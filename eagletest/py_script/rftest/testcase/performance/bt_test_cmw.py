#!/usr/bin/python
# -*- coding:utf8 -*-
from exceptions import StandardError

from rftest.rflib.wifi_lib import WIFILIB
from hal.bt_api import BTAPI
import re
from baselib.loglib.log_lib import *
from baselib.plot import *
from baselib.instrument import *
import numpy as np
import pandas as pd
import scipy
from math import *
from shutil import copy
import time
import csv
import pylab
import matplotlib.pyplot as plt
from baselib.test_channel import *
import xlrd
import sys
import random
import os
from baselib.instrument.cmw_bt import *
from baselib.instrument.spa import *
from rftest.rflib import *
from hal.common import *
from rftest.rflib.csv_report import csvreport
from rftest.rflib import rfglobal

import serial
from sys import argv
import binascii
from binascii import b2a_hex, a2b_hex
import pylink
import docx
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_TAB_ALIGNMENT,WD_TAB_LEADER #设置制表符等
from docx.shared import Inches #设置图像大小
from docx.shared import Pt #设置像素、缩进等
from docx.shared import RGBColor #设置字体颜色
from docx.shared import Length #设置宽度
import win32api
import shutil


class tester_cmw(object):
    def __init__(self, mode=0, rfport=2, cable_loss=1, bt_mode='LE1M', freq=2440, target_power=10, umargin=20, detect_mode='AUTO', burst_type='LE', phy='LE1M',repetition='SING', asyn='ON', bdaddr='050604010203', count=20, tout=10):
        if mode == 0:
            self.stx = standalone_tx('CMW')
            self.stx.reset(10)
            self.stx.rf_port(mode=mode, rfport=rfport, atten=cable_loss)
            self.stx.mode_set(mode=bt_mode)
            self.stx.analyzer_settings(enpower=target_power, umargin=umargin, freq=freq)
            self.stx.input_signal_settings(dmode=detect_mode, asyn=asyn, btype=burst_type, phy=phy, bdaddr=bdaddr)
            self.stx.measure_para(tout=tout, repetition=repetition, count=count)
            self.stx.measure_states(state=0)
        elif mode == 1:
            self.srx = standalone_rx('CMW')
            self.srx.rf_port(mode=mode, rfport=rfport, atten=cable_loss)
            self.srx.para_settings(freq=freq, level=-30)

class csv_data():
    def get_filename(self, folder, file_name, sub_folder=''):
        '''
        :folder: file store folder
        :file_name:  file name
        :sub_folder: if not need, it may be default ""
        '''
        if rfglobal.file_folder=="":
            rfdata_path = './rftest/rfdata/'
        else:
            rfdata_path = './rftest/rfdata/%s/'%rfglobal.file_folder
            if os.path.exists(rfdata_path) == False:
                os.mkdir(rfdata_path)

        data_path1 = rfdata_path+'%s/'%(folder)
        if os.path.exists(data_path1) == False:
            os.mkdir(data_path1)

        filetime = time.strftime('%Y%m%d_%H%M%S',time.localtime(time.time()));
        # mac = self.read_mac()
        mac = ''

        gen_folder = '_%s'%(filetime[0:8])
        data_path2 = data_path1 +'%s/'%(gen_folder)
        if os.path.exists(data_path2) == False:
            os.mkdir(data_path2)

        fname = '%s'%(file_name)
        outfile_name = data_path2 + fname

        if sub_folder != '':
            gen_folder = '%s_%s'%(sub_folder,filetime[0:8])
            sub_path = data_path2+'%s/'%(gen_folder)
            if os.path.exists(sub_path) == False:
                os.mkdir(sub_path)

            outfile_name = sub_path + file_name

        return outfile_name
from pygdbmi.gdbcontroller import GdbController
class gdb_load_code(object):

    def __init__(self, core='N22'):
        self.core = core
        if self.core == 'N22':
            self.gdbmin = GdbController(r'E:\project_audio\expanse\tx232_driver_src\toolchains\nds32le-elf-mculib-v5\bin\riscv32-elf-gdb', ['riscv32.elf'])
        else:
            self.gdbmin = GdbController(r'E:\project_audio\expanse\tx232_driver_src\toolchains\nds32le-elf-mculib-v5f\bin\riscv32-elf-gdb', ['riscv32.elf'])
    def close(self):
        self.gdbmin.exit()
    def open(self):
        if self.core == 'N22':
            self.gdbmin = GdbController(r'E:\project_audio\expanse\tx232_driver_src\toolchains\nds32le-elf-mculib-v5\bin\riscv32-elf-gdb', ['riscv32.elf'])
        else:
            self.gdbmin = GdbController(r'E:\project_audio\expanse\tx232_driver_src\toolchains\nds32le-elf-mculib-v5f\bin\riscv32-elf-gdb', ['riscv32.elf'])
    def load_file(self,filepath=''):
        self.close()
        self.open()
        self.gdbmin.write('monitor reset', timeout_sec=20, read_response=False)
        self.gdbmin.write('file {}'.format(filepath.replace('\\','/')),timeout_sec=5,read_response=False)
        self.gdbmin.write('load',timeout_sec=20,read_response=False)
        self.gdbmin.write('c',timeout_sec=20,read_response=False)

class gdb_server(object):
    def __init__(self,core='N22'):

        if core == 'N22':
            win32api.ShellExecute(0, 'open', 'JLinkGDBServer.exe',
                              '-select USB -device rv32 -endian little -if JTAG -speed 5000 -ir -LocalhostOnly -nologtofile -port 2335 -SWOPort 2336 -TelnetPort 2337 -jtagconf 5,1',
                              '', 1)
        else:
            win32api.ShellExecute(0, 'open', 'JLinkGDBServer.exe',
                              '-select USB -device rv32 -endian little -if JTAG -speed 5000 -ir -LocalhostOnly -nologtofile -port 2331 -SWOPort 2332 -TelnetPort 2333',
                              '', 1)


class bt_curr(object):
    def __init__(self, comport, chipv='TX232', jlink='59610138'):
        self.comport = comport
        self.chipv = chipv
        self.bt = BTAPI(self.comport,self.chipv)
        self.wifi = WIFILIB(self.comport,self.chipv)
        self.mem_ts = MEM_TS(self.comport)
        self.jlink = jlink
        self.tp = testpin(self.comport, chipv=self.chipv, jlink=self.jlink)
        self.bttest = bt_test(self.comport, chipv=self.chipv, jlink=self.jlink)

    def curr_paramp_tx_carrier(self, device_name='MY49260023', freq_list=[2402], vdd_pa_list=[3.3], cable_loss=2.6, tpmode=1, name_str='tpmode'):
        title = 'freq(MHz),vdd_pa(mv),pa_ramp,tx_carrier_pwr(dBm),curr_vddpa(mA),curr_vline3(mA)\n'
        fname = self.wifi.get_filename('ts_bt_test/', 'curr_pa_tx_carrier_{}'.format(name_str))
        fw1 = csvreport(fname, title)
        self.jlink.wrm(0xa0120080,9,9,0)
        self.jlink.wrm(0xa0120080,28,25,0)
        DM_VDDPA = dm.dm(device_name=device_name, num_of_machine=0, comm='USB')
        DM_VDDPA.device.write('VOLT {}'.format(3.3))
        DM_VDDPA.device.write('OUTPUT ON')
        # DM_VLINE3 = dm.dm(device_name='MY49260023', num_of_machine=0, comm='USB')
        myspa = Agilent()
        myspa.set_reflvl(15)
        self.jlink.wrm(0xa0120080, 9, 9, 1)
        bttest = bt_test(self.comport,chipv=self.chipv,jlink=self.jlink)
        for freq in freq_list:
            myspa.set_param(freq, 100)
            for vdd_pa in vdd_pa_list:
                DM_VDDPA.device.write('VOLT {}'.format(vdd_pa/1000.00))
                time.sleep(2)

                for ramp in range(63,0,-1):
                    bttest.tx_carrier_stop()
                    bttest.tx_carrier(freq,0xff,ramp,tpmode)
                    res = myspa.pk_search_timesleep(timesleep=1.5)
                    txp = res[0][1]+cable_loss
                    curr_vddpa = eval(DM_VDDPA.device.ask('SENS:CURR?'))*1000
                    # curr_vline3 = eval(DM_VLINE3.device.ask('MEAS:CURR?'))*1000
                    loginfo('ramp:  {}  txp:   {}  curr_vddpa:  {}  '.format(ramp,txp,curr_vddpa))
                    fw1.write_data([freq,vdd_pa,ramp,txp,curr_vddpa])
        bttest.tx_carrier_stop()

    def curr_reg_tx_carrier(self, freq=2402, vdd_pa=3.3,cable_loss=4 ):
        title = 'freq(MHz),vdd_pa(v),pa_ramp,pa_vbiasb_trim,pa_vbiasa_trim,pa_ptat_trim,pa_vblowcas_trim,pa_vbcas_trim,tx_carrier_pwr(dBm),curr_vddpa(mA),curr_vline3(mA),txp_br,' \
                'delta_f2_99,delta_f2_avg,curr_br_vddpa,curr_br_vline3,txp_edr2,DEVM_RMS,curr_edr2_vddpa,curr_edr2_vline3\n'
        fname = self.wifi.get_filename('ts_bt_test/', 'curr_reg_tx_carrier')
        fw1 = csvreport(fname, title)
        self.jlink.wrm(0xa0120080,9,9,0)
        self.jlink.wrm(0xa0120080,28,25,8)
        DM_VDDPA = dm.dm(device_name='MY50180049', num_of_machine=0, comm='USB')
        DM_VDDPA.device.write('VOLT {}'.format(vdd_pa))
        DM_VDDPA.device.write('OUTPUT ON')
        DM_VLINE3 = dm.dm(device_name='MY49260023', num_of_machine=0, comm='USB')
        myspa = Agilent()
        tester = tester_cmw(mode=0, rfport=1, cable_loss=cable_loss, bt_mode='BR', freq=freq,
                            target_power=15, burst_type='BR', asyn='OFF', bdaddr='050604010203')
        # self.BT_INIT('BR')
        bttest = bt_test(self.comport, chipv=self.chipv, jlink=self.jlink)
        bttest.cmdstop(1)
        bttest.cmdstop(0)
        bttest.tx_carrier_stop()
        myspa.set_param(freq,100)
        myspa.set_reflvl(15)
        self.jlink.wrm(0xa0120080, 9, 9, 1)

        reg_dic={
            0:  ['pa_vbiasb_trim',[28,27]],
            1: ['pa_vbiasa_trim', [26, 25]],
            2:  ['pa_ptat_trim',[7,6]],
            3: ['pa_vblowcas_trim',[4,3]],
            4: ['pa_vbcas_trim',[2,1]]
            }
        msb0 = reg_dic[0][1][0]
        lsb0 = reg_dic[0][1][1]
        init_pa_vbiasb_trim = self.jlink.rdm(0xa0421038, msb0, lsb0)
        msb1 = reg_dic[1][1][0]
        lsb1 = reg_dic[1][1][1]
        init_pa_vbiasa_trim = self.jlink.rdm(0xa0421038, msb1, lsb1)
        msb2 = reg_dic[2][1][0]
        lsb2 = reg_dic[2][1][1]
        init_pa_ptat_trim = self.jlink.rdm(0xa0421038, msb2, lsb2)
        msb3 = reg_dic[3][1][0]
        lsb3 = reg_dic[3][1][1]
        init_pa_vblowcas_trim = self.jlink.rdm(0xa0421038, msb3, lsb3)
        msb4 = reg_dic[4][1][0]
        lsb4 = reg_dic[4][1][1]
        init_pa_vbcas_trim = self.jlink.rdm(0xa0421038, msb4, lsb4)
        for ramp in [63]:

            for pa_vbiasb_trim in range(0,-1,-1):
                self.jlink.wrm(0xa0421038, msb0, lsb0, pa_vbiasb_trim)
                for pa_vbiasa_trim in range(0,-1,-1):
                    self.jlink.wrm(0xa0421038, msb1, lsb1, pa_vbiasa_trim)
                    for pa_ptat_trim in range(3,-1,-1):
                        self.jlink.wrm(0xa0421038, msb2, lsb2, pa_ptat_trim)
                        for pa_vblowcas_trim in range(3,-1,-1):
                            self.jlink.wrm(0xa0421038, msb3, lsb3, pa_vblowcas_trim)
                            for pa_vbcas_trim in range(3,-1,-1):
                                self.jlink.wrm(0xa0421038, msb4, lsb4, pa_vbcas_trim)

                                bttest.tx_carrier(freq, 0xff, ramp)
                                res = myspa.pk_search_avg_timesleep(timesleep=1.5)
                                tx_tone_pwr = res[0][1]
                                curr_tone_vddpa = eval(DM_VDDPA.device.ask('SENS:CURR?'))*1000
                                time.sleep(0.5)
                                curr_tone_vline3 = eval(DM_VLINE3.device.ask('MEAS:CURR?'))*1000
                                bttest.tx_carrier_stop()

                                bttest.BR_TX(chan=0,len=27,ptype=2,rate=1)
                                tester.stx.mode_set('BR')
                                tester.stx.input_signal_settings(btype='BR',asyn='OFF')
                                res = tester.stx.get_modulation_measure_res()
                                curr_br_vddpa = eval(DM_VDDPA.device.ask('SENS:CURR?'))*1000
                                time.sleep(0.5)
                                curr_br_vline3 = eval(DM_VLINE3.device.ask('MEAS:CURR?'))*1000
                                logdebug('{}'.format(res))

                                delta_f2_99 = eval(res[2]) / 1000.00
                                freq_accuracy = eval(res[3]) / 1000.00
                                freq_drift = eval(res[4]) / 1000.00
                                drift_rate = eval(res[5]) / 1000.00
                                delta_f2_avg = eval(res[9]) / 1000.00
                                delta_f2_min = eval(res[10]) / 1000.00
                                delta_f2_max = eval(res[11]) / 1000.00
                                txp_br = eval(res[12])
                                bttest.cmdstop(1)
                                bttest.BR_TX(chan=0, len=54, ptype=7, rate=2)
                                tester.stx.mode_set('EDR')
                                tester.stx.input_signal_settings(btype='EDR',asyn='OFF')
                                res = tester.stx.get_modulation_measure_res()
                                curr_edr2_vddpa = eval(DM_VDDPA.device.ask('SENS:CURR?'))*1000
                                time.sleep(0.5)
                                curr_edr2_vline3 = eval(DM_VLINE3.device.ask('MEAS:CURR?'))*1000
                                logdebug('{}'.format(res))

                                res = [eval(i) for i in res[0:]]

                                wi = res[2] / 1000.00
                                w0_wi = res[3] / 1000.00
                                w0_max = res[4] / 1000.00
                                DEVM_RMS = res[5]
                                DEVM_peak = res[6]
                                DEVM_P99 = res[7]
                                txp_edr2 = res[8]
                                bttest.cmdstop(1)


                                loginfo('ramp:  {}  txp:   {}  curr_vddpa:  {}  curr_vline3:    {}'.format(ramp, tx_tone_pwr, curr_tone_vddpa, curr_tone_vline3))
                                fw1.write_data([freq,vdd_pa,ramp,pa_vbiasb_trim,pa_vbiasa_trim,pa_ptat_trim,pa_vblowcas_trim,pa_vbcas_trim,tx_tone_pwr, curr_tone_vddpa,
                                                curr_tone_vline3,txp_br,delta_f2_99,delta_f2_avg,curr_br_vddpa,curr_br_vline3,txp_edr2,DEVM_RMS,curr_edr2_vddpa,curr_edr2_vline3])

        self.jlink.wrm(0xa0421038, msb0, lsb0, init_pa_vbiasb_trim)
        self.jlink.wrm(0xa0421038, msb1, lsb1, init_pa_vbiasa_trim)
        self.jlink.wrm(0xa0421038, msb2, lsb2, init_pa_ptat_trim)
        self.jlink.wrm(0xa0421038, msb3, lsb3, init_pa_vblowcas_trim)
        self.jlink.wrm(0xa0421038, msb4, lsb4, init_pa_vbcas_trim)

    def curr_lna(self,freq=2402,device_name='MY50180049'):
        title = 'freq(MHz),lna_itrim,lna_hgain_code,curr_vline1(mA)\n'
        fname = self.wifi.get_filename('ts_bt_test/', 'curr_lna')
        fw1 = csvreport(fname, title)
        DM_VLINE1 = dm.dm(device_name=device_name, num_of_machine=0, comm='USB')
        self.jlink.wrm(0xa04210a0,0,0,1)   ##agc disable
        self.jlink.wrm(0xa04210a0, 20, 17, 0xc)  ##AGC initial index
        self.jlink.wr(0xa0422014,0x3)  ##manual rxon
        for lna_itrim in range(3,-1,-1):
            self.jlink.wrm(0xa0421034, 8, 7, lna_itrim)
            for lna_hgain_code in range(6):
                self.jlink.wrm(0xa04210b4, 25, 25, 1)  ##lna_hgain_pwup
                self.jlink.wrm(0xa04210b4, 24, 19, 0x3f>>lna_hgain_code)
                curr_vline1 = eval(DM_VLINE1.device.ask('MEAS:CURR?')) * 1000
                loginfo('lna_itrim:     {}   lna_hgain_code:    {}      curr_vline1:    {}'.format(lna_itrim,0x3f>>lna_hgain_code,curr_vline1))
                fw1.write_data([freq, lna_itrim,0x3f>>lna_hgain_code,curr_vline1])
        fw1.write_string('freq(MHz),lna_itrim,lna_lgain_code,curr_vline1(mA)\n')
        for lna_itrim in range(3,-1,-1):
            self.jlink.wrm(0xa0421034, 8, 7, lna_itrim)
            for lna_lgain_code in range(3,-1,-1):
                self.jlink.wrm(0xa04210b4, 25, 25, 0)  ##lna_hgain_pwup
                self.jlink.wrm(0xa04210b4, 18, 17, lna_lgain_code)
                curr_vline1 = eval(DM_VLINE1.device.ask('MEAS:CURR?')) * 1000
                loginfo('lna_itrim:     {}   lna_lgain_code:    {}      curr_vline1:    {}'.format(lna_itrim,lna_lgain_code,curr_vline1))
                fw1.write_data([freq, lna_itrim,lna_lgain_code,curr_vline1])
        self.jlink.wrm(0xa04210b4,25,17,0x1ff)
        self.jlink.wrm(0xa04210a0, 0, 0, 0)  ##agc enable
    def xo_manual(self,en=1):
        if en!=0:
            self.jlink.wrm(0xa01200cc, 29, 29, 1)  ##bbpll_bg_pup enable
            self.jlink.wrm(0xa01200cc, 30, 30, 1)  ##bg_pup_ibg_xo
            self.jlink.wrm(0xa01200cc, 26, 25, 3)  ##manual xo pup enable
            self.jlink.wrm(0xa01200cc, 27, 27, 0)  ##manual xo pup
        else:
            self.jlink.wrm(0xa01200cc, 27, 27, 1)  ##auto xo

    def curr_xo(self, device_name='MY49260023'):
        '''
        测试的是xo与bbpll_bg合在一起的电流
        '''
        title = 'xo_ldo_trim,xo_isel,curr_vline3(mA)\n'
        fname = self.wifi.get_filename('ts_bt_test/', 'curr_xo')
        fw1 = csvreport(fname, title)
        DM_VLINE3 = dm.dm(device_name=device_name, num_of_machine=0, comm='USB')
        self.jlink.wr(0xa0422014, 0x0)
        self.jlink.wrm(0xa00000a8, 1, 0,1)  ##mclk0_sel rc48m
        self.tp.bbpll_en(0)
        self.xo_manual(1)
        init_xo_ldo = self.jlink.rdm(0xa01200cc, 5, 4)
        init_xo_isel = self.jlink.rdm(0xa01200cc, 10, 7)
        for xo_ldo_trim in range(3,-1,-1):
            self.jlink.wrm(0xa01200cc, 5, 4, xo_ldo_trim)
            for xo_isel in range(0xf,-1,-1):
                self.jlink.wrm(0xa01200cc, 10, 7, xo_isel)
                time.sleep(1)
                curr_vline3 = eval(DM_VLINE3.device.ask('MEAS:CURR?')) * 1000
                loginfo('xo_ldo_trim:   {}  xo_isel:    {}  curr_vline3:    {}'.format(xo_ldo_trim,xo_isel,curr_vline3))
                fw1.write_data([xo_ldo_trim,xo_isel,curr_vline3])
        self.jlink.wrm(0xa01200cc, 5, 4, init_xo_ldo)
        self.jlink.wrm(0xa01200cc, 10, 7, init_xo_isel)

    def vline1_init(self):
        self.jlink.wrm(0xa0421004, 9, 9, 0)  ##lna_pup
        self.jlink.wrm(0xa0421004, 12, 12, 0)  ##cbpf1_pup
        self.jlink.wrm(0xa0421000, 17, 17, 0)  ##ldo_trxhf_pup
        self.jlink.wrm(0xa0421000, 11, 11, 0)  ##bg_pup
        self.jlink.wrm(0xa0421004, 17, 17, 0)  ##txmix_pup
        self.jlink.wrm(0xa0421004, 16, 16, 0)  ##txpa_pup

    def curr_vline1(self, device_name='MY50180049'):
        title = 'mode,curr_init(mA),curr_LNA+mix_off(mA),curr_cbpf_off(mA),curr_ldo_trxhf_off(mA),curr_bg_off(mA)\n'
        fname = self.wifi.get_filename('ts_bt_test/', 'curr_vline1')
        fw1 = csvreport(fname, title)
        DM_VLINE = dm.dm(device_name=device_name, num_of_machine=0, comm='USB')
        self.jlink.wrm(0xa00000a8, 1, 0,1)  ##mclk0_sel rc48m
        self.jlink.wrm(0xa04210a0,0,0,1)   ##agc disable
        self.jlink.wrm(0xa04210a0, 20, 17, 0xc)  ##AGC initial index

        for lna_itrim in range(4):
            self.jlink.wrm(0xa0421034, 8, 7, lna_itrim)
            for cbpf_bias_trim in range(8):
                self.jlink.wrm(0xa0421034, 18, 16, cbpf_bias_trim)
                for ldo_trxhf_trim in range(4):
                    self.jlink.wrm(0xa0421028, 7, 6, ldo_trxhf_trim)

                    self.vline1_init()
                    self.jlink.wr(0xa0422014, 0x3)  ##manual rxon
                    curr_init = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
                    self.jlink.wrm(0xa0421004, 9, 9, 1)  ##lna_pup disable
                    curr_lna_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
                    self.jlink.wrm(0xa0421004, 12, 12, 1)  ##cbpf1_pupp disable
                    curr_cbpf_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
                    self.jlink.wrm(0xa0421000, 17, 17, 1)  ##ldo_trxhf_pup disable
                    curr_trxhf_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
                    self.jlink.wrm(0xa0421000, 11, 11, 1)  ##bg_pup disable
                    curr_bg_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
                    loginfo('rxon curr:  {}      {}  {}  {}  {}  {}'.format('rxon', curr_init, curr_lna_off, curr_cbpf_off, curr_trxhf_off, curr_bg_off))
                    fw1.write_data(['rxon', curr_init, curr_lna_off, curr_cbpf_off, curr_trxhf_off, curr_bg_off])


        fw1.write_string('mode,curr_init(mA),curr_txmix_off(mA),curr_txpa_off(mA),curr_ldo_trxhf_off(mA),curr_bg_off(mA)\n')
        self.vline1_init()
        self.jlink.wr(0xa0422014, 0xe)  ##manual rxon
        curr_init = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
        self.jlink.wrm(0xa0421004, 17, 17, 0)  ##txmix_pupp disable
        curr_txmix_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
        self.jlink.wrm(0xa0421004, 16, 16, 0)  ##txpa_pup disable
        curr_txpa_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
        self.jlink.wrm(0xa0421000, 17, 17, 1)  ##ldo_trxhf_pup disable
        curr_trxhf_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
        self.jlink.wrm(0xa0421000, 11, 11, 1)  ##bg_pup disable
        curr_bg_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
        loginfo('txon curr:  {}      {}  {}  {}  {}  {}'.format('rxon', curr_init, curr_txmix_off, curr_txpa_off, curr_trxhf_off, curr_bg_off))
        fw1.write_data(['txon', curr_init, curr_txmix_off, curr_txpa_off, curr_trxhf_off, curr_bg_off])

        self.vline1_init()

    def vline2_init(self):
        self.jlink.wrm(0xa0421020, 25, 25, 1)  ##iq mode
        self.jlink.wrm(0xa0421004, 3, 2, 0)  ##lo_pup_vlo_txfskdrv and lo_pup_vlo_txfsk
        self.jlink.wrm(0xa0421004, 1, 0, 0)  ##lo_pup_vlo_txdrv and lo_pup_vlo_tx
        self.jlink.wrm(0xa0421000, 31, 30, 0)  ##lo_pup_vlo_rxdrv and lo_pup_vlo_rx
        self.jlink.wrm(0xa0421000, 28, 28, 0)  ##vcodiv2_pup
        self.jlink.wrm(0xa0421000, 27, 27, 0)  ##vco_pup
        self.jlink.wrm(0xa0421000, 23, 23, 0)  ##ldo_vco_pup

    def curr_vline2(self, device_name='MY50180049'):
        title = 'mode,ldo_vco_trim,curr_init(mA),curr_vlo_rx_off(mA),curr_vcodiv2_off(mA),curr_vco_off(mA),curr_ldo_vco_off(mA)\n'
        fname = self.wifi.get_filename('ts_bt_test/', 'curr_vline2')
        fw1 = csvreport(fname, title)
        DM_VLINE2 = dm.dm(device_name=device_name, num_of_machine=0, comm='USB')

        self.jlink.wrm(0xa00000a8, 1, 0,1)  ##mclk0_sel rc48m
        for ldo_vco_trim in range(0,4):
            self.jlink.wrm(0xa0421028, 15, 14, ldo_vco_trim)
            self.vline2_init()
            self.jlink.wr(0xa0422014, 0x3)  ##manual rxon
            curr_init = eval(DM_VLINE2.device.ask('MEAS:CURR?')) * 1000
            self.jlink.wrm(0xa0421000, 31, 30, 3)  ##lo_pup_vlo_rxdrv and lo_pup_vlo_rx disable
            curr_vlo_rx_off = eval(DM_VLINE2.device.ask('MEAS:CURR?')) * 1000
            self.jlink.wrm(0xa0421000, 28, 28, 1)  ##vcodiv2_pup disable
            curr_vcodiv2_off = eval(DM_VLINE2.device.ask('MEAS:CURR?')) * 1000
            self.jlink.wrm(0xa0421000, 27, 27, 1)  ##vco_pup disable
            curr_vco_off = eval(DM_VLINE2.device.ask('MEAS:CURR?')) * 1000
            self.jlink.wrm(0xa0421000, 23, 23, 1)  ##ldo_vco_pup disable
            curr_ldo_vco_off = eval(DM_VLINE2.device.ask('MEAS:CURR?')) * 1000
            loginfo('{}     {}  {}  {}  {}  {}  {}'.format('rxon',ldo_vco_trim,curr_init,curr_vlo_rx_off,curr_vcodiv2_off,curr_vco_off,curr_ldo_vco_off))
            fw1.write_data(['rxon',ldo_vco_trim,curr_init,curr_vlo_rx_off,curr_vcodiv2_off,curr_vco_off,curr_ldo_vco_off])

        fw1.write_string('mode,ldo_vco_trim,curr_init(mA),curr_vlo_tx_off(mA),curr_vcodiv2_off(mA),curr_vco_off(mA),curr_ldo_vco_off(mA)\n')
        for ldo_vco_trim in range(0, 4):
            self.jlink.wrm(0xa0421028, 15, 14, ldo_vco_trim)
            self.vline2_init()
            self.jlink.wr(0xa0422014, 0xe)  ##manual txon
            curr_init = eval(DM_VLINE2.device.ask('MEAS:CURR?')) * 1000
            self.jlink.wrm(0xa0421004, 1, 0, 3)  ##lo_pup_vlo_txdrv and lo_pup_vlo_tx disable
            curr_vlo_tx_off = eval(DM_VLINE2.device.ask('MEAS:CURR?')) * 1000
            self.jlink.wrm(0xa0421000, 28, 28, 1)  ##vcodiv2_pup disable
            curr_vcodiv2_off = eval(DM_VLINE2.device.ask('MEAS:CURR?')) * 1000
            self.jlink.wrm(0xa0421000, 27, 27, 1)  ##vco_pup disable
            curr_vco_off = eval(DM_VLINE2.device.ask('MEAS:CURR?')) * 1000
            self.jlink.wrm(0xa0421000, 23, 23, 1)  ##ldo_vco_pup disable
            curr_ldo_vco_off = eval(DM_VLINE2.device.ask('MEAS:CURR?')) * 1000
            loginfo('{}     {}  {}  {}  {}  {}  {}'.format('rxon', ldo_vco_trim, curr_init, curr_vlo_tx_off, curr_vcodiv2_off, curr_vco_off, curr_ldo_vco_off))
            fw1.write_data(['txon', ldo_vco_trim, curr_init, curr_vlo_tx_off, curr_vcodiv2_off, curr_vco_off, curr_ldo_vco_off])

        fw1.write_string('mode,ldo_vco_trim,curr_init(mA),curr_vlo_txgfsk_off(mA),curr_vcodiv2_off(mA),curr_vco_off(mA),curr_ldo_vco_off(mA)\n')
        for ldo_vco_trim in range(0, 4):
            self.vline2_init()
            self.jlink.wr(0xa0422014, 0x0)  ##auto txrx
            self.jlink.wrm(0xa0421020, 25, 25, 0)  ##tp mode
            self.jlink.wr(0xa0422014, 0xe)  ##manual txon
            curr_init = eval(DM_VLINE2.device.ask('MEAS:CURR?')) * 1000
            self.jlink.wrm(0xa0421004, 3, 2, 3)  ##lo_pup_vlo_txdrv and lo_pup_vlo_tx disable
            curr_vlo_txgfsk_off = eval(DM_VLINE2.device.ask('MEAS:CURR?')) * 1000
            self.jlink.wrm(0xa0421000, 28, 28, 1)  ##vcodiv2_pup disable
            curr_vcodiv2_off = eval(DM_VLINE2.device.ask('MEAS:CURR?')) * 1000
            self.jlink.wrm(0xa0421000, 27, 27, 1)  ##vco_pup disable
            curr_vco_off = eval(DM_VLINE2.device.ask('MEAS:CURR?')) * 1000
            self.jlink.wrm(0xa0421000, 23, 23, 1)  ##ldo_vco_pup disable
            curr_ldo_vco_off = eval(DM_VLINE2.device.ask('MEAS:CURR?')) * 1000
            loginfo('{}     {}  {}  {}  {}  {}  {}'.format('rxon', ldo_vco_trim, curr_init, curr_vlo_txgfsk_off, curr_vcodiv2_off, curr_vco_off, curr_ldo_vco_off))
            fw1.write_data(['txon_tp', ldo_vco_trim, curr_init, curr_vlo_txgfsk_off, curr_vcodiv2_off, curr_vco_off, curr_ldo_vco_off])
        self.vline2_init()

    def vline3_init(self):
        self.jlink.wrm(0xa0421004, 15, 15, 0)  ##adc_pup
        self.jlink.wrm(0xa0421004, 6, 6, 0)  ##pfdcp_pup
        self.jlink.wrm(0xa0421004, 5, 5, 0)  ##divn_pup
        self.jlink.wrm(0xa0421000, 21, 21, 0)  ##ldo_pll_pup
        self.jlink.wrm(0xa0421000, 19, 19, 0)  ##ldo_trxlf_pup
        self.jlink.wrm(0xa0421000, 16, 16, 0)  ##ldo_lv_pup
        self.tp.bbpll_en(1)
        self.jlink.wrm(0xa01200cc,27,27,1)     ##xo
        self.jlink.wrm(0xa0421004, 18, 18, 0)  ##txlpf_pup
        self.jlink.wrm(0xa0421004, 20, 20, 0)  ##txdac_pup
    def curr_vline3(self, device_name='MY50180049'):
        title = 'mode,curr_init(mA),curr_adc_off(mA),curr_pfdcp_off(mA),curr_divn_off(mA),curr_ldo_pll_off(mA),curr_ldo_trxlf_off(mA),curr_ldo_lv_off(mA),curr_bbpll_off(mA),' \
                'curr_xo_off(mA)\n'
        fname = self.wifi.get_filename('ts_bt_test/', 'curr_vline3')
        fw1 = csvreport(fname, title)
        DM_VLINE = dm.dm(device_name=device_name, num_of_machine=0, comm='USB')
        self.jlink.wrm(0xa00000a8, 1, 0,1)  ##mclk0_sel rc48m

        self.vline3_init()
        self.jlink.wr(0xa0422014, 0x3)  ##manual rxon
        curr_init = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
        self.jlink.wrm(0xa0421004, 15, 15, 1)  ##adc_pup disable
        curr_adc_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
        self.jlink.wrm(0xa0421004, 6, 6, 1)  ##pfdcp_pup disable
        curr_pfdcp_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
        self.jlink.wrm(0xa0421004, 5, 5, 1)  ##divn_pup disable
        curr_divn_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
        self.jlink.wrm(0xa0421000, 21, 21, 1)  ##ldo_pll_pup disable
        curr_ldo_pll_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
        self.jlink.wrm(0xa0421000, 19, 19, 1)  ##ldo_trxlf_pup disable
        curr_ldo_trxlf_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
        self.jlink.wrm(0xa0421000, 16, 16, 1)  ##ldo_lv_pup disable
        curr_ldo_lv_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
        self.jlink.wrm(0xa0000144, 23, 22, 0)  ##bbpll disable
        curr_bbpll_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
        self.jlink.wrm(0xa01200cc, 27, 27, 0)  ##xo disable
        self.jlink.wrm(0xa01200cc, 29, 29, 0)  ##bbpll_bg disable
        curr_xo_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
        self.jlink.wrm(0xa01200cc, 29, 29, 1)  ##bbpll_bg enable
        self.jlink.wrm(0xa01200cc, 27, 27, 1)  ##xo enable

        loginfo('rxon curr:  {}      {}  {}  {}  {}  {}     {}  {}  {}  {}'.format('rxon', curr_init, curr_adc_off, curr_pfdcp_off, curr_divn_off, curr_ldo_pll_off,
                                                                                  curr_ldo_trxlf_off,
                                                                      curr_ldo_lv_off,curr_bbpll_off,curr_xo_off))
        fw1.write_data(['rxon', curr_init, curr_adc_off, curr_pfdcp_off, curr_divn_off, curr_ldo_pll_off,curr_ldo_trxlf_off,curr_ldo_lv_off,curr_bbpll_off,curr_xo_off])

        fw1.write_string('mode,curr_init(mA),curr_txlpf_off(mA),curr_txdac_off(mA),curr_pfdcp_off(mA),curr_divn_off(mA),curr_ldo_pll_off(mA),curr_ldo_trxlf_off(mA),'
                         'curr_ldo_lv_off(mA),curr_bbpll_off(mA),curr_xo_off(mA)\n')
        self.vline3_init()
        self.jlink.wr(0xa0422014, 0xe)  ##manual rxon
        curr_init = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
        self.jlink.wrm(0xa0421004, 18, 18, 1)  ##txlpf_pup disable
        curr_txlpf_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
        self.jlink.wrm(0xa0421004, 20, 20, 1)  ##txdac_pup disable
        curr_txdac_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
        self.jlink.wrm(0xa0421004, 6, 6, 1)  ##pfdcp_pup disable
        curr_pfdcp_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
        self.jlink.wrm(0xa0421004, 5, 5, 1)  ##divn_pup disable
        curr_divn_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
        self.jlink.wrm(0xa0421000, 21, 21, 1)  ##ldo_pll_pup disable
        curr_ldo_pll_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
        self.jlink.wrm(0xa0421000, 19, 19, 1)  ##ldo_trxlf_pup disable
        curr_ldo_trxlf_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
        self.jlink.wrm(0xa0421000, 16, 16, 1)  ##ldo_lv_pup disable
        curr_ldo_lv_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
        self.jlink.wrm(0xa0000144, 23, 22, 0)  ##bbpll disable
        curr_bbpll_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
        self.jlink.wrm(0xa01200cc, 27, 27, 0)  ##xo disable
        curr_xo_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
        self.jlink.wrm(0xa01200cc, 27, 27, 1)  ##xo enable

        loginfo('txon curr:  {}      {}  {}  {}  {}  {}     {}  {}  {}  {}'.format('txon', curr_init, curr_txlpf_off, curr_txdac_off, curr_pfdcp_off, curr_divn_off,
                                                                                   curr_ldo_pll_off,
                                                                                  curr_ldo_trxlf_off,
                                                                      curr_ldo_lv_off,curr_bbpll_off,curr_xo_off))
        fw1.write_data(['txon', curr_init, curr_txlpf_off, curr_txdac_off, curr_pfdcp_off, curr_divn_off, curr_ldo_pll_off,curr_ldo_trxlf_off,curr_ldo_lv_off,curr_bbpll_off,\
                                                                          curr_xo_off])
        self.vline3_init()


    def curr_bbpll(self, device_name='MY50180049'):
        title = 'bbpll_ldo_trim,curr_init(mA),curr_bbpll_off(mA),curr_bbpll_ldo_off(mA),curr_xo_off(mA),curr_bg_off(mA)\n'
        fname = self.wifi.get_filename('ts_bt_test/', 'curr_bbpll')
        fw1 = csvreport(fname, title)
        DM_VLINE = dm.dm(device_name=device_name, num_of_machine=0, comm='USB')
        self.jlink.wr(0xa0422014, 0x0)
        self.jlink.wrm(0xa00000a8, 1, 0, 1)  ##mclk0_sel rc48m
        self.xo_manual(1)
        bbpll_ldo_trim_init = self.jlink.rdm(0xa0000144, 14, 12)
        for bbpll_ldo_trim in range(8):
            self.jlink.wrm(0xa0000144, 14, 12, bbpll_ldo_trim)
            self.xo_manual(1)
            self.tp.bbpll_en(1)
            curr_init = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
            self.jlink.wrm(0xa0000144, 22, 22, 0)  ##bbpll_pup disable
            curr_bbpll_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
            self.jlink.wrm(0xa0000144, 23, 23, 0)  ##bbpll_ldo_pup disable
            self.jlink.wrm(0xa0000144, 25, 25, 0)  ##bbpll_bg_pup_ibg_bbpll
            curr_bbpll_ldo_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
            self.jlink.wrm(0xa01200cc, 26, 25, 0)  ##xo_pup disable
            curr_xo_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
            self.jlink.wrm(0xa01200cc, 30, 29, 0)  ##bg_pup disable
            curr_bg_off = eval(DM_VLINE.device.ask('MEAS:CURR?')) * 1000
            loginfo('bbpll curr:  {}      {}  {}  {}  {} {}'.format( bbpll_ldo_trim,curr_init, curr_bbpll_off, curr_bbpll_ldo_off,curr_xo_off,curr_bg_off))
            fw1.write_data([bbpll_ldo_trim,curr_init, curr_bbpll_off, curr_bbpll_ldo_off,curr_xo_off,curr_bg_off])
        self.xo_manual(1)
        self.tp.bbpll_en(1)
        self.jlink.wrm(0xa0000144, 14, 12, bbpll_ldo_trim_init)

class testpin(object):
    def __init__(self, comport, chipv='tx232', jlink_en=1, jlink='59610138'):
        self.comport = comport
        self.chipv = chipv
        self.mem_ts = MEM_TS(self.comport)
        if jlink_en!=0:
            self.jlink = jlink
            self.jlink.wrm(0xa0200258, 10, 6, 6)
        self.bttest = bt_test(self.comport, chipv=self.chipv, jlink_en=jlink_en,jlink=jlink)

    def bbpll_en(self, en=1):

        if self.chipv == 'TX232_MPW3':
            if en !=0:
                self.jlink.wrm(0xa01200cc, 29, 29, 1)  ##bbpll_bg_pup enbale
                self.jlink.wrm(0xa0000144, 25, 25, 1)  ##bbpll_bg_pup_ibg_bbpll
                # self.jlink.wrm(0xa0000144, 29, 29, 1)  ##bbpll_bg_pup_ibg_tbuf
                self.jlink.wrm(0xa0000144, 20, 20, 1)  ##bbpll_bg_fc on
                self.jlink.wrm(0xa01200cc, 22, 22, 1)  ##rf_aon_xo_en_clk_ana
                self.jlink.wrm(0xa01200cc, 20, 20, 1)  ##rf_aon_xo_en_clk_ref
                self.jlink.wrm(0xa0000144, 23, 23, 1)  ##bbpll_ldo_pup
                self.jlink.wrm(0xa0000144, 22, 22, 1)  ##bbpll_pup
                self.jlink.wrm(0xa0000144, 21, 21, 0)  ##bbpll_ld_pd
                self.jlink.wrm(0xa0000144, 18, 18, 1)  ##bbpll_ldo_fc on
                self.jlink.wrm(0xa0000144, 17, 17, 1)  ##bbpll_vco_fast_en on

                ##wait 40us and end bbpll fast charge
                self.jlink.wrm(0xa0000144, 20, 20, 0)  ##bbpll_bg_fc off
                self.jlink.wrm(0xa0000144, 18, 18, 0)  ##bbpll_ldo_fc off
                self.jlink.wrm(0xa0000144, 17, 17, 0)  ##bbpll_vco_fast_en off
                self.jlink.wrm(0xa0000144, 16, 16, 1)  ##bbpll_ck_dig_en
                self.jlink.wrm(0xa0000144, 15, 15, 1)  ##bbpll_ck_codec_en
            else:
                self.jlink.wrm(0xa0000144, 23, 23, 0)  ##bbpll_ldo_pup
                self.jlink.wrm(0xa0000144, 22, 22, 0)  ##bbpll_pup
                # self.jlink.wrm(0xa01200cc, 29, 29, 0)  ##bbpll_bg_pup disbale
                self.jlink.wrm(0xa0000144, 25, 25, 0)  ##bbpll_bg_pup_ibg_bbpll
        else:
            if en!=0:
                self.jlink.wrm(0xa01200a0, 25, 25, 1)  ##bbpll_bg_pup_ibg
                self.jlink.wrm(0xa01200a0, 29, 29, 1)  ##bbpll_bg_pup_ibg_tbuf
                self.jlink.wrm(0xa01200a0, 20, 20, 1)  ##bbpll_bg_fc on
                self.jlink.wrm(0xa01200cc, 22, 22, 1)  ##rf_aon_xo_en_clk_ana
                self.jlink.wrm(0xa01200cc, 20, 20, 1)  ##rf_aon_xo_en_clk_ref
                self.jlink.wrm(0xa01200a0, 23, 23, 1)  ##bbpll_ldo_pup
                self.jlink.wrm(0xa01200a0, 22, 22, 1)  ##bbpll_pup
                self.jlink.wrm(0xa01200a0, 21, 21, 0)  ##bbpll_ld_pd
                self.jlink.wrm(0xa01200a0, 18, 18, 1)  ##bbpll_ldo_fc on
                self.jlink.wrm(0xa01200a0, 17, 17, 1)  ##bbpll_vco_fast_en on

                ##wait 40us and end bbpll fast charge
                self.jlink.wrm(0xa01200a0, 20, 20, 0)  ##bbpll_bg_fc off
                self.jlink.wrm(0xa01200a0, 18, 18, 0)  ##bbpll_ldo_fc off
                self.jlink.wrm(0xa01200a0, 17, 17, 0)  ##bbpll_vco_fast_en off
                self.jlink.wrm(0xa01200a0, 16, 16, 1)  ##bbpll_ck_dig_en
                self.jlink.wrm(0xa01200a0, 15, 15, 1)  ##bbpll_ck_codec_en
            else:
                self.jlink.wrm(0xa01200a0, 23, 23, 0)  ##bbpll_ldo_pup
                self.jlink.wrm(0xa01200a0, 22, 22, 0)  ##bbpll_pup

    def frac_pll(self,en=1):
        if en!=0:
            self.jlink.wrm(0xa000014c, 1, 1, 0)  ##frac_pll_resetn
            self.jlink.wrm(0xa000014c, 0, 0, 1) ##frac_pll enable
            self.jlink.wrm(0xa000014c, 1, 1, 1) ##frac_pll_resetn
        else:
            self.jlink.wrm(0xa000014c, 1, 1, 0)  ##frac_pll_resetn
            self.jlink.wrm(0xa000014c, 0, 0, 0) ##frac_pll enable


    def mclk1_test(self,mclk1_sel=1):
        '''
        0xa0200258[10:6]不能等于0，否则PA0上拉不出来时钟
        mclk1_sel:
        00: clk_pll0_out
        01: clk_pll1_out
        10: clk_hfrc_out
        '''
        self.jlink.wrm(0xa0200258, 10, 6, 6)
        self.jlink.wrm(0xa0200240,2,0,3)    ##PA0 func set testpin_clk
        self.jlink.wrm(0xa0200258, 5, 1, 21)    ##PA0 testpin test_clk_sel  codec_adc5_clk
        self.jlink.wrm(0xa0000044, 14, 13, 3,lsb_dis=1)     ##mclk1 enable and codec_adc5_clk enable
        self.jlink.wrm(0xa0000050,1,0,mclk1_sel)

    def mclk0_test(self,mclk0_sel=1):
        '''
        mclk0_sel:
        00: clk_pll_out,divided pll clk
        01: clk_hfrc_out(48m)
        10: clk_hfxo_out(24m)
        11: doubler clock out
        '''
        self.jlink.wrm(0xa0200258, 10, 6, 6)
        self.jlink.wrm(0xa0200240,2,0,3)
        self.jlink.wrm(0xa0200258, 5, 1, 7)
        self.jlink.wrm(0xa01200cc, 31, 31, 1)   ##doubler clock enable
        self.bbpll_en(1)
        self.jlink.wrm(0xa00000a8, 1, 0, mclk0_sel)

    def sys_clk_set(self, clk=170):
        pllput_div_init = self.jlink.rdm(0xa00000a4, 31, 24)
        clkbus_div_init = self.jlink.rdm(0xa0000054, 31, 0)
        clkm4_div_init = self.jlink.rdm(0xa0000058, 31, 0)
        loginfo('pllput_div_init:   {}  clkbus_div_init:    {}  clkm4_div_init: {}'.format(hex(pllput_div_init),hex(clkbus_div_init),hex(clkm4_div_init)))

        self.jlink.wrm(0xa00000a8, 1, 0, 1)     ##mclk0 set rc48m
        self.frac_pll(1)    ##fracpll enable
        self.frac_pll_freq_set(clk)
        self.jlink.wrm(0xa00000a4, 31, 24, 0)   ##pllout div set 0
        self.jlink.wrm(0xa0000054, 31, 0, 0)  ##mclk0_g div/ahb/apb2/apb0 div set 0
        self.jlink.wrm(0xa0000058, 31, 0, 0)  ##m4_ap_system_clk/m4_system_clk div set 0

        time.sleep(0.5)
        self.jlink.wrm(0xa00000a8, 3, 3, 1)  ##mclk0 pll sel fracpll
        self.jlink.wrm(0xa00000a8, 1, 0, 0)  ##mclk0 set pll


    def frac_pll_freq_set(self, freq=150):
        '''
        freq
        unit:   MHz
        range:  144--304    MHz
        accuracy:   62.5    kHz
        '''
        self.jlink.wrm(0xa000014c, 1, 1, 0)  ##frac_pll_resetn disable
        self.jlink.wrm(0xa000014c, 2, 2, 1)  ##frac divider enable
        fn = int((freq/8)-18)
        divf =int(((freq%8)/8.000)*(2**16))
        loginfo('{}     {}'.format(fn,divf))
        self.jlink.wrm(0xa000014c, 7, 3, fn)  ##FN select
        self.jlink.wrm(0xa000014c, 31, 16, divf)  ##FD select
        self.jlink.wrm(0xa000014c, 1, 1, 1)  ##frac_pll_resetn enable

    def frac_divf_set(self, divf=0x0):
        self.jlink.wrm(0xa000014c, 1, 1, 0)  ##frac_pll_resetn disable
        self.jlink.wrm(0xa000014c, 31, 16, divf)  ##FD select
        self.jlink.wrm(0xa000014c, 1, 1, 1)  ##frac_pll_resetn enable

    def frac_pll_looptest(self, freq_list=[150,300],loop_count=1000):
        '''
        frac_clk拉到mclk1上，mclk1拉到PA0上，PA0信号耦合到RF trace来测试
        freq为非8的整数倍数
        '''
        title = 'freq_center(MHz),freq_divf-1(MHz),freq_center(MHz),freq_divf+1(MHz),freq_center_ld,freq_divf-1_ld,freq_center1_ld,freq_divf+1_ld,txp1,' \
                'txp2,txp3,txp4,divf,step\n'
        fname = self.wifi.get_filename('ts_bt_test/', 'frac_pll_looptest')
        fw1 = csvreport(fname, title)
        self.jlink.wrm(0xa0200258, 10, 6, 6)
        self.jlink.wrm(0xa0200240,2,0,3)
        self.jlink.wrm(0xa0200258, 5, 1, 8)     ##PA0 testpin test_clk_sel mclk1
        self.jlink.wrm(0xa0000044, 13, 13, 1,lsb_dis=1)
        self.jlink.wrm(0xa0000050,1,0,1)
        self.frac_pll(en=1)

        myspa = Agilent()
        myspa.set_reflvl(-10)
        myspa.set_rb_HZ(100)
        myspa.set_vb_HZ(100)
        myspa.set_span(1)

        import matplotlib.pyplot as plt
        import numpy as np
        import time
        from math import *



        for freq in freq_list:
            plt.ion()  # 开启interactive mode 成功的关键函数
            plt.figure(freq)
            per_list = []
            t = 0
            t_list = []
            self.frac_pll_freq_set(freq=freq)

            divf = int(((freq % 8) / 8.000) * (2 ** 16))
            step = (1<<10)
            myspa.set_cfreq(freq)
            for loop in range(loop_count):
                ld_list=[]
                self.frac_divf_set(divf)
                res = myspa.pk_search_timesleep(timesleep=1)
                freq_center = res[0][0]
                txp1 = res[0][1]
                ld = self.jlink.rdm(0xa0000150,3,3)
                ld_list.append(freq_center)

                self.frac_divf_set(divf-step)
                res = myspa.pk_search_timesleep(timesleep=1)
                freq_divf1 = res[0][0]
                txp2 = res[0][1]
                ld1 = self.jlink.rdm(0xa0000150, 3, 3)
                ld_list.append(freq_divf1)

                self.frac_divf_set(divf)
                res = myspa.pk_search_timesleep(timesleep=1)
                freq_center1 = res[0][0]
                txp3 = res[0][1]
                ld2 = self.jlink.rdm(0xa0000150,3,3)
                ld_list.append(freq_center1)

                self.frac_divf_set(divf+step)
                res = myspa.pk_search_timesleep(timesleep=1)
                freq_divf2 = res[0][0]
                txp4 = res[0][1]
                ld3 = self.jlink.rdm(0xa0000150, 3, 3)
                ld_list.append(freq_divf2)

                loginfo('{}'.format(freq_center,freq_divf1,freq_divf2,ld,ld1,ld2))
                fw1.write_data([freq_center,freq_divf1,freq_center1,freq_divf2,ld,ld1,ld2,ld3,txp1,txp2,txp3,txp4,divf,hex(step)],float_num=6)

                plt.clf()  # 清空画布上的所有内容

                for i in range(len(ld_list)):
                # time.sleep(time_interval)
                    t = t+1
                    # per = eval(res[1])
                    per_list.append(ld_list[i])
                    t_list.append(t)
                    plt.plot(t_list,per_list, '-r')
                    plt.pause(0.01)
                loop = loop+1

    def rw_validate(self,diag_cntl=0x29):
        '''

        '''
        self.jlink.wrm(0xa0200258,11,6,14)  ##{bt_sys_bt_dbg1,bt_sys_bt_dbg2}
        self.jlink.wrm(0xa0200240,26,3,3|(3<<3)|(3<<6)|(3<<9)|(3<<12)|(3<<15)|(3<<18)|(3<<21))  ##PA1--8 set testpin
        self.jlink.wrm(0xa0400450, 23, 23, 1)   ##bt_sys_bt_dbg2 enable
        self.jlink.wrm(0xa0400450, 22, 16, diag_cntl)   ##diag_cntl




class bt_test(object):
    def __init__(self, comport, chipv='TX232_MPW3', jlink='59610138',jlink_en=1,board_name=''):
        self.comport = comport
        self.chipv = chipv
        self.mem_ts = MEM_TS(self.comport)
        self.board_name = board_name
        self.jlink = jlink
        self.jlink_en = jlink_en
        # if jlink_en !=0:
        #     self.jlink = jlink(jlink_sn=jlink_sn)

        self.br_len_dic = {
            'DH1': 27,
            'DH3': 183,
            'DH5': 339
        }
        self.br_rate_dic = {
            'DH1': 1,
            'DH3': 4,
            'DH5': 7
        }
        self.edr_len_dic = {
            '2_DH1': 31,
            '2_DH3': 356,
            '2_DH5': 656,
            '3_DH1': 83,
            '3_DH3': 552,
            '3_DH5': 1021,
        }
        self.edr_rate_dic = {
            '2_DH1': 2,
            '2_DH3': 5,
            '2_DH5': 8,
            '3_DH1': 3,
            '3_DH3': 6,
            '3_DH5': 9,
        }
        self.le_rate_dic = {
            'LE1M': 1,
            'LE2M': 2,
            'LE500k': 3,
            'LE125K': 4
        }

    def hci_reset(self):
        self.comport.req_com(cmdstr='\x01\x03\x0c\x00',timeout=1,endstr='TS')

    def hci_read_buffer_size(self):
        self.comport.req_com(cmdstr='\x01\x05\x10\x00', timeout=1, endstr='TS')

    def hci_enable_under_test_mode(self):
        self.comport.req_com(cmdstr='\x01\x03\x18\x00', timeout=1, endstr='TS')

    def hci_enable_write_both_scan(self):
        self.comport.req_com(cmdstr='\x01\x1a\x0c\x01\x03', timeout=1, endstr='TS')

    def hci_disable_write_both_scan(self):
        self.comport.req_com(cmdstr='\x01\x1a\x0c\x01\x00', timeout=1, endstr='TS')

    def hci_auto_accept_all_connection(self):
        self.comport.req_com(cmdstr='\x01\x05\x0c\x03\x02\x00\x02', timeout=1, endstr='TS')

    def dut_signal_mode_set(self):

        res = self.comport.req_com(cmdstr='signaltest', wr_end='\r\n', timeout=5)
        logdebug(res)
        if res.find('signal test')!=-1:
            self.comport.Hexmd = 1
            self.hci_reset()
            self.hci_read_buffer_size()
            self.hci_enable_under_test_mode()
            self.hci_enable_write_both_scan()
            self.hci_auto_accept_all_connection()
            self.comport.Hexmd = 0
        else:
            logwarn('can not enter signal test mode')

    def BT_INIT(self, rate='LE1M'):
        if rate not in ['BR','EDR','LE1M','LE2M','LELR']:
            logerror('rate is wrong,must be BR,EDR,LE1M,LE2M,LELR ')
            return False
        self.mem_ts.wr(0xa0421004, 0x37)
        self.mem_ts.wr(0xa04210f8, 0x2f0fd2c)
        if rate == 'LE1M':
            self.mem_ts.wrm(0xa0400880,7,0,0X40)
            self.mem_ts.wr(0xa01200c8, 0x1802f0)
        elif rate == 'LE2M':
            self.mem_ts.wrm(0xa0400884,7,0,0X40)


    def LE_TX(self, chan=0, len=37, ptype=2, phy=2):
        '''
        chan:   0--39
        len:    1--255
        ptype:    0 ->PSEUDO_RAND_9; 1 ->11110000; 2->10101010; 3->PSEUDO_RAND_15
                4 ->11111111;      5 ->00000000; 6->00001111; 7->01010101
        phy:    1 ->1M; 2 ->2M; 3->coded S=8; 4->coded S=2
        '''
        if chan < 0 or chan > 39:
            logerror(' le channel is out of range')
            return False
        if len < 1 or len > 255:
            logerror('le packet lenght is out of range')
            return False
        # if phy == 2:
        #     self.mem_ts.wr(0xa04200e4, 0x9d8a)  ##GFSK modulation index for BLE mode
        # self.comport.req_com("bletx ch=%d len=%d pay=%d phy=%d" %(chan, len, ptype, phy), wr_end='\r\n', timeout=5)
        self.cmdstop(0)
        self.cmdstop(1)
        for i in range(5):
            res = self.comport.req_com("bletx ch=%d len=%d pay=%d phy=%d" %(chan, len, ptype, phy), wr_end='\r\n', timeout=5)
            if res.find("bletx ch=%d len=%d pay=%d phy=%d" %(chan, len, ptype, phy)) != -1:
                logdebug('le tx ok')
                break
            i = i+1
            self.cmdstop(0)
            logwarn('cmd return error')
        # self.mem_ts.wrm(0xa04210bc, 28, 26, 7)  ##bt_padrv_gain_table_gfsk_0
        # self.mem_ts.wrm(0xa04210c0, 28, 26, 7)  ##bt_padrv_gain_table_gfsk_1
        # self.mem_ts.wrm(0xa04210b4, 28, 26, 7)  ##bt_padrv_gain_table_dpsk_0
        # self.mem_ts.wrm(0xa04210b8, 28, 26, 7)
        # # self.mem_ts.wr(0xa01200c8, 0x1c02f5)
        # self.mem_ts.wr(0xa0421088, 0x00408411)  ##PA on
        # self.mem_ts.wr(0xa01200c8, 0x0018807A)  ##xtal
        # self.mem_ts.wr(0xa01200cc, 0x00000034)  ##xtal
        # self.mem_ts.wr(0xa0421024, 0x00000F07)  ##bt_padrv_setting
        # self.mem_ts.wr(0xa0421064, 0x0311FA80)  ##vco_pkd
        time.sleep(0.5)


    def BR_TX(self, chan=0, len=37, ptype=2, rate=2):
        '''
        chan:       0--78
        len:        1--1021
                    lenMax: (DH1)->27;    (2-DH1)->54;  (3-DH1)->83
                    (DH3)->183;   (2-DH3)->367; (3-DH3)->552
                    (DH5)->339;   (2-DH5)->679; (3-DH5)->1021
        ptype:      1 ->11110000; 2->10101010;  3->11111111
                    4 ->00000000; 5->00001111;  6->01010101
        type:       1 ->(DH1);    2 ->(2-DH1);  3->(3-DH1);
                    4 ->(DH3);    5 ->(2-DH3);  6->(3-DH3);
                    7 ->(DH5);    8 ->(2-DH5);  9->(3-DH5);
        '''
        self.cmdstop(1)
        self.cmdstop(0)
        if chan < 0 or chan > 78:
            logerror(' BT channel is out of range')
            return False
        if len < 1 or len > 1021:
            logerror('BT packet lenght is out of range')
            return False
        for i in range(5):
            res = self.comport.req_com("bttx ch=%d len=%d pay=%d type=%d" %(chan, len, ptype, rate), wr_end='\r\n', timeout=5)
            if res.find("bttx ch=%d len=%d pay=%d type=%d" %(chan, len, ptype, rate)) != -1:
                logdebug('bt tx ok')
                break
            i = i+1
            self.cmdstop(1)
            logwarn('cmd return error')

        # if rate in [1, 4, 7]:
        #     self.mem_ts.wrm(0xa04210bc, 28, 26, 7)  ##bt_padrv_gain_table_gfsk_0
        #     self.mem_ts.wrm(0xa04210c0, 28, 26, 7) ##bt_padrv_gain_table_gfsk_1
        #     self.mem_ts.wrm(0xa04210b4, 28, 26, 7)  ##bt_padrv_gain_table_dpsk_0
        #     self.mem_ts.wrm(0xa04210b8, 28, 26, 7)
        # # self.mem_ts.wr(0xa01200c8,0x1c02f5)
        # self.mem_ts.wr(0xa0421088, 0x00408411)  ##PA on
        # self.mem_ts.wr(0xa01200c8, 0x0018807A)  ##xtal
        # self.mem_ts.wr(0xa01200cc, 0x00000034)  ##xtal
        # self.mem_ts.wr(0xa0421024, 0x00000F07)  ##bt_padrv_setting
        # self.mem_ts.wr(0xa0421064, 0x0311FA80)  ##vco_pkd

        # self.mem_ts.wrm(0xa0421098, 12, 10, 7)  ##Manual Tx Mixer gain setting
        # self.mem_ts.wrm(0xa0421094, 11, 10, 3)
        # self.mem_ts.wrm(0xa01200cc, 14, 11, 1)

    def BR_RX(self, chan=0, len=27, ptype=5, rate=1, ltaddr=1, uap=0x48, lap=0x000080):
        if chan < 0 or chan > 78:
            logerror(' BT channel is out of range')
            return False
        if len < 1 or len > 1021:
            logerror('BT packet length is out of range')
            return False
        if ptype < 1 or ptype > 6:
            logerror('BT packet paytype is out of range')
            return False
        self.comport.req_com("btrx ch={} len={} pay={} type={} ltaddr={} uap={} lap={}".format(chan, len, ptype, rate, ltaddr, hex(uap), hex(lap)), wr_end='\r\n', timeout=5)

    def LE_RX(self, chan=0, phy=1, mod=0, countMode=0):
        if chan < 0 or chan > 39:
            logerror(' channel is out of range')
            return False
        self.comport.req_com("blerx ch={} phy={} mod={} countMode={}".format(chan, phy, mod, countMode), wr_end='\r\n', timeout=5)

    def tx_gain_set_tx232(self, en=1,  gain=63):
        self.mem_ts.wrm(0xa0421004, 22, 22, en)  ##Manual control of Tx PA target power enable
        self.mem_ts.wrm(0xa042100c, 31, 26, gain)  ##Manual control of Tx PA target power value


    def tx_carrier(self, freq=2402 , dac_gain=0xff, pa_gain=63, tp_mode=0):
        self.mem_ts.wr(0xa0422014, 0x0)  ##manual tx off
        self.mem_ts.wrm(0xa0421004, 22, 22, 1)  ##Manual control of Tx PA target power enable
        self.mem_ts.wrm(0xa042100c, 31, 26, pa_gain)  ##Manual control of Tx PA target power value
        self.mem_ts.wrm(0xa0421064, 12, 12, 1)  ##Frequency value manual enable
        self.mem_ts.wrm(0xa0421064, 11, 0, freq)  ##Manual value of frequency (unit: MHz)
        self.mem_ts.wr(0xa0422014, 0x0)  ##manual tx off
        ##chip.mem_ts.wrm(0xa0422044,10,0,0x7ff)
        ##chip.mem_ts.wr(0xa042100c,0x103e0249)  	##rf_ctrl_clk_manual
        self.mem_ts.wrm(0xa01200cc, 23, 22, 3)  ##rf_aon_xo_en_clk_adcdac_dig & rf_aon_xo_en_clk_ana_dig
        self.mem_ts.wrm(0xa01200cc, 20, 20, 1)  ##rf_aon_xo_en_clk_ref_dig
        ##IQ MODE
        if tp_mode==0:
            self.mem_ts.wrm(0xa0421020, 25, 24, 0)
            self.mem_ts.wrm(0xa04210e8, 26, 26, 1)  ##txdcoc sw mode
            self.mem_ts.wrm(0xa04210e8, 25, 18, dac_gain)  ##Manual TX DAC value for q-path gain index 7
            self.mem_ts.wrm(0xa04210e8, 17, 10, dac_gain)  ##Manual TX DAC value for i-path gain index 7
        ##TP MODE
        else:
            self.mem_ts.wrm(0xa0421020, 25, 24, 1)
            self.mem_ts.wrm(0xa0421088, 20, 20, 1)  ##txdcoc sw mode
            self.mem_ts.wrm(0xa0421088, 28, 21, dac_gain)  ##TP modulation dac code manual value

        self.mem_ts.wr(0xa0422014, 0xe)  ##manual tx on

    def tx_carrier_stop(self):
        self.mem_ts.wrm(0xa0421020, 25, 24, 0)
        self.mem_ts.wrm(0xa0421004, 22, 22, 0)  ##Manual control of Tx PA target power disable
        self.mem_ts.wrm(0xa0421064, 12, 12, 0)  ##Frequency value manual enable
        self.mem_ts.wr(0xa0422014, 0x0)  ##manual tx off
        self.mem_ts.wrm(0xa0422044, 10, 0, 0x0)
        self.mem_ts.wrm(0xa04210e8, 26, 26, 0)  ##txdcoc sw mode

    def tx_tone(self, freq=2402, symbol_code=1, pa_gain=63):
        self.mem_ts.wr(0xa0422014, 0x0)  ##manual tx off
        self.mem_ts.wrm(0xa0421004, 22, 22, 1)  ##Manual control of Tx PA target power enable
        self.mem_ts.wrm(0xa042100c, 31, 26, pa_gain)  ##Manual control of Tx PA target power value
        self.mem_ts.wrm(0xa0421064, 12, 12, 1)  ##Frequency value manual enable
        self.mem_ts.wrm(0xa0421064, 11, 0, freq)  ##Manual value of frequency (unit: MHz)
        # self.mem_ts.wrm(0xa0422044, 10, 0, 0x7ff)
        self.mem_ts.wrm(0xa04210e8, 26, 26, 1)  ##txdcoc sw mode
        self.mem_ts.wrm(0xa04210e8, 25, 18, 0xff)  ##Manual TX DAC value for q-path gain index 7
        self.mem_ts.wrm(0xa04210e8, 17, 10, 0xff)  ##Manual TX DAC value for i-path gain index 7
        self.mem_ts.wrm(0xa0422040, 2, 0, symbol_code)  ##symbol_tx，select tx 1 or 0
        self.mem_ts.wrm(0xa0422040, 8, 8, 1)  ##le mode
        self.mem_ts.wr(0xa0422014, 0xe)  ##manual tx on
        self.mem_ts.wrm(0xa0422040, 23, 23, 1)  ##tx ramp
        self.mem_ts.wrm(0xa0422040, 14, 14, 1)  ##modem tx_en
        self.mem_ts.wrm(0xa0422040, 6, 4, 1)  ##symbol_f

    def tx_tone_stop(self):
        self.mem_ts.wrm(0xa0421004, 22, 22, 0)  ##Manual control of Tx PA target power enable
        self.mem_ts.wrm(0xa0421064, 12, 12, 0)  ##Frequency value manual enable
        self.mem_ts.wr(0xa0422014, 0x0)  ##manual tx off
        self.mem_ts.wrm(0xa0422040, 2, 0, 0)  ##symbol_tx
        self.mem_ts.wrm(0xa0422040, 8, 8, 0)  ##le mode
        self.mem_ts.wrm(0xa0422040, 23, 23, 0)  ##tx ramp
        self.mem_ts.wrm(0xa0422040, 14, 14, 0)  ##modem tx_en
        self.mem_ts.wrm(0xa0422040, 6, 4, 0)  ##symbol_flag_tx
        self.mem_ts.wrm(0xa04210e8, 26, 26, 0)  ##txdcoc sw mode

    def cmdstop(self, mode=0):

        if mode == 0 or mode == 'le':
            cmdstr = 'letestend'
        else:
            cmdstr = 'bttestend'
        for i in range(10):
            time.sleep(1)
            res = self.comport.req_com(cmdstr, endstr='TS', wr_end='\r\n', timeout=1)
            logdebug(res)
            if len(res) != 0:
                if res[0] != 'cmd   head error! Send Again!   \n':
                    break
        return res

    def get_cmw_exp_pwr(self, tester=tester_cmw, target_power=30):
        self.BR_TX(chan=0, len=27, ptype=1, rate=1)
        for i in range(15):
            tester.stx.analyzer_settings(enpower=target_power, umargin=8, freq=2402)
            res = tester.stx.get_power_measure_res()
            if eval(res[0]) == 0:
                txpower = eval(res[3])
                target_power = txpower +10
                break
            else:
                target_power = target_power-5
        loginfo('target power set {}'.format(target_power))
        return target_power

    def test_br_tx(self,chan_list=[0,39,78], rfport=2, cable_loss=1, rate_list=['DH1'],  target_power=10, fig_en=1, csv_save=True, report_save=True):

        if csv_save :
            title = 'channel,type,nominal_pow(dBm),peak_pow(dBm),leakage_pow(dBm),freq_accuracy(kHz),freq_drift(kHz),drift_rate(Hz/50 μs),'
            title = title + 'delta_f1_avg(kHz),delta_f1_min(kHz),delta_f1_max(kHz),delta_f2_avg(kHz),delta_f2_min(kHz),delta_f2_max(kHz),mod_ratio,delta_f2_99(kHz),'
            title = title + 'obw(kHz),frange_l(kHz),frange_h(kHz),acp_list_21ch\n'
            fname = self.get_filename('ts_bt_test/','test_br_tx_{}'.format(self.board_name))
            fw1=csvreport(fname,title)
        if report_save:
            title2 = 'BR_tx_performance'
            fname2 = self.get_filename('ts_bt_test/','test_br_tx_report_{}'.format(self.board_name))
            fw2=csvreport(fname2,title2)
            item = ['channel', 'DH', u'Tx/01 - Output Power  - Average', u'- Peak',
                    u'@   - Leakage',
                    u' @   - Packet Timing',
                    u'Tx/04 - Spectrum - Frequency Range - Left',
                    u'- Right',
                    u'Tx/05 - Spectrum - 20 dB Bandwidth - f(L)',
                    u'- f(H)',
                    u'- f(H) - f(L)',
                    u'Tx/06 - Spectrum - Adjacent channel power - Max Power',
                    u'- ACP - >3 MHz',
                    u'- ACP -3 MHz',
                    u'- ACP -2 MHz',
                    u'- ACP -1 MHz',
                    u' - ACP +1 MHz',
                    u'- ACP +2 MHz',
                    u'- ACP +3 MHz',
                    u'- ACP + >3 MHz',
                    u'Tx/07 - Modulation Characteristics - Delta F1 Avg',
                    u'Tx/07 - Modulation Characteristics - Delta F2 Avg',
                    u'- Delta F2 Max 99%',
                    u'- Delta F2 Avg/Delta F1 Avg',
                    u'Tx/08 - Initial Carrier Frequency Tolerance - Avg',
                    u'- Max',
                    u'Tx/09 - Carrier Frequency Drift - Max.Drift DH1',
                    u' - Max.Drift DH5',
                    u' - Max.Drift',
                    u' - Max.Drift Rate/50us DH1',
                    u' - Max.Drift Rate/50us DH5',
                    u'- Max.Drift Rate/50us',
                    u'Harmonic Spurs']

            df = pd.DataFrame(item)
            df.to_csv(fw2.filename, index=False)

        tester = tester_cmw(mode=0, rfport=rfport, cable_loss=cable_loss, bt_mode='BR', freq=2402, target_power=target_power, burst_type='BR', asyn='OFF', bdaddr='050604010203')
        res_list = []
        self.cmdstop(1)
        self.cmdstop(0)
        tester.stx.trigger_settings(source='power', threshold=-20, tout=1)
        flag = True
        for chan in chan_list:
            for DH in rate_list:
                tester.stx.analyzer_settings(enpower=target_power, freq=2402 + chan)
                br_len = self.br_len_dic[DH]
                _rate = self.br_rate_dic[DH]
                self.BR_TX(chan=chan, len=br_len, ptype=1, rate=_rate)
                res = tester.stx.get_modulation_measure_res()
                logdebug('{}'.format(res))

                delta_f1_avg = eval(res[6])/1000.00
                delta_f1_min = eval(res[7])/1000.00
                delta_f1_max = eval(res[8])/1000.00

                time.sleep(0.5)
                self.cmdstop(1)
                self.BR_TX(chan=chan, len=br_len, ptype=2, rate=_rate)
                res = tester.stx.get_modulation_measure_res()

                logdebug('{}'.format(res))

                delta_f2_99 = eval(res[2])/1000.00
                freq_accuracy = eval(res[3])/1000.00
                freq_drift = eval(res[4])/1000.00
                drift_rate = eval(res[5])/1000.00
                delta_f2_avg = eval(res[9])/1000.00
                delta_f2_min = eval(res[10])/1000.00
                delta_f2_max = eval(res[11])/1000.00
                mod_ratio = delta_f2_avg/delta_f1_avg

                _res = tester.stx.get_modulation_measure_res(data_type='MAX')
                freq_drift = eval(_res[4])/1000.00
                drift_rate = eval(_res[5])/1000.00

                time.sleep(0.5)
                self.cmdstop(1)
                self.BR_TX(chan=chan, len=br_len, ptype=7, rate=_rate)
                res1 = tester.stx.get_power_measure_res()
                print(res1)
                if eval(res1[0]) != 0:
                    self.cmdstop(1)
                    self.BR_TX(chan=chan, len=br_len, ptype=7, rate=_rate)
                    res1 = tester.stx.get_power_measure_res()
                res2 = tester.stx.get_acp_res()
                print(res2)
                if eval(res2[0]) != 0:
                    self.cmdstop(1)
                    self.BR_TX(chan=chan, len=br_len, ptype=7, rate=_rate)
                    res2 = tester.stx.get_acp_res()
                res3 = tester.stx.get_obw_res()
                print(res3)
                if eval(res3[0]) != 0:
                    self.cmdstop(1)
                    self.BR_TX(chan=chan, len=br_len, ptype=7, rate=_rate)
                    res3 = tester.stx.get_obw_res()
                logdebug('{}'.format(res1))
                logdebug('{}'.format(res2))
                logdebug('{}'.format(res3))

                nominal_pow = eval(res1[2])
                peak_pow = eval(res1[3])
                leakage_pow = eval(res1[4])
                PacketTiming = eval(res1[5])
                acp_list = [eval(i) for i in res2[1:]]
                print (acp_list)

                acp_center = 10
                acp_max_pwr = acp_list[acp_center]
                acp_r4 = acp_list[acp_center + 4]
                acp_r3 = acp_list[acp_center + 3]
                acp_r2 = acp_list[acp_center + 2]
                acp_r1 = acp_list[acp_center + 1]
                acp_l1 = acp_list[acp_center - 1]
                acp_l2 = acp_list[acp_center - 2]
                acp_l3 = acp_list[acp_center - 3]
                acp_l4 = acp_list[acp_center - 4]
                obw = eval(res3[6])/1000.00
                obw_l = eval(res3[4])/1000.00
                obw_h = eval(res3[5])/1000.00


                time.sleep(0.5)
                self.cmdstop(1)
                self.BR_TX(chan=0, len=br_len, ptype=7, rate=_rate)
                tester.stx.analyzer_settings(enpower=target_power, freq=2402)
                res4 = tester.stx.get_frange_res()
                print(res4)
                if eval(res4[0]) != 0:
                    self.cmdstop(1)
                    self.BR_TX(chan=0, len=br_len, ptype=7, rate=_rate)
                    res4 = tester.stx.get_frange_res()
                logdebug('{}'.format(res4))
                frange_l = eval(res4[3])/1000.00

                self.cmdstop(1)
                self.BR_TX(chan=78, len=br_len, ptype=7, rate=_rate)
                tester.stx.analyzer_settings(enpower=target_power, freq=2480)
                res4 = tester.stx.get_frange_res()
                print(res4)
                if eval(res4[0]) != 0:
                    self.cmdstop(1)
                    self.BR_TX(chan=78, len=br_len, ptype=7, rate=_rate)
                    res4 = tester.stx.get_frange_res()
                logdebug('{}'.format(res4))
                frange_h = eval(res4[4])/1000.00
                self.cmdstop(1)

                loginfo("channel: {}    packet type: {}     payload length: {}".format(chan, DH, len))
                loginfo("nominal power: {}      peak power: {}      leakage power: {}   unit: dBm".format(nominal_pow, peak_pow, leakage_pow))
                loginfo("freq accuracy: {}      freq drift: {}      drift rate: {}  unit: Hz".format(freq_accuracy, freq_drift, drift_rate))
                loginfo("delta f1 avg: {}       delta f1 min: {}       delta f1 max: {}  unit: Hz".format(delta_f1_avg, delta_f1_min, delta_f1_max))
                loginfo("delta f2 avg: {}       delta f2 min: {}       delta f2 max: {}  unit: Hz".format(delta_f2_avg, delta_f2_min, delta_f2_max))
                loginfo("mod ratio: {}".format(mod_ratio))
                loginfo("delta f2 99.9%: {}  unit: Hz".format(delta_f2_99))
                loginfo("spectrum 20 dB bandwidth: {}".format(obw))
                loginfo("spectrum freq range: {} -- {} MHz".format(frange_l, frange_h))

                if fig_en != 0:
                    plt.ion()
                    x = [2402+chan+xi for xi in range(-10, 11, 1)]
                    fig = plt.figure('br_acp_channel{}_{}'.format(chan,DH))
                    fig.set_size_inches(13, 7)
                    ax = fig.add_subplot(111)
                    ax.bar(x, 80+np.array(acp_list), width=0.5, bottom=-80)
                    ax.set_xticks(np.array(x))
                    ax.set_xlabel('freq(MHz)')
                    ax.set_ylabel('txp(dBm)')
                    ax.set_title('BR_ACP')
                    for xi,yi in zip(x,acp_list):
                        ax.text(xi-0.25, yi+0.25, '%.2f'%yi, fontsize=8)
                    plt.savefig(fname+'br_acp_channel{}_{}_{}'.format(chan,DH,time.strftime('%Y%m%d_%H%M%S',time.localtime(time.time()))))
                    # plt.show()
                if report_save:
                    df = pd.DataFrame(pd.read_csv(fw2.filename))
                    df_value = [chan,DH,nominal_pow,peak_pow,leakage_pow,PacketTiming,frange_l,frange_h,obw_l,obw_h,obw,acp_max_pwr,acp_l4,acp_l3,
                                                           acp_l2,acp_l1,acp_r1,acp_r2,acp_r3,acp_r4,delta_f1_avg,delta_f2_avg,delta_f2_99,mod_ratio,freq_accuracy,'',freq_drift,
                                                           '','',drift_rate,'','','']
                    print (len(df_value))
                    df['channel_{}_{}'.format(chan,DH)] = df_value
                    df.to_csv(fw2.filename,index=False)
                if csv_save:
                    fw1.write_data([chan,DH,nominal_pow,peak_pow,leakage_pow,freq_accuracy,freq_drift,drift_rate,delta_f1_avg,delta_f1_min,delta_f1_max,delta_f2_avg,delta_f2_min,delta_f2_max,mod_ratio,delta_f2_99,obw,frange_l,frange_h,acp_list])
                else:
                    res_list.append([chan,DH,nominal_pow,peak_pow,leakage_pow,freq_accuracy,freq_drift,drift_rate,delta_f1_avg,delta_f1_min,delta_f1_max,delta_f2_avg,delta_f2_min,delta_f2_max,mod_ratio,delta_f2_99,obw,frange_l,frange_h,acp_list])
        return res_list
    def test_br_tx_debug(self, chan_list=[0, 39, 78], rfport=2, cable_loss=1, rate_list=['DH1'], target_power=10, fig_en=1,
                   csv_save=True, pa_gain=7, mixer_gain=7, dig_gain=0):

        if csv_save:
            title = 'channel,type,nominal_pow(dBm),peak_pow(dBm),leakage_pow(dBm),freq_accuracy(kHz),freq_drift(kHz),drift_rate(Hz/50 μs),'
            title = title + 'delta_f1_avg(kHz),delta_f1_min(kHz),delta_f1_max(kHz),delta_f2_avg(kHz),delta_f2_min(kHz),delta_f2_max(kHz),mod_ratio,delta_f2_99(kHz),'
            title = title + 'obw(kHz),frange_l(kHz),frange_h(kHz),acp_list_21ch\n'
            fname = self.get_filename('ts_bt_test/','test_br_tx_{}'.format(self.board_name))
            fw1=csvreport(fname,title)

        tester = tester_cmw(mode=0, rfport=rfport, cable_loss=cable_loss, bt_mode='BR', freq=2402,
                            target_power=target_power, burst_type='BR', asyn='OFF', bdaddr='050604010203')
        # self.BT_INIT('BR')
        self.cmdstop(1)
        self.cmdstop(0)
        res_list=[]
        self.tx_gain_set(pa_gain,mixer_gain,dig_gain)
        self.BR_TX(chan=0, len=27, ptype=1, rate=1)
        self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
        for i in range(15):
            tester.stx.measure_para(2, 'CONT', 20, MOEXception='ON')
            tester.stx.analyzer_settings(enpower=target_power, umargin=8, freq=2402)
            res = tester.stx.get_power_measure_res()
            loginfo(res)
            if eval(res[0]) == 0 or eval(res[0]) == 8:
                txpower = eval(res[3])
                target_power = txpower +10
                break
            else:
                target_power = target_power-5

        loginfo('target power set {}'.format(target_power))
        for chan in chan_list:
            for DH in rate_list:
                tester.stx.analyzer_settings(enpower=target_power, freq=2402 + chan)
                len = self.br_len_dic[DH]
                _rate = self.br_rate_dic[DH]
                self.cmdstop(1)
                # self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                self.BR_TX(chan=chan, len=len, ptype=1, rate=_rate)
                self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                res = tester.stx.get_modulation_measure_res()
                logdebug('{}'.format(res))

                delta_f1_avg = eval(res[6])/1000.00
                delta_f1_min = eval(res[7])/1000.00
                delta_f1_max = eval(res[8])/1000.00

                time.sleep(0.5)
                self.cmdstop(1)
                # self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                self.BR_TX(chan=chan, len=len, ptype=2, rate=_rate)
                res = tester.stx.get_modulation_measure_res()

                logdebug('{}'.format(res))

                delta_f2_99 = eval(res[2])/1000.00
                freq_accuracy = eval(res[3])/1000.00
                freq_drift = eval(res[4])/1000.00
                drift_rate = eval(res[5])
                delta_f2_avg = eval(res[9])/1000.00
                delta_f2_min = eval(res[10])/1000.00
                delta_f2_max = eval(res[11])/1000.00
                mod_ratio = delta_f2_avg / delta_f1_avg

                time.sleep(0.5)
                self.cmdstop(1)
                # self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                self.BR_TX(chan=chan, len=len, ptype=7, rate=_rate)
                self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                res1 = tester.stx.get_power_measure_res()

                if eval(res1[0]) != 0:
                    self.cmdstop(1)
                    # self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                    self.BR_TX(chan=chan, len=len, ptype=7, rate=_rate)
                    self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                    res1 = tester.stx.get_power_measure_res()
                res2 = tester.stx.get_acp_res()

                if eval(res2[0]) != 0:
                    self.cmdstop(1)
                    # self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                    self.BR_TX(chan=chan, len=len, ptype=7, rate=_rate)
                    self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                    res2 = tester.stx.get_acp_res()
                res3 = tester.stx.get_obw_res()

                if eval(res3[0]) != 0:
                    self.cmdstop(1)
                    # self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                    self.BR_TX(chan=chan, len=len, ptype=7, rate=_rate)
                    self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                    res3 = tester.stx.get_obw_res()
                logdebug('{}'.format(res1))
                logdebug('{}'.format(res2))
                logdebug('{}'.format(res3))

                nominal_pow = eval(res1[2])
                peak_pow = eval(res1[3])
                leakage_pow = eval(res1[4])
                acp_list = [eval(i) for i in res2[1:]]
                obw = eval(res3[6])/1000.00

                time.sleep(0.5)
                self.cmdstop(1)
                # self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                self.BR_TX(chan=0, len=len, ptype=7, rate=_rate)
                self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                tester.stx.analyzer_settings(enpower=target_power, freq=2402)
                res4 = tester.stx.get_frange_res()

                if eval(res4[0]) != 0:
                    self.cmdstop(1)
                    # self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                    self.BR_TX(chan=0, len=len, ptype=7, rate=_rate)
                    self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                    res4 = tester.stx.get_frange_res()
                logdebug('{}'.format(res4))
                frange_l = eval(res4[3])/1000.00

                self.cmdstop(1)
                # self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                self.BR_TX(chan=78, len=len, ptype=7, rate=_rate)
                self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                tester.stx.analyzer_settings(enpower=target_power, freq=2480)
                res4 = tester.stx.get_frange_res()

                if eval(res4[0]) != 0:
                    self.cmdstop(1)
                    # self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                    self.BR_TX(chan=78, len=len, ptype=7, rate=_rate)
                    self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                    res4 = tester.stx.get_frange_res()
                logdebug('{}'.format(res4))
                frange_h = eval(res4[4])/1000.00
                self.cmdstop(1)

                loginfo("channel: {}    packet type: {}     payload length: {}".format(chan, DH, len))
                loginfo("nominal power: {}      peak power: {}      leakage power: {}   unit: dBm".format(nominal_pow,
                                                                                                          peak_pow,
                                                                                                          leakage_pow))
                loginfo("freq accuracy: {}      freq drift: {}      drift rate: {}  unit: Hz".format(freq_accuracy,
                                                                                                     freq_drift,
                                                                                                     drift_rate))
                loginfo("delta f1 avg: {}       delta f1 min: {}       delta f1 max: {}  unit: Hz".format(delta_f1_avg,
                                                                                                          delta_f1_min,
                                                                                                          delta_f1_max))
                loginfo("delta f2 avg: {}       delta f2 min: {}       delta f2 max: {}  unit: Hz".format(delta_f2_avg,
                                                                                                          delta_f2_min,
                                                                                                          delta_f2_max))
                loginfo("mod ratio: {}".format(mod_ratio))
                loginfo("delta f2 99.9%: {}  unit: Hz".format(delta_f2_99))
                loginfo("spectrum 20 dB bandwidth: {}".format(obw))
                loginfo("spectrum freq range: {} -- {} MHz".format(frange_l, frange_h))

                if csv_save:
                    if fig_en != 0:
                        plt.ion()
                        x = [2402 + chan + xi for xi in range(-10, 11, 1)]
                        fig = plt.figure('br_acp_channel{}_{}'.format(chan, DH))
                        fig.set_size_inches(13, 7)
                        ax = fig.add_subplot(111)
                        ax.bar(x, 80 + np.array(acp_list), width=0.5, bottom=-80)
                        ax.set_xticks(np.array(x))
                        ax.set_xlabel('freq(MHz)')
                        ax.set_ylabel('txp(dBm)')
                        ax.set_title('BR_ACP')
                        for xi, yi in zip(x, acp_list):
                            ax.text(xi - 0.25, yi + 0.25, '%.2f' % yi, fontsize=8)
                        plt.savefig(fname + 'br_acp_channel{}_{}_{}'.format(chan, DH, time.strftime('%Y%m%d_%H%M%S',
                                                                                                    time.localtime(
                                                                                                        time.time()))))
                        # plt.show()
                    fw1.write_data([chan, DH, nominal_pow, peak_pow, leakage_pow, freq_accuracy, freq_drift, drift_rate,
                                    delta_f1_avg, delta_f1_min, delta_f1_max, delta_f2_avg, delta_f2_min, delta_f2_max,
                                    mod_ratio, delta_f2_99, obw, frange_l, frange_h, acp_list])
                else:
                    res_list.append([chan, DH, nominal_pow, peak_pow, leakage_pow, freq_accuracy, freq_drift, drift_rate,
                            delta_f1_avg, delta_f1_min, delta_f1_max, delta_f2_avg, delta_f2_min, delta_f2_max,
                            mod_ratio, delta_f2_99, obw, frange_l, frange_h, acp_list])
        return res_list

    def test_edr_tx(self,chan_list=[0,39,78], rfport=2, cable_loss=1, rate_list=['2_DH1'],  target_power=10, fig_en=1, csv_save=True, report_save=True):
        if csv_save:
            title = 'rate,channel,nominal_pwr(dBm),peak_pwr(dBm),gfsk_pwr(dBm),dpsk_pwr(dBm),dpsk_gfsk_diff_pwr,guard_period(us),'
            title = title + 'wi(KHz),w0_wi(KHz),w0_max(KHz),DEVM_RMS(%),DEVM_peak(%),DEVM_P99(%),bit_error_rate,packet0error,'
            title = title + 'PTxRef(dBm),N26ChN1Abs(dBm),N26ChP1Abs(dBm),N26ChN1Rel(dBm),N26ChP1Rel(dBm),acp_list\n'
            fname = self.get_filename('ts_bt_test/', 'test_edr_tx_{}'.format(self.board_name))
            fw1 = csvreport(fname, title)
        if report_save:
            title2 = 'BR_tx_performance'
            fname2 = self.get_filename('ts_bt_test/','test_edr_tx_report_{}'.format(self.board_name))
            fw2=csvreport(fname2,title2)
            item = ['channel', 'rate', u'Tx/10 - EDR Relative Transmit Power - PGFSK - Max Power', u' - PDPSK',
                    u'- PDPSK - PGFSK',
                    u'@   - Nomimal Power',
                    u'@   - Guard Period',
                    u'@   - Packet Timing',
                    u'Tx/11 - CFS and MA - omega i',
                    u' - omega i + omega o',
                    u'- omega o',
                    u'- DEVM RMS',
                    u'- DEVM Peak',
                    u'- DEVM 99% - Max 0.30',
                    u'Tx/12 - EDR Differential Phase Encoding',
                    u'Tx/13 - In Band Spurious Emissins - Nominal Power',
                    u' - ACPower:-3',
                    u'- ACPower:-2',
                    u'- ACPower:-1,Ptx-26dB',
                    u' - ACPower:Center Ptxref',
                    u'- ACPower:1,Ptx-26dB',
                    u'- ACPower:+2',
                    u'- ACPower:+3',
                    u'- ACPower:+- >3',]
            df = pd.DataFrame(item)
            df.to_csv(fw2.filename, index=False)
        tester = tester_cmw(mode=0, rfport=rfport, cable_loss=cable_loss, bt_mode='EDR', freq=2402, target_power=target_power, burst_type='EDR', asyn='OFF', bdaddr='050604010203')
        # self.BT_INIT('EDR')
        self.cmdstop(1)
        self.cmdstop(0)
        self.BR_TX(chan=0, len=54, ptype=1, rate=2)
        for i in range(15):
            tester.stx.analyzer_settings(enpower=target_power, umargin=8, freq=2402)
            res = tester.stx.get_power_measure_res()
            if eval(res[0]) == 0:
                txpower = eval(res[3])
                target_power = txpower + 10
                break
            else:
                target_power = target_power - 5
        loginfo('target power set {}'.format(target_power))
        for rate in rate_list:
            for chan in chan_list:
                tester.stx.analyzer_settings(enpower=target_power, freq=2402+chan)
                edr_len = self.edr_len_dic[rate]
                _rate = self.edr_rate_dic[rate]
                self.cmdstop(1)
                self.BR_TX(chan=chan, len=edr_len, ptype=7, rate=_rate)
                res = tester.stx.get_modulation_measure_res()
                logdebug('{}'.format(res))

                res = [eval(i) for i in res[0:]]

                wi = res[2]/1000.00
                w0_wi = res[3]/1000.00
                w0_max = res[4]/1000.00
                DEVM_RMS = res[5]
                DEVM_peak = res[6]
                DEVM_P99 = res[7]

                res1 = tester.stx.get_power_measure_res()
                logdebug('{}'.format(res1))
                res1 = [eval(i) for i in res1[0:]]
                nominal_pwr = res1[2]
                gfsk_pwr = res1[3]
                dpsk_pwr = res1[4]
                dpsk_gfsk_diff_pwr = res1[5]
                guard_period = res1[6]
                packet_timing = res1[7]
                peak_pwr = res1[8]

                res2 = tester.stx.get_diff_phase_encoding_res()
                logdebug('{}'.format(res2))
                res2 = [eval(i) for i in res2[0:]]
                bit_error_rate = res2[2]
                packet0error = res2[3]

                [res3, res4] = tester.stx.get_acp_res_edr()
                logdebug('{}'.format(res3))
                logdebug('{}'.format(res4))
                res3 = [eval(i) for i in res3[0:]]
                res4 = [eval(i) for i in res4[0:]]
                acp_nominal_pwr = res3[2]
                PTxRef = res3[4]
                N26ChN1Abs = res3[5]
                N26ChP1Abs = res3[6]
                N26ChN1Rel = res3[7]
                N26ChP1Rel = res3[8]

                acp_list = res4[1:]
                acp_center = len(acp_list)/2
                acp_max_pwr = acp_list[acp_center]
                acp_r4 = acp_list[acp_center + 4]
                acp_r3 = acp_list[acp_center + 3]
                acp_r2 = acp_list[acp_center + 2]
                acp_r1 = acp_list[acp_center + 1]
                acp_l1 = acp_list[acp_center - 1]
                acp_l2 = acp_list[acp_center - 2]
                acp_l3 = acp_list[acp_center - 3]
                acp_l4 = acp_list[acp_center - 4]

                loginfo("channel: {}    packet type: {}     payload length: {}".format(chan, rate, len))
                loginfo("nominal power: {}      peak power: {}      gfsk_pwr: {}    dpsk_pwr: {}".format(nominal_pwr, peak_pwr, gfsk_pwr,dpsk_pwr))
                loginfo("dpsk_gfsk_diff_pwr: {}      guard_period: {}      ".format(dpsk_gfsk_diff_pwr, guard_period))
                loginfo("wi: {}       w0_wi: {}       w0_max: {}".format(wi, w0_wi, w0_max))
                loginfo("DEVM_RMS: {}       DEVM_peak: {}       DEVM_P99: {}".format(DEVM_RMS, DEVM_peak, DEVM_P99))
                loginfo("bit_error_rate: {}     packet0error: {}".format(bit_error_rate, packet0error))
                loginfo("PTxRef: {} N26ChN1Abs: {}  N26ChP1Abs: {}  N26ChN1Rel: {}  N26ChP1Rel:{}".format(PTxRef, N26ChN1Abs, N26ChP1Abs, N26ChN1Rel, N26ChP1Rel))

                if fig_en != 0:
                    plt.ion()
                    x = [2402+chan+xi for xi in range(-10, 11, 1)]
                    fig = plt.figure('edr_acp_channel{}_{}'.format(chan,rate))
                    fig.set_size_inches(13, 7)
                    ax = fig.add_subplot(111)
                    ax.bar(x, 80+np.array(acp_list), width=0.5, bottom=-80)
                    ax.set_xticks(np.array(x))
                    ax.set_xlabel('freq(MHz)')
                    ax.set_ylabel('txp(dBm)')
                    ax.set_title('%s_ACP'%rate)
                    for xi,yi in zip(x,acp_list):
                        ax.text(xi-0.25, yi+0.25, '%.2f'%yi, fontsize=8)
                    plt.savefig(fname + 'edr_acp_channel{}_{}_{}'.format(chan, rate, time.strftime('%Y%m%d_%H%M%S',time.localtime(time.time()))))
                        # plt.show()
                if csv_save:
                    fw1.write_data([rate,chan,nominal_pwr,peak_pwr,gfsk_pwr,dpsk_pwr,dpsk_gfsk_diff_pwr,guard_period,wi,w0_wi,w0_max,DEVM_RMS,DEVM_peak,DEVM_P99,bit_error_rate,packet0error,PTxRef,N26ChN1Abs,N26ChP1Abs,N26ChN1Rel,N26ChP1Rel,acp_list])
                if report_save:
                    df = pd.DataFrame(pd.read_csv(fw2.filename))
                    df_value = [chan,rate,peak_pwr,dpsk_pwr,dpsk_gfsk_diff_pwr,nominal_pwr,guard_period,packet_timing,wi,w0_wi,w0_max,DEVM_RMS,DEVM_peak,DEVM_P99,'',
                                acp_nominal_pwr,acp_l3,acp_l2,N26ChN1Rel,PTxRef,N26ChP1Rel,acp_r2,acp_r3,max([acp_l4,acp_r4])]

                    df['channel_{}_{}'.format(chan,rate)] = df_value
                    df.to_csv(fw2.filename,index=False)
                self.cmdstop(1)

    def test_edr_tx_debug(self,chan_list=[0,39,78], rfport=2, cable_loss=1, rate_list=['2_DH1'],  target_power=10, fig_en=1, csv_save=True, pa_gain=7, mixer_gain=7, dig_gain=0):
        if csv_save:
            title = 'rate,channel,nominal_pwr(dBm),peak_pwr(dBm),gfsk_pwr(dBm),dpsk_pwr(dBm),dpsk_gfsk_diff_pwr,guard_period(us),'
            title = title + 'wi(KHz),w0_wi(KHz),w0_max(KHz),DEVM_RMS(%),DEVM_peak(%),DEVM_P99(%),bit_error_rate,packet0error,'
            title = title + 'PTxRef(dBm),N26ChN1Abs(dBm),N26ChP1Abs(dBm),N26ChN1Rel(dBm),N26ChP1Rel(dBm),acp_list\n'
            fname = self.get_filename('ts_bt_test/', 'test_edr_tx_{}'.format(self.board_name))
            fw1 = csvreport(fname, title)

        tester = tester_cmw(mode=0, rfport=rfport, cable_loss=cable_loss, bt_mode='EDR', freq=2402, target_power=target_power, burst_type='EDR', asyn='OFF', bdaddr='050604010203')
        # self.BT_INIT('EDR')
        self.cmdstop(1)
        self.cmdstop(0)
        res_list=[]
        self.tx_gain_set(pa_gain,mixer_gain,dig_gain)
        self.BR_TX(chan=0, len=54, ptype=1, rate=2)
        self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
        for i in range(15):
            tester.stx.measure_para(10, 'CONT', 20, MOEXception='ON')
            tester.stx.analyzer_settings(enpower=target_power, umargin=8, freq=2402)
            res = tester.stx.get_power_measure_res()
            if eval(res[0]) == 0 or eval(res[0]) == 8:
                txpower = eval(res[3])
                target_power = txpower +10
                break
            else:
                target_power = target_power-5
        loginfo('target power set {}'.format(target_power))
        for rate in rate_list:
            for chan in chan_list:
                tester.stx.analyzer_settings(enpower=target_power, freq=2402+chan)
                len = self.edr_len_dic[rate]
                _rate = self.edr_rate_dic[rate]
                self.cmdstop(1)
                # self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                self.BR_TX(chan=chan, len=len, ptype=7, rate=_rate)
                res = tester.stx.get_modulation_measure_res()

                res = [eval(i) for i in res[0:]]

                wi = res[2]/1000
                w0_wi = res[3]/1000
                w0_max = res[4]/1000
                DEVM_RMS = res[5]
                DEVM_peak = res[6]
                DEVM_P99 = res[7]

                res1 = tester.stx.get_power_measure_res()
                logdebug('{}'.format(res1))
                res1 = [eval(i) for i in res1[0:]]
                nominal_pwr = res1[2]
                gfsk_pwr = res1[3]
                dpsk_pwr = res1[4]
                dpsk_gfsk_diff_pwr = res1[5]
                guard_period = res1[6]
                # packet_timing = res1[7]
                peak_pwr = res1[8]

                res2 = tester.stx.get_diff_phase_encoding_res()
                logdebug('{}'.format(res2))
                res2 = [eval(i) for i in res2[0:]]
                bit_error_rate = res2[2]
                packet0error = res2[3]

                [res3, res4] = tester.stx.get_acp_res_edr()
                logdebug('{}'.format(res3))
                logdebug('{}'.format(res4))
                res3 = [eval(i) for i in res3[0:]]
                res4 = [eval(i) for i in res4[0:]]
                PTxRef = res3[4]
                N26ChN1Abs = res3[5]
                N26ChP1Abs = res3[6]
                N26ChN1Rel = res3[7]
                N26ChP1Rel = res3[8]

                acp_list = res4[1:]

                loginfo("channel: {}    packet type: {}     payload length: {}".format(chan, rate, len))
                loginfo("nominal power: {}      peak power: {}      gfsk_pwr: {}    dpsk_pwr: {}".format(nominal_pwr, peak_pwr, gfsk_pwr,dpsk_pwr))
                loginfo("dpsk_gfsk_diff_pwr: {}      guard_period: {}      ".format(dpsk_gfsk_diff_pwr, guard_period))
                loginfo("wi: {}       w0_wi: {}       w0_max: {}".format(wi, w0_wi, w0_max))
                loginfo("DEVM_RMS: {}       DEVM_peak: {}       DEVM_P99: {}".format(DEVM_RMS, DEVM_peak, DEVM_P99))
                loginfo("bit_error_rate: {}     packet0error: {}".format(bit_error_rate, packet0error))
                loginfo("PTxRef: {} N26ChN1Abs: {}  N26ChP1Abs: {}  N26ChN1Rel: {}  N26ChP1Rel:{}".format(PTxRef, N26ChN1Abs, N26ChP1Abs, N26ChN1Rel, N26ChP1Rel))

                if csv_save:
                    if fig_en != 0:
                        plt.ion()
                        x = [2402 + chan + xi for xi in range(-10, 11, 1)]
                        fig = plt.figure('edr_acp_channel{}_{}'.format(chan, rate))
                        fig.set_size_inches(13, 7)
                        ax = fig.add_subplot(111)
                        ax.bar(x, 80 + np.array(acp_list), width=0.5, bottom=-80)
                        ax.set_xticks(np.array(x))
                        ax.set_xlabel('freq(MHz)')
                        ax.set_ylabel('txp(dBm)')
                        ax.set_title('%s_ACP' % rate)
                        for xi, yi in zip(x, acp_list):
                            ax.text(xi - 0.25, yi + 0.25, '%.2f' % yi, fontsize=8)
                        plt.savefig(fname + 'edr_acp_channel{}_{}_{}'.format(chan, rate, time.strftime('%Y%m%d_%H%M%S',
                                                                                                       time.localtime(
                                                                                                           time.time()))))
                        # plt.show()
                    fw1.write_data([rate,chan,nominal_pwr,peak_pwr,gfsk_pwr,dpsk_pwr,dpsk_gfsk_diff_pwr,guard_period,wi,w0_wi,w0_max,DEVM_RMS,DEVM_peak,DEVM_P99,bit_error_rate,packet0error,PTxRef,N26ChN1Abs,N26ChP1Abs,N26ChN1Rel,N26ChP1Rel,acp_list])
                else:
                    res_list.append([rate,chan,nominal_pwr,peak_pwr,gfsk_pwr,dpsk_pwr,dpsk_gfsk_diff_pwr,guard_period,wi,w0_wi,w0_max,DEVM_RMS,DEVM_peak,DEVM_P99,bit_error_rate,packet0error,PTxRef,N26ChN1Abs,N26ChP1Abs,N26ChN1Rel,N26ChP1Rel,acp_list])

                self.cmdstop(1)
        return res_list
    def test_le_tx(self,chan_list=[0,19,38], rfport=2, cable_loss=1, rate_list=['LE1M'],  target_power=10, fig_en=1, csv_save=True, report_save=True):

        if csv_save:
            title = 'channel,type,nominal_pow(dBm),peak_pow(dBm),leakage_pow(dBm),freq_accuracy(KHz),freq_drift(KHz),drift_rate(Hz/50us),'
            title = title + 'delta_f1_avg(KHz),delta_f1_min(KHz),delta_f1_max(KHz),delta_f2_avg(KHz),delta_f2_min(KHz),delta_f2_max(KHz),mod_ratio,delta_f1_99(KHz),delta_f2_99(KHz),'
            title = title + 'acp_list_21ch\n'
            fname = self.get_filename('ts_bt_test/','test_le_tx_{}'.format(self.board_name))
            fw1=csvreport(fname,title)
        if report_save:
            title2 = 'BR_tx_performance'
            fname2 = self.get_filename('ts_bt_test/','test_le_tx_report_{}'.format(self.board_name))
            fw2=csvreport(fname2,title2)
            item = ['channel', 'rate', u'TP/TRM-LE/CA/BV-01-C - Output power  - Average Power', u' - Peak Power',
                    u'- Peak - Average Power',
                    u'- leakage Power',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -10',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -9',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -8',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -7',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -6',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -5',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -4',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -3',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -2',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -1',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -    Ptx',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   1',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   2',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   3',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   4',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   5',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   6',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   7',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   8',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   9',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   10',
                    u'TP/TRM-LE/CA/BV-05-C - Modulation Characteristics - Fre Accuracy',
                    u'- Delta F1 Avg',
                    u' - Delta F1 Min',
                    u' - Delta F1Max',
                    u'- Delta F2 Avg',
                    u'- Delta F2 Min',
                    u'- Delta F2Max',
                    u' - Delta F2 Max 99%',
                    u'- Delta F2 Avg/Delta F1 Avg',
                    u'- Frequency Accuracy',
                    u' - Frequency Offset',
                    u'- Frequency Drift',
                    u'- Max Drift Rate - 50us',
                    u'- Intial Frequency Drift',
                    ]
            df = pd.DataFrame(item)
            df.to_csv(fw2.filename, index=False)
        tester = tester_cmw(mode=0, rfport=rfport, cable_loss=cable_loss, bt_mode='LE1M', freq=2402, target_power=target_power, burst_type='LE', asyn='ON', bdaddr='050604010203')
        tester.stx.trigger_settings(source='power', threshold=-20, tout=1)
        self.cmdstop(0)
        self.cmdstop(1)
        res_lsit=[]
        for chan in chan_list:
            tester.stx.analyzer_settings(enpower=target_power, umargin=8, freq=2402+chan*2)
            for rate in rate_list:
                # self.BT_INIT(rate)
                _rate = self.le_rate_dic[rate]
                tester.stx.input_signal_settings(btype='LE', phy=rate)
                time.sleep(0.5)
                delta_f1_99,delta_f2_99,delta_f2_avg,delta_f2_min,delta_f2_max,mod_ratio = -999,-999,-999,-999,-999,-999
                if _rate > 2:
                    self.LE_TX(chan=chan, len=255, ptype=1, phy=_rate)
                    self.stx.mode = 'LRANge'
                    res = tester.stx.get_modulation_measure_res()
                    if eval(res[0]) != 0:
                        self.cmdstop(0)
                        self.LE_TX(chan=chan, len=255, ptype=1, phy=_rate)
                        res = tester.stx.get_modulation_measure_res()
                    logdebug('{}'.format(res))

                    delta_f1_99 = eval(res[2])/1000.00
                    freq_accuracy = eval(res[3])/1000.00
                    freq_drift = eval(res[4])/1000.00
                    drift_rate = eval(res[5])
                    delta_f1_avg = eval(res[6])/1000.00
                    delta_f1_min = eval(res[7])/1000.00
                    delta_f1_max = eval(res[8])/1000.00
                else:
                    self.LE_TX(chan=chan, len=255, ptype=1, phy=_rate)
                    tester.stx.mode = rate
                    res = tester.stx.get_modulation_measure_res()
                    if eval(res[0]) != 0:
                        self.cmdstop(0)
                        self.LE_TX(chan=chan, len=255, ptype=1, phy=_rate)
                        res = tester.stx.get_modulation_measure_res()
                    logdebug('{}'.format(res))

                    delta_f1_avg = eval(res[6])/1000.00
                    delta_f1_min = eval(res[7])/1000.00
                    delta_f1_max = eval(res[8])/1000.00

                    time.sleep(0.5)
                    self.cmdstop(0)
                    self.LE_TX(chan=chan, len=255, ptype=2, phy=_rate)
                    res = tester.stx.get_modulation_measure_res()
                    if eval(res[0]) != 0:
                        self.cmdstop(0)
                        self.LE_TX(chan=chan, len=255, ptype=2, phy=_rate)
                        res = tester.stx.get_modulation_measure_res()
                    logdebug('{}'.format(res))

                    delta_f2_99 = eval(res[2])/1000.00
                    freq_accuracy = eval(res[3])/1000.00
                    freq_drift = eval(res[4])/1000.00
                    drift_rate = eval(res[5])
                    delta_f2_avg = eval(res[9])/1000.00
                    delta_f2_min = eval(res[10])/1000.00
                    delta_f2_max = eval(res[11])/1000.00
                    mod_ratio = delta_f2_avg/delta_f1_avg
                    freq_offset = eval(res[14])/1000.00
                    init_drift = eval(res[15]) / 1000.00

                time.sleep(0.5)
                self.cmdstop(0)
                self.LE_TX(chan=chan, len=255, ptype=0, phy=_rate)
                res1 = tester.stx.get_power_measure_res()
                if eval(res1[0]) != 0:
                    self.cmdstop(0)
                    self.LE_TX(chan=chan, len=255, ptype=0, phy=_rate)
                    res1 = tester.stx.get_power_measure_res()
                res2 = tester.stx.get_acp_res()
                if eval(res2[0]) != 0:
                    self.cmdstop(0)
                    self.LE_TX(chan=chan, len=255, ptype=0, phy=_rate)
                    res2 = tester.stx.get_acp_res()
                logdebug('{}'.format(res1))
                logdebug('{}'.format(res2))

                nominal_pow = eval(res1[2])
                peak_pow = eval(res1[3])
                peak_avg_pwr = eval(res1[5])
                leakage_pow = eval(res1[4])
                acp_list = [eval(i) for i in res2[1:]]
                acp_center = len(acp_list)/2
                acp_max_pwr = acp_list[acp_center]
                acp_r10 = acp_list[acp_center + 10]
                acp_r9 = acp_list[acp_center + 9]
                acp_r8 = acp_list[acp_center + 8]
                acp_r7 = acp_list[acp_center + 7]
                acp_r6 = acp_list[acp_center + 6]
                acp_r5 = acp_list[acp_center + 5]
                acp_r4 = acp_list[acp_center + 4]
                acp_r3 = acp_list[acp_center + 3]
                acp_r2 = acp_list[acp_center + 2]
                acp_r1 = acp_list[acp_center + 1]
                acp_l1 = acp_list[acp_center - 1]
                acp_l2 = acp_list[acp_center - 2]
                acp_l3 = acp_list[acp_center - 3]
                acp_l4 = acp_list[acp_center - 4]
                acp_l5 = acp_list[acp_center - 5]
                acp_l6 = acp_list[acp_center - 6]
                acp_l7 = acp_list[acp_center - 7]
                acp_l8 = acp_list[acp_center - 8]
                acp_l9 = acp_list[acp_center - 9]
                acp_l10 = acp_list[acp_center - 10]

                loginfo("channel: {}    packet type: {} ".format(chan, rate))
                loginfo("nominal power: {} dBm      peak power: {} dBm      leakage power: {} dBm ".format(nominal_pow, peak_pow, leakage_pow))
                loginfo("freq accuracy: {} Hz     freq drift: {} Hz     drift rate: {} Hz".format(freq_accuracy, freq_drift, drift_rate))
                loginfo("delta f1 avg: {}  Hz     delta f1 min: {} Hz      delta f1 max: {} Hz".format(delta_f1_avg, delta_f1_min, delta_f1_max))
                if _rate < 3:
                    loginfo("delta f2 avg: {} Hz      delta f2 min: {} Hz      delta f2 max: {} Hz".format(delta_f2_avg, delta_f2_min, delta_f2_max))
                    loginfo("mod ratio: {}".format(mod_ratio))
                    loginfo("delta f2 99.9%: {} Hz".format(delta_f2_99))
                else:
                    loginfo("delta f1 99.9%: {} Hz".format(delta_f1_99))


                if fig_en != 0:
                    plt.ion()
                    x = [2402+chan*2+xi for xi in range(-10, 11, 1)]
                    fig = plt.figure('le_acp_channel{}_{}'.format(chan,rate))
                    fig.set_size_inches(13, 7)
                    ax = fig.add_subplot(111)
                    ax.bar(x, 80+np.array(acp_list), width=0.5, bottom=-80)
                    ax.set_xticks(np.array(x))
                    ax.set_xlabel('freq(MHz)')
                    ax.set_ylabel('txp(dBm)')
                    ax.set_title('%s_ACP'%rate)
                    for xi,yi in zip(x,acp_list):
                        ax.text(xi-0.25, yi+0.25, '%.2f'%yi, fontsize=8)
                    plt.savefig(fname + 'le_acp_channel{}_{}_{}'.format(chan, rate, time.strftime('%Y%m%d_%H%M%S',time.localtime(time.time()))))
                    # plt.show()
                if report_save:
                    df = pd.DataFrame(pd.read_csv(fw2.filename))
                    df_value = [chan,rate,nominal_pow,peak_pow,peak_avg_pwr,leakage_pow,acp_l10,acp_l9,acp_l8,acp_l7,acp_l6,acp_l5,acp_l4,acp_l3,acp_l2,acp_l1,acp_max_pwr,
                                acp_r1,acp_r2,acp_r3,acp_r4,acp_r5,acp_r6,acp_r7,acp_r8,acp_r9,acp_r10,freq_accuracy,delta_f1_avg,delta_f1_min,delta_f1_max,delta_f2_avg,
                                delta_f2_min,delta_f2_max,delta_f1_99,mod_ratio,freq_accuracy,freq_offset,freq_drift,drift_rate,init_drift]

                    df['channel_{}_{}'.format(chan,rate)] = df_value
                    df.to_csv(fw2.filename,index=False)
                if csv_save:
                    fw1.write_data([chan,rate,nominal_pow,peak_pow,leakage_pow,freq_accuracy,freq_drift,drift_rate,delta_f1_avg,delta_f1_min,delta_f1_max,delta_f2_avg,delta_f2_min,delta_f2_max,mod_ratio,delta_f1_99,delta_f2_99,acp_list])
                else:
                    res_lsit.append([chan,rate,nominal_pow,peak_pow,leakage_pow,freq_accuracy,freq_drift,drift_rate,delta_f1_avg,delta_f1_min,delta_f1_max,delta_f2_avg,delta_f2_min,delta_f2_max,mod_ratio,delta_f1_99,delta_f2_99,acp_list])

                self.cmdstop(0)

        return res_lsit

    def test_le_tx_debug(self,chan_list=[0,19,38], rfport=2, cable_loss=1, rate_list=['LE1M'],  target_power=10, fig_en=1, csv_save=True, pa_gain=7, mixer_gain=7, dig_gain=0):

        if csv_save:
            title = 'channel,type,nominal_pow(dBm),peak_pow(dBm),leakage_pow(dBm),freq_accuracy(KHz),freq_drift(KHz),drift_rate(Hz/50us),'
            title = title + 'delta_f1_avg(KHz),delta_f1_min(KHz),delta_f1_max(KHz),delta_f2_avg(KHz),delta_f2_min(KHz),delta_f2_max(KHz),mod_ratio,delta_f1_99(KHz),delta_f2_99(KHz),'
            title = title + 'acp_list_21ch\n'
            fname = self.get_filename('ts_bt_test/','test_le_tx_{}'.format(self.board_name))
            fw1=csvreport(fname,title)

        tester = tester_cmw(mode=0, rfport=rfport, cable_loss=cable_loss, bt_mode='LE1M', freq=2402, target_power=target_power, burst_type='LE', asyn='ON', bdaddr='050604010203')

        self.cmdstop(0)
        self.cmdstop(1)
        res_lsit=[]
        self.tx_gain_set(pa_gain,mixer_gain,dig_gain)
        self.LE_TX(chan=0, len=37, ptype=1, phy=1)
        for i in range(20):
            tester.stx.measure_para(2, 'CONT', 20, MOEXception='ON')
            tester.stx.analyzer_settings(enpower=target_power, umargin=8, freq=2402)
            res = tester.stx.get_power_measure_res()
            if eval(res[0]) == 0 or eval(res[0]) == 8:
                txpower = eval(res[3])
                target_power = txpower +10
                break
            else:
                target_power = target_power-5
        loginfo('target power set {}'.format(target_power))
        for chan in chan_list:
            tester.stx.analyzer_settings(enpower=target_power, umargin=8, freq=2402+chan*2)
            for rate in rate_list:
                # self.BT_INIT(rate)
                _rate = self.le_rate_dic[rate]
                tester.stx.input_signal_settings(btype='LE', phy=rate)
                time.sleep(0.5)
                delta_f1_99,delta_f2_99,delta_f2_avg,delta_f2_min,delta_f2_max,mod_ratio = -999,-999,-999,-999,-999,-999
                if _rate > 2:
                    self.cmdstop(0)
                    self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                    self.LE_TX(chan=chan, len=37, ptype=1, phy=_rate)
                    self.stx.mode = 'LRANge'
                    res = tester.stx.get_modulation_measure_res()
                    if eval(res[0]) != 0:
                        self.cmdstop(0)
                        self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                        self.LE_TX(chan=chan, len=37, ptype=1, phy=_rate)
                        res = tester.stx.get_modulation_measure_res()
                    logdebug('{}'.format(res))

                    delta_f1_99 = eval(res[2])/1000.00
                    freq_accuracy = eval(res[3])/1000.00
                    freq_drift = eval(res[4])/1000.00
                    drift_rate = eval(res[5])
                    delta_f1_avg = eval(res[6])/1000.00
                    delta_f1_min = eval(res[7])/1000.00
                    delta_f1_max = eval(res[8])/1000.00
                else:
                    self.cmdstop(0)
                    self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                    self.LE_TX(chan=chan, len=37, ptype=1, phy=_rate)
                    tester.stx.mode = rate
                    res = tester.stx.get_modulation_measure_res()
                    if eval(res[0]) != 0:
                        self.cmdstop(0)
                        self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                        self.LE_TX(chan=chan, len=37, ptype=1, phy=_rate)
                        res = tester.stx.get_modulation_measure_res()
                    logdebug('{}'.format(res))

                    delta_f1_avg = eval(res[6])/1000.00
                    delta_f1_min = eval(res[7])/1000.00
                    delta_f1_max = eval(res[8])/1000.00

                    time.sleep(0.5)
                    self.cmdstop(0)
                    self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                    self.LE_TX(chan=chan, len=37, ptype=2, phy=_rate)
                    res = tester.stx.get_modulation_measure_res()
                    if eval(res[0]) != 0:
                        self.cmdstop(0)
                        self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                        self.LE_TX(chan=chan, len=37, ptype=2, phy=_rate)
                        res = tester.stx.get_modulation_measure_res()
                    logdebug('{}'.format(res))

                    delta_f2_99 = eval(res[2])/1000.00
                    freq_accuracy = eval(res[3])/1000.00
                    freq_drift = eval(res[4])/1000.00
                    drift_rate = eval(res[5])
                    delta_f2_avg = eval(res[9])/1000.00
                    delta_f2_min = eval(res[10])/1000.00
                    delta_f2_max = eval(res[11])/1000.00
                    mod_ratio = delta_f2_avg/delta_f1_avg

                time.sleep(0.5)
                self.cmdstop(0)
                self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                self.LE_TX(chan=chan, len=37, ptype=0, phy=_rate)
                res1 = tester.stx.get_power_measure_res()
                if eval(res1[0]) != 0:
                    self.cmdstop(0)
                    self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                    self.LE_TX(chan=chan, len=37, ptype=0, phy=_rate)
                    res1 = tester.stx.get_power_measure_res()
                res2 = tester.stx.get_acp_res()
                if eval(res2[0]) != 0:
                    self.cmdstop(0)
                    self.tx_gain_set(pa_gain, mixer_gain, dig_gain)
                    self.LE_TX(chan=chan, len=37, ptype=0, phy=_rate)
                    res2 = tester.stx.get_acp_res()
                logdebug('{}'.format(res1))
                logdebug('{}'.format(res2))

                nominal_pow = eval(res1[2])
                peak_pow = eval(res1[3])
                leakage_pow = eval(res1[4])
                acp_list = [eval(i) for i in res2[1:]]

                loginfo("channel: {}    packet type: {} ".format(chan, rate))
                loginfo("nominal power: {} dBm      peak power: {} dBm      leakage power: {} dBm ".format(nominal_pow, peak_pow, leakage_pow))
                loginfo("freq accuracy: {} Hz     freq drift: {} Hz     drift rate: {} Hz".format(freq_accuracy, freq_drift, drift_rate))
                loginfo("delta f1 avg: {}  Hz     delta f1 min: {} Hz      delta f1 max: {} Hz".format(delta_f1_avg, delta_f1_min, delta_f1_max))
                if _rate < 3:
                    loginfo("delta f2 avg: {} Hz      delta f2 min: {} Hz      delta f2 max: {} Hz".format(delta_f2_avg, delta_f2_min, delta_f2_max))
                    loginfo("mod ratio: {}".format(mod_ratio))
                    loginfo("delta f2 99.9%: {} Hz".format(delta_f2_99))
                else:
                    loginfo("delta f1 99.9%: {} Hz".format(delta_f1_99))

                if csv_save:
                    if fig_en != 0:
                        plt.ion()
                        x = [2402 + chan * 2 + xi for xi in range(-10, 11, 1)]
                        fig = plt.figure('le_acp_channel{}_{}'.format(chan, rate))
                        fig.set_size_inches(13, 7)
                        ax = fig.add_subplot(111)
                        ax.bar(x, 80 + np.array(acp_list), width=0.5, bottom=-80)
                        ax.set_xticks(np.array(x))
                        ax.set_xlabel('freq(MHz)')
                        ax.set_ylabel('txp(dBm)')
                        ax.set_title('%s_ACP' % rate)
                        for xi, yi in zip(x, acp_list):
                            ax.text(xi - 0.25, yi + 0.25, '%.2f' % yi, fontsize=8)
                        plt.savefig(fname + 'le_acp_channel{}_{}_{}'.format(chan, rate, time.strftime('%Y%m%d_%H%M%S', time.localtime(time.time()))))
                        # plt.show()
                    fw1.write_data([chan,rate,nominal_pow,peak_pow,leakage_pow,freq_accuracy,freq_drift,drift_rate,delta_f1_avg,delta_f1_min,delta_f1_max,delta_f2_avg,delta_f2_min,delta_f2_max,mod_ratio,delta_f1_99,delta_f2_99,acp_list])
                else:
                    res_lsit.append([chan,rate,nominal_pow,peak_pow,leakage_pow,freq_accuracy,freq_drift,drift_rate,delta_f1_avg,delta_f1_min,delta_f1_max,delta_f2_avg,delta_f2_min,delta_f2_max,mod_ratio,delta_f1_99,delta_f2_99,acp_list])
                self.cmdstop(0)

        return res_lsit

    def tx_gain_set(self, pa_gain=7, mixer_gain=7, dig_gain=0):
        self.mem_ts.wrm(0xa0421080, 31, 30, 0)  ##PA、TXM gain manual
        self.mem_ts.wrm(0xa0421098, 12, 10, mixer_gain)  ##Manual Tx Mixer gain setting
        self.mem_ts.wrm(0xa0421094, 11, 10, pa_gain)  ##PA drive ability control bits
        time.sleep(1)

    def tx_gain_set_k140t(self, pa_gain=7, mixer_gain=7, dig_gain=0):
        gain = pa_gain + (mixer_gain<<3) + (dig_gain<<6)
        self.mem_ts.rfwr(0x83,gain)
        time.sleep(1)

    def test_tx_br_gain(self, chan_list=[0,39,78], rfport=2, cable_loss=1, rate_list=['DH1'], target_power=10, fig_en=0):

        title = 'tx_pa_gain,tx_mixer_gain,tx_dig_gain,channel,type,nominal_pow(dBm),peak_pow(dBm),leakage_pow(dBm),freq_accuracy(kHz),freq_drift(kHz),drift_rate(Hz/50 μs),'
        title = title + 'delta_f1_avg(kHz),delta_f1_min(kHz),delta_f1_max(kHz),delta_f2_avg(kHz),delta_f2_min(kHz),delta_f2_max(kHz),mod_ratio,delta_f2_99(kHz),'
        title = title + 'obw(kHz),frange_l(kHz),frange_h(kHz),acp_list_21ch\n'
        fname = self.get_filename('ts_bt_test/','test_br_tx_gain_{}'.format(self.board_name))
        fw1=csvreport(fname,title)

        for pa_gain in range(2, 8):
            res = self.test_br_tx_debug(chan_list=chan_list, rfport=rfport, cable_loss=cable_loss, rate_list=rate_list,
                                        target_power=target_power - 30 + (pa_gain - 2) * 6, fig_en=fig_en,
                                        csv_save=False, pa_gain=pa_gain, mixer_gain=7, dig_gain=0)
            for i in range(len(res)):
                fw1.write_data([pa_gain, 7, 0] + res[i])

        for mixer_gain in range(0, 7):
            res = self.test_br_tx_debug(chan_list=chan_list, rfport=rfport, cable_loss=cable_loss, rate_list=rate_list,
                                        target_power=target_power - 14 + mixer_gain * 2, fig_en=fig_en, csv_save=False,
                                        pa_gain=7, mixer_gain=mixer_gain, dig_gain=0)
            for i in range(len(res)):
                fw1.write_data([7, mixer_gain, 0] + res[i])

        for dig_gain in range(1, 31):
            res = self.test_br_tx_debug(chan_list=chan_list, rfport=rfport, cable_loss=cable_loss, rate_list=rate_list,
                                        target_power=target_power - dig_gain / 2, fig_en=fig_en, csv_save=False,
                                        pa_gain=7, mixer_gain=7, dig_gain=dig_gain)
            for i in range(len(res)):
                fw1.write_data([7, 7, dig_gain] + res[i])

    def test_tx_br_gain_tx231(self, chan_list=[0,39,78], rfport=2, cable_loss=1, rate_list=['DH1'], target_power=10, fig_en=0):

        title = 'tx_pa_gain,tx_mixer_gain,tx_dig_gain,channel,type,nominal_pow(dBm),peak_pow(dBm),leakage_pow(dBm),freq_accuracy(kHz),freq_drift(kHz),drift_rate(Hz/50 μs),'
        title = title + 'delta_f1_avg(kHz),delta_f1_min(kHz),delta_f1_max(kHz),delta_f2_avg(kHz),delta_f2_min(kHz),delta_f2_max(kHz),mod_ratio,delta_f2_99(kHz),'
        title = title + 'obw(kHz),frange_l(kHz),frange_h(kHz),acp_list_21ch\n'
        fname = self.get_filename('ts_bt_test/','test_br_tx_gain_{}'.format(self.board_name))
        fw1=csvreport(fname,title)

        for pa_gain in range(0, 4):
            for mixer_gain in range(0,8):
                res = self.test_br_tx_debug(chan_list=chan_list, rfport=rfport, cable_loss=cable_loss, rate_list=rate_list,target_power=target_power - (3-pa_gain) * 3-(7-mixer_gain)*2, fig_en=fig_en,
                                            csv_save=False, pa_gain=pa_gain, mixer_gain=mixer_gain, dig_gain=0)
                for i in range(len(res)):
                    fw1.write_data([pa_gain, mixer_gain, 0] + res[i])

    def test_tx_br_gain_tx232(self):

        title = 'ramp_tgt_val,channel,type,nominal_pow(dBm),peak_pow(dBm),leakage_pow(dBm),freq_accuracy(kHz),freq_drift(kHz),drift_rate(Hz/50 μs),'
        title = title + 'delta_f1_avg(kHz),delta_f1_min(kHz),delta_f1_max(kHz),delta_f2_avg(kHz),delta_f2_min(kHz),delta_f2_max(kHz),mod_ratio,delta_f2_99(kHz),'
        title = title + 'obw(kHz),frange_l(kHz),frange_h(kHz),acp_list_21ch\n'
        fname = self.get_filename('ts_bt_test/','test_tx_br_gain_tx232_{}'.format(self.board_name))
        fw1=csvreport(fname,title)
        self.mem_ts.wrm(0xa0421004, 22, 22, 1)  ##Manual control of Tx PA target power enable
        for gain in range(63,0,-1):
            self.mem_ts.wrm(0xa042100c, 31, 26, gain)  ##Manual control of Tx PA target power value
            if gain >30:
                target_power=15
            elif gain <30 and gain >10:
                target_power = 10
            elif  gain <10:
                target_power = 0
            res = self.test_br_tx(chan_list=[0], rfport=1, cable_loss=1, rate_list=['DH1'],  target_power=target_power, fig_en=0, csv_save=False, report_save=False)
            for i in range(len(res)):
                fw1.write_data([gain] + res[i])
        self.mem_ts.wrm(0xa0421004, 22, 22, 0)

    def test232_tp_gain_cal(self,mode=1, freq=2441):
        '''
        mode:
        0:  le1m mode
        other:  le2m mode

        freq:   MHz
        '''
        if self.jlink_en==0:
            if mode != 0:
                self.mem_ts.wrm(0xa0421064, 12, 12, 1)  ##Frequency value manual enable
                self.mem_ts.wrm(0xa0421064, 11, 0, freq)  ##Manual value of frequency (unit: MHz)
                self.mem_ts.wrm(0xa0421010, 3, 3, 0)
                self.mem_ts.wrm(0xa0421024, 21, 18, 0xf)
                self.mem_ts.wrm(0xa0422014, 3, 0, 0x0)
                self.mem_ts.wrm(0xa0422014, 3, 0, 0xe)
                time.sleep(0.2)
                res = self.mem_ts.rdm(0xa04210f8, 10, 0)
                tp_gian_cal_le2m = res & 0x7ff
                logdebug('tp_gian_cal_le2m: {}'.format(tp_gian_cal_le2m))
            else:
                self.mem_ts.wrm(0xa0421064, 12, 12, 1)  ##Frequency value manual enable
                self.mem_ts.wrm(0xa0421064, 11, 0, freq)  ##Manual value of frequency (unit: MHz)
                self.mem_ts.wrm(0xa0421010, 3, 3, 0)
                self.mem_ts.wrm(0xa0421024, 21, 18, 0x0)
                self.mem_ts.wrm(0xa0422014, 3, 0, 0x0)
                self.mem_ts.wrm(0xa0422014, 3, 0, 0xe)
                time.sleep(0.2)
                res = self.mem_ts.rdm(0xa04210f8, 10, 0)
                tp_gian_cal_le1m = res & 0x7ff
                logdebug('tp_gian_cal_le1m: {}'.format(tp_gian_cal_le1m))
            self.mem_ts.wrm(0xa0422014, 3, 0, 0x0)
            self.mem_ts.wrm(0xa0421010, 3, 3, 1)
            self.mem_ts.wrm(0xa0421024, 21, 18, 0x0)
            self.mem_ts.wrm(0xa0421064, 12, 12, 0)  ##Frequency value auto
        else:
            if mode !=0:
                self.jlink.wrm(0xa0421064, 12, 12, 1)  ##Frequency value manual enable
                self.jlink.wrm(0xa0421064, 11, 0, freq)  ##Manual value of frequency (unit: MHz)
                self.jlink.wrm(0xa0421010, 3, 3, 0)
                self.jlink.wrm(0xa0421024,21,18,0xf)
                self.jlink.wrm(0xa0422014,3,0,0x0)
                self.jlink.wrm(0xa0422014, 3, 0, 0xe)
                time.sleep(0.2)
                res = self.jlink.rdm(0xa04210f8,10,0)
                tp_gian_cal_le2m = res & 0x7ff
                logdebug('tp_gian_cal_le2m: {}'.format(tp_gian_cal_le2m))
            else:
                self.jlink.wrm(0xa0421064, 12, 12, 1)  ##Frequency value manual enable
                self.jlink.wrm(0xa0421064, 11, 0, freq)  ##Manual value of frequency (unit: MHz)
                self.jlink.wrm(0xa0421010, 3, 3, 0)
                self.jlink.wrm(0xa0421024,21,18,0x0)
                self.jlink.wrm(0xa0422014,3,0,0x0)
                self.jlink.wrm(0xa0422014, 3, 0, 0xe)
                time.sleep(0.2)
                res = self.jlink.rdm(0xa04210f8,10,0)
                tp_gian_cal_le1m = res & 0x7ff
                logdebug('tp_gian_cal_le1m: {}'.format(tp_gian_cal_le1m))
            self.jlink.wrm(0xa0422014, 3, 0, 0x0)
            self.jlink.wrm(0xa0421010, 3, 3, 1)
            self.jlink.wrm(0xa0421024, 21, 18, 0x0)
            self.jlink.wrm(0xa0421064, 12, 12, 0)  ##Frequency value manual enable
        return res

    def test232_tp_gain_cal_scan(self):
        title = 'mode,channel,tp_gain_cal_value\n'
        fname = self.get_filename('ts_bt_test/', 'test232_tp_gain_cal_scan_{}'.format(self.board_name))
        fw1 = csvreport(fname, title)
        for mode in [0,1]:
            if mode == 0:
                rate = 'LE1M'
            else:
                rate = 'LE2M'
            for chan in range(0,40):
                tp_gain_cal_value = self.test232_tp_gain_cal(mode=mode,freq=2402+chan*2)
                logdebug('tp_gain_cal_value:    {}'.format(tp_gain_cal_value))
                fw1.write_data([rate,chan,tp_gain_cal_value])

    def test232_dig_gain_set(self, mode=0, gfsk_gain=0x7f,dpsk_gain=0x98):
        '''
        mode:
        0:  BR/LE
        other:  EDR
        '''
        if mode != 0:
            self.jlink.wrm(0xa0420018, 7, 0, gfsk_gain)
        else:
            self.jlink.wrm(0xa0420010, 7, 0, gfsk_gain)
        self.jlink.wrm(0xa042000c, 15, 8, dpsk_gain)

    def test232_bbpll_en(self, en=1):
        if en!=0:
            if self.chipv == 'TX232_MPW3':
                self.mem_ts.wrm(0xa00000a8, 1, 0, 1)
                self.mem_ts.wrm(0xa0000144, 25, 25, 1)  ##bbpll_bg_pup_ibg_bbpll
                self.mem_ts.wrm(0xa0000144, 29, 29, 1)  ##bbpll_bg_pup_ibg_tbuf
                self.mem_ts.wrm(0xa0000144, 20, 20, 1)  ##bbpll_bg_fc on
                self.mem_ts.wrm(0xa0000144, 22, 22, 1)  ##rf_aon_xo_en_clk_ana
                self.mem_ts.wrm(0xa0000144, 20, 20, 1)  ##rf_aon_xo_en_clk_ref
                self.mem_ts.wrm(0xa0000144, 23, 23, 1)  ##bbpll_ldo_pup
                self.mem_ts.wrm(0xa0000144, 22, 22, 1)  ##bbpll_pup
                self.mem_ts.wrm(0xa0000144, 21, 21, 0)  ##bbpll_ld_pd
                self.mem_ts.wrm(0xa0000144, 18, 18, 1)  ##bbpll_ldo_fc on
                self.mem_ts.wrm(0xa0000144, 17, 17, 1)  ##bbpll_vco_fast_en on

                ##wait 40us and end bbpll fast charge
                self.mem_ts.wrm(0xa0000144, 20, 20, 0)  ##bbpll_bg_fc off
                self.mem_ts.wrm(0xa0000144, 18, 18, 0)  ##bbpll_ldo_fc off
                self.mem_ts.wrm(0xa0000144, 17, 17, 0)  ##bbpll_vco_fast_en off
                self.mem_ts.wrm(0xa0000144, 16, 16, 1)  ##bbpll_ck_dig_en
                self.mem_ts.wrm(0xa0000144, 15, 15, 1)  ##bbpll_ck_codec_en
            else:
                self.mem_ts.wrm(0xa01200a0, 25, 25, 1)  ##bbpll_bg_pup_ibg
                self.mem_ts.wrm(0xa01200a0, 29, 29, 1)  ##bbpll_bg_pup_ibg_tbuf
                self.mem_ts.wrm(0xa01200a0, 20, 20, 1)  ##bbpll_bg_fc on
                self.mem_ts.wrm(0xa01200cc, 22, 22, 1)  ##rf_aon_xo_en_clk_ana
                self.mem_ts.wrm(0xa01200cc, 20, 20, 1)  ##rf_aon_xo_en_clk_ref
                self.mem_ts.wrm(0xa01200a0, 23, 23, 1)  ##bbpll_ldo_pup
                self.mem_ts.wrm(0xa01200a0, 22, 22, 1)  ##bbpll_pup
                self.mem_ts.wrm(0xa01200a0, 21, 21, 0)  ##bbpll_ld_pd
                self.mem_ts.wrm(0xa01200a0, 18, 18, 1)  ##bbpll_ldo_fc on
                self.mem_ts.wrm(0xa01200a0, 17, 17, 1)  ##bbpll_vco_fast_en on

                ##wait 40us and end bbpll fast charge
                self.mem_ts.wrm(0xa01200a0, 20, 20, 0)  ##bbpll_bg_fc off
                self.mem_ts.wrm(0xa01200a0, 18, 18, 0)  ##bbpll_ldo_fc off
                self.mem_ts.wrm(0xa01200a0, 17, 17, 0)  ##bbpll_vco_fast_en off
                self.mem_ts.wrm(0xa01200a0, 16, 16, 1)  ##bbpll_ck_dig_en
                self.mem_ts.wrm(0xa01200a0, 15, 15, 1)  ##bbpll_ck_codec_en
        else:
            if self.chipv == 'TX232_MPW3':
                self.mem_ts.wrm(0xa0000144, 23, 23, 0)  ##bbpll_ldo_pup
                self.mem_ts.wrm(0xa0000144, 22, 22, 0)  ##bbpll_pup
            else:
                self.mem_ts.wrm(0xa01200a0, 23, 23, 0)  ##bbpll_ldo_pup
                self.mem_ts.wrm(0xa01200a0, 22, 22, 0)  ##bbpll_pup


    def test232_bbpll_diag(self, mode=0, diag_code=1):
        '''
        mode    0:  bbpll ck test
                1:  bbpll locked voltage test
                2:  bbpll ldo output test
        diag_code is useful for mode2
        '''
        self.test232_bbpll_en(1)
        if self.chipv == 'TX232_MPW3':
            self.mem_ts.wrm(0xa0000148, 0, 0, 0)  ##bbpll_ck6p144m_test_en
            self.mem_ts.wrm(0xa0000148, 3, 3, 0)  ##bbpll_ck_test_en
            self.mem_ts.wrm(0xa0000148, 1, 1, 0)  ##bbpll_vctl_test_en
            self.mem_ts.wrm(0xa0000148, 2, 2, 0)  ##bbpll_dc_test_en
            self.mem_ts.wrm(0xa0000148, 7, 7, 0)  ##bbpll_ldo_diag_sel
            self.mem_ts.wrm(0xa0000148, 6, 4, 0)  ##bbpll_diag_code
            if mode == 0:
                self.mem_ts.wrm(0xa0000148, 0, 0, 1)  ##bbpll_ck6p144m_test_en
                self.mem_ts.wrm(0xa0000148, 3, 3, 1)  ##bbpll_ck_test_en
            elif mode == 1:
                self.mem_ts.wrm(0xa0000148, 1, 1, 1)  ##bbpll_vctl_test_en
                self.mem_ts.wrm(0xa0000148, 2, 2, 1)  ##bbpll_dc_test_en
            else:
                self.mem_ts.wrm(0xa0000148, 7, 7, 1)  ##bbpll_ldo_diag_sel
                self.mem_ts.wrm(0xa0000148, 6, 4, diag_code)  ##bbpll_diag_code
        else:
            self.mem_ts.wrm(0xa01200e0, 0, 0, 0)  ##bbpll_ck6p144m_test_en
            self.mem_ts.wrm(0xa01200e0, 11, 11, 0)  ##bbpll_ck_test_en
            self.mem_ts.wrm(0xa01200e0, 1, 1, 0)  ##bbpll_vctl_test_en
            self.mem_ts.wrm(0xa01200e0, 10, 10, 0)  ##bbpll_dc_test_en
            self.mem_ts.wrm(0xa01200e0, 16, 16, 0)  ##bbpll_ldo_diag_sel
            self.mem_ts.wrm(0xa01200e0, 15, 13, 0)  ##bbpll_diag_code
            if mode == 0:
                self.mem_ts.wrm(0xa01200e0, 0, 0, 1)  ##bbpll_ck6p144m_test_en
                self.mem_ts.wrm(0xa01200e0, 11, 11, 1)  ##bbpll_ck_test_en
            elif mode == 1:
                self.mem_ts.wrm(0xa01200e0, 1, 1, 1)  ##bbpll_vctl_test_en
                self.mem_ts.wrm(0xa01200e0, 10, 10, 1)  ##bbpll_dc_test_en
            else:
                self.mem_ts.wrm(0xa01200e0, 16, 16, 1)  ##bbpll_ldo_diag_sel
                self.mem_ts.wrm(0xa01200e0, 15, 13, diag_code)  ##bbpll_diag_code

    def test232_ldo_bbpll_trim_diag(self, device_name='MY50180049'):
        title = 'bbpll_bg_trim,ldo_bbpll_trim,diag_code,volt(v)\n'
        fname = self.get_filename('ts_bt_test/', 'test232_ldo_bbpll_trim_diag_{}'.format(self.board_name))
        fw1 = csvreport(fname, title)
        mydm = dm.dm(device_name=device_name, num_of_machine=0, comm='USB')
        self.test232_bbpll_en()
        self.test232_bbpll_diag(2,1)
        if self.chipv == 'TX232_MPW3':
            init_value = self.mem_ts.rdm(0xa0000144, 14, 12)
            bg_trim_init = self.mem_ts.rdm(0xa01200c8, 7, 5)
            for bg_trim in range(8):
                self.mem_ts.wrm(0xa01200c8, 7, 5, bg_trim)
                for value in range(0, 8):
                    self.mem_ts.wrm(0xa0000144, 14, 12, value)
                    for diag_code in range(1, 2):
                    # for diag_code in range(1, 8):
                        self.mem_ts.wrm(0xa0000148, 6, 4, diag_code)
                        time.sleep(0.2)
                        dd = mydm.device.ask('MEAS:VOLT? 10, 0.0001')
                        res = eval(dd)

                        loginfo('ldo_bbpll_trim:{}     diag_code:{}       volt:{}'.format(value, diag_code, res))
                        fw1.write_data([bg_trim, value, diag_code, res], float_num=4)
            self.mem_ts.wrm(0xa0000144, 14, 12, init_value)
            self.mem_ts.wrm(0xa01200c8, 7, 5, bg_trim_init)
        else:
            init_value = self.mem_ts.rdm(0xa01200a0, 13, 12)
            bg_trim_init = self.mem_ts.rdm(0xa01200a0, 11, 9)
            for bg_trim in range(8):
                self.mem_ts.wrm(0xa01200a0, 11, 9, bg_trim)
                for value in range(0, 4):
                    self.mem_ts.wrm(0xa01200a0, 13, 12, value)
                    for diag_code in range(1, 8):
                        self.mem_ts.wrm(0xa01200e0, 15, 13, diag_code)
                        time.sleep(5)
                        dd = mydm.device.ask('MEAS:VOLT? 10, 0.0001')
                        res = eval(dd)

                        loginfo('ldo_bbpll_trim:{}     diag_code:{}       volt:{}'.format(value, diag_code, res))
                        fw1.write_data([bg_trim,value, diag_code, res], float_num=4)
            self.mem_ts.wrm(0xa01200a0, 13, 12, init_value)
            self.mem_ts.wrm(0xa01200a0, 11, 9, bg_trim_init)

    def test232_rf_diag(self, mode=0, diag_code=1):
        '''
        mode    2:  xo diag select
                3:  xo LDO diag select
                4:  Bandgap diag select
                5:  LV LDO diag select
                6:  TRXHF LDO diag select
                7:  TRXLF LDO diag select
                8:  PLL LDO diag select
                9:  VCO LDO diag select
                10:  Rx LNA diag select
                11:  Rx CBPF diag select
                12:  Tx PA diag select
                13:  Tx Mixer diag select
                14:  Tx LPF diag select
                15:  Tx DAC diag select
                16:  PLL diag select
        diag_code
        '''
        diag_dic = {
            2:  'xo diag select',
            3:  'xo LDO diag select',
            4:  'Bandgap diag select',
            5: 'LV LDO diag select',
            6: 'TRXHF LDO diag select',
            7: 'TRXLF LDO diag select',
            8: 'PLL LDO diag select',
            9: 'VCO LDO diag select',
            10: 'Rx LNA diag select',
            11: 'Rx CBPF diag select',
            12: 'Tx PA diag select',
            13: 'Tx Mixer diag select',
            14: 'Tx LPF diag select',
            15: 'Tx DAC diag select',
            16: 'PLL diag select',
            }
        self.mem_ts.wrm(0xa0421024, 17, 17, 1)  ##RF diag  soc enable
        self.mem_ts.wrm(0xa0421024, 3, 3, 0)  ##RF diag  enable
        self.mem_ts.wrm(0xa0421024, 16, 4, 0)
        self.mem_ts.wrm(0xa01200cc, 3, 2, 0)
        if mode < 4:
            self.mem_ts.wrm(0xa01200cc, mode, mode, 1)
        else:
            self.mem_ts.wrm(0xa0421024, mode, mode, 1)
        self.mem_ts.wrm(0xa0421024, 2, 0, diag_code)  ##RF diag code select
        loginfo("{}     diag code:  {}".format(diag_dic[mode],diag_code))

    def test232_pkd_trim_diag(self, device_name='MY49260023',cableloss=1.5):
        title = 'pkd_ref,VDD,VREF(V),VREF1(V),VREF2(V)\n'
        fname = self.get_filename('ts_bt_test/', 'test232_pkd_trim_diag{}'.format(self.board_name))
        fw1 = csvreport(fname, title)
        mydm = dm.dm(device_name=device_name, num_of_machine=0, comm='USB')
        init_value = self.mem_ts.rdm(0xa0421028, 25, 19)

        self.mem_ts.wrm(0xa0421024, 3, 3, 0)
        self.mem_ts.wrm(0xa0421024, 17, 17, 1)
        self.mem_ts.wrm(0xa0421028, 31, 31, 1)      ##diag select pkd
        self.mem_ts.wrm(0xa0422014, 7,0,0x3)             ##rxon enable

        for value in range(0,2**7):
            self.mem_ts.wrm(0xa0421028, 25, 19, value)
            res_list=[]
            for diag_code in [2,4,5,6]:
                self.mem_ts.wrm(0xa0421024, 2, 0, diag_code)
                time.sleep(2)
                dd = mydm.device.ask('MEAS:VOLT? 10, 0.0001')
                res = eval(dd)
                res_list.append(res)

                loginfo('pkd_ref:{}     diag_code:{}       volt:{}'.format(value, diag_code, res))
            fw1.write_data([value]+res_list, float_num=4)
        self.mem_ts.wrm(0xa0421028, 25, 19, init_value)

        fw1.write_string('pkd_ref,rx_gain_index,input_level(dbm),VIN(V)\n')
        txg = mxg.MXG()
        # txg.arb_waveform(rate='LE_1M')
        txg.trriger_para_set(type='CONTinuous', count=1000)
        txg.output_state(1, 0)
        # txg.arb_state(1)
        txg.set_cfreq(2475)
        self.mem_ts.wrm(0xa0421064,12,12,1)
        self.mem_ts.wrm(0xa0421064, 11, 0, 2475)    ##set freq
        self.mem_ts.wrm(0xa04210A0, 0, 0, 0x1)  ##AGC OFF
        pkd_ref = self.mem_ts.rdm(0xa0421028, 25, 19)
        rx_gain_index = self.mem_ts.rdm(0xa04210a0, 20, 17)
        self.mem_ts.wrm(0xa0421024, 2, 0, 3)
        for rxlevel in range(-990,10,5):
            self.mem_ts.wrm(0xa0422014, 7,0,0x0)
            txg.set_power(rxlevel/10.0+cableloss)
            time.sleep(0.2)
            self.mem_ts.wrm(0xa0422014, 7,0,0x3)
            dd = mydm.device.ask('MEAS:VOLT? 10, 0.0001')
            res = eval(dd)
            loginfo('pkd_ref:{}     rx_gain_index:{}   rxlevel:{}    VIN:{}'.format(pkd_ref, rx_gain_index,rxlevel, res))
            fw1.write_data([pkd_ref,rx_gain_index,rxlevel,res], float_num=4)

        self.mem_ts.wrm(0xa04210A0, 0, 0, 0x0)  ##AGC ON
        self.mem_ts.wrm(0xa0421064, 12, 12, 0)  ##freq auto
        self.mem_ts.wrm(0xa0422014, 7,0,0x0)


    def test232_bg_trim_diag(self):
        title = 'bg_trim,diag_code,volt(v)\n'
        fname = self.get_filename('ts_bt_test/', 'test232_bg_trim_diag_{}'.format(self.board_name))
        fw1 = csvreport(fname, title)
        mydm = dm.dm(device_name='MY50180049', num_of_machine=0, comm='USB')
        self.test232_rf_diag(4,1)
        init_value = self.mem_ts.rdm(0xa0421028, 2, 0)
        self.mem_ts.wr(0xa0422014, 0x3)
        for value in range(0,8):
            self.mem_ts.wrm(0xa0421028, 2, 0, value)
            for diag_code in range(1,8):
                self.mem_ts.wrm(0xa0421024, 2, 0, diag_code)
                time.sleep(5)
                dd = mydm.device.ask('MEAS:VOLT? 10, 0.0001')
                res = eval(dd)

                loginfo('bg trim:{}     diag_code:{}       volt:{}'.format(value, diag_code, res))
                fw1.write_data([value, diag_code, res], float_num=4)
        self.mem_ts.wrm(0xa0421028, 2, 0, init_value)
        self.mem_ts.wr(0xa0422014, 0x0)

    def test232_ldo_trxhf_trim_diag(self):
        title = 'ldo_trxhf_trim,diag_code,volt(v)\n'
        fname = self.get_filename('ts_bt_test/', 'test232_ldo_trxhf_trim_diag_{}'.format(self.board_name))
        fw1 = csvreport(fname, title)
        mydm = dm.dm(device_name='MY50180049', num_of_machine=0, comm='USB')
        self.test232_rf_diag(6, 1)
        init_value = self.mem_ts.rdm(0xa0421028, 7, 6)
        self.mem_ts.wr(0xa0422014, 0x3)
        for value in range(0,4):
            self.mem_ts.wrm(0xa0421028, 7, 6, value)
            for diag_code in range(1,8):
                self.mem_ts.wrm(0xa0421024, 2, 0, diag_code)
                time.sleep(5)
                dd = mydm.device.ask('MEAS:VOLT? 10, 0.0001')
                res = eval(dd)

                loginfo('ldo_trxhf_trim:{}     diag_code:{}       volt:{}'.format(value, diag_code, res))
                fw1.write_data([value, diag_code, res], float_num=4)
        self.mem_ts.wrm(0xa0421028, 7, 6, init_value)
        self.mem_ts.wr(0xa0422014, 0x0)

    def test232_ldo_trxlf_trim_diag(self):
        title = 'ldo_trxlf_trim,diag_code,volt(v)\n'
        fname = self.get_filename('ts_bt_test/', 'test232_ldo_trxlf_trim_diag_{}'.format(self.board_name))
        fw1 = csvreport(fname, title)
        mydm = dm.dm(device_name='MY50180049', num_of_machine=0, comm='USB')
        self.test232_rf_diag(7, 1)
        init_value = self.mem_ts.rdm(0xa0421034, 26, 24)
        self.mem_ts.wr(0xa0422014, 0x3)
        for value in [0,1,3,7]:
            self.mem_ts.wrm(0xa0421034, 26, 24, value)
            for diag_code in [1]:
                self.mem_ts.wrm(0xa0421024, 2, 0, diag_code)
                time.sleep(5)
                dd = mydm.device.ask('MEAS:VOLT? 10, 0.0001')
                res = eval(dd)

                loginfo('ldo_trxlf_trim:{}     diag_code:{}       volt:{}'.format(value, diag_code, res))
                fw1.write_data([value, diag_code, res], float_num=4)
        self.mem_ts.wrm(0xa0421034, 26, 24, init_value)
        self.mem_ts.wr(0xa0422014, 0x0)

    def test232_ldo_rfpll_trim_diag(self):
        title = 'ldo_rfpll_trim,diag_code,volt(v)\n'
        fname = self.get_filename('ts_bt_test/', 'test232_ldo_rfpll_trim_diag_{}'.format(self.board_name))
        fw1 = csvreport(fname, title)
        mydm = dm.dm(device_name='MY50180049', num_of_machine=0, comm='USB')
        self.test232_rf_diag(8, 1)
        init_value = self.mem_ts.rdm(0xa0421028, 12, 11)
        self.mem_ts.wr(0xa0422014, 0x3)
        for value in range(0,4):
            self.mem_ts.wrm(0xa0421028, 12, 11, value)
            for diag_code in range(1,8):
                self.mem_ts.wrm(0xa0421024, 2, 0, diag_code)
                time.sleep(5)
                dd = mydm.device.ask('MEAS:VOLT? 10, 0.0001')
                res = eval(dd)

                loginfo('ldo_rfpll_trim:{}     diag_code:{}       volt:{}'.format(value, diag_code, res))
                fw1.write_data([value, diag_code, res], float_num=4)
        self.mem_ts.wrm(0xa0421028, 12, 11, init_value)
        self.mem_ts.wr(0xa0422014, 0x0)

    def test232_ldo_vco_trim_diag(self):
        title = 'ldo_vco_trim,diag_code,volt(v)\n'
        fname = self.get_filename('ts_bt_test/', 'test232_ldo_vco_trim_diag_{}'.format(self.board_name))
        fw1 = csvreport(fname, title)
        mydm = dm.dm(device_name='MY50180049', num_of_machine=0, comm='USB')
        self.test232_rf_diag(9, 1)
        init_value = self.mem_ts.rdm(0xa0421028, 15, 14)
        self.mem_ts.wr(0xa0422014, 0x3)
        for value in range(0,4):
            self.mem_ts.wrm(0xa0421028, 15, 14, value)
            for diag_code in range(1,8):
                self.mem_ts.wrm(0xa0421024, 2, 0, diag_code)
                time.sleep(5)
                dd = mydm.device.ask('MEAS:VOLT? 10, 0.0001')
                res = eval(dd)

                loginfo('ldo_vco_trim:{}     diag_code:{}       volt:{}'.format(value, diag_code, res))
                fw1.write_data([value, diag_code, res], float_num=4)
        self.mem_ts.wrm(0xa0421028, 15, 14, init_value)
        self.mem_ts.wr(0xa0422014, 0x0)

    def test232_ldo_xo_trim_diag(self, device_name='MY50180049'):
        '''
        TX232_MPW3 xo使用bbpll bg，MPW2使用的是rfpll bg
        '''
        title = 'bg_trim,ldo_xo_trim,diag_code,volt(v)\n'
        fname = self.get_filename('ts_bt_test/', 'test232_ldo_xo_trim_diag_{}'.format(self.board_name))
        fw1 = csvreport(fname, title)
        mydm = dm.dm(device_name=device_name, num_of_machine=0, comm='USB')
        self.test232_rf_diag(3, 1)
        init_value = self.mem_ts.rdm(0xa01200cc, 5, 4)
        if self.chipv == 'TX232_MPW3':
            bg_trim_init = self.mem_ts.rdm(0xa01200c8, 7, 5)
        else:
            bg_trim_init = self.mem_ts.rdm(0xa0421028, 2, 0)
        self.mem_ts.wr(0xa0422014, 0x3)
        for bg_trim in range(8):
            if self.chipv == 'TX232_MPW3':
                self.mem_ts.wrm(0xa01200c8, 7, 5, bg_trim)
            else:
                self.mem_ts.wrm(0xa0421028, 2, 0, bg_trim)
            for value in range(0,4):
                self.mem_ts.wrm(0xa01200cc, 5, 4, value)
                for diag_code in range(1,6):
                    self.mem_ts.wrm(0xa0421024, 2, 0, diag_code)
                    time.sleep(0.2)
                    dd = mydm.device.ask('MEAS:VOLT? 10, 0.001')
                    res = eval(dd)

                    loginfo('ldo xo trim:{}     diag_code:{}       volt:{}'.format(value, diag_code, res))
                    fw1.write_data([bg_trim, value, diag_code, res], float_num=4)
        self.mem_ts.wrm(0xa01200cc, 5, 4, init_value)
        if self.chipv == 'TX232_MPW3':
            bg_trim_init = self.mem_ts.wrm(0xa01200c8, 7, 5, bg_trim_init)
        else:
            bg_trim_init = self.mem_ts.wrm(0xa0421028, 2, 0, bg_trim_init)
        self.mem_ts.wr(0xa0422014, 0x0)

    def test232_rf_ldo_trim_diag(self, device_name='MY50180049'):
        ldo_diag_dic = {
                    # 3:  'LDO_XO',
                    # 4:  'Bandgap',
                    5:  'LDO_LV',
                    6:  'LDO_TRXHF',
                    7:  'LDO_TRXLF',
                    8:  'LDO_PLL',
                    9:  'LDO_VCO',
            }
        ldo_trim_dic = {
            # 3: '0xa01200cc 5 4',
            # 4: '0xa0421028 2 0',
            5: '0xa0421028 4 4',
            6: '0xa0421028 7 6',
            7: '0xa0421034 26 24',
            8: '0xa0421028 12 11',
            9: '0xa0421028 15 14',
            }
        title = 'item,bg_trim,trim value,diag_code,volt(v)\n'
        fname = self.get_filename('ts_bt_test/', 'test232_rf_ldo_trim_diag_{}'.format(self.board_name))
        fw1 = csvreport(fname, title)
        mydm = dm.dm(device_name=device_name, num_of_machine=0, comm='USB')
        bg_trim_init = self.mem_ts.rdm(0xa0421028,2,0)
        self.mem_ts.wr(0xa0422014, 0x3)
        for val in range(5,10):
            self.test232_rf_diag(val, 1)
            ldo_type = ldo_diag_dic[val]
            ldo_trim_str = ldo_trim_dic[val]
            ldo_trim_strr = ldo_trim_str.split()
            ldo_trim_reg_addr = eval(ldo_trim_strr[0])
            ldo_trim_reg_msb = eval(ldo_trim_strr[1])
            ldo_trim_reg_lsb = eval(ldo_trim_strr[2])
            init_value = self.mem_ts.rdm(ldo_trim_reg_addr, ldo_trim_reg_msb, ldo_trim_reg_lsb)
            for bg_trim in range(8):
                self.mem_ts.wrm(0xa0421028, 2, 0, bg_trim)
                for trim_value in range(0,2**(ldo_trim_reg_msb-ldo_trim_reg_lsb+1)):
                    self.mem_ts.wrm(ldo_trim_reg_addr, ldo_trim_reg_msb, ldo_trim_reg_lsb, trim_value)
                    for diag_code in range(1, 2):
                    # for diag_code in range(1, 8):
                        self.mem_ts.wrm(0xa0421024, 2, 0, diag_code)
                        time.sleep(0.2)
                        dd = mydm.device.ask('MEAS:VOLT? 10, 0.001')
                        res = eval(dd)

                        loginfo('{} trim     trim_value:{}       volt:{}'.format(ldo_type, trim_value, res))
                        fw1.write_data([ldo_type, bg_trim, trim_value, diag_code, res], float_num=4)
            self.mem_ts.wrm(ldo_trim_reg_addr, ldo_trim_reg_msb, ldo_trim_reg_lsb, init_value)
        self.mem_ts.wr(0xa0422014, 0x0)

        # self.test232_ldo_bbpll_trim_diag(device_name=device_name)
        self.test232_ldo_xo_trim_diag(device_name=device_name)


    def test232_rfpll_vtune(self):
        title = 'freq(MHz),rfpll_vtune(v)\n'
        fname = self.get_filename('ts_bt_test/', 'test232_rfpll_vtune')
        fw1 = csvreport(fname, title)
        mydm=dm.dm(device_name='MY50180049', num_of_machine=0, comm='USB')
        # self.mem_ts.wrm(0xa0421024, 17, 17, 1)  ##RF diag  soc enable
        # self.mem_ts.wrm(0xa0421024, 16, 16, 1)  ##RF PLL diag select
        # self.mem_ts.wrm(0xa0421024, 3, 3, 0)  ##RF diag  enable
        # self.mem_ts.wrm(0xa0421024, 2, 0, 1)  ##RF diag code select
        self.test232_rf_diag(16,1)
        self.mem_ts.wrm(0xa0421064, 12, 12, 1)  ##Frequency value manual enable
        for channel in range(0,79):
            self.mem_ts.wrm(0xa0421064, 11, 0, 2402+channel)
            self.mem_ts.wr(0xa0422014, 0x0)
            time.sleep(2)
            self.mem_ts.wr(0xa0422014, 0xe)
            time.sleep(5)
            dd = mydm.device.ask('MEAS:VOLT?')
            res = eval(dd)

            loginfo('channel:{}     vtune:{}'.format(channel,res))
            fw1.write_data([channel,res],float_num=4)

    def test232_rf_buf(self, mode='rx'):
        self.mem_ts.wrm(0xa0120038, 11, 8, 0)
        self.mem_ts.wrm(0xa012005c, 31, 16, 0)  ##
        self.mem_ts.wrm(0xa0120060, 7, 0, 0)  ##PC8/9/10/11 高阻
        self.mem_ts.wrm(0xa01200a0, 29, 29, 1)  ##bbpll_bg_pup_ibg_tbuf
        self.mem_ts.wrm(0xa01200c8, 4, 4, 1)  ##test buf en
        if mode == 'tx':
            self.mem_ts.wrm(0xa01200c8,3,2,2)	##tx IF test buf
        else:
            self.mem_ts.wrm(0xa01200c8, 3, 2, 1)  ##rx IF test buf

    def test232_cfo_est_test(self,chan_list=[0],signal_freq_oft=400):
        title = 'channel,signal_freq(khz),cfo_est, cfo_est_step\n'
        fname = self.get_filename('ts_bt_test/','test232_cfo_est_test_{}'.format(self.board_name))
        fw1 = csvreport(fname,title)
        txg = mxg.MXG()
        txg.arb_waveform(rate='LE_1M')
        txg.arb_state(1)
        txg.trriger_para_set(type='CONTinuous', count=1000)
        txg.output_state(1, 1)
        txg.set_power(-40)
        for chan in chan_list:
            self.cmdstop(0)
            time.sleep(0.5)
            self.LE_RX(chan=chan,phy=1,mod=0)
            signal_freq_base = 2402+2*chan
            cfo_est_init = 0
            for freq_oft in range(-signal_freq_oft,signal_freq_oft+1):
                while 1:
                    recv_pkt1 = self.jlink.rdm(0xa04008d8, 15, 0)
                    time.sleep(0.1)
                    recv_pkt2 = self.jlink.rdm(0xa04008d8, 15, 0)
                    recv_pkt = recv_pkt2 - recv_pkt1
                    if recv_pkt != 0:
                        freq = signal_freq_base+freq_oft/1000.000
                        txg.set_cfreq(freq)
                        time.sleep(0.5)
                        cfo_est = self.jlink.rdm(0xa04202f0,15,0)
                        if cfo_est>>15 ==1:
                            cfo_est = cfo_est - 65535
                        cfo_est_step = cfo_est - cfo_est_init
                        cfo_est_init = cfo_est
                        loginfo('chan:{}    signal_freq:{}      cfo_est:{}      step:{}'.format(chan, freq, cfo_est, cfo_est_step))
                        fw1.write_data([chan,freq,cfo_est,cfo_est_step])
                        break
                    else:
                        self.cmdstop(0)
                        time.sleep(0.5)
                        self.LE_RX(chan=chan, phy=1, mod=0)


    def test232_rxdcoc_code_rd(self):
        res1_i = self.mem_ts.rdm(0xa04210f4,11,6)
        res1_q = self.mem_ts.rdm(0xa04210f4, 17, 12)
        res2_i = self.mem_ts.rdm(0xa04210f4,23,18)
        res2_q = self.mem_ts.rdm(0xa04210f4, 29, 24)
        loginfo("rxdcoc1_i_code: {}     rxdcoc1_q_code: {}     rxdcoc2_i_code: {}     rxdcoc2_q_code: {}     ".format(res1_i,res1_q,res2_i,res2_q))
        return [res1_i,res1_q,res2_i,res2_q]

    def test232_rxdcoc_volt_rd(self):
        mydm = dm.dm(device_name='MY50180049', num_of_machine=0, comm='USB')
        self.mem_ts.wrm(0xa0421024, 17, 17, 1)  ##RF diag  soc enable
        self.mem_ts.wrm(0xa0421024, 3, 3, 0)  ##RF diag  enable
        self.mem_ts.wrm(0xa0421024, 11, 11, 1)

        self.mem_ts.wr(0xa0422014, 0x0)
        time.sleep(0.5)
        self.mem_ts.wr(0xa0422014, 0x3)
        time.sleep(0.5)
        self.mem_ts.wrm(0xa0421024, 2, 0, 1)  ##RF diag code select
        res_qn = eval(mydm.device.ask('MEAS:VOLT?'))
        time.sleep(0.5)
        self.mem_ts.wrm(0xa0421024, 2, 0, 2)  ##RF diag code select
        res_qp = eval(mydm.device.ask('MEAS:VOLT?'))
        time.sleep(0.5)
        self.mem_ts.wrm(0xa0421024, 2, 0, 3)  ##RF diag code select
        res_in = eval(mydm.device.ask('MEAS:VOLT?'))
        time.sleep(0.5)
        self.mem_ts.wrm(0xa0421024, 2, 0, 4)  ##RF diag code select
        res_ip = eval(mydm.device.ask('MEAS:VOLT?'))

        loginfo("rxdcoc_qn: {}     rxdcoc_qp: {}     rxdcoc_in: {}     rxdcoc_ip: {}     ".format(res_qn,res_qp,res_in,res_ip))
        return [res_qn,res_qp,res_in,res_ip]

    def test232_rxdcoc_offset_wr(self,i_offset=0,q_offset=0):
        if self.jlink_en != 0:
            self.jlink.wrm(0xa0421058, 12, 12, 1)
            self.jlink.wrm(0xa0421010, 6, 5, 3)
            self.jlink.wrm(0xa0421058, 18, 13, i_offset)
            self.jlink.wrm(0xa0421058, 24, 19, q_offset)
        else:
            self.mem_ts.wrm(0xa0421058,12,12,1)
            self.mem_ts.wrm(0xa0421010, 6, 5, 3)
            self.mem_ts.wrm(0xa0421058, 18, 13, i_offset)
            self.mem_ts.wrm(0xa0421058, 24, 19, q_offset)
        loginfo("wr rxdcoc offset:  i_offset {}     q_offset {}".format(i_offset,q_offset))

    def test232_rxdcoc_cal_bypass(self,en=1):
        self.mem_ts.wrm(0xa0421058, 12, 12, 0)
        if en == 0:
            self.mem_ts.wrm(0xa0421010, 6, 5, 0)
        else:
            self.mem_ts.wrm(0xa0421010, 6, 5, 3)

    def test232_rxdcoc_scan_lna_gain(self, device_name='MY50180049',idac_code=30,qdac_code=28):
        title = 'idac_code,qdac_code,hgain_code,lgain_code,atten,rxdcoc_qn(v),rxdcoc_qp(v),rxdcoc_in(v),rxdcoc_ip(v)\n'
        fname = self.get_filename('ts_bt_test/', 'test232_rxdcoc_scan_lna_gain_{}'.format(self.board_name))
        fw1 = csvreport(fname, title)
        mydm = dm.dm(device_name=device_name, num_of_machine=0, comm='USB')
        self.mem_ts.wrm(0xa0421024, 17, 17, 1)  ##RF diag  soc enable
        self.mem_ts.wrm(0xa0421024, 3, 3, 0)  ##RF diag  enable
        self.mem_ts.wrm(0xa0421024, 11, 11, 1)
        self.mem_ts.wrm(0xa04210a0, 0, 0, 1)  ##agc off
        self.mem_ts.wrm(0xa04210a0, 20, 17, 0xc)  ##agc index set 0
        self.mem_ts.wrm(0xa04210cc, 14, 13, 3)  ##cbpf gain1 on
        self.test232_rxdcoc_cal_bypass(1)
        self.test232_rxdcoc_offset_wr(idac_code, qdac_code)
        self.mem_ts.wrm(0xa0422014, 3, 0, 3)

        for hgain_code in [0x3f,0x20,0x10,0x8,0x4,0x2]:
            self.mem_ts.wrm(0xa04210b4, 24, 19, hgain_code)
            lgain_code = self.mem_ts.rdm(0xa04210b4, 18, 17)
            atten = self.mem_ts.rdm(0xa04210b4, 16, 15)
            self.mem_ts.wrm(0xa0421024, 2, 0, 1)  ##RF diag code select
            res_qn = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 2)  ##RF diag code select
            res_qp = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 3)  ##RF diag code select
            res_in = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 4)  ##RF diag code select
            res_ip = eval(mydm.device.ask('MEAS:VOLT?'))

            loginfo("rxdcoc_qn: {}     rxdcoc_qp: {}     rxdcoc_in: {}     rxdcoc_ip: {}     ".format(res_qn, res_qp, res_in, res_ip))
            rxdcoc_qn = res_qn
            rxdcoc_qp = res_qp
            rxdcoc_in = res_in
            rxdcoc_ip = res_ip

            fw1.write_data([idac_code, qdac_code, hgain_code, lgain_code, atten, rxdcoc_qn, rxdcoc_qp, rxdcoc_in, rxdcoc_ip ])
        for lgain_code in range(3,-1,-1):
            self.mem_ts.wrm(0xa04210b4, 25, 25, 0)
            self.mem_ts.wrm(0xa04210b4, 18, 17, lgain_code)
            hgain_code = self.mem_ts.rdm(0xa04210b4, 24, 19)
            atten = self.mem_ts.rdm(0xa04210b4, 16, 15)

            self.mem_ts.wrm(0xa0421024, 2, 0, 1)  ##RF diag code select
            res_qn = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 2)  ##RF diag code select
            res_qp = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 3)  ##RF diag code select
            res_in = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 4)  ##RF diag code select
            res_ip = eval(mydm.device.ask('MEAS:VOLT?'))

            loginfo("rxdcoc_qn: {}     rxdcoc_qp: {}     rxdcoc_in: {}     rxdcoc_ip: {}     ".format(res_qn, res_qp, res_in, res_ip))
            rxdcoc_qn = res_qn
            rxdcoc_qp = res_qp
            rxdcoc_in = res_in
            rxdcoc_ip = res_ip

            fw1.write_data([idac_code, qdac_code, hgain_code, lgain_code, atten, rxdcoc_qn, rxdcoc_qp, rxdcoc_in, rxdcoc_ip ])

        for atten in range(4):
            self.mem_ts.wrm(0xa04210b4, 16, 15, atten)
            hgain_code = self.mem_ts.rdm(0xa04210b4, 24, 19)
            lgain_code = self.mem_ts.rdm(0xa04210b4, 18, 17)

            # res = self.test232_rxdcoc_volt_rd()
            self.mem_ts.wrm(0xa0421024, 2, 0, 1)  ##RF diag code select
            res_qn = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 2)  ##RF diag code select
            res_qp = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 3)  ##RF diag code select
            res_in = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 4)  ##RF diag code select
            res_ip = eval(mydm.device.ask('MEAS:VOLT?'))

            loginfo("rxdcoc_qn: {}     rxdcoc_qp: {}     rxdcoc_in: {}     rxdcoc_ip: {}     ".format(res_qn, res_qp, res_in, res_ip))
            rxdcoc_qn = res_qn
            rxdcoc_qp = res_qp
            rxdcoc_in = res_in
            rxdcoc_ip = res_ip

            fw1.write_data([idac_code, qdac_code, hgain_code, lgain_code, atten, rxdcoc_qn, rxdcoc_qp, rxdcoc_in, rxdcoc_ip ])
        self.mem_ts.wrm(0xa04210b4, 16, 15, 0)
        self.mem_ts.wrm(0xa04210b4, 25, 17, 0x1ff)
    def test232_rxdcoc_cal_diff(self, device_name='MY50180049',board_name=''):
        mydm = dm.dm(device_name=device_name, num_of_machine=0, comm='USB')
        self.mem_ts.wrm(0xa0421024, 17, 17, 1)  ##RF diag  soc enable
        self.mem_ts.wrm(0xa0421024, 3, 3, 0)  ##RF diag  enable
        self.mem_ts.wrm(0xa0421024, 11, 11, 1)
        self.mem_ts.wrm(0xa04210a0, 0, 0, 1)  ##agc off
        self.mem_ts.wrm(0xa04210a0, 20, 17, 0)  ##agc index set 0
        self.mem_ts.wrm(0xa04210cc, 14, 13, 3)  ##cbpf gain1&2 on

        title = 'cbpf2_en,cbpf1_idac_code,cbpf1_qdac_code,cbpf2_idac_code,cbpf2_qdac_code,rxdcoc_qn(v),rxdcoc_qp(v),rxdcoc_in(v),rxdcoc_ip(v)\n'
        fname = self.get_filename('ts_bt_test/', 'test232_rxdcoc_cal_diff_{}'.format(board_name))
        fw1 = csvreport(fname, title)
        for cbpf_en in [0,1]:
            self.mem_ts.wrm(0xa0422014, 3, 0, 0)
            self.mem_ts.wrm(0xa0421020, 20, 20, cbpf_en)
            self.mem_ts.wrm(0xa0421010, 6, 5, 3)
            self.mem_ts.wrm(0xa0421058, 12, 12, 1)  ##cbpf1 rxdcoc code manual
            self.mem_ts.wrm(0xa042105c, 12, 12, 1)  ##cbpf2 rxdcoc code manual
            self.mem_ts.wrm(0xa042105c, 18, 13, 32)
            self.mem_ts.wrm(0xa042105c, 24, 19, 32)
            self.mem_ts.wrm(0xa0421058, 18, 13, 32)
            self.mem_ts.wrm(0xa0421058, 24, 19, 32)
            self.mem_ts.wrm(0xa0422014, 3, 0, 3)

            self.mem_ts.wrm(0xa0421024, 2, 0, 1)  ##RF diag code select
            res_qn = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 2)  ##RF diag code select
            res_qp = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 3)  ##RF diag code select
            res_in = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 4)  ##RF diag code select
            res_ip = eval(mydm.device.ask('MEAS:VOLT?'))

            loginfo("rxdcoc_qn: {}     rxdcoc_qp: {}     rxdcoc_in: {}     rxdcoc_ip: {}     ".format(res_qn, res_qp, res_in, res_ip))
            rxdcoc_qn = res_qn
            rxdcoc_qp = res_qp
            rxdcoc_in = res_in
            rxdcoc_ip = res_ip

            fw1.write_data([cbpf_en, 32, 32, 32, 32, rxdcoc_qn, rxdcoc_qp, rxdcoc_in, rxdcoc_ip])

            self.mem_ts.wrm(0xa0421058, 12, 12, 0)  ##cbpf1 rxdcoc code auto
            self.mem_ts.wrm(0xa042105c, 12, 12, 0)  ##cbpf2 rxdcoc code auto
            time.sleep(2)
            self.mem_ts.wrm(0xa0421010, 6, 5, 0)  ##rxdcoc cal bypass disable
            self.mem_ts.wrm(0xa0422014, 3, 0, 0)
            self.mem_ts.wrm(0xa0422014, 3, 0, 3)
            time.sleep(0.5)
            ro_rxdcoc1_q_code = self.mem_ts.rdm(0xa04210f4,17,12)
            ro_rxdcoc1_i_code = self.mem_ts.rdm(0xa04210f4,11,6)
            ro_rxdcoc2_q_code = self.mem_ts.rdm(0xa04210f4,23,18)
            ro_rxdcoc2_i_code = self.mem_ts.rdm(0xa04210f4,29,24)
            self.mem_ts.wrm(0xa0421024, 2, 0, 1)  ##RF diag code select
            res_qn = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 2)  ##RF diag code select
            res_qp = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 3)  ##RF diag code select
            res_in = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 4)  ##RF diag code select
            res_ip = eval(mydm.device.ask('MEAS:VOLT?'))

            loginfo("rxdcoc_qn: {}     rxdcoc_qp: {}     rxdcoc_in: {}     rxdcoc_ip: {}     ".format(res_qn, res_qp, res_in, res_ip))
            rxdcoc_qn = res_qn
            rxdcoc_qp = res_qp
            rxdcoc_in = res_in
            rxdcoc_ip = res_ip
            fw1.write_data([cbpf_en, ro_rxdcoc1_i_code, ro_rxdcoc1_q_code, ro_rxdcoc2_i_code, ro_rxdcoc2_q_code, rxdcoc_qn, rxdcoc_qp, rxdcoc_in, rxdcoc_ip])

    def test232_rxdcoc_cal_diff_loop(self, device_name='MY50180049',board_name='',loop=100):
        mydm = dm.dm(device_name=device_name, num_of_machine=0, comm='USB')
        self.mem_ts.wrm(0xa0421024, 17, 17, 1)  ##RF diag  soc enable
        self.mem_ts.wrm(0xa0421024, 3, 3, 0)  ##RF diag  enable
        self.mem_ts.wrm(0xa0421024, 11, 11, 1)
        self.mem_ts.wrm(0xa04210a0, 0, 0, 1)  ##agc off
        self.mem_ts.wrm(0xa04210a0, 20, 17, 0)  ##agc index set 0
        self.mem_ts.wrm(0xa04210cc, 14, 13, 3)  ##cbpf gain1&2 on
        self.mem_ts.wrm(0xa0421020, 20, 20, 1)

        title = 'cbpf2_en,cbpf1_idac_code,cbpf1_qdac_code,cbpf2_idac_code,cbpf2_qdac_code,rxdcoc_qn(v),rxdcoc_qp(v),rxdcoc_in(v),rxdcoc_ip(v)\n'
        fname = self.get_filename('ts_bt_test/', 'test232_rxdcoc_cal_diff_{}'.format(board_name))
        fw1 = csvreport(fname, title)
        for i in range(loop):
            # self.mem_ts.wrm(0xa0421010, 6, 5, 3)
            # self.mem_ts.wrm(0xa0421058, 12, 12, 1)  ##cbpf1 rxdcoc code manual
            # self.mem_ts.wrm(0xa042105c, 12, 12, 1)  ##cbpf2 rxdcoc code manual
            # self.mem_ts.wrm(0xa042105c, 18, 13, 32)
            # self.mem_ts.wrm(0xa042105c, 24, 19, 32)
            # self.mem_ts.wrm(0xa0421058, 18, 13, 32)
            # self.mem_ts.wrm(0xa0421058, 24, 19, 32)
            self.mem_ts.wrm(0xa0422014, 3, 0, 0)
            self.mem_ts.wrm(0xa0421058, 12, 12, 0)  ##cbpf1 rxdcoc code auto
            self.mem_ts.wrm(0xa042105c, 12, 12, 0)  ##cbpf2 rxdcoc code auto
            time.sleep(0.1)
            # self.mem_ts.wrm(0xa0421010, 6, 5, 0)  ##rxdcoc cal bypass disable
            self.mem_ts.wrm(0xa0422014, 3, 0, 3)
            time.sleep(0.5)
            ro_rxdcoc1_q_code = self.mem_ts.rdm(0xa04210f4,17,12)
            ro_rxdcoc1_i_code = self.mem_ts.rdm(0xa04210f4,11,6)
            ro_rxdcoc2_q_code = self.mem_ts.rdm(0xa04210f4,23,18)
            ro_rxdcoc2_i_code = self.mem_ts.rdm(0xa04210f4,29,24)
            self.mem_ts.wrm(0xa0421024, 2, 0, 1)  ##RF diag code select
            res_qn = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 2)  ##RF diag code select
            res_qp = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 3)  ##RF diag code select
            res_in = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 4)  ##RF diag code select
            res_ip = eval(mydm.device.ask('MEAS:VOLT?'))

            loginfo("rxdcoc_qn: {}     rxdcoc_qp: {}     rxdcoc_in: {}     rxdcoc_ip: {}     ".format(res_qn, res_qp, res_in, res_ip))
            rxdcoc_qn = res_qn
            rxdcoc_qp = res_qp
            rxdcoc_in = res_in
            rxdcoc_ip = res_ip
            fw1.write_data([1, ro_rxdcoc1_i_code, ro_rxdcoc1_q_code, ro_rxdcoc2_i_code, ro_rxdcoc2_q_code, rxdcoc_qn, rxdcoc_qp, rxdcoc_in, rxdcoc_ip])


    def test232_rxdcoc_scan(self, device_name='MY50180049', cbpf1_en=1, cbpf2_en=1):

        mydm = dm.dm(device_name=device_name, num_of_machine=0, comm='USB')
        self.mem_ts.wrm(0xa0421024, 17, 17, 1)  ##RF diag  soc enable
        self.mem_ts.wrm(0xa0421024, 3, 3, 0)  ##RF diag  enable
        self.mem_ts.wrm(0xa0421024, 11, 11, 1)
        self.mem_ts.wrm(0xa04210a0, 0, 0, 1)  ##agc off
        self.mem_ts.wrm(0xa04210a0, 20, 17, 0)  ##agc index set 0
        self.mem_ts.wrm(0xa04210cc, 14, 13, 2)  ##cbpf gain1 on
        self.test232_rxdcoc_cal_bypass(1)
        self.test232_rxdcoc_offset_wr(32, 32)
        self.mem_ts.wrm(0xa0422014, 3, 0, 3)

        if cbpf1_en!=0:
            title = 'idac_code,qdac_code,rxdcoc_qn(v),rxdcoc_qp(v),rxdcoc_in(v),rxdcoc_ip(v)\n'
            fname = self.get_filename('ts_bt_test/', 'test232_rxdcoc_scan_cbpf1_{}'.format(self.board_name))
            fw1 = csvreport(fname, title)
            self.mem_ts.wrm(0xa0421020, 20, 20, 0)  ##cbpf2 disable
            for idac_code in range(64):
                for qdac_code in range(64):
                    self.test232_rxdcoc_offset_wr(idac_code, qdac_code)
                    # res = self.test232_rxdcoc_volt_rd()
                    self.mem_ts.wrm(0xa0421024, 2, 0, 1)  ##RF diag code select
                    res_qn = eval(mydm.device.ask('MEAS:VOLT?'))

                    self.mem_ts.wrm(0xa0421024, 2, 0, 2)  ##RF diag code select
                    res_qp = eval(mydm.device.ask('MEAS:VOLT?'))

                    self.mem_ts.wrm(0xa0421024, 2, 0, 3)  ##RF diag code select
                    res_in = eval(mydm.device.ask('MEAS:VOLT?'))

                    self.mem_ts.wrm(0xa0421024, 2, 0, 4)  ##RF diag code select
                    res_ip = eval(mydm.device.ask('MEAS:VOLT?'))

                    loginfo("rxdcoc_qn: {}     rxdcoc_qp: {}     rxdcoc_in: {}     rxdcoc_ip: {}     ".format(res_qn, res_qp, res_in, res_ip))
                    rxdcoc_qn = res_qn
                    rxdcoc_qp = res_qp
                    rxdcoc_in = res_in
                    rxdcoc_ip = res_ip

                    fw1.write_data([idac_code, qdac_code, rxdcoc_qn, rxdcoc_qp, rxdcoc_in, rxdcoc_ip, ])
        if cbpf2_en!=0:
            self.mem_ts.wrm(0xa0421020, 20, 20, 0)  ##cbpf2 disable
            self.mem_ts.wrm(0xa0422014, 3, 0, 0)
            self.test232_rxdcoc_cal_bypass(0)
            self.mem_ts.wrm(0xa0421058, 12, 12, 0)
            self.mem_ts.wrm(0xa04210cc, 14, 13, 3)  ##cbpf gain1&2 on
            self.mem_ts.wrm(0xa0422014, 3, 0, 3)    ##rxon
            time.sleep(0.5)
            ro_rxdcoc1_q_code = self.mem_ts.rdm(0xa04210f4,17,12)
            ro_rxdcoc1_i_code = self.mem_ts.rdm(0xa04210f4,11,6)
            ro_rxdcoc2_q_code = self.mem_ts.rdm(0xa04210f4,23,18)
            ro_rxdcoc2_i_code = self.mem_ts.rdm(0xa04210f4,29,24)
            self.test232_rxdcoc_offset_wr(ro_rxdcoc1_i_code, ro_rxdcoc1_q_code)
            self.mem_ts.wrm(0xa0422014, 3, 0, 0)
            self.mem_ts.wrm(0xa0421020, 20, 20, 1)  ##cbpf2 enable
            self.mem_ts.wrm(0xa042105c, 12, 12, 1)  ##cbpf2 rxdcoc code manual
            self.mem_ts.wrm(0xa0422014, 3, 0, 3)  ##rxon

            title = 'ro_rxdcoc1_q_code={},ro_rxdcoc1_i_code={},ro_rxdcoc2_q_code={},ro_rxdcoc2_i_code={}\n'.format(ro_rxdcoc1_q_code,ro_rxdcoc1_i_code,ro_rxdcoc2_q_code,ro_rxdcoc2_i_code)
            title = title + 'idac_code,qdac_code,rxdcoc_qn(v),rxdcoc_qp(v),rxdcoc_in(v),rxdcoc_ip(v)\n'
            fname = self.get_filename('ts_bt_test/', 'test232_rxdcoc_scan_cbpf2')
            fw1 = csvreport(fname, title)
            for idac_code in range(64):
                self.mem_ts.wrm(0xa042105c, 18, 13, idac_code)
                for qdac_code in range(64):
                    self.mem_ts.wrm(0xa042105c, 24, 19, qdac_code)
                    # res = self.test232_rxdcoc_volt_rd()
                    self.mem_ts.wrm(0xa0421024, 2, 0, 1)  ##RF diag code select
                    res_qn = eval(mydm.device.ask('MEAS:VOLT?'))

                    self.mem_ts.wrm(0xa0421024, 2, 0, 2)  ##RF diag code select
                    res_qp = eval(mydm.device.ask('MEAS:VOLT?'))

                    self.mem_ts.wrm(0xa0421024, 2, 0, 3)  ##RF diag code select
                    res_in = eval(mydm.device.ask('MEAS:VOLT?'))

                    self.mem_ts.wrm(0xa0421024, 2, 0, 4)  ##RF diag code select
                    res_ip = eval(mydm.device.ask('MEAS:VOLT?'))

                    loginfo("rxdcoc_qn: {}     rxdcoc_qp: {}     rxdcoc_in: {}     rxdcoc_ip: {}     ".format(res_qn, res_qp, res_in, res_ip))
                    rxdcoc_qn = res_qn
                    rxdcoc_qp = res_qp
                    rxdcoc_in = res_in
                    rxdcoc_ip = res_ip

                    fw1.write_data([idac_code, qdac_code, rxdcoc_qn, rxdcoc_qp, rxdcoc_in, rxdcoc_ip, ])
            self.mem_ts.wrm(0xa042105c, 18, 13, ro_rxdcoc2_i_code)
            self.mem_ts.wrm(0xa042105c, 24, 19, ro_rxdcoc2_q_code)
            self.mem_ts.wrm(0xa0421020, 20, 20, 0)  ##cbpf2 disable
            self.mem_ts.wrm(0xa042105c, 12, 12, 0)  ##cbpf2 rxdcoc code manual


    def test232_rxdcoc_manual_cal(self, device_name='MY50180049'):
        title = 'idac_code,rxdcoc_qn(v),rxdcoc_qp(v),rxdcoc_in(v),rxdcoc_ip(v),cbpf_outi,cbpf_outq\n'
        fname = self.get_filename('ts_bt_test/', 'test232_rxdcoc_manual_cal_{}'.format(self.board_name))
        fw1 = csvreport(fname, title)
        mydm = dm.dm(device_name=device_name, num_of_machine=0, comm='USB')
        self.mem_ts.wrm(0xa0421024, 17, 17, 1)  ##RF diag  soc enable
        self.mem_ts.wrm(0xa0421024, 3, 3, 0)  ##RF diag  enable
        self.mem_ts.wrm(0xa0421024, 11, 11, 1)
        self.mem_ts.wrm(0xa04210a0, 0, 0, 1)    ##agc off
        self.mem_ts.wrm(0xa04210a0, 20, 17, 0)  ##agc index set 0
        self.mem_ts.wrm(0xa04210cc, 14, 13, 2)  ##cbpf gain1 on
        self.test232_rxdcoc_cal_bypass(1)
        self.test232_rxdcoc_offset_wr(32,32)
        self.mem_ts.wrm(0xa0422014, 3, 0, 3)

        for idac_code in range(64):
            self.test232_rxdcoc_offset_wr(idac_code, 32)
            # res = self.test232_rxdcoc_volt_rd()
            self.mem_ts.wrm(0xa0421024, 2, 0, 1)  ##RF diag code select
            res_qn = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 2)  ##RF diag code select
            res_qp = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 3)  ##RF diag code select
            res_in = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 4)  ##RF diag code select
            res_ip = eval(mydm.device.ask('MEAS:VOLT?'))

            loginfo("rxdcoc_qn: {}     rxdcoc_qp: {}     rxdcoc_in: {}     rxdcoc_ip: {}     ".format(res_qn, res_qp, res_in, res_ip))
            rxdcoc_qn = res_qn
            rxdcoc_qp = res_qp
            rxdcoc_in = res_in
            rxdcoc_ip = res_ip
            daci_outi = rxdcoc_ip-rxdcoc_in
            daci_outq = rxdcoc_qp-rxdcoc_qn
            fw1.write_data([idac_code,rxdcoc_qn,rxdcoc_qp,rxdcoc_in,rxdcoc_ip,daci_outi,daci_outq])
        fw1.write_string('qdac_code,rxdcoc_qn(v),rxdcoc_qp(v),rxdcoc_in(v),rxdcoc_ip(v),cbpf_outi,cbpf_outq\n')
        for qdac_code in range(64):
            self.test232_rxdcoc_offset_wr(32, qdac_code)
            # res = self.test232_rxdcoc_volt_rd()
            self.mem_ts.wrm(0xa0421024, 2, 0, 1)  ##RF diag code select
            res_qn = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 2)  ##RF diag code select
            res_qp = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 3)  ##RF diag code select
            res_in = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 4)  ##RF diag code select
            res_ip = eval(mydm.device.ask('MEAS:VOLT?'))

            loginfo("rxdcoc_qn: {}     rxdcoc_qp: {}     rxdcoc_in: {}     rxdcoc_ip: {}     ".format(res_qn, res_qp, res_in, res_ip))
            rxdcoc_qn = res_qn
            rxdcoc_qp = res_qp
            rxdcoc_in = res_in
            rxdcoc_ip = res_ip
            dacq_outi = rxdcoc_ip-rxdcoc_in
            dacq_outq = rxdcoc_qp-rxdcoc_qn
            fw1.write_data([qdac_code,rxdcoc_qn,rxdcoc_qp,rxdcoc_in,rxdcoc_ip,dacq_outi,dacq_outq])


    def test232_txdcoc_rd(self):
        res1_i = self.mem_ts.rdm(0xa042103c,25,20)
        res1_q = self.mem_ts.rdm(0xa042103c, 31, 26)
        loginfo("txdcoc_i_code: {}     txdcoc_q_code: {}".format(res1_i,res1_q))
        return [res1_i,res1_q]

    def test232_txdcoc_wr(self, manual_en=1, icode=32, qcode=32):
        self.mem_ts.wrm(0xa0421040, 6, 6, manual_en)
        self.mem_ts.wrm(0xa0421040, 14, 14, manual_en)
        self.mem_ts.wrm(0xa0421040, 5, 0, icode)
        self.mem_ts.wrm(0xa0421040, 13, 8, qcode)
        self.mem_ts.wr(0xa0422014, 0x0)
        time.sleep(1)
        self.mem_ts.wr(0xa0422014, 0xe)
        self.test232_txdcoc_rd()

    def test232_txdcoc_allchannel_rd(self):
        title = 'channel,txdcoc_i_code,txdcoc_q_code\n'
        fname = self.get_filename('ts_bt_test/', 'test232_txdcoc_allchannel_rd_{}'.format(self.board_name))
        fw1 = csvreport(fname, title)
        self.mem_ts.wrm(0xa0421064, 12, 12, 1)  ##Frequency value manual enable
        for channel in range(0,79):
            self.mem_ts.wrm(0xa0421064, 11, 0, 2402 + channel)
            self.mem_ts.wr(0xa0422014, 0x0)
            time.sleep(1)
            self.mem_ts.wr(0xa0422014, 0xe)
            time.sleep(1)
            res = self.test232_txdcoc_rd()
            fw1.write_data([channel, res[0],res[1]])

    def test232_tx_gain_scan(self, csv_save=True):
        if csv_save:
            title = 'tx pa tpr value,tx carrier power(dbm)\n'
            fname = self.get_filename('ts_bt_test/','test232_tx_gain_scan_{}'.format(self.board_name))
            fw1 = csvreport(fname,title)

        spa = Agilent()
        spa.set_param(2402,0.5,3,3)
        spa.set_reflvl(15)
        self.tx_carrier(freq=2402, dac_gain=0x9f, pa_gain=63)
        self.mem_ts.wrm(0xa0421004, 22, 22, 1)  ##Manual control of Tx PA target power enable

        for gain in range(63,0,-1):
            self.mem_ts.wrm(0xa042100c, 31, 26, gain)  ##Manual control of Tx PA target power value
            self.mem_ts.wr(0xa0422014, 0x0)
            time.sleep(0.2)
            self.mem_ts.wr(0xa0422014, 0xe)
            res = spa.pk_search_avg_timesleep(timesleep=1.5)
            pwr = res[0][1]
            loginfo('gain:  {}  tx carrier power:   {}'.format(gain,pwr))
            fw1.write_data([gain,pwr])
        self.tx_carrier_stop()

    def test_tx_edr_gain(self, chan_list=[0,39,78], rfport=2, cable_loss=1, rate_list=['DH1'], target_power=10, fig_en=0):

        title = 'tx_pa_gain,tx_mixer_gain,tx_dig_gain,rate,channel,nominal_pwr(dBm),peak_pwr(dBm),gfsk_pwr(dBm),dpsk_pwr(dBm),dpsk_gfsk_diff_pwr,guard_period(us),'
        title = title + 'wi(KHz),w0_wi(KHz),w0_max(KHz),DEVM_RMS(%),DEVM_peak(%),DEVM_P99(%),bit_error_rate,packet0error,'
        title = title + 'PTxRef(dBm),N26ChN1Abs(dBm),N26ChP1Abs(dBm),N26ChN1Rel(dBm),N26ChP1Rel(dBm),acp_list\n'
        fname = self.get_filename('ts_bt_test/', 'test_edr_tx_gain_{}'.format(self.board_name))
        fw1 = csvreport(fname, title)

        # for pa_gain in range(2,8):
        #     res = self.test_edr_tx_debug(chan_list=chan_list, rfport=rfport, cable_loss=cable_loss, rate_list=rate_list, target_power=target_power-30+(
        #             pa_gain-2)*6, fig_en=fig_en, csv_save=False, pa_gain=pa_gain, mixer_gain=7, dig_gain=0)
        #     for i in range(len(res)):
        #         fw1.write_data([pa_gain,7,0]+res[i])
        #
        # for mixer_gain in range(0,7):
        #     res = self.test_edr_tx_debug(chan_list=chan_list, rfport=rfport, cable_loss=cable_loss, rate_list=rate_list,
        #                                target_power=target_power-14+mixer_gain*2, fig_en=fig_en, csv_save=False, pa_gain=7, mixer_gain=mixer_gain, dig_gain=0)
        #     for i in range(len(res)):
        #         fw1.write_data([7, mixer_gain, 0] + res[i])

        for dig_gain in range(6,31):
            res = self.test_edr_tx_debug(chan_list=chan_list, rfport=rfport, cable_loss=cable_loss, rate_list=rate_list, target_power=target_power-dig_gain/2,
                                       fig_en=fig_en, csv_save=False, pa_gain=7, mixer_gain=7, dig_gain=dig_gain)
            for i in range(len(res)):
                fw1.write_data([7, 7, dig_gain] + res[i])

    def test_tx_edr_gain_tx231(self, chan_list=[0,39,78], rfport=2, cable_loss=1, rate_list=['DH1'], target_power=10, fig_en=0):

        title = 'tx_pa_gain,tx_mixer_gain,tx_dig_gain,rate,channel,nominal_pwr(dBm),peak_pwr(dBm),gfsk_pwr(dBm),dpsk_pwr(dBm),dpsk_gfsk_diff_pwr,guard_period(us),'
        title = title + 'wi(KHz),w0_wi(KHz),w0_max(KHz),DEVM_RMS(%),DEVM_peak(%),DEVM_P99(%),bit_error_rate,packet0error,'
        title = title + 'PTxRef(dBm),N26ChN1Abs(dBm),N26ChP1Abs(dBm),N26ChN1Rel(dBm),N26ChP1Rel(dBm),acp_list\n'
        fname = self.get_filename('ts_bt_test/', 'test_edr_tx_gain_{}'.format(self.board_name))
        fw1 = csvreport(fname, title)

        for pa_gain in range(0,4):
            for mixer_gain in range(0,8):
                res = self.test_edr_tx_debug(chan_list=chan_list, rfport=rfport, cable_loss=cable_loss, rate_list=rate_list, target_power=target_power - (3-pa_gain) * 3-(7-mixer_gain)*2, fig_en=fig_en, csv_save=False, pa_gain=pa_gain, mixer_gain=7, dig_gain=0)
                for i in range(len(res)):
                    fw1.write_data([pa_gain,mixer_gain,0]+res[i])


    def test_sel(self):
        self.test_tx_br_gain(chan_list=[0,38,78], rfport=1, cable_loss=1, rate_list=['DH1'], target_power=10, fig_en=0)
        time.sleep(1)
        self.test_tx_edr_gain(chan_list=[0, 38, 78], rfport=1, cable_loss=1, rate_list=['2_DH1','3_DH1'], target_power=10, fig_en=0)

    def test_tx_le_gain(self, chan_list=[0,19,39], rfport=2, cable_loss=1, rate_list=['LE1M'], target_power=10, fig_en=0):


        title = 'tx_pa_gain,tx_mixer_gain,tx_dig_gain,channel,type,nominal_pow(dBm),peak_pow(dBm),leakage_pow(dBm),freq_accuracy(KHz),freq_drift(KHz),drift_rate(Hz/50us),'
        title = title + 'delta_f1_avg(KHz),delta_f1_min(KHz),delta_f1_max(KHz),delta_f2_avg(KHz),delta_f2_min(KHz),delta_f2_max(KHz),mod_ratio,delta_f1_99(KHz),delta_f2_99(KHz),'
        title = title + 'acp_list_21ch\n'
        fname = self.get_filename('ts_bt_test/','test_le_tx_gain_{}'.format(self.board_name))
        fw1 = csvreport(fname, title)

        for pa_gain in range(2,8):
            res = self.test_le_tx_debug(chan_list=chan_list, rfport=rfport, cable_loss=cable_loss, rate_list=rate_list, target_power=target_power-30+(pa_gain-2)*6, fig_en=fig_en, csv_save=False, pa_gain=pa_gain, mixer_gain=7, dig_gain=0)
            for i in range(len(res)):
                fw1.write_data([pa_gain,7,0]+res[i])

        for mixer_gain in range(0,7):
            res = self.test_le_tx_debug(chan_list=chan_list, rfport=rfport, cable_loss=cable_loss, rate_list=rate_list, target_power=target_power-14+mixer_gain*2, fig_en=fig_en, csv_save=False, pa_gain=7, mixer_gain=mixer_gain, dig_gain=0)
            for i in range(len(res)):
                fw1.write_data([7, mixer_gain, 0] + res[i])

        for dig_gain in range(1,31):
            res = self.test_le_tx_debug(chan_list=chan_list, rfport=rfport, cable_loss=cable_loss, rate_list=rate_list, target_power=target_power-dig_gain/2, fig_en=fig_en, csv_save=False, pa_gain=7, mixer_gain=7, dig_gain=dig_gain)
            for i in range(len(res)):
                fw1.write_data([7, 7, dig_gain] + res[i])
    def test_tx_le_gain_tx231(self, chan_list=[0,19,39], rfport=2, cable_loss=1, rate_list=['LE1M'], target_power=10, fig_en=0):


        title = 'tx_pa_gain,tx_mixer_gain,tx_dig_gain,channel,type,nominal_pow(dBm),peak_pow(dBm),leakage_pow(dBm),freq_accuracy(KHz),freq_drift(KHz),drift_rate(Hz/50us),'
        title = title + 'delta_f1_avg(KHz),delta_f1_min(KHz),delta_f1_max(KHz),delta_f2_avg(KHz),delta_f2_min(KHz),delta_f2_max(KHz),mod_ratio,delta_f1_99(KHz),delta_f2_99(KHz),'
        title = title + 'acp_list_21ch\n'
        fname = self.get_filename('ts_bt_test/','test_le_tx_gain_{}'.format(self.board_name))
        fw1 = csvreport(fname, title)

        for pa_gain in range(0,4):
            for mixer_gain in  range(0,8):
                res = self.test_le_tx_debug(chan_list=chan_list, rfport=rfport, cable_loss=cable_loss, rate_list=rate_list, target_power=target_power-(3-pa_gain)*3-(
                        7-mixer_gain)*2,fig_en=fig_en, csv_save=False, pa_gain=pa_gain, mixer_gain=mixer_gain, dig_gain=0)
                for i in range(len(res)):
                    fw1.write_data([pa_gain,mixer_gain,0]+res[i])

    def rw_per(self, rfport=2, cable_loss=5, chan=0, rate='1M_DH1', rxpwr_range=[-98,-20], pkt_num=1000, dirty_en=0, csv_save=True, device='CMW', num_of_machine=1):
        if csv_save:
            title = 'channel,rate,rxpwr,rev_pkg,total_pkg,err_bit,total_bit,err_bit_ratio(%)\n'
            fname = self.get_filename('ts_bt_test/','rw_per_{}_{}'.format(rate,chan))
            fw1 = csvreport(fname,title)
        perform_list = []
        rxpwr_max = max(rxpwr_range)
        rxpwr_min = min(rxpwr_range)

        if device == 'CMW':
            tester = tester_cmw(mode=1, rfport=rfport, cable_loss=cable_loss)
            tester.srx.gen_wave(repeat=pkt_num, data_rate=rate, dirty_en=dirty_en)
        else:
            tester = mxg.MXG(device, num_of_machine)
            tester.arb_waveform(rate=rate)
            tester.trriger_para_set(type='SINGle', count=pkt_num)
            tester.output_state(1, 1)
            tester.arb_state(1)

        for rxpwr in  range(rxpwr_min, rxpwr_max+1):
            tester.srx.para_settings(freq=2402 + chan, level=rxpwr)
            self.BR_RX()
            tester.srx.gen_switch('ON')
            time.sleep(pkt_num*2e-3)
            res = self.cmdstop(1)
            rev_pkg = res[1]
            err_bit = res[2]
            total_bit = res[3]
            err_bit_ratio = (100*err_bit+0.01)/total_bit
            if csv_save:
                fw1.write_data([chan,rate,rxpwr,rev_pkg,pkt_num,err_bit,total_bit,err_bit_ratio])
            perform_list.append([chan,rate,rxpwr,rev_pkg,pkt_num,err_bit,total_bit,err_bit_ratio])
        return perform_list

    def le_per(self, rfport=1, cable_loss=5, chan_list=[0], rate_list=['LE_1M'], rxpwr_range=[-98,-20], pkt_num=2000, dirty_en=0, csv_save=True, device='CMW',
               num_of_machine=1):
        self.cmdstop(0)
        self.cmdstop(1)
        # self.mem_ts.wrm(0xa0420004,30,16,0x1c4)
        if csv_save:
            title = 'channel,rate,rxpwr,total_pkg,rev_pkg,per(%)\n'
            fname = self.get_filename('ts_bt_test/','le_per_{}'.format(self.board_name))
            fw1 = csvreport(fname,title)
        perform_list = []
        rxpwr_max = max(rxpwr_range)
        rxpwr_min = min(rxpwr_range)
        for chan in chan_list:
            for rate in rate_list:
                if device == 'CMW':
                    tester = tester_cmw(mode=1, rfport=rfport, cable_loss=cable_loss)
                    tester.srx.gen_wave(repeat=pkt_num, data_rate=rate, dirty_en=dirty_en)
                else:
                    tester = mxg.MXG(device, num_of_machine)
                    tester.arb_waveform(rate=rate)
                    tester.trriger_para_set(type='SINGle', count=pkt_num)
                    tester.output_state(1,1)
                    tester.arb_state(1)

                if rate.find('LE_1M') != -1:
                    phy = 1
                else:
                    phy = 2
                for rxpwr in  range(rxpwr_min, rxpwr_max+1):
                    self.LE_RX(chan=chan, phy=phy, mod=0, countMode=0)
                    if device == 'CMW':
                        tester.srx.para_settings(freq=2402 + chan*2, level=rxpwr)
                        tester.srx.gen_switch('ON')
                    else:
                        tester.para_set(freq=2402+chan*2, power=rxpwr+cable_loss)
                        tester.trigger_on()
                    time.sleep(pkt_num*3e-3)
                    res = self.cmdstop(0)
                    print(res)
                    res = res[0].split('bleRxCount=')[1].strip('\r\n')
                    logdebug('rev_pkt:  {}'.format(res))
                    rev_pkg = eval(res)
                    per = (pkt_num-rev_pkg)*100.0/pkt_num
                    loginfo('{}     {}      {}      {}      {}      {}'.format(chan,rate,rxpwr,pkt_num,rev_pkg,per))
                    if csv_save:
                        fw1.write_data([chan,rate,rxpwr,pkt_num,rev_pkg,per])

                    perform_list.append([chan,rate,rxpwr,pkt_num,rev_pkg,per])
        return perform_list

    def le_rx_ci(self, rfport=1, cable_loss=5, chan_list=[0], rate_list=['LE_1M'], pkt_num=2000, dirty_en=0, csv_save=True, device_signal='N5182B', rxgain='0xfb'):
        self.dict_pwr_inter_freqoffset = {
            0: range(-90, -60),
            1: range(-80, -40),
            2: range(-52, -20),
            3: range(-41, -10),
            -1: range(-80, -40),
            -2: range(-52, -20),
            -3: range(-41, -10)
            }
        self.cmdstop(0)
        self.cmdstop(1)
        # self.mem_ts.wrm(0xa0420004,30,16,0x1c4)
        if csv_save:
            title = 'signal_channel,rate,freq_interference,rxpwr_interference,total_pkg,rev_pkg,per(%)\n'
            fname = self.get_filename('ts_bt_test/','le_rx_ci_{}'.format(self.board_name))
            fw1 = csvreport(fname,title)
        for chan in chan_list:
            for rate in rate_list:
                if device_signal == 'CMW':
                    tester_signal = tester_cmw(mode=1, rfport=rfport, cable_loss=cable_loss)
                    tester_signal.srx.gen_wave(repeat=pkt_num, data_rate=rate, dirty_en=dirty_en)
                    tester_signal.srx.para_settings(freq=2402 + chan*2, level=-67)
                    tester_interference = mxg.MXG()
                else:
                    tester_signal = mxg.MXG('N5182B',1)
                    tester_signal.arb_waveform(rate=rate)
                    tester_signal.trriger_para_set(type='SINGle', count=pkt_num)
                    tester_signal.output_state(1,1)
                    tester_signal.arb_state(1)
                    tester_signal.para_set(freq=2402 + chan*2, power=-67 + cable_loss)

                    tester_interference = mxg.MXG('N5182B',2)
                tester_interference.arb_waveform(rate=rate+'_PN15')
                tester_interference.trriger_para_set(type='CONT', count=pkt_num)
                tester_interference.output_state(1, 1)
                tester_interference.arb_state(1)

                if rate.find('LE_1M') != -1:
                    phy = 1
                    freq_inter_offset_list = [0,1,-1,2,-2,3,-3]
                else:
                    phy = 2
                    freq_inter_offset_list = [0, 2, -2, 4, -4, 6, -6]
                for freq_inter_offset in freq_inter_offset_list:
                    freq_inter = freq_inter_offset+2402+chan*2
                    pwr_list = self.dict_pwr_inter_freqoffset[freq_inter_offset/phy]
                    for power_inter in pwr_list:
                        tester_interference.para_set(freq=freq_inter, power=power_inter + cable_loss)
                        time.sleep(0.2)
                        loginfo('freq_inter:    {}      pwr_inter:  {}'.format(freq_inter,power_inter))
                        self.LE_RX(chan=chan, phy=phy, mod=0, countMode=0)
                        if device_signal == 'CMW':
                            tester_signal.srx.gen_switch('ON')
                        else:
                            tester_signal.trigger_on()
                        if phy == 2:
                            time.sleep(pkt_num * 3e-3)
                        else:
                            time.sleep(pkt_num * 1.5e-3)
                        res = self.cmdstop(0)
                        res = res[0].split('bleRxCount=')[1].strip('\r\n')
                        logdebug('rev_pkt:  {}'.format(res))
                        rev_pkg = eval(res)
                        per = (pkt_num - rev_pkg) * 100.0 / pkt_num
                        loginfo('{}     {}      {}      {}      {}      {}      {}'.format(chan, rate, freq_inter, power_inter, pkt_num, rev_pkg, per))
                        fw1.write_data([chan, rate, freq_inter, power_inter, pkt_num, rev_pkg, per])

    def le_rx_ci_rxgain(self, rxgain_list=[0xff,0xfd,0xfb,0xf9], rfport=1, cable_loss=5, chan_list=[0], rate_list=['LE_1M'], pkt_num=2000, dirty_en=0, csv_save=True,
                        device_signal='N5182B'):
        self.dict_pwr_inter_freqoffset = {
            0: range(-90, -60),
            1: range(-85, -50),
            2: range(-52, -30),
            3: range(-41, -18),
            -1: range(-85, -50),
            -2: range(-52, -30),
            -3: range(-41, -18)
            }
        self.cmdstop(0)
        self.cmdstop(1)
        self.mem_ts.wrm(0xa0420004,30,16,0x1c4)
        if csv_save:
            title = 'rxgain,signal_channel,rate,freq_interference,rxpwr_interference,total_pkg,rev_pkg,per(%)\n'
            fname = self.get_filename('ts_bt_test/','le_rx_ci_{}'.format(self.board_name))
            fw1 = csvreport(fname,title)
        for rxgain in rxgain_list:
            for chan in chan_list:
                for rate in rate_list:
                    if device_signal == 'CMW':
                        tester_signal = tester_cmw(mode=1, rfport=rfport, cable_loss=cable_loss)
                        tester_signal.srx.gen_wave(repeat=pkt_num, data_rate=rate, dirty_en=dirty_en)
                        tester_signal.srx.para_settings(freq=2402 + chan*2, level=-67)
                        tester_interference = mxg.MXG()
                    else:
                        tester_signal = mxg.MXG('N5182B',1)
                        tester_signal.arb_waveform(rate=rate)
                        tester_signal.trriger_para_set(type='SINGle', count=pkt_num)
                        tester_signal.output_state(1,1)
                        tester_signal.arb_state(1)
                        tester_signal.para_set(freq=2402 + chan*2, power=-67 + cable_loss)

                        tester_interference = mxg.MXG('N5182B',2)
                    tester_interference.arb_waveform(rate=rate+'_PN15')
                    tester_interference.trriger_para_set(type='CONT', count=pkt_num)
                    tester_interference.output_state(1, 1)
                    tester_interference.arb_state(1)

                    if rate.find('LE_1M') != -1:
                        phy = 1
                    else:
                        phy = 2
                    for freq_inter_offset in[0,1,-1,2,-2,3,-3]:
                        freq_inter = freq_inter_offset+2402+chan*2
                        pwr_list = self.dict_pwr_inter_freqoffset[freq_inter_offset]
                        for power_inter in pwr_list:
                            tester_interference.para_set(freq=freq_inter, power=power_inter + cable_loss)
                            time.sleep(0.2)
                            loginfo('freq_inter:    {}      pwr_inter:  {}'.format(freq_inter,power_inter))
                            self.mem_ts.wrm(0xa042205c, 7, 0, rxgain)
                            self.LE_RX(chan=chan, phy=phy, mod=0, countMode=0)
                            if device_signal == 'CMW':
                                tester_signal.srx.gen_switch('ON')
                            else:
                                tester_signal.trigger_on()
                            time.sleep(pkt_num * 1e-3)
                            res = self.cmdstop(0)
                            res = res[0].split('bleRxCount=')[1].strip('\r\n')
                            logdebug('rev_pkt:  {}'.format(res))
                            rev_pkg = eval(res)
                            per = (pkt_num - rev_pkg) * 100.0 / pkt_num
                            loginfo('{}     {}      {}      {}      {}      {}      {}      {}'.format(rxgain, chan, rate, freq_inter, power_inter, pkt_num, rev_pkg, per))
                            fw1.write_data([rxgain, chan, rate, freq_inter, power_inter, pkt_num, rev_pkg, per])

    def le_rx_blocking(self, rfport=1, cable_loss=5,  rate_list=['LE_1M'], inter_power_list=[], pkt_num=2000, dirty_en=0, csv_save=True, device_signal='N5182B'):
        self.cmdstop(0)
        self.cmdstop(1)
        # self.mem_ts.wrm(0xa0420004,30,16,0x1c4)
        if csv_save:
            title = 'rate,freq_blocking_1,freq_blocking_2,per(%)\n'
            fname = self.get_filename('ts_bt_test/','le_rx_blocking_{}'.format(self.board_name))
            fw1 = csvreport(fname,title)
        freq_blocking_1_list = []
        freq_blocking_2_list = []
        for rate in rate_list:
            if device_signal == 'CMW':
                tester_signal = tester_cmw(mode=1, rfport=rfport, cable_loss=cable_loss)
                tester_signal.srx.gen_wave(repeat=pkt_num, data_rate=rate, dirty_en=dirty_en)
                tester_signal.srx.para_settings(freq=2426, level=-67)
                tester_interference = mxg.MXG()
            else:
                tester_signal = mxg.MXG('N5182B', 1)
                tester_signal.arb_waveform(rate=rate)
                tester_signal.trriger_para_set(type='SINGle', count=pkt_num)
                tester_signal.output_state(1, 1)
                tester_signal.arb_state(1)
                tester_signal.para_set(freq=2426, power=-67 + cable_loss)

                tester_interference = mxg.MXG('N5182B', 2)
            tester_interference.output_state(1, 0)

            if rate.find('LE_1M') != -1:
                phy = 1
            else:
                phy = 2
            freq_range= range(30,2001,10)+range(2003,2400,3)+range(2484,2998,3)+range(3000,6000,25)
            freq_blocking_2 = 'NA'
            for inter_pwr in inter_power_list:
                for freq_inter in freq_range:
                    if freq_inter < 2001 or freq_inter > 2999:
                        tester_interference.para_set(freq=freq_inter, power=inter_pwr + cable_loss)
                    else:
                        tester_interference.para_set(freq=freq_inter, power=inter_pwr - 5 + cable_loss)
                    loginfo('freq_inter:    {}'.format(freq_inter))
                    # self.mem_ts.wrm(0xa042205c, 7, 0, rxgain)
                    self.LE_RX(chan=12, phy=phy, mod=0, countMode=0)
                    if device_signal == 'CMW':
                        tester_signal.srx.gen_switch('ON')
                    else:
                        tester_signal.trigger_on()
                    time.sleep(pkt_num * 1e-3)
                    res = self.cmdstop(0)
                    res = res[0].split('bleRxCount=')[1].strip('\r\n')
                    logdebug('rev_pkt:  {}'.format(res))
                    rev_pkg = eval(res)
                    per = (pkt_num - rev_pkg) * 100.0 / pkt_num
                    if per > 31:
                        freq_blocking_1 = freq_inter
                        freq_blocking_1_list.append(freq_blocking_1)
                        fw1.write_data([rate,freq_blocking_1,freq_blocking_2,per])

                freq_blocking_1 = 'NA'
                for freq_inter in freq_range:
                    tester_interference.para_set(freq=freq_inter, power=inter_pwr - 20 + cable_loss)
                    loginfo('freq_inter:    {}'.format(freq_inter))
                    # self.mem_ts.wrm(0xa042205c, 7, 0, rxgain)
                    self.LE_RX(chan=12, phy=phy, mod=0, countMode=0)
                    if device_signal == 'CMW':
                        tester_signal.srx.gen_switch('ON')
                    else:
                        tester_signal.trigger_on()
                    time.sleep(pkt_num * 1e-3)
                    res = self.cmdstop(0)
                    res = res[0].split('bleRxCount=')[1].strip('\r\n')
                    logdebug('rev_pkt:  {}'.format(res))
                    rev_pkg = eval(res)
                    per = (pkt_num - rev_pkg) * 100.0 / pkt_num
                    if per > 31:
                        freq_blocking_2 = freq_inter
                        freq_blocking_2_list.append(freq_blocking_2)
                        fw1.write_data([rate,freq_blocking_1,freq_blocking_2,per])
        return freq_blocking_1_list + [len(freq_blocking_1_list)], freq_blocking_2_list + [len(freq_blocking_2_list)]

    def test_bt_rx(self,rfport=2, cable_loss=5, chan_list=[0], rate_list=['1M_DH1'], rxpwr_range=[-98,-20], pkt_num=1000, dirty_en=0):
        tester = tester_cmw(mode=1, rfport=rfport, cable_loss=cable_loss)
        for chan in chan_list:
            for rate in rate_list:
                perform_list = self.rw_per(rfport=rfport, cable_loss=cable_loss, chan=chan, rate= rate, rxpwr_range=rxpwr_range, pkt_num=pkt_num, dirty_en=dirty_en)

    def rxgain_scan_tx232(self, channel_list=[0], csv_save=True, rf_cableloss=1, if_cableloss=1):
        self.cmdstop(0)
        self.cmdstop(1)
        if csv_save:
            title = 'channel,lna_itrim,cbpf_bias_trim,LNA_HGAIN,LNA_LGAIN,LNA_ATTEN,FLT_GAIN1,FLT_GAIN2,RF_pwr(dbm),IF_pwr(dbm),dalta\n'
            fname = self.get_filename('ts_bt_test/','rxgain_scan_tx232_{}'.format(self.board_name))
            fw1 = csvreport(fname,title)

        spa = Agilent()
        spa.set_param(1,0.5,1,3)
        txg = mxg.MXG()
        # txg.arb_waveform(rate='LE_1M')
        txg.trriger_para_set(type='CONTinuous', count=1000)
        txg.output_state(1, 0)
        # txg.arb_state(1)
        self.mem_ts.wrm(0xa04210A0, 0, 0, 0x1)  ##AGC OFF
        self.mem_ts.wrm(0xa0421020, 20, 20, 0x0)
        self.mem_ts.wrm(0xa0120038, 11, 8, 0)
        self.mem_ts.wrm(0xa012005c, 31, 16, 0)  ##
        self.mem_ts.wrm(0xa0200240, 17, 6, 0)  ##PA5/4/3/2 高阻
        self.mem_ts.wrm(0xa01200a0, 29, 29, 1)  ##bbpll_bg_pup_ibg_tbuf
        self.mem_ts.wrm(0xa01200c8, 4, 4, 1)  ##test buf en
        ##self.mem_ts.wrm(0xa01200c8,3,2,2)	##tx IF test buf
        self.mem_ts.wrm(0xa01200c8, 3, 2, 1)  ##rx IF test buf
        lna_itrim_init = self.mem_ts.rdm(0xa0421034,8,7)
        cbpf_bias_trim_init = self.mem_ts.rdm(0xa0421034,18,16)
        for channel in channel_list:
            for lna_itrim in range(4):
                self.mem_ts.wrm(0xa0421034, 8, 7,lna_itrim)
                for cbpf_bias_trim in range(8):
                    self.mem_ts.wrm(0xa0421034, 18, 16, cbpf_bias_trim)
                    self.mem_ts.wrm(0xa04210b4, 25, 17, 0x1ff)
                    self.mem_ts.wrm(0xa04210b4, 16, 13, 0x3)
                    txg.para_set(freq=2402 + channel , power=-67)
                    self.mem_ts.wr(0xa0422014, 0x0)
                    self.mem_ts.wrm(0xa0421064, 12, 12, 1)  ##Frequency value manual enable
                    self.mem_ts.wrm(0xa0421064, 11, 0, 2402+channel)
                    self.mem_ts.wr(0xa0422014, 0x3)  ##manual rx on

                    for lna_hgain in (0x3f,0x20,0x10,0x8,0x4,0x2):

                        self.mem_ts.wrm(0xa04210b4, 24, 19, lna_hgain)
                        lna_lgain = 3
                        lna_atten = 0
                        flt_gain1 = 1
                        flt_gain2 = 1

                        pwr = -40
                        while 1:
                            txg.set_power(pwr+rf_cableloss)
                            time.sleep(0.2)
                            res = spa.pk_search_avg_timesleep(timesleep=1.5)
                            if_pwr = res[0][1]+if_cableloss
                            loginfo(if_pwr)
                            if if_pwr > 0:
                                pwr = pwr - 5
                            else:
                                break
                        loginfo('RF pwer:{}    IF power:{}'.format(pwr, if_pwr))
                        fw1.write_data([channel, lna_itrim, cbpf_bias_trim, lna_hgain, lna_lgain, lna_atten,  flt_gain1, flt_gain2, pwr, if_pwr, if_pwr - pwr])

                    self.mem_ts.wrm(0xa04210b4, 25, 25, 0)
                    for lna_lgain in (3,2):
                        self.mem_ts.wrm(0xa04210b4, 18, 17, lna_lgain)
                        lna_hgain = 2
                        lna_atten = 0
                        flt_gain1 = 1
                        flt_gain2 = 1

                        pwr = -40
                        while 1:
                            txg.set_power(pwr+rf_cableloss)
                            time.sleep(0.2)
                            res = spa.pk_search_avg_timesleep(timesleep=1.5)
                            if_pwr = res[0][1]+if_cableloss
                            print(if_pwr)
                            if if_pwr > 0:
                                pwr = pwr - 5
                            else:
                                break
                        loginfo('RF pwer:{}    IF power:{}'.format(pwr, if_pwr))
                        fw1.write_data([channel, lna_itrim, cbpf_bias_trim, lna_hgain, lna_lgain, lna_atten,  flt_gain1, flt_gain2, pwr, if_pwr, if_pwr - pwr])
                    self.mem_ts.wrm(0xa04210b4, 18, 17, 2)
                    for lna_atten in (1,2,3):
                        self.mem_ts.wrm(0xa04210b4, 16, 15, lna_atten)
                        lna_hgain = 2
                        lna_lgain = 2
                        flt_gain1 = 1
                        flt_gain2 = 1

                        pwr = -40
                        while 1:
                            txg.set_power(pwr+rf_cableloss)
                            time.sleep(0.2)
                            res = spa.pk_search_avg_timesleep(timesleep=1.5)
                            if_pwr = res[0][1]+if_cableloss
                            print(if_pwr)
                            if if_pwr > 0:
                                pwr = pwr - 5
                            else:
                                break
                        loginfo('RF pwer:{}    IF power:{}'.format(pwr, if_pwr))
                        fw1.write_data([channel, lna_itrim, cbpf_bias_trim, lna_hgain, lna_lgain, lna_atten,  flt_gain1, flt_gain2, pwr, if_pwr, if_pwr - pwr])
                    self.mem_ts.wrm(0xa04210b4, 16, 15, 3)
                    self.mem_ts.wrm(0xa04210b4, 14, 14, 0)
                    for flt_gain2 in (1,0) :
                        self.mem_ts.wrm(0xa04210b4, 13, 13, flt_gain2)
                        lna_hgain = 0
                        lna_lgain = 0
                        flt_gain1 = 0
                        lna_atten = 3

                        pwr = -40
                        while 1:
                            txg.set_power(pwr+rf_cableloss)
                            time.sleep(0.2)
                            res = spa.pk_search_avg_timesleep(timesleep=1.5)
                            if_pwr = res[0][1]+if_cableloss
                            print(if_pwr)
                            if if_pwr > 0:
                                pwr = pwr - 5
                            else:
                                break
                        loginfo('RF pwer:{}    IF power:{}'.format(pwr, if_pwr))
                        fw1.write_data([channel, lna_itrim, cbpf_bias_trim, lna_hgain, lna_lgain, lna_atten,  flt_gain1, flt_gain2, pwr, if_pwr, if_pwr - pwr])
        txg.output_state(0, 0)
        self.mem_ts.wrm(0xa0421034, 8, 7, lna_itrim_init)
        self.mem_ts.wrm(0xa0421034, 18, 16, cbpf_bias_trim_init)
        self.mem_ts.wrm(0xa04210A0, 0, 0, 0x0)
        self.mem_ts.wrm(0xa04210b4, 25, 13, 0x1ff3)
        self.mem_ts.wrm(0xa0421020, 20, 20, 0x0)


    def rxgain_scan_tx231(self, channel_list=[0], csv_save=True):
        self.cmdstop(0)
        self.cmdstop(1)
        if csv_save:
            title = 'channel,LNA_GAIN,LNA_ATTEN_R,LNA_ATTEN_C,FLT_GAIN1,FLT_GAIN2,RF_pwr(dbm),IF_pwr(dbm),dalta\n'
            fname = self.get_filename('ts_bt_test/','rxgain_scan_tx231_{}'.format(self.board_name))
            fw1 = csvreport(fname,title)

        spa = Agilent()
        spa.set_param(1,1,10,30)
        txg = mxg.MXG()
        # txg.arb_waveform(rate='LE_1M')
        txg.trriger_para_set(type='CONTinuous', count=1000)
        txg.output_state(1, 0)
        # txg.arb_state(1)
        for channel in channel_list:
            txg.para_set(freq=2402 + channel * 2, power=-67)
            self.mem_ts.wrm(0xa042107c,30,26,0x1f)  ##rx gain manual
            time.sleep(.5)
            self.mem_ts.wrm(0xa0421064, 0, 0, 0x1)  ##AGC OFF
            time.sleep(.5)
            self.mem_ts.wrm(0xa04210a8, 6, 6, 1)  ##Tx DAC output test enable signal
            time.sleep(.5)
            self.mem_ts.wrm(0xa04210a8, 0, 0, 1)  ##Test buffer enable signal
            time.sleep(.5)
            self.LE_RX(chan=channel,phy=1,mod=0,countMode=0)
            for lna_gain in range(3,-1,-1):
                if lna_gain == 0:
                    for lna_atten_c in range(0,4):
                        for lna_atten_r in range(0, 4):
                            for flt_gain1 in range(3, -1, -1):
                                for flt_gain2 in range(3, -1, -1):
                                    self.mem_ts.wrm(0xa042107c, 21, 20, lna_atten_r)
                                    time.sleep(.5)
                                    self.mem_ts.wrm(0xa042107c, 17, 16, lna_gain)
                                    time.sleep(.5)
                                    self.mem_ts.wrm(0xa042107c, 19, 18, lna_atten_c)
                                    time.sleep(.5)
                                    self.mem_ts.wrm(0xa042107c, 23, 22, flt_gain1)
                                    time.sleep(.5)
                                    self.mem_ts.wrm(0xa042107c, 25, 24, flt_gain2)
                                    pwr = -40
                                    while 1:
                                        txg.set_power(pwr)
                                        time.sleep(0.2)
                                        res = spa.pk_search_avg_timesleep(timesleep=1.5)
                                        if_pwr = res[0][1]
                                        print(if_pwr)
                                        if if_pwr > 0:
                                            pwr = pwr - 5
                                        else:
                                            break
                                    loginfo('RF pwer:{}    IF power:{}'.format(pwr, if_pwr))
                                    fw1.write_data([channel, lna_gain, lna_atten_r, lna_atten_c, flt_gain1, flt_gain2, pwr, if_pwr, if_pwr - pwr])
                else:
                    lna_atten_r = 0
                    lna_atten_c = 0
                    for flt_gain1 in range(3,-1,-1):
                        for flt_gain2 in range(3, -1, -1):
                            self.mem_ts.wrm(0xa042107c, 21, 20, lna_atten_r)
                            time.sleep(.5)
                            self.mem_ts.wrm(0xa042107c, 17, 16, lna_gain)
                            time.sleep(.5)
                            self.mem_ts.wrm(0xa042107c, 19, 18, lna_atten_c)
                            time.sleep(.5)
                            self.mem_ts.wrm(0xa042107c, 23, 22, flt_gain1)
                            time.sleep(.5)
                            self.mem_ts.wrm(0xa042107c, 25, 24, flt_gain2)
                            pwr = -40

                            while 1:
                                txg.set_power(pwr)
                                time.sleep(0.2)
                                res = spa.pk_search_avg_timesleep(timesleep=1.5)
                                if_pwr = res[0][1]
                                print(if_pwr)
                                if if_pwr > 0:
                                    pwr = pwr - 5
                                else:
                                    break
                            loginfo('RF pwer:{}    IF power:{}'.format(pwr,if_pwr))
                            fw1.write_data([channel,lna_gain,lna_atten_r,lna_atten_c,flt_gain1,flt_gain2,pwr,if_pwr,if_pwr-pwr])
            # for lna_gain in range(0,4):
            #     lna_atten_r = 0
            #     lna_atten_c = 0
            #     flt_gain1 = 3
            #     flt_gain2 = 3
            #     self.mem_ts.wrm(0xa042107c, 21, 20, lna_atten_r)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 17, 16, lna_gain)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 19, 18, lna_atten_c)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 23, 22, flt_gain1)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 25, 24, flt_gain2)
            #     pwr = -40
            #
            #     while 1:
            #         txg.set_power(pwr)
            #         time.sleep(0.2)
            #         res = spa.pk_search_avg_timesleep(timesleep=1.5)
            #         if_pwr = res[0][1]
            #         print(if_pwr)
            #         if if_pwr > 0:
            #             pwr = pwr - 5
            #         else:
            #             break
            #     loginfo('RF pwer:{}    IF power:{}'.format(pwr,if_pwr))
            #     fw1.write_data([channel,lna_gain,lna_atten_r,lna_atten_c,flt_gain1,flt_gain2,pwr,if_pwr,if_pwr-pwr])
            # for lna_atten_c in range(0,4):
            #     lna_gain = 0
            #     lna_atten_r = 0
            #     flt_gain1 = 3
            #     flt_gain2 = 3
            #     self.mem_ts.wrm(0xa042107c, 21, 20, lna_atten_r)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 17, 16, lna_gain)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 19, 18, lna_atten_c)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 23, 22, flt_gain1)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 25, 24, flt_gain2)
            #     pwr = -30
            #
            #     while 1:
            #         txg.set_power(pwr)
            #         time.sleep(0.2)
            #         res = spa.pk_search_avg_timesleep(timesleep=1.5)
            #         if_pwr = res[0][1]
            #         print(if_pwr)
            #         if if_pwr > 0 :
            #             pwr = pwr - 5
            #         else:
            #             break
            #     loginfo('RF pwer:{}    IF power:{}'.format(pwr,if_pwr))
            #     fw1.write_data([channel,lna_gain,lna_atten_r,lna_atten_c,flt_gain1,flt_gain2,pwr,if_pwr,if_pwr-pwr])
            # for lna_atten_r in range(0,4):
            #     lna_gain = 0
            #     lna_atten_c = 3
            #     flt_gain1 = 3
            #     flt_gain2 = 3
            #
            #     self.mem_ts.wrm(0xa042107c, 21, 20, lna_atten_r)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 17, 16, lna_gain)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 19, 18, lna_atten_c)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 23, 22, flt_gain1)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 25, 24, flt_gain2)
            #     pwr = -20
            #
            #     while 1:
            #         txg.set_power(pwr)
            #         time.sleep(0.2)
            #         res = spa.pk_search_avg_timesleep(timesleep=1.5)
            #         if_pwr = res[0][1]
            #         print(if_pwr)
            #         if if_pwr > 0 :
            #             pwr = pwr - 5
            #         else:
            #             break
            #     loginfo('RF pwer:{}    IF power:{}'.format(pwr,if_pwr))
            #     fw1.write_data([channel,lna_gain,lna_atten_r,lna_atten_c,flt_gain1,flt_gain2,pwr,if_pwr,if_pwr-pwr])
            # for flt_gain1 in range(0,4):
            #     lna_gain = 0
            #     lna_atten_r = 3
            #     lna_atten_c = 3
            #     flt_gain2 = 3
            #
            #     self.mem_ts.wrm(0xa042107c, 21, 20, lna_atten_r)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 17, 16, lna_gain)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 19, 18, lna_atten_c)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 23, 22, flt_gain1)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 25, 24, flt_gain2)
            #     pwr = -10
            #
            #     while 1:
            #         txg.set_power(pwr)
            #         time.sleep(0.2)
            #         res = spa.pk_search_avg_timesleep(timesleep=2)
            #         if_pwr = res[0][1]
            #         print(if_pwr)
            #         if if_pwr > 5 :
            #             pwr = pwr - 5
            #         else:
            #             break
            #     loginfo('RF pwer:{}    IF power:{}'.format(pwr,if_pwr))
            #     fw1.write_data([channel,lna_gain,lna_atten_r,lna_atten_c,flt_gain1,flt_gain2,pwr,if_pwr,if_pwr-pwr])
            # for flt_gain2 in range(0,4):
            #     lna_gain = 0
            #     lna_atten_r = 3
            #     lna_atten_c = 3
            #     flt_gain1 = 0
            #
            #     self.mem_ts.wrm(0xa042107c, 21, 20, lna_atten_r)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 17, 16, lna_gain)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 19, 18, lna_atten_c)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 23, 22, flt_gain1)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 25, 24, flt_gain2)
            #     pwr = -10
            #
            #     while 1:
            #         txg.set_power(pwr)
            #         time.sleep(0.2)
            #         res = spa.pk_search_avg_timesleep(timesleep=2)
            #         if_pwr = res[0][1]
            #         print(if_pwr)
            #         if if_pwr > 0 :
            #             pwr = pwr - 5
            #         else:
            #             break
            #     loginfo('RF pwer:{}    IF power:{}'.format(pwr,if_pwr))
            #     fw1.write_data([channel,lna_gain,lna_atten_r,lna_atten_c,flt_gain1,flt_gain2,pwr,if_pwr,if_pwr-pwr])

    def test231_agc_scan(self,  flt_pkd_list=[0x40], agc_tgt_sat_list=[[407,580]], signal_cableloss=7, interference_cableloss=7,  if_cableloss=1, signal_freq=2440,
                         signal_level=-67, interference_en=1,
                         interference_offset_freq=3):

        title = 'flt_pkd_value,agc_tgt,agc_sat,interference_pwr(dbm)@offset {}MHz,signal_pwr(dbm),IF_pwr(dbm),dalta\n'.format(interference_offset_freq)
        fname = self.get_filename('ts_bt_test/','test231_agc_scan_{}'.format(self.board_name))
        fw1 = csvreport(fname,title)

        spa = Agilent()
        spa.set_param(1,1,10,30)
        txg = mxg.MXG()
        if interference_en == 0:
            txg_freq = signal_freq
        else:
            txg_freq = signal_freq + interference_offset_freq
            tester = tester_cmw(mode=1, rfport=1, cable_loss=signal_cableloss)
            tester.srx.trig_cw()
            tester.srx.para_settings(freq=signal_freq, level=signal_level)
            tester.srx.gen_switch('ON')
        txg.trriger_para_set(type='CONTinuous', count=1000)
        txg.output_state(1, 0)
        self.mem_ts.wrm(0xa04210a8, 6, 6, 1)  ##Tx DAC output test enable signal
        time.sleep(.5)
        self.mem_ts.wrm(0xa04210a8, 0, 0, 1)  ##Test buffer enable signal
        time.sleep(.5)
        self.mem_ts.wrm(0xa0421060, 12, 12, 1)
        self.mem_ts.wrm(0xa0421060, 11, 0, signal_freq)

        loginfo('txg_freq:  {}'.format(txg_freq))
        for flt_pkd_value in flt_pkd_list:
            self.mem_ts.wrm(0xa0421090, 20, 14, flt_pkd_value)
            for agc_tgt_sat in agc_tgt_sat_list:
                agc_tgt = agc_tgt_sat[0]
                agc_sat = agc_tgt_sat[1]
                self.mem_ts.wrm(0xa0421068, 9, 0, agc_tgt)
                self.mem_ts.wrm(0xa042106c, 25, 16, agc_sat)
                for rfpower in range(-80,10,3):
                    txg_rfpower = rfpower + interference_cableloss
                    txg.para_set(freq=txg_freq, power=txg_rfpower)
                    self.mem_ts.wr(0xa0422014, 0x0)
                    time.sleep(1)
                    self.mem_ts.wr(0xa0422014, 0x3)
                    res = spa.pk_search_avg_timesleep(timesleep=0.5)
                    if_pwr = res[0][1]
                    loginfo('flt_pkd_value:{}  agc_tgt:{}  agc_sat:{}  txg_freq:{}  RF pwer:{}  IF power:{}  gain:{}'.format(flt_pkd_value,agc_tgt,agc_sat,txg_freq,
                                                                                                                                          rfpower,
                                                                                                                                          if_pwr,
                                                                                                                                   if_pwr-rfpower))
                    if interference_en == 0:
                        fw1.write_data([flt_pkd_value,agc_tgt, agc_sat, -999, rfpower, if_pwr + if_cableloss, if_pwr - rfpower + if_cableloss])
                    else:
                        fw1.write_data([flt_pkd_value,agc_tgt, agc_sat, rfpower, signal_level, if_pwr + if_cableloss, if_pwr - signal_level + if_cableloss])
                    self.mem_ts.wr(0xa0422014, 0x0)

    def test232_agc_gain_table_set(self, file_path=r'E:\project_audio\TX232\report\NPW2\TX232_QFN68_A1\rxgain\tx232_rxgain_table_v1.xls'):
        fd = pd.read_excel(file_path,index_col=None,header=None)
        for i in range(len(fd.index)):
            reg_addr = eval(fd[0][i])
            addr_msb = fd[1][i]
            addr_lsb = fd[2][i]
            reg_value = eval(fd[7][i])
            self.mem_ts.wrm(reg_addr,addr_msb,addr_lsb,reg_value)


    def test232_agc_scan(self, agc_tgt_sat_list=[[407,580]], signal_cableloss=7, interference_cableloss=7, if_cableloss=1, signal_freq=2440, signal_level=-67,
                         interference_en=0,
                         interference_offset_freq=3):

        title = 'agc_tgt,agc_sat,interference_pwr(dbm)@offset {}MHz,signal_pwr(dbm),IF_pwr(dbm),dalta,agc_gain_idx\n'.format(interference_offset_freq)
        fname = self.get_filename('ts_bt_test/','test232_agc_scan_{}M'.format(interference_offset_freq))
        fw1 = csvreport(fname,title)
        self.test232_agc_gain_table_set()
        spa = Agilent()
        spa.set_param(1,1,10,30)
        txg = mxg.MXG()
        tester = tester_cmw(mode=1, rfport=1, cable_loss=signal_cableloss)
        if interference_en == 0:
            txg_freq = signal_freq
        else:
            txg_freq = signal_freq + interference_offset_freq
            tester.srx.trig_cw()
            tester.srx.para_settings(freq=signal_freq, level=signal_level)
            tester.srx.gen_switch('ON')
        txg.trriger_para_set(type='CONTinuous', count=1000)
        txg.output_state(1, 0)
        self.test232_rf_buf('rx')

        self.mem_ts.wrm(0xa0421064, 12, 12, 1)
        self.mem_ts.wrm(0xa0421064, 11, 0, signal_freq)

        for agc_tgt_sat in agc_tgt_sat_list:
            agc_tgt = agc_tgt_sat[0]
            agc_sat = agc_tgt_sat[1]
            self.mem_ts.wrm(0xa04210ac, 9, 0, agc_tgt)
            self.mem_ts.wrm(0xa04210a8, 11, 2, agc_sat)
            for rfpower in range(-90,10,3):
                txg_rfpower = rfpower + interference_cableloss
                txg.para_set(freq=txg_freq, power=txg_rfpower)
                self.mem_ts.wr(0xa0422014, 0x0)
                time.sleep(1)
                self.mem_ts.wr(0xa0422014, 0x3)
                res = spa.pk_search_avg_timesleep(timesleep=1.5)
                if_pwr = res[0][1]
                agc_gain_idx = self.mem_ts.rdm(0xa04210dc, 11, 8)
                loginfo('agc_tgt:{}  agc_sat:{}  txg_freq:{}  RF pwer:{}  IF power:{}  gain:{}  agc_gain_idx:{}'.format(agc_tgt,agc_sat,txg_freq,
                                                                                                                                      rfpower,
                                                                                                                                      if_pwr,
                                                                                                                               if_pwr-rfpower,agc_gain_idx))
                if interference_en == 0:
                    fw1.write_data([agc_tgt, agc_sat,-999, rfpower, if_pwr+if_cableloss, if_pwr - rfpower  + if_cableloss,agc_gain_idx])
                else:
                    fw1.write_data([agc_tgt, agc_sat, rfpower, signal_level, if_pwr + if_cableloss, if_pwr - signal_level + if_cableloss,agc_gain_idx])
                self.mem_ts.wr(0xa0422014, 0x0)
        txg.output_state(0, 0)
        tester.srx.gen_switch('OFF')

    def test232_agc_vs_gainindex(self, channel_list=[0], rf_cableloss=7):
        title = 'mode,channel,signal_pwr(dbm),agc_gain_idx\n'
        fname = self.get_filename('ts_bt_test/','test232_agc_vs_gainindex_{}'.format(self.board_name))
        fw1 = csvreport(fname,title)
        txg = mxg.MXG()
        txg.arb_waveform(rate='LE_1M')
        txg.arb_state(1)
        txg.trriger_para_set(type='CONTinuous', count=1000)
        txg.output_state(1, 1)
        for channel in channel_list:
            self.cmdstop(0)
            self.cmdstop(1)
            self.LE_RX(chan=channel, phy=1, mod=0, countMode=0)
            txg.para_set(freq=2402 + 2*channel, power=-67)
            for signal_pwr in range(-97,10):
                txg.set_power(signal_pwr + rf_cableloss)
                time.sleep(1)
                gain_index = self.jlink.rdm(0xa04210dc, 11, 8)

                fw1.write_data(['LE1M',channel,signal_pwr,gain_index])






    def le_rx_sniff(self, chan=0, rxpwr=-60):
        import matplotlib.pyplot as plt
        import numpy as np
        import time
        from math import *
        plt.ion()  # 开启interactive mode 成功的关键函数
        plt.figure(1)

        self.cmdstop(0)
        self.cmdstop(1)
        tester_signal = tester_cmw(mode=1, rfport=1, cable_loss=4.6)
        tester_signal.srx.gen_wave(repeat=1500, data_rate='LE_1M', dirty_en=0)
        tester_signal.srx.para_settings(freq=2402+chan*2, level=rxpwr)
        time.sleep(1)
        per_list=[]
        t=0
        t_list=[]
        while 1 :
            plt.clf()  # 清空画布上的所有内容
            self.LE_RX(chan=chan, phy=1, mod=0, countMode=0)
            tester_signal.srx.gen_switch('ON')
            time.sleep(1500 * 1e-3)
            res = self.cmdstop(0)
            res = res[0].split('bleRxCount=')[1].strip('\r\n')
            logdebug('rev_pkt:  {}'.format(res))
            rev_pkg = eval(res)
            per = (1500 - rev_pkg) * 100.0 / 1500
            # time.sleep(time_interval)
            t = t+1
            # per = eval(res[1])
            per_list.append(per)
            t_list.append(t)
            plt.plot(t_list,per_list, '-r')
            plt.pause(0.01)

    def test232_rssi_print(self):
        self.comport.ser.flushInput()
        self.comport.ser.flushOutput()
        self.rssi_print_end = 0
        while 1:

            self.rssi_print = self.comport.ser.readline()
            loginfo('{}'.format(self.rssi_print))
            if self.rssi_print_end == 1:
                break
            # elif self.rssi_print == '':
            #     break

    def test232_rssi_scan(self, freq=2476, mode=0, cable_loss=7, name_str=''):
        '''
        freq:   range2402--2480;unit:MHz
        mode:   0: br dh1;  1: LE1M
        '''

        title = 'rate,freq(MHz),instrument_tx_level(dBm),rssi/agc_gain_index\n'
        fname = self.get_filename('ts_bt_test/','test232_rssi_scan_{}'.format(name_str))
        fw1 = csvreport(fname,title)
        tester_signal = mxg.MXG('N5182B', 1)
        if mode!=0:
            tester_signal.arb_waveform(rate='LE_1M')
            rate = 'LE_1M'
        else:
            tester_signal.arb_waveform(rate='1M_DH1')
            rate = '1M_DH1'
        tester_signal.trriger_para_set(type='SINGle', count=100)
        tester_signal.output_state(1, 1)
        tester_signal.arb_state(1)
        tester_signal.para_set(freq=freq, power=-67 + cable_loss)
        time.sleep(2)

        # rssi_rd = threading.Thread(target=self.test232_rssi_print)
        # rssi_rd.start()
        for pwr in range(-97,0):
            tester_signal.set_power(pwr+cable_loss)
            time.sleep(1)
            if mode != 0:
                self.LE_RX(chan=(freq - 2402) / 2, phy=1, mod=0, countMode=0)
            else:
                self.BR_RX(chan=freq - 2402, len=27, ptype=5, rate=1, ltaddr=1, uap=0x0, lap=0x8)
            self.comport.ser.readlines()
            for num in range(10):
                tester_signal.trigger_on()

                time.sleep(1)
                rssi_print = self.comport.ser.readlines()
                # loginfo('{}'.format(rssi_print))
                if rssi_print != []:
                    break

            rssi_list=[]
            for i in range(1,len(rssi_print)):
                if rssi_print[i].find('rssi=')!=-1:
                    rssi_val = rssi_print[i].split('rssi=')[1].split(',')[0]
                    loginfo('{}'.format(rssi_val))
                    agc_gain_index = rssi_print[i].split('agc_gain_index=')[1]
                    loginfo('agc_gain_index:    {}'.format(agc_gain_index))
                    rssi_list.append(eval(rssi_val))
                    rssi_list.append(eval(agc_gain_index))

            # agc_gain_index = self.jlink.rdm(0xa04210dc,11,8)

            fw1.write_data([rate,freq,pwr]+rssi_list)

            if mode != 0:
                self.cmdstop(0)
            else:
                self.cmdstop(1)

    def rf_reg_read_all(self,file_path=''):
        # import docx
        # from docx import Document
        # title = 'reg_addr,reg_msb,reg_lsb,RW,reg_name,description,value_default,value_present_{}\n'.format(time.strftime("%Y-%m-%d %H:%M:%S", time.localtime()))
        # fname = self.get_filename('ts_bt_test/', 'rf_reg_read_all')
        # fw1 = csvreport(fname, title)
        # document = Document(file_path)
        # tables = document.tables
        # for i in range(1,len(tables)):
        #     table = tables[i]
        #     reg_addr = 0xa0421000 + (i-1)*4
        #     for table_row in range(1,len(table.rows)):
        #         reg_bit = table.cell(table_row,0).text.encode('utf-8')
        #         reg_bit_num = reg_bit.split(':')
        #         msb = reg_bit_num[0]
        #         if len(reg_bit_num) > 1 :
        #             lsb = reg_bit_num[1]
        #         elif len(reg_bit_num) > 2 :
        #             logerror('{} bit read fail'.format(reg_addr))
        #         else:
        #             lsb = msb
        #         value = self.mem_ts.rdm(reg_addr,eval(msb),eval(lsb))
        #
        #         reg_rw = table.cell(table_row,1).text.encode('utf-8')
        #         reg_value_default = table.cell(table_row,2).text.encode('utf-8')
        #         reg_name = table.cell(table_row,3).text.encode('utf-8')
        #         reg_description = table.cell(table_row,4).text.encode('utf-8')
        #         reg_description = reg_description.replace('\n', '; ').replace('\r', '; ')
        #         reg_description = r'{}'.format(reg_description)
        #         print(msb,lsb,reg_rw,reg_name,reg_description,reg_value_default)
        #         fw1.write_data([str(hex(reg_addr)),msb,lsb,reg_rw,reg_name,reg_description,reg_value_default,str(hex(value))],float_num=0)

        df=pd.read_excel(file_path,header=None,names=['a','b','c','d','e','f'])
        reg_addr_offset = 0
        for i in df['a'].index:
            offset = df.iloc[i,0]
            logdebug(offset)
            if str(offset).find('Offset:')!=-1:
                reg_addr_offset = offset.split('Offset:')[1]
                reg_addr_offset = eval(reg_addr_offset)
            reg_addr = 0xa0421000 + reg_addr_offset
            logdebug(reg_addr)
            if df.iloc[i,2] == 'RW' or df.iloc[i,2] == 'R':
                reg_bit = df.iloc[i,0]
                if reg_bit.find(':')!=-1:
                    bit_msb = reg_bit.split(':')[0]
                    bit_lsb = reg_bit.split(':')[1]
                else:
                    bit_msb = reg_bit
                    bit_lsb = reg_bit
                reg_value_pre = self.mem_ts.rdm(reg_addr,eval(bit_msb),eval(bit_lsb))
                df.iloc[i, 5] = hex(reg_value_pre)
        df.to_excel(r'E:\360MoveData\Users\123\Desktop\bt_rf_mpw3.xlsx',header=False,index=False)


    def rf_top_reg_to_code(self, file_path=''):
        import docx
        from docx import Document
        document = Document(file_path)
        tables = document.tables
        paragraph_list = []
        fname = self.get_filename('ts_bt_test/', 'rf_top_reg_to_code')
        logtime = time.strftime('%Y%m%d_%H%M%S', time.localtime(time.time()))
        with open('{}_{}.txt'.format(fname,logtime),'w+') as f:
            f.write('/*\n * DEFINES\n *****************************************************************************************\n */\n')
            # f.write('#define    RF_TOP_BASE_ADDR                0xA0421000\n')
            for paragraph in document.paragraphs:
                logdebug('{}'.format(paragraph.text))
                if paragraph.text.find('(offset:') != -1:
                    stru_name = paragraph.text.split('(offset:')[0].split(' ')[1]
                    offset = paragraph.text.split('(offset:')[1].split(')')[0].replace(' ','')
                    reg_addr = 0xa0421000 + eval(offset)
                    w_str = '#define    {}_REG'.format(stru_name)
                    w_str = w_str + (40-len(w_str))*' ' + '{}\n'.format(hex(reg_addr).replace('L',''))
                    f.write(w_str)


            for paragraph in document.paragraphs:
                logdebug('{}'.format(paragraph.text))
                if paragraph.text.find('(offset:') != -1:
                    stru_name = paragraph.text.split('(offset:')[0].split(' ')[1]
                    offset = paragraph.text.split('(offset:')[1].split(')')[0].replace(' ','')
                    table_num = eval(offset)/4 +1
                    logdebug('table_num: {}'.format(table_num))
                    table = tables[table_num]
                    struct_wstr1 = 'typedef union{\n    struct{\n'
                    struct_wstr2_list = ''
                    struct_reset_list = ''
                    for table_row in range(len(table.rows)-1,0,-1):
                        logdebug('lens: {} ;table_row: {}'.format(len(table.rows),table_row))
                        reg_bit = table.cell(table_row,0).text.encode('utf-8')
                        reg_bit_num = reg_bit.split(':')
                        msb = eval(reg_bit_num[0])
                        if len(reg_bit_num) > 1 :
                            lsb = eval(reg_bit_num[1])
                        elif len(reg_bit_num) > 2 :
                            logerror('{} bit read fail'.format(stru_name))
                        else:
                            lsb = msb
                        reg_rw = table.cell(table_row,1).text.encode('utf-8')
                        reg_value_default = table.cell(table_row,2).text.encode('utf-8')
                        reg_name = table.cell(table_row,3).text.encode('utf-8')
                        reg_description = table.cell(table_row,4).text.encode('utf-8')
                        reg_description = reg_description.replace('\n', '; ').replace('\r', '; ')
                        reg_description = r'{}'.format(reg_description)
                        struct_wstr2 = '        unsigned int {} : {};'.format(reg_name,(msb-lsb+1))
                        struct_wstr2 = struct_wstr2 + (60-len(struct_wstr2))*' ' + '/* bit[{}:{}]; {}; {} */\n'.format(lsb,msb,reg_rw,reg_description)
                        struct_wstr2_list = struct_wstr2_list + struct_wstr2
                        if reg_rw == 'RW':
                            struct_reset = '    {}->{}_BIT.{} = {};\n'.format(stru_name,stru_name,reg_name,reg_value_default)
                            struct_reset_list = struct_reset_list + struct_reset
                    # struct_wstr3 = '    }'+'{}_BIT;\n    uint32_t    {}_BYTE;\n'.format(stru_name,stru_name)+'}'+'_{};\n'.format(stru_name)
                    struct_wstr3 = '    }' + 'BF;\n    uint32_t    U32;\n' + '}' + '_{};\n'.format(stru_name)
                    struct_wstr4 = 'volatile _{} *{} = (_{}*){}_REG;\n'.format(stru_name,stru_name,stru_name,stru_name)
                    struct_wstr5 = 'extern void {}_reset(void)\n'.format(stru_name.lower()) +'{\n'+ struct_reset_list + '}\n'
                    f.write(struct_wstr1+struct_wstr2_list+struct_wstr3+struct_wstr4+struct_wstr5)

    def excel_to_c(self,filepath=''):
        import xlrd
        from glob import glob
        '''
        获取要转换的excel文件路径
        '''
        # path_name_list=[]
        # if os.path.isdir(filepath):
        #     path_names = os.listdir(filepath)
        #     for _path_name in path_names:
        #         if _path_name.find('bt_') !=-1:
        #             _path_name = os.path.join(filepath,_path_name)
        #             path_name_list.append(_path_name)
        path_name_list = glob(filepath+'\\bt_*.xlsx')
        '''
        读取excel表格并转换为c代码
        '''
        for path_name in path_name_list:
            datafile = xlrd.open_workbook(filename=path_name)
            names = datafile.sheet_names()
            for name in names:
                logdebug(name)
                fname = self.get_filename('ts_bt_test/', name)
                with open('{}.h'.format(fname), 'w+') as f:
                    # f.write('struct {}'.format(name.upper())+'{\n')
                    f.write('#ifndef __{}_H__\n'.format(name.upper()))
                    f.write('#define __{}_H__\n \n'.format(name.upper()))
                    f.write('#include "chip_public.h"\n \n')
                    if name == 'bt_modem':
                        table = datafile.sheet_by_name(name)
                        nrows = table.nrows
                        logdebug(nrows)
                        struct_name_list = []
                        for i in range(nrows):
                            if table.cell_value(i,0).find('0x')!=-1:            ##addr offset所在行查找
                                addr_offset = table.cell_value(i,0)
                                struct_name = table.cell_value(i,1)
                                struct_name_list.append([struct_name,addr_offset])
                                logdebug('addr_offset: {};   struct_name:   {}'.format(addr_offset,struct_name))
                                w_str = '//offset address: {}\n'.format(addr_offset)+'  union REG_{}\n'.format(struct_name.upper())+'  {\n'+'    DRV_IOM uint32_t raw;\n'+'    ' \
                                                                                                                                                                       'struct'+'{\n'
                                for j in range(40):
                                    if table.cell_value(i+j,0).find(':0]')!=-1 or table.cell_value(i+j,0) == '0]':          ##一个32位寄存器所有bit定义的起始行查找
                                        w_bit_row_start = i+j
                                        w_bit_row_stop = i+1
                                        logdebug('w_bit_row_start:  {}; w_bit_row_stop: {}'.format(w_bit_row_start,w_bit_row_stop))
                                        break
                                bit_num_check = 0
                                for row in range(w_bit_row_start,i,-1):                         ##bit位定义，bit位数确定以及定义，描述
                                    reg_bit = table.cell_value(row,0).replace(']','').replace('[','')
                                    logdebug('reg_bit:  {}'.format(reg_bit))
                                    if reg_bit.find(':')!=-1:
                                        _reg_bit = reg_bit.split(':')
                                        bit_num = eval(_reg_bit[0])-eval(_reg_bit[1])+1
                                    elif reg_bit == '':
                                        continue
                                    else:
                                        bit_num = 1
                                    logdebug('bit_num:  {}'.format(bit_num))
                                    bit_num_check = bit_num_check + bit_num
                                    reg_name = table.cell_value(row,1).replace('reserved','  ')
                                    reg_name = reg_name.split('[')[0]
                                    reg_desc = table.cell_value(row,4).replace('\n',';  ')
                                    logdebug('reg_name:{}    reg_desc:{}'.format(reg_name,reg_desc))
                                    w_bit = '      DRV_IOM uint32_t {}'.format(reg_name)+(40-len(reg_name))*' '+':{} ;      //{}\n'.format(bit_num,reg_desc)
                                    w_str = w_str + w_bit
                                if bit_num_check != 32:
                                    logerror('{}    {}      bit total is {},please check table!'.format(name,struct_name,bit_num_check))
                                    raise
                                w_str = w_str + '         }' +'bits;\n' + '  }' + ';\n'+'\n'
                                f.write(w_str)

                        w_str = 'typedef struct {}_reg\n'.format(name) + '{\n'
                        f.write(w_str)
                        addr_offset_pre = 0
                        addr_step_flag = 0
                        for struct_name_t in struct_name_list:
                            addr_offset_now = eval(struct_name_t[1])
                            addr_step = addr_offset_now - addr_offset_pre
                            if addr_step >4:
                                w_str_addr_offset = 'DRV_IOM uint32_t  rev_{}[({} - {}) / 4 - 1];\n'.format(addr_step_flag,addr_offset_now,addr_offset_pre)
                                addr_step_flag = addr_step_flag+1
                                f.write(w_str_addr_offset)
                            w_str_name_t = 'union REG_{}'.format(struct_name_t[0].upper()) + (40-len(struct_name_t[0]))*' ' + '{};\n'.format(struct_name_t[0].upper())
                            f.write(w_str_name_t)
                            addr_offset_pre = addr_offset_now
                        f.write('}'+'{}_reg_t;\n \n'.format(name))
                        add_str = "#define {}_READ_REG(_BA, _REG)                             (_BA->_REG.raw)\n".format(name.upper())
                        add_str = add_str + "#define {}_WRITE_REG(_BA, _REG, _V)                        (_BA->_REG.raw = _V)\n".format(name.upper())
                        add_str = add_str + "#define {}_READ_BITS(_BA, _REG, _FLD)".format(name.upper())+"                      ({union REG_##_REG _r;_r.raw = " \
                                                                                                         "_BA->_REG.raw;_r.bits._FLD;})\n"
                        add_str = add_str + "#define {}_WRITE_BITS(_BA, _REG, _FLD, _V)".format(name.upper()) + "                 ({union REG_##_REG _r;_r.raw = # _BA->_REG.raw;_r.bits._FLD = _V;_BA->_REG.raw = _r.raw;})\n"
                        add_str = add_str + "#define {}_READ_MULTI_BITS(_BA, _REG, ELEMENT, _FLD)".format(name.upper()) + "       ({union REG_##_REG _r;_r.raw = _BA->_REG.raw;_r.ELEMENT._FLD;})\n"
                        add_str = add_str + "#define {}_WRITE_MULTI_BITS(_BA, _REG, ELEMENT, _FLD, _V)".format(name.upper()) + "  ({union REG_##_REG _r;_r.raw = _BA->_REG.raw;_r.ELEMENT._FLD = " \
                                            "_V;_BA->_REG.raw = _r.raw;})\n"
                        add_str = add_str + "#define {}_WRITE_BITS_MSK(_BA, _REG, _FLD, _V)".format(name.upper()) + "             ({union REG_##_REG _r;_r.raw = _BA->_REG.raw;_r.bits._FLD = " \
                                            "_V;_r.bits._FLD##_msk = 1;_BA->_REG.raw = _r.raw;})\n"

                        f.write(add_str)
                        f.write('#define {}                                (({}_reg_t *) {}_BASE)\n \n#endif'.format(name.upper(),name,name.upper()))
                        f.close()
                    # elif name == 'bt_rf':
                    #     table = datafile.sheet_by_name(name)
                    #     nrows = table.nrows
                    #     logdebug(nrows)
                    #     struct_name_list = []
                    #     for i in range(nrows):
                    #         if type(table.cell_value(i, 1)) == int or type(table.cell_value(i, 1)) == float:
                    #             aa = str(table.cell_value(i, 1))
                    #         else:
                    #             aa = table.cell_value(i, 1)
                    #         logdebug(aa.encode('utf-8'))
                    #         if aa.encode('utf-8').find('Offset:') != -1:
                    #             addr_offset = table.cell_value(i, 0).split(':')[1]
                    #             if name == 'bt_core_reg':
                    #                 struct_name = table.cell_value(i - 1, 0).split('(')[0].replace(' ', '')
                    #             else:
                    #                 struct_name = table.cell_value(i - 1, 0).split('(')[1].replace(')', '')
                    #             logdebug('addr_offset: {};   struct_name:   {}'.format(addr_offset, struct_name))
                    #             struct_name_list.append([struct_name, addr_offset])
                    #             w_str = '//offset address: {}\n'.format(addr_offset) + '  union REG_{}\n'.format(
                    #                 struct_name.upper()) + '  {\n' + '    DRV_IOM uint32_t raw;\n' + '    ' \
                    #                                                                                  'struct' + '{\n'
                    #             for j in range(1, 40):
                    #                 if str(table.cell_value(i + j, 0)).find(':0') != -1 or str(table.cell_value(i + j, 0)) == '0' or str(table.cell_value(i + j, 0)) == '0.0':
                    #                     w_bit_row_start = i + j
                    #                     w_bit_row_stop = i + 1
                    #                     logdebug('w_bit_row_start:  {}; w_bit_row_stop: {}'.format(w_bit_row_start, w_bit_row_stop))
                    #                     break
                    #             bit_num_check = 0
                    #             for row in range(w_bit_row_start, i + 1, -1):
                    #                 reg_bit = str(table.cell_value(row, 0)).replace(']', '').replace('[', '')
                    #                 logdebug('reg_bit:  {}'.format(reg_bit))
                    #                 if reg_bit.find(':') != -1:
                    #                     _reg_bit = reg_bit.split(':')
                    #                     bit_num = eval(_reg_bit[0]) - eval(_reg_bit[1]) + 1
                    #                 elif reg_bit == '':
                    #                     continue
                    #                 else:
                    #                     bit_num = 1
                    #                 logdebug('bit_num:  {}'.format(bit_num))
                    #                 bit_num_check = bit_num_check + bit_num
                    #                 reg_name = table.cell_value(row, 1).replace('NA', '  ').replace('reserved', '  ')
                    #                 # reg_name = table.cell_value(row, 2).replace('reserved', '  ')
                    #                 reg_name = reg_name.split('[')[0]
                    #                 reg_desc = table.cell_value(row, 4).replace('-', '').replace('\n', ';  ').encode('utf-8')
                    #                 logdebug('reg_name:{}    reg_desc:{}'.format(reg_name, reg_desc))
                    #                 w_bit = '      DRV_IOM uint32_t {}'.format(reg_name) + (40 - len(reg_name)) * ' ' + ':{} ;      //{}\n'.format(bit_num, reg_desc)
                    #                 w_str = w_str + w_bit
                    #             if bit_num_check != 32:
                    #                 logerror('{}    {}      bit total is {},please check table!'.format(name, struct_name, bit_num_check))
                    #                 raise
                    #             w_str = w_str + '         }' + 'bits;\n' + '  }' + ';\n' + '\n'
                    #             f.write(w_str)
                    #
                    #     w_str = 'typedef struct {}_reg\n'.format(name) + '{\n'
                    #     f.write(w_str)
                    #     addr_offset_pre = 0
                    #     addr_step_flag = 0
                    #     for struct_name_t in struct_name_list:
                    #         addr_offset_now = eval(struct_name_t[1])
                    #         addr_step = addr_offset_now - addr_offset_pre
                    #         if addr_step > 4:
                    #             w_str_addr_offset = 'DRV_IOM uint32_t  rev_{}[({} - {}) / 4 - 1];\n'.format(addr_step_flag, addr_offset_now, addr_offset_pre)
                    #             addr_step_flag = addr_step_flag + 1
                    #             f.write(w_str_addr_offset)
                    #         w_str_name_t = 'union REG_{}'.format(struct_name_t[0].upper()) + (40 - len(struct_name_t[0])) * ' ' + '{};\n'.format(struct_name_t[0].upper())
                    #         f.write(w_str_name_t)
                    #         addr_offset_pre = addr_offset_now
                    #     f.write('}' + '{}_reg_t;\n \n'.format(name))
                    #     f.write('#define {}                                (({}_reg_t *) {}_BASE)\n \n#endif'.format(name.upper(), name, name.upper()))
                    #     f.close()
                    else:
                        table = datafile.sheet_by_name(name)
                        nrows = table.nrows
                        logdebug(nrows)
                        struct_name_list = []
                        for i in range(nrows):
                            if type(table.cell_value(i, 0)) == int or type(table.cell_value(i, 0)) == float :
                                aa = str(table.cell_value(i, 0))
                            else:
                                aa = table.cell_value(i, 0)
                            logdebug(aa.encode('utf-8'))
                            if aa.encode('utf-8').find('Offset:') != -1:
                                addr_offset = table.cell_value(i,0).split(':')[1]
                                if name == 'bt_core_reg':
                                    struct_name = table.cell_value(i - 1, 0).split('(')[0].replace(' ', '')
                                else:
                                    struct_name = table.cell_value(i-1,0).split('(')[1].replace(')','')
                                logdebug('addr_offset: {};   struct_name:   {}'.format(addr_offset,struct_name))
                                struct_name_list.append([struct_name,addr_offset])
                                w_str = '//offset address: {}\n'.format(addr_offset)+'  union REG_{}\n'.format(struct_name.upper())+'  {\n'+'    DRV_IOM uint32_t raw;\n'+'    ' \
                                                                                                                                                                       'struct'+'{\n'
                                for j in range(1,40):
                                    if str(table.cell_value(i+j,0)).find(':0')!=-1 or str(table.cell_value(i+j,0)) == '0' or str(table.cell_value(i+j,0)) == '0.0':
                                        w_bit_row_start = i+j
                                        w_bit_row_stop = i+1
                                        logdebug('w_bit_row_start:  {}; w_bit_row_stop: {}'.format(w_bit_row_start,w_bit_row_stop))
                                        break
                                bit_num_check = 0
                                for row in range(w_bit_row_start,i+1,-1):
                                    reg_bit = str(table.cell_value(row,0)).replace(']','').replace('[','')
                                    logdebug('reg_bit:  {}'.format(reg_bit))
                                    if reg_bit.find(':')!=-1:
                                        _reg_bit = reg_bit.split(':')
                                        bit_num = eval(_reg_bit[0])-eval(_reg_bit[1])+1
                                    elif reg_bit == '':
                                        continue
                                    else:
                                        bit_num = 1
                                    logdebug('bit_num:  {}'.format(bit_num))
                                    bit_num_check = bit_num_check + bit_num
                                    reg_name = table.cell_value(row,1).replace('NA','  ').replace('reserved', '  ')
                                    # reg_name = table.cell_value(row, 1).replace('reserved', '  ')
                                    reg_name = reg_name.split('[')[0]
                                    reg_desc = table.cell_value(row,4).replace('-','').replace('\n',';  ').encode('utf-8')
                                    logdebug('reg_name:{}    reg_desc:{}'.format(reg_name,reg_desc))
                                    w_bit = '      DRV_IOM uint32_t {}'.format(reg_name)+(40-len(reg_name))*' '+':{} ;      //{}\n'.format(bit_num,reg_desc)
                                    w_str = w_str + w_bit
                                if bit_num_check != 32:
                                    logerror('{}    {}      bit total is {},please check table!'.format(name,struct_name,bit_num_check))
                                    raise
                                w_str = w_str + '         }' + 'bits;\n' + '  }' + ';\n' + '\n'
                                f.write(w_str)

                        w_str = 'typedef struct {}_reg\n'.format(name) + '{\n'
                        f.write(w_str)
                        addr_offset_pre = 0
                        addr_step_flag = 0
                        for struct_name_t in struct_name_list:
                            addr_offset_now = eval(struct_name_t[1])
                            addr_step = addr_offset_now - addr_offset_pre
                            if addr_step >4:
                                w_str_addr_offset = 'DRV_IOM uint32_t  rev_{}[({} - {}) / 4 - 1];\n'.format(addr_step_flag,addr_offset_now,addr_offset_pre)
                                addr_step_flag = addr_step_flag+1
                                f.write(w_str_addr_offset)
                            w_str_name_t = 'union REG_{}'.format(struct_name_t[0].upper()) + (40-len(struct_name_t[0]))*' ' + '{};\n'.format(struct_name_t[0].upper())
                            f.write(w_str_name_t)
                            addr_offset_pre = addr_offset_now
                        f.write('}'+'{}_reg_t;\n \n'.format(name))
                        add_str = "#define {}_READ_REG(_BA, _REG)                             (_BA->_REG.raw)\n".format(name.upper())
                        add_str = add_str + "#define {}_WRITE_REG(_BA, _REG, _V)                        (_BA->_REG.raw = _V)\n".format(name.upper())
                        add_str = add_str + "#define {}_READ_BITS(_BA, _REG, _FLD)".format(name.upper())+"                      ({union REG_##_REG _r;_r.raw = " \
                                                                                                         "_BA->_REG.raw;_r.bits._FLD;})\n"
                        add_str = add_str + "#define {}_WRITE_BITS(_BA, _REG, _FLD, _V)".format(name.upper()) + "                 ({union REG_##_REG _r;_r.raw = # _BA->_REG.raw;_r.bits._FLD = _V;_BA->_REG.raw = _r.raw;})\n"
                        add_str = add_str + "#define {}_READ_MULTI_BITS(_BA, _REG, ELEMENT, _FLD)".format(name.upper()) + "       ({union REG_##_REG _r;_r.raw = _BA->_REG.raw;_r.ELEMENT._FLD;})\n"
                        add_str = add_str + "#define {}_WRITE_MULTI_BITS(_BA, _REG, ELEMENT, _FLD, _V)".format(name.upper()) + "  ({union REG_##_REG _r;_r.raw = _BA->_REG.raw;_r.ELEMENT._FLD = " \
                                            "_V;_BA->_REG.raw = _r.raw;})\n"
                        add_str = add_str + "#define {}_WRITE_BITS_MSK(_BA, _REG, _FLD, _V)".format(name.upper()) + "             ({union REG_##_REG _r;_r.raw = _BA->_REG.raw;_r.bits._FLD = " \
                                            "_V;_r.bits._FLD##_msk = 1;_BA->_REG.raw = _r.raw;})\n"

                        f.write(add_str)
                        f.write('#define {}                                (({}_reg_t *) {}_BASE)\n \n#endif'.format(name.upper(),name,name.upper()))
                        f.close()

    def get_filename(self, folder, file_name, sub_folder=''):
        '''
        :folder: file store folder
        :file_name:  file name
        :sub_folder: if not need, it may be default ""
        '''
        if rfglobal.file_folder=="":
            rfdata_path = './rftest/rfdata/'
        else:
            rfdata_path = './rftest/rfdata/%s/'%rfglobal.file_folder
            if os.path.exists(rfdata_path) == False:
                os.mkdir(rfdata_path)

        data_path1 = rfdata_path+'%s/'%(folder)
        if os.path.exists(data_path1) == False:
            os.mkdir(data_path1)

        filetime = time.strftime('%Y%m%d_%H%M%S',time.localtime(time.time()));
        # mac = self.read_mac()
        mac = ''

        gen_folder = '_%s'%(filetime[0:8])
        data_path2 = data_path1 +'%s/'%(gen_folder)
        if os.path.exists(data_path2) == False:
            os.mkdir(data_path2)

        fname = '%s'%(file_name)
        outfile_name = data_path2 + fname

        if sub_folder != '':
            gen_folder = '%s_%s'%(sub_folder,filetime[0:8])
            sub_path = data_path2+'%s/'%(gen_folder)
            if os.path.exists(sub_path) == False:
                os.mkdir(sub_path)

            outfile_name = sub_path + file_name

        return outfile_name

class bt_signaling(object):

    def __init__(self, comport, chipv='TX212',board_name='', jlink='', rfport=1, atten=4.6, le_com_baudrate='B500K'):
        self.rfport = rfport
        self.atten = atten
        self.comport = comport
        self.chipv = chipv
        self.board_name = board_name
        self.jlink = jlink
        self.comport.Hexmd=1
        self.le_com_baudrate = le_com_baudrate
        # if jlink_en !=0:
        #     self.jlink = jlink(jlink_sn=jlink_sn)
            # self.jlink = pylink.JLink()
            # self.jlink_sn = jlink_sn
            # self.jlink.open()
            # self.jlink.set_tif(pylink.enums.JLinkInterfaces.SWD)
            # self.jlink.connect('ARMCM4_FP')

        cmw_bt().reset()
        self.csp = combined_signal_path()
        self.csp.signaling_switch(1)
        self.csp.rf_port(mode='signaling', rfport=self.rfport, atten=self.atten)
        self.br_rate_list = ['DH1','DH3','DH5']
        self.edr_rate_list = ['2_DH1','2_DH3','2_DH5','3_DH1','3_DH3','3_DH5']
        self.dic_edr_rate = {
            '2_DH1':    'E21P',
            '2_DH3':    'E23P',
            '2_DH5':    'E25P',
            '3_DH1':    'E31P',
            '3_DH3':    'E33P',
            '3_DH5':    'E35P'
            }

    # def jlink_connect(self):
    #     self.jlink.jlink_connect()
        # self.jlink.open(self.jlink_sn)
        # self.jlink.set_tif(pylink.enums.JLinkInterfaces.SWD)
        # self.jlink.connect('ARMCM4_FP')

    # def jlink_mem_rd(self, addr=0xa04210a0):
    #     res = self.jlink.memory_read(addr,4)
    #     res_int = 0x1000000*res[3] + 0x10000*res[2] + 0x100*res[1] + res[0]
    #     res = '0x%x'%(res_int)
    #     loginfo('addr:  {}  value:  {}'.format(hex(addr), res))
    #     return res
    #
    # def jlink_mem_rdm(self, addr=0xa04210a0, msb=1, lsb=0):
    #     result = self.jlink_mem_rd(addr=addr)
    #     mask = (1<<(msb+1)) - (1<<lsb);
    #     try:
    #         result = (eval(result) & mask) >> lsb;
    #         # logdebug("resm:" + "0x%x"%result)
    #     except:
    #         logerror("mem reads " + "%s"%result)
    #     return result
    #
    # def jlink_mem_wr(self, addr=0xa04210a0, value=0x18c619):
    #
    #     value_list = [value&0xff,(value&0xff00)>>8,(value&0xff0000)>>16,(value&0xff000000)>>24]
    #     self.jlink.memory_write(addr,value_list)
    #     res = self.jlink_mem_rd(addr)
    #
    # def jlink_mem_wrm(self, addr, msb, lsb, value):
    #     result = self.jlink_mem_rd(addr=addr)
    #     mask = (1<<(msb+1)) - (1<<lsb);
    #     try:
    #         result = (eval(result) & ~mask) | ((value << lsb) & mask);
    #         return self.jlink_mem_wr(addr,result);
    #     except:
    #         logerror("mem write fail, reads  %s"%result)
    #         return

    def hci_reset(self):
        self.comport.req_com(cmdstr='\x01\x03\x0c\x00',timeout=1,endstr='TS')

    def hci_read_buffer_size(self):
        self.comport.req_com(cmdstr='\x01\x05\x10\x00', timeout=1, endstr='TS')

    def hci_enable_under_test_mode(self):
        self.comport.req_com(cmdstr='\x01\x03\x18\x00', timeout=1, endstr='TS')

    def hci_enable_write_both_scan(self):
        self.comport.req_com(cmdstr='\x01\x1a\x0c\x01\x03', timeout=1, endstr='TS')

    def hci_auto_accept_all_connection(self):
        self.comport.req_com(cmdstr='\x01\x05\x0c\x03\x02\x00\x02', timeout=1, endstr='TS')

    def dut_signal_mode_set(self):
        res = self.comport.req_com(cmdstr='signaltest', wr_end='\r\n', timeout=5)
        if res.find('signal test')!=-1:
            self.comport.Hexmd = 1
            self.hci_reset()
            self.hci_read_buffer_size()
            self.hci_enable_under_test_mode()
            self.hci_enable_write_both_scan()
            self.hci_auto_accept_all_connection()
            self.comport.Hexmd = 0
        else:
            logwarn('can not enter signal test mode')

    def config_per_le_connect_usb(self, rate='LE1M', packet_len=37):
        self.csp.reset()
        self.csp.clean()
        time.sleep(1)
        self.csp.signaling_switch(1)
        time.sleep(2)
        self.csp.rf_port(mode='signaling', rfport=self.rfport, atten=self.atten)
        self.csp.sig_opmode(mode='RFT')
        self.csp.sig_std(std='LESignaling')
        self.csp.sig_btype('LE')
        self.csp.sig_le_phy(phy=rate)
        self.csp.RF_Frequency_Settings_rx(mode='DTM', ch_tx=0)
        self.csp.RF_Power_settings(rx_level=-40, tx_power=10,margin=8)
        self.csp.RF_Power_settings_autoranging(1)
        self.csp.config_connect_le_packet_len(rate=rate, len=packet_len)
        self.csp.config_connect_le_packet_pattern(rate=rate, pattern='PRBS9')
        self.csp.config_hwinterface('RS232')
        self.csp.config_comport()
        self.csp.config_comport_baudrate(self.le_com_baudrate)
        self.csp.check_le_connect_usb()
        time.sleep(1)

        self.csp.config_rxq_repetion(rep='SING')
        self.csp.config_rxq_le_packets(rate=rate, num=1500)
        self.csp.per_meas_state(0)
        time.sleep(1)

    def config_classic_ber(self, rate='BR', chan=0, rx_level=-60, tx_power=10):
        '''
        配置经典蓝牙连接CMW流程：
        复位仪器，打开信令模式，配置好RF参数和蓝牙包的参数，inquire EUT的BD，paging EUT并连接EUT。
        '''
        if rate not in ('BR','EDR'):
            raise StandardError('rate command wrong')
        # self.dut_signal_mode_set()  #设置dut进入BT信令测试模式
        self.csp.reset()    #复位cmw设置
        self.csp.clean()
        self.csp.signaling_switch(1)    #1：仪器的蓝牙信令模式打开，0：关闭
        self.csp.rf_port(mode='signaling', rfport=self.rfport, atten=self.atten)    #设置信令模式下 rfport和线损
        self.csp.sig_opmode(mode='RFT')
        self.csp.sig_std(std='CLAS')
        self.csp.sig_btype(btype=rate)
        self.csp.config_sig_testmode(testmode='LOOP')
        self.csp.RF_Frequency_Settings_rx(mode='LOOP', ch_tx=chan)
        self.csp.RF_Power_settings(rx_level=rx_level, tx_power=tx_power, margin=8)
        # self.csp.RF_Power_settings_autoranging()
        if rate == 'BR':
            self.csp.config_connect_br_packet_pattern(pattern='PRBS9')  #设置BR包 payload 数据类型
            self.csp.config_connect_br_packet_ptype(ptype='DH1')    #设置BR包类型
            # self.csp.config_connect_br_packet_len(len=packet_len)   #设置BR包长
        else:
            self.csp.config_connect_edr_packet_pattern(pattern='PRBS9') #设置EDR包 payload 数据类型
            self.csp.config_connect_edr_packet_ptype(ptype='E21P')  #设置EDR的包类型，E21P：2-DH1，E31P：3-DH1
            # self.csp.config_connect_edr_packet_len(len=packet_len)  #设置EDR包长
        self.csp.config_paging_classic_NOResponses(1)   #设置stop inquire的条件，当inquire到的BD数量为设定值时则stop inquire
        self.csp.bt_connect_action(action='INQuire')    #开始inquire命令
        time.sleep(2)
        for i in range(5):
            self.csp.bt_connect_action(action='INQuire')  # 开始inquire命令
            time.sleep(2)
            target = self.csp.get_paging_classic_target()   #返回inquire到的BD 列表
            self.csp.config_paging_classic_target(target=1) #设置inquire到的第一个BD用来paging
            self.csp.bt_connect_action(action='TMConnect')  #
            res = self.csp.get_bt_connect_state()
            if res == 'TCON':
                break

        time.sleep(1)

    def cmd_br_connect(self):
        for i in range(5):
            self.csp.bt_connect_action(action='INQuire')  # 开始inquire命令
            time.sleep(2)
            target = self.csp.get_paging_classic_target()   #返回inquire到的BD 列表
            self.csp.config_paging_classic_target(target=1) #设置inquire到的第一个BD用来paging
            self.csp.bt_connect_action(action='TMConnect')  #
            res = self.csp.get_bt_connect_state()
            if res == 'TCON':
                loginfo('*********** connect sucess ************')
                time.sleep(2)
                break
            else:
                logwarn('connect fail, state: {} , please try again'.format(res))

    def le_tx_tp_debug(self, chan_list=[0,19,39]):
        '''

        '''
        title = 'type,tp_gain_value,channel,nominal_pow(dBm),'
        title = title + 'delta_f1_avg(KHz),delta_f1_avg_max(KHz),delta_f2_avg(KHz),delta_f2_avg_max(KHz),mod_ratio,delta_f2_99(KHz),'
        title = title + '\n'
        fname = self.get_filename('ts_bt_test/','le_tx_tp_debug_{}'.format(self.board_name))
        fw1=csvreport(fname,title)
        self.config_per_le_connect_usb(rate='LE1M', packet_len=37)
        self.csp.config_rx_level(rxpwr=-67)
        self.csp.tx_measure_para(tout=5, repetition='SINGleshot', count=1000, MOEXception='OFF')
        for rate in ['LE1M','LE2M']:
            if rate == 'LE500k':
                _rate = 'LELR'
                coding = 'S2'
                mode = 'LRANge'
                self.csp.sig_le_lr_coding(coding=coding)
            elif rate == 'LE125K':
                _rate = 'LELR'
                coding = 'S8'
                mode = 'LRANge'
                self.csp.sig_le_lr_coding(coding=coding)
            else:
                mode = rate
                _rate = rate

            self.csp.mode_set(mode=mode)
            self.csp.sig_le_phy(phy=_rate)
            bttest = bt_test(self.comport,self.chipv,self.jlink)
            self.jlink.wrm(0xa0421020,25,24,0x1)
            for tp_gain_val_chan in [0,19,39]:
                if rate == 'LE1M':
                    tp_gain_cal_value = bttest.test232_tp_gain_cal(0,2402+tp_gain_val_chan*2)
                else:
                    tp_gain_cal_value = bttest.test232_tp_gain_cal(1, 2402 + tp_gain_val_chan * 2)
                for chan in chan_list:
                    self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan)
                    self.csp.config_connect_le_packet_pattern(rate=_rate, pattern='P44')
                    res = self.csp.get_modulation_measure_res(rate=mode,data_type='AVER')
                    logdebug('{}'.format(res))
                    delta_f1_avg = eval(res[6]) / 1000.00
                    nominal_pow = eval(res[12])

                    res = self.csp.get_modulation_measure_res(rate=mode, data_type='MAX')
                    logdebug('{}'.format(res))
                    delta_f1_avg_max = eval(res[6]) / 1000.00


                    time.sleep(0.5)

                    self.csp.config_connect_le_packet_pattern(rate=_rate, pattern='P11')
                    res = self.csp.get_modulation_measure_res(rate=mode,data_type='AVER')
                    logdebug('{}'.format(res))

                    delta_f2_99 = eval(res[2]) / 1000.00
                    delta_f2_avg = eval(res[9]) / 1000.00
                    mod_ratio = delta_f2_avg / delta_f1_avg

                    res = self.csp.get_modulation_measure_res(rate=mode, data_type='MAX')
                    logdebug('{}'.format(res))
                    delta_f2_avg_max = eval(res[9]) / 1000.00

                    fw1.write_data([rate,tp_gain_cal_value,chan,nominal_pow,delta_f1_avg,delta_f1_avg_max,delta_f2_avg,delta_f2_avg_max,mod_ratio,delta_f2_99])



    def le_tx_meas(self, chan_list=[0], rate_list=['LE1M'], packet_len=37, csv_save=True, report_save=True, fig_en=0):
        '''
        rate_list:  LE1M、LE2M、LE500K、LE125K
        '''
        if csv_save:
            title = 'channel,type,nominal_pow(dBm),peak_pow(dBm),leakage_pow(dBm),freq_accuracy(KHz),freq_drift(KHz),drift_rate(Hz/50us),'
            title = title + 'delta_f1_avg(KHz),delta_f1_min(KHz),delta_f1_max(KHz),delta_f2_avg(KHz),delta_f2_min(KHz),delta_f2_max(KHz),mod_ratio,delta_f1_99(KHz),delta_f2_99(KHz),'
            title = title + 'acp_list_21ch\n'
            fname = self.get_filename('ts_bt_test/','test_le_tx_{}'.format(self.board_name))
            fw1=csvreport(fname,title)
        if report_save:
            title2 = 'BR_tx_performance'
            fname2 = self.get_filename('ts_bt_test/','test_le_tx_report_{}'.format(self.board_name))
            fw2=csvreport(fname2,title2)
            item = ['channel', 'rate', u'TP/TRM-LE/CA/BV-01-C - Output power  - Average Power', u' - Peak Power',
                    u'- Peak - Average Power',
                    u'- leakage Power',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -10',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -9',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -8',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -7',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -6',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -5',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -4',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -3',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -2',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -1',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -    Ptx',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   1',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   2',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   3',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   4',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   5',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   6',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   7',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   8',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   9',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   10',
                    u'TP/TRM-LE/CA/BV-05-C - Modulation Characteristics - Fre Accuracy',
                    u'- Delta F1 Avg',
                    u' - Delta F1 Min',
                    u' - Delta F1Max',
                    u' - Fre Accuracy',
                    u'- Delta F2 Avg',
                    u'- Delta F2 Min',
                    u'- Delta F2Max',
                    u' - Delta F2 Max 99%',
                    u'- Delta F2 Avg/Delta F1 Avg',
                    u'- Frequency Accuracy',
                    u' - Frequency Offset',
                    u'- Frequency Drift',
                    u'- Max Drift Rate - 50us',
                    u'- Intial Frequency Drift',
                    ]
            df = pd.DataFrame(item)
            df.to_csv(fw2.filename, index=False)
        res_list=[]
        self.config_per_le_connect_usb(rate='LE1M', packet_len=37)
        self.csp.config_rx_level(rxpwr=-67)
        self.csp.tx_measure_para(tout=5, repetition='SINGleshot', count=20, MOEXception='OFF')
        for chan in chan_list:
            self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan)
            for rate in rate_list:
                if rate == 'LE500k':
                    _rate = 'LELR'
                    coding = 'S2'
                    mode = 'LRANge'
                    self.csp.sig_le_lr_coding(coding=coding)
                elif rate == 'LE125K':
                    _rate = 'LELR'
                    coding = 'S8'
                    mode = 'LRANge'
                    self.csp.sig_le_lr_coding(coding=coding)
                else:
                    mode = rate
                    _rate = rate

                self.csp.mode_set(mode=mode)
                self.csp.sig_le_phy(phy=_rate)
                # self.config_per_le_connect_usb(rate=_rate, packet_len=packet_len)
                # self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan)
                # self.csp.config_rx_level(rxpwr=-67)
                time.sleep(1)
                delta_f1_99, delta_f2_99, delta_f1_avg, delta_f1_min, delta_f1_max,delta_f2_avg, delta_f2_min, delta_f2_max, mod_ratio = -999, -999, -999, -999, -999, -999, -999, -999, -999
                self.csp.config_connect_le_packet_pattern(rate=_rate,pattern='P44')
                res = self.csp.get_modulation_measure_res(rate=mode)
                logdebug('{}'.format(res))

                if mode == 'LRANge':
                    freq_accuracy = eval(res[3]) / 1000.00
                    freq_drift = eval(res[4]) / 1000.00
                    drift_rate = eval(res[5]) / 1000.00
                else:
                    delta_f1_99 = eval(res[2]) / 1000.00
                    freq_accuracy = eval(res[3]) / 1000.00
                    freq_drift = eval(res[4]) / 1000.00
                    drift_rate = eval(res[5]) / 1000.00
                    delta_f1_avg = eval(res[6]) / 1000.00
                    delta_f1_min = eval(res[7]) / 1000.00
                    delta_f1_max = eval(res[8]) / 1000.00

                time.sleep(0.5)

                self.csp.config_connect_le_packet_pattern(rate=_rate,pattern='P11')
                res = self.csp.get_modulation_measure_res(rate=mode)
                logdebug('{}'.format(res))

                if mode == 'LRANge':
                    freq_accuracy = eval(res[3]) / 1000.00
                    freq_drift = eval(res[4]) / 1000.00
                    drift_rate = eval(res[5]) / 1000.00
                else:
                    delta_f2_99 = eval(res[2]) / 1000.00
                    freq_accuracy = eval(res[3]) / 1000.00
                    freq_drift = eval(res[4]) / 1000.00
                    drift_rate = eval(res[5]) / 1000.00
                    delta_f2_avg = eval(res[9]) / 1000.00
                    delta_f2_min = eval(res[10]) / 1000.00
                    delta_f2_max = eval(res[11]) / 1000.00
                    mod_ratio = delta_f2_avg / delta_f1_avg
                    freq_offset = eval(res[14])/1000.00
                    init_drift = eval(res[15]) / 1000.00
                time.sleep(0.5)

                self.csp.config_connect_le_packet_pattern(rate=_rate,pattern='PRBS9')
                res1 = self.csp.get_power_measure_res(rate=mode)
                res2 = self.csp.get_acp_res()
                logdebug('{}'.format(res1))
                logdebug('{}'.format(res2))

                nominal_pow = eval(res1[2])
                peak_pow = eval(res1[3])
                leakage_pow = eval(res1[4])
                peak_avg_pwr = eval(res1[5])
                acp_list = [eval(i) for i in res2[1:]]
                acp_center = len(acp_list)/2
                acp_max_pwr = acp_list[acp_center]
                acp_r10 = acp_list[acp_center + 10]
                acp_r9 = acp_list[acp_center + 9]
                acp_r8 = acp_list[acp_center + 8]
                acp_r7 = acp_list[acp_center + 7]
                acp_r6 = acp_list[acp_center + 6]
                acp_r5 = acp_list[acp_center + 5]
                acp_r4 = acp_list[acp_center + 4]
                acp_r3 = acp_list[acp_center + 3]
                acp_r2 = acp_list[acp_center + 2]
                acp_r1 = acp_list[acp_center + 1]
                acp_l1 = acp_list[acp_center - 1]
                acp_l2 = acp_list[acp_center - 2]
                acp_l3 = acp_list[acp_center - 3]
                acp_l4 = acp_list[acp_center - 4]
                acp_l5 = acp_list[acp_center - 5]
                acp_l6 = acp_list[acp_center - 6]
                acp_l7 = acp_list[acp_center - 7]
                acp_l8 = acp_list[acp_center - 8]
                acp_l9 = acp_list[acp_center - 9]
                acp_l10 = acp_list[acp_center - 10]
                loginfo("channel: {}    packet type: {} ".format(chan, rate))
                loginfo("nominal power: {} dBm      peak power: {} dBm      leakage power: {} dBm ".format(nominal_pow, peak_pow, leakage_pow))
                loginfo("freq accuracy: {} Hz     freq drift: {} Hz     drift rate: {} Hz".format(freq_accuracy, freq_drift, drift_rate))
                loginfo("delta f1 avg: {}  Hz     delta f1 min: {} Hz      delta f1 max: {} Hz".format(delta_f1_avg, delta_f1_min, delta_f1_max))

                if fig_en != 0:
                    plt.ion()
                    x = [2402+chan*2+xi for xi in range(-10, 11, 1)]
                    fig = plt.figure('le_acp_channel{}_{}'.format(chan,rate))
                    fig.set_size_inches(13, 7)
                    ax = fig.add_subplot(111)
                    ax.bar(x, 80+np.array(acp_list), width=0.5, bottom=-80)
                    ax.set_xticks(np.array(x))
                    ax.set_xlabel('freq(MHz)')
                    ax.set_ylabel('txp(dBm)')
                    ax.set_title('%s_ACP'%rate)
                    for xi,yi in zip(x,acp_list):
                        ax.text(xi-0.25, yi+0.25, '%.2f'%yi, fontsize=8)
                    plt.savefig(fname + 'le_acp_channel{}_{}_{}'.format(chan, rate, time.strftime('%Y%m%d_%H%M%S',time.localtime(time.time()))))
                    # plt.show()
                if report_save:
                    df = pd.DataFrame(pd.read_csv(fw2.filename))
                    df_value = [chan,rate,nominal_pow,peak_pow,peak_avg_pwr,leakage_pow,acp_l10,acp_l9,acp_l8,acp_l7,acp_l6,acp_l5,acp_l4,acp_l3,acp_l2,acp_l1,acp_max_pwr,
                                acp_r1,acp_r2,acp_r3,acp_r4,acp_r5,acp_r6,acp_r7,acp_r8,acp_r9,acp_r10,freq_accuracy,delta_f1_avg,delta_f1_min,delta_f1_max,freq_accuracy,
                                delta_f2_avg,delta_f2_min,delta_f2_max,delta_f2_99,mod_ratio,freq_accuracy,freq_offset,freq_drift,drift_rate,init_drift]

                    df['channel_{}_{}'.format(chan,rate)] = df_value
                    df.to_csv(fw2.filename,index=False)
                if csv_save:
                    fw1.write_data([chan,rate,nominal_pow,peak_pow,leakage_pow,freq_accuracy,freq_drift,drift_rate,delta_f1_avg,delta_f1_min,delta_f1_max,delta_f2_avg,delta_f2_min,delta_f2_max,mod_ratio,delta_f1_99,delta_f2_99,acp_list])
                else:
                    res_list.append([chan,rate,nominal_pow,peak_pow,leakage_pow,freq_accuracy,freq_drift,drift_rate,delta_f1_avg,delta_f1_min,delta_f1_max,delta_f2_avg,delta_f2_min,delta_f2_max,mod_ratio,delta_f1_99,delta_f2_99,acp_list])

        return res_list

    def br_tx_meas(self, chan_list=[0], rate_list=['DH1'], rx_level=-50, tx_level=10, csv_save=True, report_save=True, fig_en=0):
        '''
        rate_list:  DH1\DH3\DH5
        '''
        if csv_save :
            title = 'channel,type,nominal_pow(dBm),peak_pow(dBm),leakage_pow(dBm),freq_accuracy(kHz),freq_drift(kHz),drift_rate(Hz/50 μs),'
            title = title + 'delta_f1_avg(kHz),delta_f1_min(kHz),delta_f1_max(kHz),delta_f2_avg(kHz),delta_f2_min(kHz),delta_f2_max(kHz),mod_ratio,delta_f2_99(kHz),'
            title = title + 'obw(kHz),frange_l(kHz),frange_h(kHz),acp_list_21ch\n'
            fname = self.get_filename('ts_bt_test/','test_br_tx_{}'.format(self.board_name))
            fw1=csvreport(fname,title)
        if report_save:
            title2 = 'BR_tx_performance'
            fname2 = self.get_filename('ts_bt_test/','test_br_tx_report_{}'.format(self.board_name))
            fw2=csvreport(fname2,title2)
            item = ['channel', 'DH', u'Tx/01 - Output Power  - Average', u'- Peak',
                    u'@   - Leakage',
                    u' @   - Packet Timing',
                    u'Tx/04 - Spectrum - Frequency Range - Left',
                    u'- Right',
                    u'Tx/05 - Spectrum - 20 dB Bandwidth - f(L)',
                    u'- f(H)',
                    u'- f(H) - f(L)',
                    u'Tx/06 - Spectrum - Adjacent channel power - Max Power',
                    u'- ACP - >3 MHz',
                    u'- ACP -3 MHz',
                    u'- ACP -2 MHz',
                    u'- ACP -1 MHz',
                    u' - ACP +1 MHz',
                    u'- ACP +2 MHz',
                    u'- ACP +3 MHz',
                    u'- ACP + >3 MHz',
                    u'Tx/07 - Modulation Characteristics - Delta F1 Avg',
                    u'Tx/07 - Modulation Characteristics - Delta F2 Avg',
                    u'- Delta F2 Max 99%',
                    u'- Delta F2 Avg/Delta F1 Avg',
                    u'Tx/08 - Initial Carrier Frequency Tolerance - Avg',
                    u'- Max',
                    u'Tx/09 - Carrier Frequency Drift - Max.Drift DH1',
                    u' - Max.Drift DH5',
                    u' - Max.Drift',
                    u' - Max.Drift Rate/50us DH1',
                    u' - Max.Drift Rate/50us DH5',
                    u'- Max.Drift Rate/50us',
                    u'Harmonic Spurs']

            df = pd.DataFrame(item)
            df.to_csv(fw2.filename, index=False)
        res_list=[]
        self.config_classic_ber(rate='BR', chan=1, rx_level=rx_level, tx_power=tx_level)
        self.csp.tx_measure_para(tout=10, repetition='SINGleshot', count=20, MOEXception='OFF')
        for chan in chan_list:
            # self.config_classic_ber(rate='BR', chan=1, rx_level=rx_level, tx_power=tx_level)
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan)
            self.csp.config_rxq_br_tout(tout=10)
            for rate in rate_list:
                self.csp.config_connect_br_packet_ptype(ptype=rate)  # 设置BR包类型

                time.sleep(1)
                delta_f1_99, delta_f2_99, delta_f2_avg, delta_f2_min, delta_f2_max, mod_ratio = -999, -999, -999, -999, -999, -999
                self.csp.config_connect_br_packet_pattern(pattern='P44')  # 设置BR包 payload 数据类型
                res = self.csp.get_modulation_measure_res()
                logdebug('{}'.format(res))

                delta_f1_99 = eval(res[2]) / 1000.00
                freq_accuracy = eval(res[3]) / 1000.00
                freq_drift = eval(res[4]) / 1000.00
                drift_rate = eval(res[5])
                delta_f1_avg = eval(res[6]) / 1000.00
                delta_f1_min = eval(res[7]) / 1000.00
                delta_f1_max = eval(res[8]) / 1000.00
                time.sleep(0.5)

                self.csp.config_connect_br_packet_pattern(pattern='P11')  # 设置BR包 payload 数据类型
                res = self.csp.get_modulation_measure_res()
                logdebug('{}'.format(res))

                delta_f2_99 = eval(res[2]) / 1000.00
                freq_accuracy = eval(res[3]) / 1000.00
                freq_drift = eval(res[4]) / 1000.00
                drift_rate = eval(res[5]) / 1000.00
                delta_f2_avg = eval(res[9]) / 1000.00
                delta_f2_min = eval(res[10]) / 1000.00
                delta_f2_max = eval(res[11]) / 1000.00
                mod_ratio = delta_f2_avg / delta_f1_avg
                time.sleep(0.5)

                self.csp.config_connect_br_packet_pattern(pattern='PRBS9')  # 设置BR包 payload 数据类型
                res1 = self.csp.get_power_measure_res()
                time.sleep(1)
                res2 = self.csp.get_acp_res()
                # time.sleep(1)
                res3 = self.csp.get_obw_res()
                time.sleep(1)
                logdebug('{}'.format(res1))
                logdebug('{}'.format(res2))

                nominal_pow = eval(res1[2])
                peak_pow = eval(res1[3])
                leakage_pow = eval(res1[4])
                PacketTiming = eval(res1[5])
                acp_list = [eval(i) for i in res2[1:]]
                # obw = eval(res3[6]) / 1000.00
                # obw = -999
                acp_center = 10
                acp_max_pwr = acp_list[acp_center]
                acp_r4 = acp_list[acp_center + 4]
                acp_r3 = acp_list[acp_center + 3]
                acp_r2 = acp_list[acp_center + 2]
                acp_r1 = acp_list[acp_center + 1]
                acp_l1 = acp_list[acp_center - 1]
                acp_l2 = acp_list[acp_center - 2]
                acp_l3 = acp_list[acp_center - 3]
                acp_l4 = acp_list[acp_center - 4]
                obw = eval(res3[6])/1000.00
                obw_l = eval(res3[4])/1000.00
                obw_h = eval(res3[5])/1000.00

                self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=0)
                res4 = self.csp.get_frange_res()
                frange_l = eval(res4[3]) / 1000000.00
                self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=78)
                res4 = self.csp.get_frange_res()
                frange_h = eval(res4[4]) / 1000000.00
                self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan)
                loginfo("channel: {}    packet type: {}   ".format(chan, rate))
                loginfo("nominal power: {}      peak power: {}      leakage power: {}   unit: dBm".format(nominal_pow, peak_pow, leakage_pow))
                loginfo("freq accuracy: {}      freq drift: {}      drift rate: {}  unit: Hz".format(freq_accuracy, freq_drift, drift_rate))
                loginfo("delta f1 avg: {}       delta f1 min: {}       delta f1 max: {}  unit: Hz".format(delta_f1_avg, delta_f1_min, delta_f1_max))
                loginfo("delta f2 avg: {}       delta f2 min: {}       delta f2 max: {}  unit: Hz".format(delta_f2_avg, delta_f2_min, delta_f2_max))
                loginfo("mod ratio: {}".format(mod_ratio))
                loginfo("delta f2 99.9%: {}  unit: Hz".format(delta_f2_99))
                loginfo("spectrum 20 dB bandwidth: {}".format(obw))

                if fig_en != 0:
                    plt.ion()
                    x = [2402 + chan + xi for xi in range(-10, 11, 1)]
                    fig = plt.figure('br_acp_channel{}_{}'.format(chan, rate))
                    fig.set_size_inches(13, 7)
                    ax = fig.add_subplot(111)
                    ax.bar(x, 80 + np.array(acp_list), width=0.5, bottom=-80)
                    ax.set_xticks(np.array(x))
                    ax.set_xlabel('freq(MHz)')
                    ax.set_ylabel('txp(dBm)')
                    ax.set_title('BR_ACP')
                    for xi, yi in zip(x, acp_list):
                        ax.text(xi - 0.25, yi + 0.25, '%.2f' % yi, fontsize=8)
                    plt.savefig(fname + 'br_acp_channel{}_{}_{}'.format(chan, rate, time.strftime('%Y%m%d_%H%M%S', time.localtime(time.time()))))
                    # plt.show()
                if report_save:
                    df = pd.DataFrame(pd.read_csv(fw2.filename))
                    df_value = [chan,rate,nominal_pow,peak_pow,leakage_pow,PacketTiming,frange_l,frange_h,obw_l,obw_h,obw,acp_max_pwr,acp_l4,acp_l3,
                                                           acp_l2,acp_l1,acp_r1,acp_r2,acp_r3,acp_r4,delta_f1_avg,delta_f2_avg,delta_f2_99,mod_ratio,freq_accuracy,'',freq_drift,
                                                           '','',drift_rate,'','','']
                    print (len(df_value))
                    df['channel_{}_{}'.format(chan,rate)] = df_value
                    df.to_csv(fw2.filename,index=False)
                if csv_save:
                    fw1.write_data(
                        [chan, rate, nominal_pow, peak_pow, leakage_pow, freq_accuracy, freq_drift, drift_rate, delta_f1_avg, delta_f1_min, delta_f1_max, delta_f2_avg,
                         delta_f2_min,
                         delta_f2_max, mod_ratio, delta_f2_99, obw, acp_list])
                else:
                    res_list.append(
                        [chan, rate, nominal_pow, peak_pow, leakage_pow, freq_accuracy, freq_drift, drift_rate, delta_f1_avg, delta_f1_min, delta_f1_max, delta_f2_avg,
                         delta_f2_min,
                         delta_f2_max, mod_ratio, delta_f2_99, obw,  acp_list])
        return res_list

    def edr_tx_meas(self, chan_list=[0], rate_list=['2_DH1'], rx_level=-50, tx_level=10, csv_save=True, report_save=True, fig_en=0):
        '''
        rate_list:  2_DH1\2_DH3\2_DH5\3_DH1\3_DH3\3_DH5
        '''
        if csv_save:
            title = 'rate,channel,nominal_pwr(dBm),peak_pwr(dBm),gfsk_pwr(dBm),dpsk_pwr(dBm),dpsk_gfsk_diff_pwr,guard_period(us),'
            title = title + 'wi(KHz),w0_wi(KHz),w0_max(KHz),DEVM_RMS(%),DEVM_peak(%),DEVM_P99(%),bit_error_rate,packet0error,'
            title = title + 'PTxRef(dBm),N26ChN1Abs(dBm),N26ChP1Abs(dBm),N26ChN1Rel(dBm),N26ChP1Rel(dBm),acp_list\n'
            fname = self.get_filename('ts_bt_test/', 'test_edr_tx_{}'.format(self.board_name))
            fw1 = csvreport(fname, title)
        if report_save:
            title2 = 'BR_tx_performance'
            fname2 = self.get_filename('ts_bt_test/','test_edr_tx_report_{}'.format(self.board_name))
            fw2=csvreport(fname2,title2)
            item = ['channel', 'rate', u'Tx/10 - EDR Relative Transmit Power - PGFSK - Max Power', u' - PDPSK',
                    u'- PDPSK - PGFSK',
                    u'@   - Nomimal Power',
                    u'@   - Guard Period',
                    u'@   - Packet Timing',
                    u'Tx/11 - CFS and MA - omega i',
                    u' - omega i + omega o',
                    u'- omega o',
                    u'- DEVM RMS',
                    u'- DEVM Peak',
                    u'- DEVM 99% - Max 0.30',
                    u'Tx/12 - EDR Differential Phase Encoding',
                    u'Tx/13 - In Band Spurious Emissins - Nominal Power',
                    u' - ACPower:-3',
                    u'- ACPower:-2',
                    u'- ACPower:-1,Ptx-26dB',
                    u' - ACPower:Center Ptxref',
                    u'- ACPower:1,Ptx-26dB',
                    u'- ACPower:+2',
                    u'- ACPower:+3',
                    u'- ACPower:+- >3',]
            df = pd.DataFrame(item)
            df.to_csv(fw2.filename, index=False)
        res_list=[]
        self.config_classic_ber(rate='EDR', chan=1, rx_level=rx_level, tx_power=tx_level)
        self.csp.tx_measure_para(tout=5, repetition='SINGleshot', count=20, MOEXception='OFF')
        for chan in chan_list:
            # self.config_classic_ber(rate='EDR', chan=1, rx_level=rx_level, tx_power=tx_level)
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan)
            self.csp.config_connect_edr_packet_pattern(pattern='PRBS9')  # 设置EDR包 payload 数据类型
            for rate in rate_list:
                _rate = self.dic_edr_rate[rate]
                self.csp.config_connect_edr_packet_ptype(ptype=_rate)  # 设置EDR包类型

                res = self.csp.get_modulation_measure_res()
                logdebug('{}'.format(res))
                res = [eval(i) for i in res[0:]]
                wi = res[2]/1000
                w0_wi = res[3]/1000
                w0_max = res[4]/1000
                DEVM_RMS = res[5]
                DEVM_peak = res[6]
                DEVM_P99 = res[7]

                res1 = self.csp.get_power_measure_res()
                logdebug('{}'.format(res1))
                res1 = [eval(i) for i in res1[0:]]
                nominal_pwr = res1[2]
                gfsk_pwr = res1[3]
                dpsk_pwr = res1[4]
                dpsk_gfsk_diff_pwr = res1[5]
                guard_period = res1[6]
                packet_timing = res1[7]
                peak_pwr = res1[8]

                self.csp.RF_Frequency_Settings_tx(mode='TXT', ch_tx=chan)
                res2 = self.csp.get_diff_phase_encoding_res()
                logdebug('{}'.format(res2))
                res2 = [eval(i) for i in res2[0:]]
                bit_error_rate = res2[2]
                packet0error = res2[3]

                self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan)
                self.csp.tx_measure_para(tout=5, repetition='SINGleshot', count=100, MOEXception='OFF')
                self.csp.RF_Power_settings(rx_level=rx_level, tx_power=nominal_pwr+3, margin=8)
                [res3, res4] = self.csp.get_acp_res_edr()
                self.csp.tx_measure_para(tout=5, repetition='SINGleshot', count=20, MOEXception='OFF')
                logdebug('{}'.format(res3))
                logdebug('{}'.format(res4))
                res3 = [eval(i) for i in res3[0:]]
                res4 = [eval(i) for i in res4[0:]]
                acp_nominal_pwr = res3[2]
                PTxRef = res3[4]
                N26ChN1Abs = res3[5]
                N26ChP1Abs = res3[6]
                N26ChN1Rel = res3[7]
                N26ChP1Rel = res3[8]

                acp_list = res4[1:]
                acp_center = len(acp_list)/2
                acp_max_pwr = acp_list[acp_center]
                acp_r4 = acp_list[acp_center + 4]
                acp_r3 = acp_list[acp_center + 3]
                acp_r2 = acp_list[acp_center + 2]
                acp_r1 = acp_list[acp_center + 1]
                acp_l1 = acp_list[acp_center - 1]
                acp_l2 = acp_list[acp_center - 2]
                acp_l3 = acp_list[acp_center - 3]
                acp_l4 = acp_list[acp_center - 4]

                loginfo("channel: {}    packet type: {}     payload length: ".format(chan, rate))
                loginfo("nominal power: {}      peak power: {}      gfsk_pwr: {}    dpsk_pwr: {}".format(nominal_pwr, peak_pwr, gfsk_pwr,dpsk_pwr))
                loginfo("dpsk_gfsk_diff_pwr: {}      guard_period: {}      ".format(dpsk_gfsk_diff_pwr, guard_period))
                loginfo("wi: {}       w0_wi: {}       w0_max: {}".format(wi, w0_wi, w0_max))
                loginfo("DEVM_RMS: {}       DEVM_peak: {}       DEVM_P99: {}".format(DEVM_RMS, DEVM_peak, DEVM_P99))
                loginfo("bit_error_rate: {}     packet0error: {}".format(bit_error_rate, packet0error))
                loginfo("PTxRef: {} N26ChN1Abs: {}  N26ChP1Abs: {}  N26ChN1Rel: {}  N26ChP1Rel:{}".format(PTxRef, N26ChN1Abs, N26ChP1Abs, N26ChN1Rel, N26ChP1Rel))
                if report_save:
                    df = pd.DataFrame(pd.read_csv(fw2.filename))
                    df_value = [chan,rate,peak_pwr,dpsk_pwr,dpsk_gfsk_diff_pwr,nominal_pwr,guard_period,packet_timing,wi,w0_wi,w0_max,DEVM_RMS,DEVM_peak,DEVM_P99,packet0error,
                                acp_nominal_pwr,acp_l3,acp_l2,N26ChN1Rel,PTxRef,N26ChP1Rel,acp_r2,acp_r3,min([acp_l4,acp_r4])]

                    df['channel_{}_{}'.format(chan,rate)] = df_value
                    df.to_csv(fw2.filename,index=False)
                if csv_save:
                    if fig_en != 0:
                        plt.ion()
                        x = [2402 + chan + xi for xi in range(-10, 11, 1)]
                        fig = plt.figure('edr_acp_channel{}_{}'.format(chan, rate))
                        fig.set_size_inches(13, 7)
                        ax = fig.add_subplot(111)
                        ax.bar(x, 80 + np.array(acp_list), width=0.5, bottom=-80)
                        ax.set_xticks(np.array(x))
                        ax.set_xlabel('freq(MHz)')
                        ax.set_ylabel('txp(dBm)')
                        ax.set_title('%s_ACP' % rate)
                        for xi, yi in zip(x, acp_list):
                            ax.text(xi - 0.25, yi + 0.25, '%.2f' % yi, fontsize=8)
                        plt.savefig(fname + 'edr_acp_channel{}_{}_{}'.format(chan, rate, time.strftime('%Y%m%d_%H%M%S',
                                                                                                       time.localtime(
                                                                                                           time.time()))))
                        # plt.show()
                    fw1.write_data([rate,chan,nominal_pwr,peak_pwr,gfsk_pwr,dpsk_pwr,dpsk_gfsk_diff_pwr,guard_period,wi,w0_wi,w0_max,DEVM_RMS,DEVM_peak,DEVM_P99,bit_error_rate,packet0error,PTxRef,N26ChN1Abs,N26ChP1Abs,N26ChN1Rel,N26ChP1Rel,acp_list])
                else:
                    res_list.append([rate,chan,nominal_pwr,peak_pwr,gfsk_pwr,dpsk_pwr,dpsk_gfsk_diff_pwr,guard_period,wi,w0_wi,w0_max,DEVM_RMS,DEVM_peak,DEVM_P99,bit_error_rate,packet0error,PTxRef,N26ChN1Abs,N26ChP1Abs,N26ChN1Rel,N26ChP1Rel,acp_list])

        return res_list


    def bedr_rx_sens(self, chan_list=[0,38,78], rate_list=['DH1'], rx_level=-60, tx_power=10, packet_num=2000, step=0.5, csv_save=True):
        if csv_save:
            title = 'rate,channel,ber,per,NAK,hec_err,crc_err,packet_type_err,pay_err,sens\n'
            fname = self.get_filename('ts_bt_test/', 'br_rx_sens_{}'.format(self.board_name))
            fw1 = csvreport(fname, title)
        for chan in chan_list:
            self.config_classic_ber(rate='BR',chan=chan, rx_level=rx_level, tx_power=tx_power)
            self.csp.config_rx_level(rxpwr=rx_level)
            self.csp.config_rxq_br_search_packets(num=packet_num)
            self.csp.config_rxq_br_search_step(step=step)
            # self.csp.config_rxq_br_search_tout(tout=20)
            for rate in rate_list:
                if rate in self.br_rate_list:
                    btype = 'BR'
                    self.csp.config_connect_br_packet_ptype(ptype=rate)  # 设置BR包类型
                elif rate in self.edr_rate_list:
                    btype = 'EDR'
                    rate = self.dic_edr_rate[rate]
                    self.csp.config_connect_edr_packet_ptype(ptype=rate)  # 设置EDR的包类型，E21P：2-DH1，E31P：3-DH1
                self.csp.sig_btype(btype=btype)
                self.csp.ber_search_meas_state(state=0)
                while 1:
                    state_res = self.csp.get_ber_search_meas_state()    #ber search 完成则退出循环
                    # logdebug(state_res)
                    if state_res == 'RDY':
                        loginfo('search complete')
                        break
                    elif state_res == 'OFF':
                        self.cmd_br_connect()
                        self.csp.ber_search_meas_state(state=0)
                    else:
                        logdebug('searching')
                    time.sleep(5)

                res = self.csp.meas_bt_ber_search_res(cmd_type='FETCh')
                ber = eval(res[1])
                per = eval(res[2])
                NAK = eval(res[5])
                hec_err = eval(res[6])
                crc_err = eval(res[7])
                packet_type_err = eval(res[8])
                pay_err = eval(res[9])
                sens = eval(res[11])
                loginfo('{} {}  {}  {}  {}  {}  {}  {}'.format(ber,per,NAK,hec_err,crc_err,packet_type_err,pay_err,sens))
                if csv_save:
                    fw1.write_data([rate,chan,ber,per,NAK,hec_err,crc_err,packet_type_err,pay_err,sens])

    def bedr_rx_range(self, chan_list=[0,38,78], rate_list=['DH1'], rx_level=-60, tx_power=10, packet_num=2000, rxpwr_range=range(-98,0),  csv_save=True):
        if csv_save:
            title = 'rate,channel,ber,per,NAK,hec_err,crc_err,packet_type_err,pay_err,sens\n'
            fname = self.get_filename('ts_bt_test/', 'bedr_rx_range_{}'.format(self.board_name))
            fw1 = csvreport(fname, title)

        self.config_classic_ber(rate='BR', chan=1, rx_level=rx_level, tx_power=tx_power)
        self.csp.config_rx_level(rxpwr=rx_level)
        self.csp.config_rxq_br_packets(num=packet_num)
        self.csp.config_rxq_br_tout(tout=10)
        for chan in chan_list:
            self.csp.config_rx_level(rxpwr=rx_level)
            self.csp.RF_Frequency_Settings_rx(ch_tx=chan)
            for rate in rate_list:
                if rate in self.br_rate_list:
                    btype = 'BR'
                    self.csp.config_connect_br_packet_ptype(ptype=rate)  # 设置BR包类型
                elif rate in self.edr_rate_list:
                    btype = 'EDR'
                    rate = self.dic_edr_rate[rate]
                    self.csp.config_connect_edr_packet_ptype(ptype=rate)  # 设置EDR的包类型，E21P：2-DH1，E31P：3-DH1
                self.csp.sig_btype(btype=btype)
                self.csp.ber_meas_state(0)
                for rxpwr in rxpwr_range:
                    self.csp.config_rx_level(rxpwr=rxpwr)
                    # while 1:
                    #     state_res = self.csp.get_ber_search_meas_state()
                    #     if state_res == 'RDY':
                    #         logdebug(state_res)
                    #         break
                    #     time.sleep(3)
                    # delay = 0.001*packet_num*(rx_level+100)/step
                    # time.sleep(delay)
                    while 1:
                        res = self.csp.meas_bt_ber_res(cmd_type='READ')
                        if eval(res[0]) == 0:
                            break
                        else:
                            if self.csp.get_bt_connect_state() == 'SBY':
                                self.cmd_br_connect()

                    ber = eval(res[1])
                    per = eval(res[2])
                    NAK = eval(res[5])
                    hec_err = eval(res[6])
                    crc_err = eval(res[7])
                    packet_type_err = eval(res[8])
                    pay_err = eval(res[9])

                    time.sleep(0.5)
                    loginfo('{} {}  {}  {}  {}  {}  {}  {}'.format(ber,per,NAK,hec_err,crc_err,packet_type_err,pay_err,rxpwr))
                    if csv_save:
                        fw1.write_data([rate,chan,ber,per,NAK,hec_err,crc_err,packet_type_err,pay_err,rxpwr])

    def bedr_rx_ci(self, chan_list=[0,38,78], rate_list=['DH1'], cable_loss=4.6,rx_level=-60, tx_power=10, packet_num=2000,  csv_save=True):
        self.dict_pwr_inter_freqoffset = {
            0: range(-85, -40),
            1: range(-65, -10),
            2: range(-50, 0),
            3: range(-50, 0),
            4: range(-50, 0),
            -1: range(-65, -10),
            -2: range(-60, 0),
            -3: range(-50, 0),
            -4: range(-50, 0),
            -5: range(-50, 0),
            -6: range(-50, 0),
            -7: range(-50, 0),
            -8: range(-50, 0)
            }

        if csv_save:
            title = 'signal_channel,rate,freq_interference,rxpwr_interference,ber(%), per(%),agc_gain_index,agc_tgt,agc_sat\n'
            fname = self.get_filename('ts_bt_test/','bedr_rx_ci_data_{}'.format(self.board_name))
            fw1 = csvreport(fname,title)

            title1 = 'signal_channel,rate,freq_interference,rxpwr_interference(dbm),ci(db),agc_tgt,agc_sat\n'
            fname1 = self.get_filename('ts_bt_test/','bedr_rx_ci_report_{}'.format(self.board_name))
            fw2 = csvreport(fname1,title1)

        self.config_classic_ber(rate='BR', chan=1, rx_level=rx_level, tx_power=tx_power)
        self.csp.config_rx_level(rxpwr=rx_level)
        self.csp.config_rxq_br_packets(num=packet_num)
        self.csp.config_rxq_br_tout(tout=10)

        for chan in chan_list:
            self.csp.config_rx_level(rxpwr=-60)
            self.csp.RF_Frequency_Settings_rx(mode='LOOP', ch_tx=chan)
            for rate in rate_list:
                self.tester_inter = mxg.MXG(device_name='N5182B', num_of_machine=1)
                self.tester_inter.arb_waveform(rate=rate+'_PN15')
                self.tester_inter.trriger_para_set(type='CONT', count=2000)
                self.tester_inter.output_state(1, 1)
                self.tester_inter.arb_state(1)

                if rate in self.br_rate_list:
                    btype = 'BR'
                    self.csp.config_connect_br_packet_ptype(ptype=rate)  # 设置BR包类型
                elif rate in self.edr_rate_list:
                    btype = 'EDR'
                    _rate = self.dic_edr_rate[rate]
                    self.csp.config_connect_edr_packet_ptype(ptype=_rate)  # 设置EDR的包类型，E21P：2-DH1，E31P：3-DH1
                self.csp.sig_btype(btype=btype)
                for freq_inter_offset in[0,1,-1,2,-2,3,-3,4,-4]:
                    freq_inter = freq_inter_offset+2402+chan
                    if freq_inter_offset == 0 and btype == 'EDR':
                        self.tester_inter.arb_waveform(rate=rate + '_PN15')
                    else:
                        self.tester_inter.arb_waveform(rate='DH1_PN15')
                    if -3 < freq_inter_offset < 3:
                        self.csp.config_rx_level(rxpwr=-60)
                    else:
                        self.csp.config_rx_level(rxpwr=-67)
                    pwr_list = self.dict_pwr_inter_freqoffset[freq_inter_offset]
                    for power_inter in pwr_list:
                        self.tester_inter.para_set(freq=freq_inter, power=power_inter + cable_loss)
                        # time.sleep(1)
                        loginfo('freq_inter:    {}      pwr_inter:  {}'.format(freq_inter,power_inter))

                        while 1:
                            res = self.csp.meas_bt_ber_res(cmd_type='READ')

                            if eval(res[0]) == 0 or eval(res[0]) == 1:
                                break
                            else:
                                if self.csp.get_bt_connect_state() == 'SBY':
                                    self.tester_inter.para_set(freq=freq_inter, power=power_inter - 1 + cable_loss)
                                    self.cmd_br_connect()
                                    time.sleep(10)
                                    res = [-999,-999,-999,-999,-999,-999,-999,-999,-999,-999,-999,-999]
                                    break
                        ber = eval(res[1])
                        per = eval(res[2])
                        NAK = eval(res[5])
                        hec_err = eval(res[6])
                        crc_err = eval(res[7])
                        packet_type_err = eval(res[8])
                        pay_err = eval(res[9])

                        loginfo('{} {}  {}  {}  {}  {}  {} '.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err))

                        fw1.write_data([chan, rate, freq_inter, power_inter,  ber, per])
                        ber_min = 0.1
                        if btype == 'EDR':
                            ber_min = 0.01
                        if ber > ber_min or per > 90 or ber == -999:
                            if -3 < freq_inter_offset < 3:
                                ci = -60 - power_inter +1
                            else:
                                ci = -67 - power_inter + 1
                            fw2.write_data([chan, rate, freq_inter, power_inter-1, ci])
                            break
        self.tester_inter.output_state(0, 0)

    def bedr_rx_ci_debug(self, chan_list=[0,38,78], rate_list=['DH1'], cable_loss=4.6,rx_level=-60, tx_power=10, packet_num=2000,  csv_save=True, name_str='cbpf1',
                   cbpf2_en=0,agc_tgt_sat_list=[[300,400]], loopmax=1):
        self.dict_pwr_inter_freqoffset = {
            0: range(-85, -40),
            1: range(-65, -10),
            2: range(-50, 0),
            3: range(-50, 0),
            4: range(-50, 0),
            -1: range(-65, -10),
            -2: range(-60, 0),
            -3: range(-50, 0),
            -4: range(-50, 0),
            -5: range(-50, 0),
            -6: range(-50, 0),
            -7: range(-50, 0),
            -8: range(-50, 0)
            }

        if csv_save:
            title = 'signal_channel,rate,freq_interference,rxpwr_interference,ber(%), per(%),agc_gain_index,agc_tgt,agc_sat\n'
            fname = self.get_filename('ts_bt_test/','bedr_rx_ci_data_{}'.format(name_str))
            fw1 = csvreport(fname,title)

            title1 = 'signal_channel,rate,freq_interference,rxpwr_interference(dbm),ci(db),agc_tgt,agc_sat\n'
            fname1 = self.get_filename('ts_bt_test/','bedr_rx_ci_report_{}'.format(name_str))
            fw2 = csvreport(fname1,title1)
        if self.chipv.find('TX232') != -1:
            if cbpf2_en !=0:
                self.jlink.wrm(0xa0421020,20,20,1)
            else:
                self.jlink.wrm(0xa0421020, 20, 20, 0)

        self.config_classic_ber(rate='BR', chan=1, rx_level=rx_level, tx_power=tx_power)
        self.csp.config_rx_level(rxpwr=rx_level)
        self.csp.config_rxq_br_packets(num=packet_num)
        self.csp.config_rxq_br_tout(tout=10)
        for loop in range(0,loopmax):
            loop = loop+1
            for cbpf_bw in range(3,-1,-1):
                self.jlink.wrm(0xa0421034, 13, 12, cbpf_bw)
                for cbpf_bias_trim in range(0,8):
                    self.jlink.wrm(0xa0421034, 18, 16, cbpf_bias_trim)
                    for chan in chan_list:
                        self.csp.config_rx_level(rxpwr=-60)
                        self.csp.RF_Frequency_Settings_rx(mode='LOOP', ch_tx=chan)
                        for rate in rate_list:
                            self.tester_inter = mxg.MXG(device_name='N5182B', num_of_machine=1)
                            self.tester_inter.arb_waveform(rate=rate+'_PN15')
                            self.tester_inter.trriger_para_set(type='CONT', count=2000)
                            self.tester_inter.output_state(1, 1)
                            self.tester_inter.arb_state(1)

                            if rate in self.br_rate_list:
                                btype = 'BR'
                                self.csp.config_connect_br_packet_ptype(ptype=rate)  # 设置BR包类型
                            elif rate in self.edr_rate_list:
                                btype = 'EDR'
                                _rate = self.dic_edr_rate[rate]
                                self.csp.config_connect_edr_packet_ptype(ptype=_rate)  # 设置EDR的包类型，E21P：2-DH1，E31P：3-DH1
                            self.csp.sig_btype(btype=btype)
                            for agc_tgt_sat in agc_tgt_sat_list:
                                agc_tgt = agc_tgt_sat[0]
                                agc_sat = agc_tgt_sat[1]
                                if self.chipv == 'TX232':
                                    self.jlink.wrm(0xa04210ac,9,0,agc_tgt)
                                    self.jlink.wrm(0xa04210a8,11,2,agc_sat)

                                for freq_inter_offset in[0,1,-1,2,-2,3,-3,4,-4]:
                                    freq_inter = freq_inter_offset+2402+chan
                                    if freq_inter_offset == 0 and btype == 'EDR':
                                        self.tester_inter.arb_waveform(rate=rate + '_PN15')
                                    else:
                                        self.tester_inter.arb_waveform(rate='DH1_PN15')
                                    if -3 < freq_inter_offset < 3:
                                        self.csp.config_rx_level(rxpwr=-60)
                                    else:
                                        self.csp.config_rx_level(rxpwr=-67)
                                    pwr_list = self.dict_pwr_inter_freqoffset[freq_inter_offset]
                                    for power_inter in pwr_list:
                                        self.tester_inter.para_set(freq=freq_inter, power=power_inter + cable_loss)
                                        # time.sleep(1)
                                        loginfo('freq_inter:    {}      pwr_inter:  {}'.format(freq_inter,power_inter))
                                        agc_gain_index = self.jlink.rdm(0xa04210dc, 11, 8)
                                        while 1:
                                            res = self.csp.meas_bt_ber_res(cmd_type='READ')

                                            if eval(res[0]) == 0 or eval(res[0]) == 1:
                                                break
                                            else:
                                                if self.csp.get_bt_connect_state() == 'SBY':
                                                    self.tester_inter.para_set(freq=freq_inter, power=power_inter - 1 + cable_loss)
                                                    self.cmd_br_connect()
                                                    time.sleep(10)
                                                    res = [-999,-999,-999,-999,-999,-999,-999,-999,-999,-999,-999,-999]
                                                    break
                                        ber = eval(res[1])
                                        per = eval(res[2])
                                        NAK = eval(res[5])
                                        hec_err = eval(res[6])
                                        crc_err = eval(res[7])
                                        packet_type_err = eval(res[8])
                                        pay_err = eval(res[9])

                                        loginfo('{} {}  {}  {}  {}  {}  {} '.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err))

                                        fw1.write_data([chan, rate, freq_inter, power_inter,  ber, per, agc_gain_index,agc_tgt,agc_sat])
                                        ber_min = 0.1
                                        if btype == 'EDR':
                                            ber_min = 0.01
                                        if ber > ber_min or per > 90 or ber == -999:
                                            if -3 < freq_inter_offset < 3:
                                                ci = -60 - power_inter +1
                                            else:
                                                ci = -67 - power_inter + 1
                                            fw2.write_data([chan, rate, freq_inter, power_inter-1, ci,agc_tgt,agc_sat])
                                            break
        self.tester_inter.output_state(0, 0)

    def bedr_rx_blocking(self, rate_list=['DH1'],  packet_num=1500, inter_loss=11, inter_rxpwr=[-30], freq_range=[]):
        title = 'rate,blocking_level,freq_blocking_1,freq_blocking_2,ber,per(%)\n'
        fname = self.get_filename('ts_bt_test/', 'bedr_rx_blocking_{}'.format(self.board_name))
        fw1 = csvreport(fname, title)
        self.tester_inter = mxg.MXG(device_name='N5182B',num_of_machine=1)
        self.tester_inter.para_set(30,-30+inter_loss)
        self.tester_inter.output_state(1, 0)

        self.config_classic_ber(rate='BR', chan=58, rx_level=-60, tx_power=10)
        self.csp.config_rxq_br_packets(num=packet_num)
        self.csp.config_rxq_br_tout(tout=10)
        self.csp.config_rx_level(rxpwr=-67)
        self.csp.RF_Frequency_Settings_rx(mode='LOOP', ch_tx=58)

        for rate in rate_list:
            self.csp.config_rx_level(rxpwr=-67)
            if rate in self.br_rate_list:
                btype = 'BR'
                self.csp.config_connect_br_packet_ptype(ptype=rate)  # 设置BR包类型
            elif rate in self.edr_rate_list:
                btype = 'EDR'
                _rate = self.dic_edr_rate[rate]
                self.csp.config_connect_edr_packet_ptype(ptype=_rate)  # 设置EDR的包类型，E21P：2-DH1，E31P：3-DH1
            self.csp.sig_btype(btype=btype)

            # freq_range= range(30,2401,1)+range(2500,6000,1)

            # for rxpwr in inter_rxpwr:
            rxpwr = inter_rxpwr
            freq_inter_list = []
            freq_inter_list2 = []
            for freq_inter in freq_range:
                freq_blocking_2 = 'NA'
                freq_blocking_1 = 'NA'

                if freq_inter < 2001 or freq_inter > 2999:
                    pwr = rxpwr + 17 + inter_loss
                else:
                    pwr = rxpwr + inter_loss
                self.csp.config_rx_level(rxpwr=-67)
                connect_num = 0
                while 1:
                    self.tester_inter.para_set(freq=freq_inter, power=pwr)

                    loginfo('freq_inter:    {}'.format(freq_inter))

                    res = self.csp.meas_bt_ber_res(cmd_type='READ')
                    if eval(res[0]) == 0 or eval(res[0]) == 1:
                        break
                    else:
                        connect_num = connect_num+1
                        loginfo('reconnect num {}'.format(connect_num))
                        if self.csp.get_bt_connect_state() == 'SBY':
                            self.cmd_br_connect()
                            time.sleep(10)
                            if connect_num > 3:
                                res = ['0','100','100','100','100','100','100','100','100','100','100']
                                break
                ber = eval(res[1])
                per = eval(res[2])
                NAK = eval(res[5])
                hec_err = eval(res[6])
                crc_err = eval(res[7])
                packet_type_err = eval(res[8])
                pay_err = eval(res[9])
                loginfo('freq_blocking_1:   {}      level:  {}      ber:    {}      per:    {}'.format(freq_inter, inter_rxpwr, ber, per))
                if ber > 0.1 or per > 90:
                    loginfo('freq_blocking_1:   {}      level:  {}      ber:    {}      per:    {}'.format(freq_inter,inter_rxpwr,ber,per))
                    freq_blocking_1 = freq_inter
                    freq_inter_list.append(freq_inter)
                    fw1.write_data([rate,pwr-inter_loss,freq_blocking_1,freq_blocking_2,ber,per])


            for freq_inter2 in freq_inter_list:
                freq_blocking_2 = 'NA'
                freq_blocking_1 = 'NA'

                if freq_inter2 < 2001 or freq_inter2 > 2999:
                    pwr = rxpwr + 17 + inter_loss
                else:
                    pwr = rxpwr + inter_loss
                self.csp.config_rx_level(rxpwr=-50)
                connect_num = 0
                while 1:
                    self.tester_inter.para_set(freq=freq_inter2, power=pwr)

                    loginfo('freq_inter:    {}'.format(freq_inter2))

                    res = self.csp.meas_bt_ber_res(cmd_type='READ')
                    if eval(res[0]) == 0 or eval(res[0]) == 1:
                        break
                    else:
                        connect_num = connect_num+1
                        loginfo('reconnect num {}'.format(connect_num))
                        if self.csp.get_bt_connect_state() == 'SBY':
                            self.cmd_br_connect()
                            time.sleep(10)
                            if connect_num > 3:
                                res = ['0','100','100','100','100','100','100','100','100','100','100']
                                break
                ber = eval(res[1])
                per = eval(res[2])
                NAK = eval(res[5])
                hec_err = eval(res[6])
                crc_err = eval(res[7])
                packet_type_err = eval(res[8])
                pay_err = eval(res[9])
                loginfo('freq_blocking_1:   {}      level:  {}      ber:    {}      per:    {}'.format(freq_inter2, inter_rxpwr, ber, per))
                if ber > 0.1 or per > 90:
                    loginfo('freq_blocking_1:   {}      level:  {}      ber:    {}      per:    {}'.format(freq_inter2,inter_rxpwr,ber,per))
                    freq_blocking_2 = freq_inter2
                    freq_inter_list2.append(freq_inter2)

                    fw1.write_data([rate,pwr-inter_loss,freq_blocking_1,freq_blocking_2,ber,per])
            #     self.csp.config_rx_level(rxpwr=-50)
            #     connect_num = 0
            #     while 1:
            #         self.tester_inter.para_set(freq=freq_inter, power=pwr)
            #
            #         loginfo('freq_inter:    {}'.format(freq_inter))
            #
            #         res = self.csp.meas_bt_ber_res(cmd_type='READ')
            #         if eval(res[0]) == 0 or eval(res[0]) == 1:
            #             break
            #         else:
            #             connect_num = connect_num+1
            #             loginfo('reconnect num {}'.format(connect_num))
            #             if self.csp.get_bt_connect_state() == 'SBY':
            #                 self.cmd_br_connect()
            #                 if connect_num > 3:
            #                     res = [0, 100, 100, 100, 100, 100, 100, 100, 100, 100, 100]
            #                     break
            #     ber = eval(res[1])
            #     per = eval(res[2])
            #     NAK = eval(res[5])
            #     hec_err = eval(res[6])
            #     crc_err = eval(res[7])
            #     packet_type_err = eval(res[8])
            #     pay_err = eval(res[9])
            #     loginfo('freq_blocking_2:   {}      level:  {}      ber:    {}      per:    {}'.format(freq_inter, pwr-inter_loss, ber, per))
            #     if ber > 0.1 or per > 90:
            #         loginfo('freq_blocking_2:   {}      level:  {}      ber:    {}      per:    {}'.format(freq_inter,pwr-inter_loss,ber,per))
            #         freq_blocking_2 = freq_inter
            #         # if freq_inter < 2001 or freq_inter > 2999:
            #         #     pwr = rxpwr + 17 + inter_loss
            #         fw1.write_data([rate,pwr,freq_blocking_1,freq_blocking_2,ber,per])
            #         count1 = count1 + 1
            #     # fw1.write_data([rate, pwr - inter_loss, freq_blocking_1, freq_blocking_2, ber, per])
            # if count > 24 or count1 > 5:
            #     self.tester_inter.output_state(0, 0)
            #     break
            # if freq_blocking_1 == 'NA':
            #     fw1.write_data([rate, pwr-inter_loss, freq_blocking_1, freq_blocking_2, ber,per])

        self.tester_inter.output_state(0, 0)
        return freq_inter_list+freq_inter_list2

    def le_rx_blocking(self, rate_list=['LE1M'], packet_len=37, packet_num=1500, inter_loss=11, inter_rxpwr=-10, freq_range_le=[]):
        title = 'rate,blocking_level,freq_blocking_1,freq_blocking_2,per(%)\n'
        fname = self.get_filename('ts_bt_test/', 'le_rx_blocking_{}'.format(self.board_name))
        fw1 = csvreport(fname, title)
        self.tester_inter = mxg.MXG(device_name='N5182B',num_of_machine=1)
        self.tester_inter.para_set(30,-30+inter_loss)
        self.tester_inter.output_state(1, 0)
        for rate in rate_list:
            if rate in ['LE_1M','LE1M']:
                _rate = 'LE1M'
            else:
                _rate = 'LE2M'
            self.config_per_le_connect_usb(rate=_rate,packet_len=packet_len)
            self.csp.RF_Frequency_Settings_rx(mode='DTM', ch_tx=19)
            self.csp.config_rxq_le_packets(rate=_rate, num=packet_num)
            self.csp.config_rx_level(rxpwr=-67)
            # freq_range= range(30,2001,10)+range(2003,2400,1)+range(2484,2998,1)+range(3000,6000,25)
            freq_range = freq_range_le
            freq_blocking_2 = 'NA'
            freq_blocking_1 = 'NA'
            freq_inter_list = []
            # for rxpwr in inter_rxpwr:
            count = 0
            count1 = 0
            for freq_inter in freq_range:
                freq_blocking_2 = 'NA'
                freq_blocking_1 = 'NA'
                # if freq_inter < 2001 or freq_inter > 2999:
                #     pwr =  + 5  + inter_loss
                #     # self.tester_inter.para_set(freq=freq_inter, power=rxpwr + inter_loss)
                # else:
                #     pwr = rxpwr + inter_loss
                #     # self.tester_inter.para_set(freq=freq_inter, power=rxpwr-5 + inter_loss)
                pwr = inter_rxpwr + inter_loss
                self.tester_inter.para_set(freq=freq_inter, power=pwr)
                loginfo('freq_inter:    {}'.format(freq_inter))
                self.csp.config_rx_level(rxpwr=-67)
                self.csp.per_meas_state(0)
                time.sleep(1)
                res = self.csp.meas_le_per_res(rate=_rate)
                per = eval(res[1])
                if per > 30.8:
                    loginfo('freq_blocking_1:   {}      level:  {}      per:    {}'.format(freq_inter,inter_rxpwr,per))
                    # freq_blocking_1 = freq_inter
                    # if freq_inter > 1999 or freq_inter < 3001:
                    #     # pwr = rxpwr - 5 + inter_loss
                    #     rxpwr = rxpwr - 5
                    freq_inter_list.append(freq_inter)
                    # fw1.write_data([rate,rxpwr,freq_blocking_1,freq_blocking_2,per])
                    count = count + 1
            for freq_inter_check in freq_inter_list:
                pwr = inter_rxpwr + inter_loss
                self.tester_inter.para_set(freq=freq_inter_check, power=pwr)
                loginfo('freq_inter:    {}'.format(freq_inter_check))
                self.csp.per_meas_state(0)
                time.sleep(1)
                res = self.csp.meas_le_per_res(rate=_rate)
                per = eval(res[1])
                if per > 30.8:
                    loginfo('freq_blocking_1:   {}      level:  {}      per:    {}'.format(freq_inter_check,inter_rxpwr,per))
                    freq_blocking_1 = freq_inter_check
                    fw1.write_data([rate, inter_rxpwr, freq_blocking_1, freq_blocking_2, per])
                    # self.csp.config_rx_level(rxpwr=-50)
                    # self.csp.per_meas_state(0)
                    # time.sleep(1)
                    # res = self.csp.meas_le_per_res(rate=_rate)
                    # per = eval(res[1])
                    # if per > 30.8:
                    #     loginfo('freq_blocking_2:   {}      level:  {}      per:    {}'.format(freq_inter_check,rxpwr,per))
                    #     freq_blocking_2 = freq_inter
                    #     # pwr = rxpwr
                    #     # if freq_inter > 1999 or freq_inter < 3001:
                    #     #     rxpwr = rxpwr - 5
                    #     #     pwr = rxpwr - 5 + inter_loss
                    #     fw1.write_data([rate,rxpwr,freq_blocking_1,freq_blocking_2,per])
                    #     count1 = count1 + 1

                # if count > 10 or count1 > 3:
                #     self.tester_inter.output_state(0, 0)
                #     break
            # if freq_blocking_1 == 'NA':
            #     fw1.write_data([rate, pwr-inter_loss, freq_blocking_1, freq_blocking_2, per])


    def le_rx_ci(self,  cable_loss=5, chan_list=[0], rate_list=['LE_1M'], packet_len=37, pkt_num=2000,  csv_save=True, cbpf2_en=0,agc_tgt_sat_list=[[300,400]]):
        self.dict_pwr_inter_freqoffset = {
            0: range(-80, -50),
            1: range(-75, -20),
            2: range(-52, -10),
            3: range(-41, 0),
            -1: range(-75, -20),
            -2: range(-52, -10),
            -3: range(-41, 0)
            }

        if csv_save:
            title = 'signal_channel,rate,freq_interference,rxpwr_interference,per(%),agc_gain_index,agc_tgt,agc_sat\n'
            fname = self.get_filename('ts_bt_test/','le_rx_ci_{}'.format(self.board_name))
            fw1 = csvreport(fname,title)

            title1 = 'signal_channel,rate,freq_interference,rxpwr_interference(dbm),ci(db),agc_tgt,agc_sat\n'
            fname1 = self.get_filename('ts_bt_test/','le_rx_ci_report_{}'.format(self.board_name))
            fw2 = csvreport(fname1,title1)
        if self.chipv == 'TX232':
            if cbpf2_en !=0:
                self.jlink.wrm(0xa0421020,20,20,1)
            else:
                self.jlink.wrm(0xa0421020, 20, 20, 0)
        for chan in chan_list:
            for rate in rate_list:
                self.tester_inter = mxg.MXG(device_name='N5182B', num_of_machine=1)
                self.tester_inter.arb_waveform(rate=rate+'_PN15')
                self.tester_inter.trriger_para_set(type='CONT', count=pkt_num)
                self.tester_inter.output_state(1, 1)
                self.tester_inter.arb_state(1)

                if rate == 'LE_1M':
                    _rate = 'LE1M'
                else:
                    _rate = 'LE2M'
                self.config_per_le_connect_usb(rate=_rate, packet_len=packet_len)
                self.csp.RF_Frequency_Settings_rx(mode='DTM', ch_tx=chan)
                self.csp.config_rxq_le_packets(rate=_rate, num=pkt_num)
                self.csp.config_rx_level(rxpwr=-67)
                time.sleep(1)
                for agc_tgt_sat in agc_tgt_sat_list:
                    agc_tgt = agc_tgt_sat[0]
                    agc_sat = agc_tgt_sat[1]
                    if self.chipv == 'TX232':
                        self.jlink.wrm(0xa04210ac,9,0,agc_tgt)
                        self.jlink.wrm(0xa04210a8,11,2,agc_sat)
                    for freq_inter_offset in[0,1,-1,2,-2,3,-3]:
                        if rate == 'LE_2M':
                            freq_inter = freq_inter_offset*2 + 2402 + chan * 2
                        else:
                            freq_inter = freq_inter_offset+2402+chan*2
                        pwr_list = self.dict_pwr_inter_freqoffset[freq_inter_offset]
                        for power_inter in pwr_list:
                            self.tester_inter.para_set(freq=freq_inter, power=power_inter + cable_loss)
                            time.sleep(0.2)
                            loginfo('freq_inter:    {}      pwr_inter:  {}'.format(freq_inter,power_inter))
                            agc_gain_index = self.jlink.rdm(0xa04210dc,11,8)
                            per_list=[]
                            for loop in range(10):
                                res = self.csp.meas_le_per_res(rate=_rate)
                                # time.sleep(1)
                                per = eval(res[1])
                                per_list.append(per)
                                time.sleep(0.5)
                            loginfo(per_list)
                            per = max(per_list)
                            loginfo('{}     {}      {}      {}      {}'.format(chan, rate, freq_inter, power_inter,  per))
                            fw1.write_data([chan, rate, freq_inter, power_inter,  per, agc_gain_index, agc_tgt, agc_sat])

                            if per > 30:
                                ci = -67 - power_inter +1
                                fw2.write_data([chan, rate, freq_inter, power_inter, ci, agc_tgt, agc_sat])
                                break


        self.tester_inter.output_state(0, 0)

    def test_le_sens(self, chan_list=[], rate_list=[], rx_pwr_range=[-99,-90], cable_loss=5, csv_save=True):
        if csv_save:
            title = 'rate,channel,sensitivity,per(%)\n'
            fname = self.get_filename('ts_bt_test/','test_le_per_{}'.format(self.board_name))
            fw1 = csvreport(fname,title)
        for rate in rate_list:
            if rate == 'LE_1M' or rate == 'LE1M':
                _rate = 'LE1M'
            else:
                _rate = 'LE2M'
            self.config_per_le_connect_usb(rate=_rate, packet_len=37)
            self.csp.RF_Frequency_Settings_rx(mode='DTM', ch_tx=0)
            self.csp.config_rxq_le_packets(rate=_rate, num=1500)
            self.csp.config_rx_level(rxpwr=-67)
            self.csp.RF_Power_settings_autoranging(0)  # 打开autorange会导致接收per 概率性100%
            time.sleep(1)
            for chan in chan_list:
                self.csp.RF_Frequency_Settings_rx(mode='DTM', ch_tx=chan)
                for rxpwr in range(rx_pwr_range[0],rx_pwr_range[1],1):
                    self.csp.config_rx_level(rxpwr=rxpwr)

                    per_list = []
                    for loop in range(5):
                        res = self.csp.meas_le_per_res(rate=_rate)
                        time.sleep(1)
                        per = eval(res[1])
                        per_list.append(per)

                    per = max(per_list)
                    if per < 30.8:
                        break
                fw1.write_data([rate,chan,rxpwr,per])



    def le_rx_intermod(self,  cable_loss=5, chan_list=[0], rate_list=['LE_1M'], packet_len=37, pkt_num=2000,  csv_save=True):
        self.dict_pwr_inter_freqoffset = {
            0: range(-90, -60),
            1: range(-85, -50),
            2: range(-52, -30),
            3: range(-41, -18),
            -1: range(-85, -50),
            -2: range(-52, -30),
            -3: range(-41, -18)
            }

        if csv_save:
            title = 'signal_channel,rate,freq_interference_tone,freq_interference_mod,rxpwr_interference,per(%)\n'
            fname = self.get_filename('ts_bt_test/','le_rx_ci_{}'.format(self.board_name))
            fw1 = csvreport(fname,title)
        for chan in chan_list:
            for rate in rate_list:
                if rate == 'LE_1M':
                    _rate = 'LE1M'
                else:
                    _rate = 'LE2M'
                self.tester_inter1 = mxg.MXG(device_name='N5182B', num_of_machine=1)
                self.tester_inter1.arb_waveform(rate=rate+'_PN15')
                self.tester_inter1.trriger_para_set(type='CONT', count=pkt_num)
                self.tester_inter1.output_state(1, 1)
                self.tester_inter1.arb_state(1)

                self.tester_inter2 = mxg.MXG(device_name='N5182B', num_of_machine=2)

                self.tester_inter2.output_state(1, 0)

                self.config_per_le_connect_usb(rate=_rate, packet_len=packet_len)
                self.csp.RF_Frequency_Settings_rx(mode='DTM', ch_tx=chan)
                self.csp.config_rxq_le_packets(rate=_rate, num=pkt_num)
                self.csp.config_rx_level(rxpwr=-67)

                freq_inter_offset_range = range(3,13) + range(-3,-13,-1)
                for freq_inter_offset in freq_inter_offset_range:
                    freq_inter2 = freq_inter_offset + 2402 + chan * 2
                    freq_inter1 = freq_inter_offset * 2 + 2402 + chan * 2
                    pwr_list = self.dict_pwr_inter_freqoffset[freq_inter_offset]
                    for power_inter in pwr_list:
                        self.tester_inter1.para_set(freq=freq_inter1, power=power_inter + cable_loss)
                        self.tester_inter2.para_set(freq=freq_inter2, power=power_inter + cable_loss)
                        time.sleep(0.2)
                        loginfo('freq_interference_tone:    {}      pwr_inter:  {}'.format(freq_inter2,power_inter))
                        per_list=[]
                        for loop in range(10):
                            res = self.csp.meas_le_per_res(rate=_rate)
                            per = eval(res[1])
                            per_list.append(per)
                            time.sleep(0.5)

                        per = max(per_list)
                        loginfo('{}     {}      {}      {}      {}      {}'.format(chan, rate, freq_inter2, freq_inter1, power_inter + cable_loss,  per))
                        fw1.write_data([chan, rate, freq_inter2, freq_inter1,  power_inter + cable_loss,  per])


    def le_rx_sniff(self, chan=0, rxpwr=-60, tx_power=15, time_interval=0.5):
        import matplotlib.pyplot as plt
        import numpy as np
        import time
        from math import *
        plt.ion()  # 开启interactive mode 成功的关键函数
        plt.figure(1)

        self.config_per_le_connect_usb(rate='LE1M', packet_len=37)
        self.csp.RF_Frequency_Settings_rx(mode='DTM', ch_tx=chan)
        self.csp.config_rxq_le_packets(rate='LE1M', num=1500)
        self.csp.RF_Power_settings(rx_level=rxpwr, tx_power=tx_power,margin=8)
        self.csp.RF_Power_settings_autoranging(0)   #打开autorange会导致接收per 概率性100%
        time.sleep(1)
        per_list=[]
        t=0
        t_list=[]
        while 1 :
            # plt.clf()  # 清空画布上的所有内容
            res = self.csp.meas_le_per_res(rate='LE1M')
            time.sleep(time_interval)
            t = t+time_interval
            per = eval(res[1])
            rxpkt_count = self.jlink.rd(0xa04008d8)
            loginfo('rxpkt_count:{}'.format(rxpkt_count))
            # per_list.append(per)
            # t_list.append(t)
            # plt.plot(t_list,per_list, '.r')
            plt.plot(t, per, '.r')
            plt.draw()
            plt.pause(0.01)

    def br_rx_sniff(self, chan=0, rxpwr=-60, time_interval=0.5, count=10000):
        import matplotlib.pyplot as plt
        import numpy as np
        import time
        from math import *
        plt.ion()  # 开启interactive mode 成功的关键函数
        plt.figure(200)
        plt.subplot(211)
        plt.subplot(212)


        self.config_classic_ber(rate='BR', chan=chan, rx_level=rxpwr, tx_power=15)
        self.csp.config_rx_level(rxpwr=rxpwr)
        self.csp.config_rxq_br_packets(num=1000)
        self.csp.config_rxq_br_tout(tout=10)
        time.sleep(1)
        ber_list=[]
        per_list=[]
        t=0
        t_list=[]
        for i in range(0,count) :
            # plt.clf()  # 清空画布上的所有内容
            while 1:
                res = self.csp.meas_bt_ber_res(cmd_type='READ')
                if eval(res[0]) == 0 or eval(res[0]) == 1:
                    break
                else:
                    if self.csp.get_bt_connect_state() == 'SBY':
                        self.cmd_br_connect()
                        time.sleep(10)
            time.sleep(time_interval)
            t = t+1
            ber = eval(res[1])
            per = eval(res[2])
            loginfo('ber  {}   per  {}'.format(ber,per))
            # ber_list.append(ber)
            # per_list.append(per)
            # t_list.append(t)
            plt.subplot(211)
            plt.plot(t,ber, '.r')
            plt.subplot(212)
            plt.plot( t, per, '.b')
            plt.pause(0.01)

    def get_filename(self, folder, file_name, sub_folder=''):
        '''
        :folder: file store folder
        :file_name:  file name
        :sub_folder: if not need, it may be default ""
        '''
        if rfglobal.file_folder=="":
            rfdata_path = './rftest/rfdata/'
        else:
            rfdata_path = './rftest/rfdata/%s/'%rfglobal.file_folder
            if os.path.exists(rfdata_path) == False:
                os.mkdir(rfdata_path)

        data_path1 = rfdata_path+'%s/'%(folder)
        if os.path.exists(data_path1) == False:
            os.mkdir(data_path1)

        filetime = time.strftime('%Y%m%d_%H%M%S',time.localtime(time.time()));
        # mac = self.read_mac()
        mac = ''

        gen_folder = '_%s'%(filetime[0:8])
        data_path2 = data_path1 +'%s/'%(gen_folder)
        if os.path.exists(data_path2) == False:
            os.mkdir(data_path2)

        fname = '%s'%(file_name)
        outfile_name = data_path2 + fname

        if sub_folder != '':
            gen_folder = '%s_%s'%(sub_folder,filetime[0:8])
            sub_path = data_path2+'%s/'%(gen_folder)
            if os.path.exists(sub_path) == False:
                os.mkdir(sub_path)

            outfile_name = sub_path + file_name

        return outfile_name

class BQB_autotest(bt_signaling):
    # def __init__(self, comport, chipv='TX212', rfport=1, atten=4.6):
    #     # bt_signaling(comport, chipv=chipv, rfport=rfport, atten=atten)
    #     self.rfport = rfport
    #     self.atten = atten
    #     self.comport = comport
    #     self.chipv = chipv
        # self.comport.Hexmd=1

    def pass_verdict(self, argu, meas):
        low_limit = str(argu.cells[1].text)
        upper_limit = str(argu.cells[2].text)
        loginfo(low_limit,upper_limit)
        argu.cells[5].text = ''
        if low_limit == '' and upper_limit != '':
            if meas < eval(upper_limit):
                run = argu.cells[5].paragraphs[0].add_run('Pass')
                run.font.color.rgb = RGBColor(0, 200, 0)
                res = 'Pass'
            else:
                run = argu.cells[5].paragraphs[0].add_run('Fail')
                run.font.color.rgb = RGBColor(200, 0, 0)
                res = 'Fail'
        elif upper_limit == '' and low_limit != '':
            if meas > eval(low_limit):
                run = argu.cells[5].paragraphs[0].add_run('Pass')
                run.font.color.rgb = RGBColor(0, 200, 0)
                res = 'Pass'
            else:
                run = argu.cells[5].paragraphs[0].add_run('Fail')
                run.font.color.rgb = RGBColor(200, 0, 0)
                res = 'Fail'
        elif low_limit != '' and upper_limit != '':
            if meas >= eval(low_limit) and meas <= eval(upper_limit):
                run = argu.cells[5].paragraphs[0].add_run('Pass')
                run.font.color.rgb = RGBColor(0, 200, 0)
                res = 'Pass'
            else:
                run = argu.cells[5].paragraphs[0].add_run('Fail')
                run.font.color.rgb = RGBColor(200, 0, 0)
                res = 'Fail'
        else:
            run = argu.cells[5].paragraphs[0].add_run('Pass')
            run.font.color.rgb = RGBColor(0, 200, 0)
            res = 'Pass'
        return res

    def bqb_autotest(self, inter_cableloss=5, signal_cableloss=6, dut_name='', temp=25):
        # title = 'reg_addr,reg_msb,reg_lsb,RW,reg_name,description,value_default,value_present_{}\n'.format(time.strftime("%Y-%m-%d %H:%M:%S", time.localtime()))
        # fname = self.get_filename('BQB_autotest', 'BQB_autotest')
        # fw1 = csvreport(fname, title)
        # shutil.copyfile(r'E:\chip\eagletest\py_script\rftest\test_script\BQB_autotest\BQB_RF_PHY_test_report.docx',os.path.join(fname,'BQB_RF_PHY_test_report_{}_{}.docx'.format(
        #     dut_name,temp)))
        self.document = Document(r'E:\chip\eagletest\py_script\rftest\test_script\BQB_autotest\BQB_RF_PHY_test_report_tws_L_1.docx')
        fname = 'E:/chip/eagletest/py_script/rftest/test_script/BQB_autotest/'
        # self.document = Document(r'E:\chip\eagletest\py_script\rftest\rfdata\BQB_autotest\_20210629\BQB_RF_PHY_test_report_QFN56_ECO_25.docx')

        self.tables = self.document.tables

        self.fname = fname
        self.dut_name = dut_name
        self.temp = temp
        self.cableloss = inter_cableloss
        self.atten =signal_cableloss

        ##BR/EDR test
        # self.csp.sig_std('CLAS')
        # self.config_classic_ber(rate='BR', chan=1, rx_level=-50, tx_power=10)
        # self.csp.tx_measure_para(tout=10, repetition='SINGleshot', count=10, MOEXception='OFF')
        # self.BR_Output_Power()
        # self.BR_Power_Density()
        # self.BR_Power_Control()
        # self.BR_Frequency_Range()
        # self.BR_20dB_Bandwidth()
        # self.BR_ACP()
        # self.BR_Modulation_Characteristics()
        # self.BR_Initial_Carrier_Frequency_Tolerance()
        # self.BR_Carrier_Frequency_Drift()
        # self.EDR_Relative_Transmit_Power()
        # self.EDR_Modulation_Accuracy()
        # self.EDR_Differential_Phase_Encoding()
        # self.EDR_Inband_Spurious_Emissions()
        # self.EDR_Power_Control()
        # self.BR_Sensitivity_single_slot_packets()
        # self.BR_Sensitivity_multi_slot_packets()
        # self.BR_CI_Performance()
        # self.BR_Blocking_Performance()
        # self.BR_Maximum_Input_Level()
        # self.EDR_Sensitivity()
        # self.EDR_BER_Floor_Performance()
        # self.EDR_CI_Performance()
        # self.EDR_Maximum_Input_Level()
        #
        # ##LE test
        # self.csp.sig_std('LE')
        # self.config_per_le_connect_usb(rate='LE1M', packet_len=37)
        # self.csp.config_rx_level(rxpwr=-67)
        # self.csp.trigger_settings(source='power', threshold=-35, tout=1)
        #
        # self.LE1M_Output_power()
        # self.LE1M_Inband_emissions()
        # self.LE1M_Modulation_Characteristics()
        # self.LE1M_Carrier_frequency_offset_and_drift()
        # self.LE1M_Receiver_sensitivity()
        # self.LE1M_Maximum_input_signal_level()
        # self.LE1M_PER_Report_Integrity()
        # self.LE1M_CI_Performance()
        # # self.LE1M_Blocking_Performance()

        # self.LE2M_Output_power()
        # self.LE2M_Inband_emissions()
        # self.LE2M_Modulation_Characteristics()
        # self.LE2M_Carrier_frequency_offset_and_drift()
        # self.LE2M_Receiver_sensitivity()
        # self.LE2M_Maximum_input_signal_level()
        # self.LE2M_PER_Report_Integrity()
        # self.LE2M_CI_Performance()
        # # self.LE2M_Blocking_Performance()

        # for i in range(3,len(self.tables)):
        #     logdebug(len(self.tables))
        #     res_list=[]
        #     for row in range(2,len(self.tables[i].rows)):
        #         self.tables[i].cell(row, 3).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
        #         self.tables[i].cell(row, 5).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
        #         res_list.append(self.tables[i].cell(row,5).text)
        #
        #     loginfo(res_list)
        #     for res in res_list:
        #         if res == 'Fail':
        #             sum_res = 'Fail'
        #             break
        #         elif res == 'Pass':
        #             sum_res = 'Pass'
        #         else:
        #             sum_res = ''
        #     if i <= 15:
        #         self.tables[2].cell(i - 2, 2).text = ''
        #         run = self.tables[2].cell(i - 2, 2).paragraphs[0].add_run(sum_res)
        #     elif i ==16:
        #         if self.tables[2].cell(i - 3, 2).text != 'Fail':
        #             self.tables[2].cell(i - 3, 2).text = ''
        #             run = self.tables[2].cell(i - 3, 2).paragraphs[0].add_run(sum_res)
        #     elif i > 16 and i <= 26:
        #         self.tables[2].cell(i - 3, 2).text = ''
        #         run = self.tables[2].cell(i - 3, 2).paragraphs[0].add_run(sum_res)
        #     elif i == 27:
        #         if self.tables[2].cell(i - 4, 2).text != 'Fail':
        #             self.tables[2].cell(i - 4, 2).text = ''
        #             run = self.tables[2].cell(i - 4, 2).paragraphs[0].add_run(sum_res)
        #     else:
        #         self.tables[2].cell(i - 4, 2).text = ''
        #         run = self.tables[2].cell(i - 4, 2).paragraphs[0].add_run(sum_res)
        #     if sum_res == 'Fail':
        #         run.font.color.rgb = RGBColor(200, 0, 0)
        #     else:
        #         run.font.color.rgb = RGBColor(0, 200, 0)
        # for i in range(2,len(self.tables[2].rows)):
        #     self.tables[2].cell(i,2).paragraphs[0].alginment = WD_TABLE_ALIGNMENT.CENTER
        #
        # self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))


    def BR_Output_Power(self):
        self.csp.mode_set(mode='BR')
        self.csp.sig_btype(btype='BR')
        self.csp.config_connect_br_packet_ptype(ptype='DH1')  # 设置BR包类型
        self.csp.config_connect_br_packet_pattern(pattern='PRBS9')  # 设置BR包 payload 数据类型
        self.csp.trigger_settings(source='power', threshold=-20, tout=1)
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan*39)
            res1 = self.csp.get_power_measure_res()
            nominal_pow = eval(res1[2])
            peak_pow = eval(res1[3])
            self.tables[3].cell(3+chan*3, 3).text = str(nominal_pow)
            self.tables[3].cell(4+chan*3, 3).text = str(peak_pow)
            self.pass_verdict(self.tables[3].rows[3+chan*3], nominal_pow)
            self.pass_verdict(self.tables[3].rows[4+chan*3], peak_pow)
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))
        self.csp.trigger_settings(source='power', threshold=-30, tout=1)

    def BR_Power_Density(self):
        self.csp.mode_set(mode='BR')
        self.csp.sig_btype(btype='BR')
        self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=18)
        self.csp.config_connect_br_packet_ptype(ptype='DH1')  # 设置BR包类型
        self.csp.config_connect_br_packet_pattern(pattern='PRBS9')  # 设置BR包 payload 数据类型
        self.csp.trigger_settings(source='power', threshold=-20, tout=1)
        res1 = self.csp.get_power_measure_res()
        nominal_pow = eval(res1[2])
        peak_pow = eval(res1[3])
        self.tables[4].cell(2, 3).text = str(peak_pow)
        self.pass_verdict(self.tables[4].rows[2], peak_pow)
        self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))
        self.csp.trigger_settings(source='power', threshold=-30, tout=1)

    def BR_Power_Control(self):
        self.csp.mode_set(mode='BR')
        self.csp.sig_btype(btype='BR')
        self.csp.config_connect_br_packet_ptype(ptype='DH1')  # 设置BR包类型
        self.csp.config_connect_br_packet_pattern(pattern='PRBS9')  # 设置BR包 payload 数据类型
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan*39)

            self.csp.config_power_control(para='MAX')
            time.sleep(5)
            res1 = self.csp.get_power_measure_res()
            nominal_pow_max = eval(res1[2])
            peak_pow = eval(res1[3])
            self.tables[5].cell(3+chan*14, 3).text = str(nominal_pow_max)
            self.pass_verdict(self.tables[5].rows[3+chan*14], nominal_pow_max)
            nominal_pow = nominal_pow_max
            for i in range(4,9):
                self.csp.config_power_control(para='DOWN')
                self.csp.trigger_settings(source='power', threshold=-20-i, tout=1)
                res1 = self.csp.get_power_measure_res()
                nominal_pow_d = eval(res1[2])
                power_step  = nominal_pow - nominal_pow_d
                nominal_pow = nominal_pow_d
                self.tables[5].cell(i+chan*14, 3).text = str(power_step)
                self.pass_verdict(self.tables[5].rows[i+chan*14], power_step)

            for i in range(10):
                self.csp.config_power_control(para='DOWN')
                res = self.csp.config_power_control_state()
                res = res.split(',')[1]
                logdebug(res)
                if res == 'MIN':
                    time.sleep(5)
                    res1 = self.csp.get_power_measure_res()
                    nominal_pow_min = eval(res1[2])
                    break
            self.tables[5].cell(9+chan*14, 3).text = str(nominal_pow_min)
            self.pass_verdict(self.tables[5].rows[9+chan*14], nominal_pow_min)
            nominal_pow = nominal_pow_min

            for i in range(10,15):
                self.csp.config_power_control(para='UP')
                self.csp.trigger_settings(source='power', threshold=-35 + i, tout=1)
                res1 = self.csp.get_power_measure_res()
                nominal_pow_d = eval(res1[2])
                power_step = nominal_pow_d - nominal_pow
                nominal_pow = nominal_pow_d
                self.tables[5].cell(i+chan*14, 3).text = str(power_step)
                self.pass_verdict(self.tables[5].rows[i+chan*14], power_step)

            self.csp.config_power_control(para='MAX')
            res1 = self.csp.get_power_measure_res()
            nominal_pow_max = eval(res1[2])
            peak_pow = eval(res1[3])
            self.tables[5].cell(15+chan*14, 3).text = str(nominal_pow_max)
            self.pass_verdict(self.tables[5].rows[15+chan*14], nominal_pow_max)
            self.csp.trigger_settings(source='power', threshold=-30, tout=1)
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))


    def BR_Frequency_Range(self):
        self.csp.mode_set(mode='BR')
        self.csp.sig_btype(btype='BR')
        self.csp.config_connect_br_packet_pattern(pattern='PRBS9')  # 设置BR包 payload 数据类型
        self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=0)
        res4 = self.csp.get_frange_res()
        frange_l = eval(res4[3]) / 1000000.00
        self.tables[6].cell(2, 3).text = str(frange_l)
        self.pass_verdict(self.tables[6].rows[2], frange_l)
        self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=78)
        res4 = self.csp.get_frange_res()
        frange_h = eval(res4[4]) / 1000000.00
        self.tables[6].cell(3, 3).text = str(frange_h)
        self.pass_verdict(self.tables[6].rows[3], frange_h)
        self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def BR_20dB_Bandwidth(self):
        self.csp.mode_set(mode='BR')
        self.csp.sig_btype(btype='BR')
        self.csp.config_connect_br_packet_pattern(pattern='PRBS9')  # 设置BR包 payload 数据类型
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan*39)
            res3 = self.csp.get_obw_res()
            obw = eval(res3[6]) / 1000.00
            obw_l = eval(res3[4]) / 1000.00
            obw_h = eval(res3[5]) / 1000.00
            self.tables[7].cell(3+chan*4, 3).text = str(obw_l)
            self.pass_verdict(self.tables[7].rows[3+chan*4], obw_l)
            self.tables[7].cell(4+chan*4, 3).text = str(obw_h)
            self.pass_verdict(self.tables[7].rows[4+chan*4], obw_h)
            self.tables[7].cell(5+chan*4, 3).text = str(obw)
            self.pass_verdict(self.tables[7].rows[5+chan*4], obw)
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def BR_ACP(self):
        self.csp.mode_set(mode='BR')
        self.csp.sig_btype(btype='BR')
        self.csp.config_connect_br_packet_pattern(pattern='PRBS9')  # 设置BR包 payload 数据类型
        self.csp.acp_meas_settings(opt='CH79')
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=3+chan*36)
            self.csp.tx_measure_states(state=0)
            res2 = self.csp.get_acp_res('READ')
            time.sleep(1)
            acp_list = [eval(i) for i in res2[1:]]
            loginfo(acp_list)
            for i in range(len(acp_list)):
                self.tables[8].cell(3 + i + chan * 80, 3).text = str(acp_list[i])
                self.pass_verdict(self.tables[8].rows[3 + i + chan * 80], acp_list[i])
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def BR_Modulation_Characteristics(self):
        self.csp.mode_set(mode='BR')
        self.csp.sig_btype(btype='BR')
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan * 39)
            delta_f1_99, delta_f2_99, delta_f2_avg, delta_f2_min, delta_f2_max, mod_ratio = -999, -999, -999, -999, -999, -999
            self.csp.config_connect_br_packet_pattern(pattern='P44')  # 设置BR包 payload 数据类型
            res = self.csp.get_modulation_measure_res()
            logdebug('{}'.format(res))

            delta_f1_99 = eval(res[2]) / 1000.00
            freq_accuracy = eval(res[3]) / 1000.00
            freq_drift = eval(res[4]) / 1000.00
            drift_rate = eval(res[5])
            delta_f1_avg = eval(res[6]) / 1000.00
            delta_f1_min = eval(res[7]) / 1000.00
            delta_f1_max = eval(res[8]) / 1000.00
            time.sleep(0.5)

            self.csp.config_connect_br_packet_pattern(pattern='P11')  # 设置BR包 payload 数据类型
            res = self.csp.get_modulation_measure_res()
            logdebug('{}'.format(res))

            delta_f2_99 = eval(res[2]) / 1000.00
            freq_accuracy = eval(res[3]) / 1000.00
            freq_drift = eval(res[4]) / 1000.00
            drift_rate = eval(res[5]) / 1000.00
            delta_f2_avg = eval(res[9]) / 1000.00
            delta_f2_min = eval(res[10]) / 1000.00
            delta_f2_max = eval(res[11]) / 1000.00
            mod_ratio = delta_f2_avg / delta_f1_avg
            time.sleep(0.5)

            self.tables[9].cell(3 + chan * 4, 3).text = str(delta_f1_avg)
            self.pass_verdict(self.tables[9].rows[3 + chan * 4], delta_f1_avg)

            self.tables[9].cell(4 + chan * 4, 3).text = str(delta_f2_99)
            self.pass_verdict(self.tables[9].rows[4 + chan * 4], delta_f2_99)

            self.tables[9].cell(5 + chan * 4, 3).text = str(mod_ratio)
            self.pass_verdict(self.tables[9].rows[5 + chan * 4], mod_ratio)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def BR_Initial_Carrier_Frequency_Tolerance(self):
        self.csp.mode_set(mode='BR')
        self.csp.sig_btype(btype='BR')
        self.csp.config_connect_br_packet_pattern(pattern='P11')  # 设置BR包 payload 数据类型
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan * 39)
            res = self.csp.get_modulation_measure_res()
            freq_accuracy = eval(res[3]) / 1000.00
            self.tables[10].cell(3 + chan * 2, 3).text = str(freq_accuracy)
            self.pass_verdict(self.tables[10].rows[3 + chan * 2], freq_accuracy)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def BR_Carrier_Frequency_Drift(self):
        self.csp.mode_set(mode='BR')
        self.csp.sig_btype(btype='BR')
        self.csp.config_connect_br_packet_pattern(pattern='P11')  # 设置BR包 payload 数据类型
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan * 39)
            self.csp.config_connect_br_packet_ptype(ptype='DH1')  # 设置BR包类型
            res = self.csp.get_modulation_measure_res()
            freq_drift = eval(res[4]) / 1000.00
            drift_rate = eval(res[5]) / 1000.00
            self.tables[11].cell(3 + chan * 7, 3).text = str(freq_drift)
            self.pass_verdict(self.tables[11].rows[3 + chan * 7], freq_drift)
            self.tables[11].cell(6 + chan * 7, 3).text = str(freq_drift)
            self.pass_verdict(self.tables[11].rows[6 + chan * 7], freq_drift)

            self.csp.config_connect_br_packet_ptype(ptype='DH3')  # 设置BR包类型
            res = self.csp.get_modulation_measure_res()
            freq_drift = eval(res[4]) / 1000.00
            drift_rate = eval(res[5]) / 1000.00
            self.tables[11].cell(4 + chan * 7, 3).text = str(freq_drift)
            self.pass_verdict(self.tables[11].rows[4 + chan * 7], freq_drift)
            self.tables[11].cell(7 + chan * 7, 3).text = str(freq_drift)
            self.pass_verdict(self.tables[11].rows[7 + chan * 7], freq_drift)

            self.csp.config_connect_br_packet_ptype(ptype='DH5')  # 设置BR包类型
            res = self.csp.get_modulation_measure_res()
            freq_drift = eval(res[4]) / 1000.00
            drift_rate = eval(res[5]) / 1000.00
            self.tables[11].cell(5 + chan * 7, 3).text = str(freq_drift)
            self.pass_verdict(self.tables[11].rows[5 + chan * 7], freq_drift)
            self.tables[11].cell(8 + chan * 7, 3).text = str(freq_drift)
            self.pass_verdict(self.tables[11].rows[8 + chan * 7], freq_drift)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def EDR_Relative_Transmit_Power(self):
        self.csp.mode_set(mode='EDR')
        self.csp.sig_btype(btype='EDR')
        self.csp.config_connect_edr_packet_pattern(pattern='PRBS9')  # 设置EDR包 payload 数据类型
        self.csp.trigger_settings(source='power', threshold=-20, tout=1)
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan * 39)
            self.csp.config_connect_edr_packet_ptype(ptype='E25P')  # 设置EDR包类型
            res1 = self.csp.get_power_measure_res()
            logdebug('{}'.format(res1))
            res1 = [eval(i) for i in res1[0:]]
            nominal_pwr = res1[2]
            gfsk_pwr = res1[3]
            dpsk_pwr = res1[4]
            dpsk_gfsk_diff_pwr = res1[5]
            guard_period = res1[6]
            packet_timing = res1[7]
            peak_pwr = res1[8]

            self.tables[12].cell(3 + chan * 7, 3).text = str(dpsk_gfsk_diff_pwr)
            self.pass_verdict(self.tables[12].rows[3 + chan * 7], dpsk_gfsk_diff_pwr)
            self.tables[12].cell(4 + chan * 7, 3).text = str(gfsk_pwr)
            self.pass_verdict(self.tables[12].rows[4 + chan * 7], gfsk_pwr)
            self.tables[12].cell(5 + chan * 7, 3).text = str(dpsk_pwr)
            self.pass_verdict(self.tables[12].rows[5 + chan * 7], dpsk_pwr)

            self.csp.config_connect_edr_packet_ptype(ptype='E35P')  # 设置EDR包类型
            res1 = self.csp.get_power_measure_res()
            logdebug('{}'.format(res1))
            res1 = [eval(i) for i in res1[0:]]
            nominal_pwr = res1[2]
            gfsk_pwr = res1[3]
            dpsk_pwr = res1[4]
            dpsk_gfsk_diff_pwr = res1[5]
            guard_period = res1[6]
            packet_timing = res1[7]
            peak_pwr = res1[8]

            self.tables[12].cell(6 + chan * 7, 3).text = str(dpsk_gfsk_diff_pwr)
            self.pass_verdict(self.tables[12].rows[6 + chan * 7], dpsk_gfsk_diff_pwr)
            self.tables[12].cell(7 + chan * 7, 3).text = str(gfsk_pwr)
            self.pass_verdict(self.tables[12].rows[7 + chan * 7], gfsk_pwr)
            self.tables[12].cell(8 + chan * 7, 3).text = str(dpsk_pwr)
            self.pass_verdict(self.tables[12].rows[8 + chan * 7], dpsk_pwr)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))
        self.csp.trigger_settings(source='power', threshold=-30, tout=1)

    def EDR_Modulation_Accuracy(self):
        self.csp.mode_set(mode='EDR')
        self.csp.sig_btype(btype='EDR')
        self.csp.config_connect_edr_packet_pattern(pattern='PRBS9')  # 设置EDR包 payload 数据类型
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan * 39)
            self.csp.config_connect_edr_packet_ptype(ptype='E25P')  # 设置EDR包类型
            for i in range(1):
                self.csp.config_power_control(para='DOWN')
            res = self.csp.get_modulation_measure_res()
            logdebug('{}'.format(res))
            res = [eval(i) for i in res[0:]]
            wi = res[2] / 1000
            w0_wi = res[3] / 1000
            w0_max = res[4] / 1000
            DEVM_RMS = res[5]*100.00
            DEVM_peak = res[6]*100.00
            DEVM_P99 = res[7]*100.00

            self.tables[13].cell(3 + chan * 13, 3).text = str(wi)
            self.pass_verdict(self.tables[13].rows[3 + chan * 13], wi)
            self.tables[13].cell(4 + chan * 13, 3).text = str(w0_wi)
            self.pass_verdict(self.tables[13].rows[4 + chan * 13], w0_wi)
            self.tables[13].cell(5 + chan * 13, 3).text = str(w0_max)
            self.pass_verdict(self.tables[13].rows[5 + chan * 13], w0_max)
            self.tables[13].cell(6 + chan * 13, 3).text = str(DEVM_RMS)
            self.pass_verdict(self.tables[13].rows[6 + chan * 13], DEVM_RMS)
            self.tables[13].cell(7 + chan * 13, 3).text = str(DEVM_peak)
            self.pass_verdict(self.tables[13].rows[7 + chan * 13], DEVM_peak)
            self.tables[13].cell(8 + chan * 13, 3).text = str(DEVM_P99)
            self.pass_verdict(self.tables[13].rows[8 + chan * 13], DEVM_P99)

            self.csp.config_connect_edr_packet_ptype(ptype='E35P')  # 设置EDR包类型
            for i in range(1):
                self.csp.config_power_control(para='DOWN')
            res = self.csp.get_modulation_measure_res()
            logdebug('{}'.format(res))
            res = [eval(i) for i in res[0:]]
            wi = res[2] / 1000
            w0_wi = res[3] / 1000
            w0_max = res[4] / 1000
            DEVM_RMS = res[5]*100.00
            DEVM_peak = res[6]*100.00
            DEVM_P99 = res[7]*100.00

            self.tables[13].cell(9 + chan * 13, 3).text = str(wi)
            self.pass_verdict(self.tables[13].rows[9 + chan * 13], wi)
            self.tables[13].cell(10 + chan * 13, 3).text = str(w0_wi)
            self.pass_verdict(self.tables[13].rows[10 + chan * 13], w0_wi)
            self.tables[13].cell(11 + chan * 13, 3).text = str(w0_max)
            self.pass_verdict(self.tables[13].rows[11 + chan * 13], w0_max)
            self.tables[13].cell(12 + chan * 13, 3).text = str(DEVM_RMS)
            self.pass_verdict(self.tables[13].rows[12 + chan * 13], DEVM_RMS)
            self.tables[13].cell(13 + chan * 13, 3).text = str(DEVM_peak)
            self.pass_verdict(self.tables[13].rows[13 + chan * 13], DEVM_peak)
            self.tables[13].cell(14 + chan * 13, 3).text = str(DEVM_P99)
            self.pass_verdict(self.tables[13].rows[14 + chan * 13], DEVM_P99)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def EDR_Differential_Phase_Encoding(self):
        self.csp.mode_set(mode='EDR')
        self.csp.sig_btype(btype='EDR')
        self.csp.config_connect_edr_packet_pattern(pattern='PRBS9')
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='TXT', ch_tx=chan * 39)
            self.csp.config_connect_edr_packet_ptype(ptype='E21P')  # 设置EDR包类型
            res2 = self.csp.get_diff_phase_encoding_res()
            logdebug('{}'.format(res2))
            res2 = [eval(i) for i in res2[0:]]
            bit_error_rate = res2[2]
            packet0error = res2[3]
            self.tables[14].cell(3 + chan * 3, 3).text = str(packet0error)
            self.pass_verdict(self.tables[14].rows[3 + chan * 3], packet0error)

            self.csp.config_connect_edr_packet_ptype(ptype='E31P')  # 设置EDR包类型
            res2 = self.csp.get_diff_phase_encoding_res()
            logdebug('{}'.format(res2))
            res2 = [eval(i) for i in res2[0:]]
            bit_error_rate = res2[2]
            packet0error = res2[3]
            self.tables[14].cell(4 + chan * 3, 3).text = str(packet0error)
            self.pass_verdict(self.tables[14].rows[4 + chan * 3], packet0error)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def EDR_Inband_Spurious_Emissions(self):
        self.csp.mode_set(mode='EDR')
        self.csp.sig_btype(btype='EDR')
        self.csp.config_connect_edr_packet_pattern(pattern='PRBS9')
        self.csp.acp_meas_settings(opt='CH79')
        self.csp.config_connect_edr_packet_ptype(ptype='E21P')  # 设置EDR包类型
        self.csp.trigger_settings(source='power', threshold=-40, tout=1)
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=3+chan * 36)
            time.sleep(1)
            self.csp.config_power_control(para='MAX')
            self.csp.RF_Power_settings(rx_level=-40, tx_power=15, margin=8)

            res1 = self.csp.get_power_measure_res()
            logdebug('{}'.format(res1))
            res1 = [eval(i) for i in res1[0:]]
            nominal_pwr = res1[2]
            self.csp.RF_Power_settings(rx_level=-40, tx_power=nominal_pwr+3, margin=8)
            # time.sleep(2)
            # # for i in range(1):
            # self.csp.config_power_control(para='DOWN')
            # self.csp.config_power_control(para='DOWN')
            time.sleep(1)
            [res3, res4] = self.csp.get_acp_res_edr('READ')
            logdebug('{}'.format(res3))
            logdebug('{}'.format(res4))
            res3 = [eval(i) for i in res3[0:]]
            res4 = [eval(i) for i in res4[0:]]
            acp_nominal_pwr = res3[2]
            PTxRef = res3[4]
            N26ChN1Abs = res3[5]
            N26ChP1Abs = res3[6]
            N26ChN1Rel = res3[7]
            N26ChP1Rel = res3[8]

            acp_list = res4[1:]
            for i in range(len(acp_list)):
                self.tables[15].cell(4 + i + chan * 80, 3).text = str(acp_list[i])
                self.pass_verdict(self.tables[15].rows[4 + i + chan * 80], acp_list[i])

            self.tables[15].cell(6 + chan*36 + chan * 80, 3).text = str(N26ChN1Abs)
            self.pass_verdict(self.tables[15].rows[6 + chan*36 + chan * 80], N26ChN1Abs)
            self.tables[15].cell(7 + chan*36 + chan * 80, 3).text = str(PTxRef)
            self.pass_verdict(self.tables[15].rows[7 + chan*36 + chan * 80], PTxRef)
            self.tables[15].cell(8 + chan*36 + chan * 80, 3).text = str(N26ChP1Abs)
            self.pass_verdict(self.tables[15].rows[8 + chan*36 + chan * 80], N26ChP1Abs)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

        self.csp.config_connect_edr_packet_ptype(ptype='E31P')  # 设置EDR包类型

        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=3+chan * 36)
            time.sleep(1)
            self.csp.config_power_control(para='MAX')
            # time.sleep(2)
            # # for i in range(1):
            # self.csp.config_power_control(para='DOWN')
            # self.csp.config_power_control(para='DOWN')
            time.sleep(1)
            [res3, res4] = self.csp.get_acp_res_edr('READ')
            logdebug('{}'.format(res3))
            logdebug('{}'.format(res4))
            res3 = [eval(i) for i in res3[0:]]
            res4 = [eval(i) for i in res4[0:]]
            acp_nominal_pwr = res3[2]
            PTxRef = res3[4]
            N26ChN1Abs = res3[5]
            N26ChP1Abs = res3[6]
            N26ChN1Rel = res3[7]
            N26ChP1Rel = res3[8]

            acp_list = res4[1:]
            for i in range(len(acp_list)):
                self.tables[16].cell(4 + i + chan * 80, 3).text = str(acp_list[i])
                self.pass_verdict(self.tables[16].rows[4 + i + chan * 80], acp_list[i])

            self.tables[16].cell(6 + chan*36 + chan * 80, 3).text = str(N26ChN1Abs)
            self.pass_verdict(self.tables[16].rows[6 + chan*36 + chan * 80], N26ChN1Abs)
            self.tables[16].cell(7 + chan*36 + chan * 80, 3).text = str(PTxRef)
            self.pass_verdict(self.tables[16].rows[7 + chan*36 + chan * 80], PTxRef)
            self.tables[16].cell(8 + chan*36 + chan * 80, 3).text = str(N26ChP1Abs)
            self.pass_verdict(self.tables[16].rows[8 + chan*36 + chan * 80], N26ChP1Abs)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))
        self.csp.trigger_settings(source='power', threshold=-30, tout=1)

    def EDR_Power_Control(self):
        self.csp.mode_set(mode='EDR')
        self.csp.sig_btype(btype='EDR')
        self.csp.config_connect_edr_packet_pattern(pattern='PRBS9')
        self.csp.config_connect_edr_packet_ptype(ptype='E25P')  # 设置EDR包类型
        self.csp.RF_Power_settings_autoranging(1)
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan*39)
            self.csp.config_power_control(para='MAX')
            res1 = self.csp.get_power_measure_res()
            logdebug('{}'.format(res1))
            res1 = [eval(i) for i in res1[0:]]
            nominal_pow_max = res1[2]
            self.tables[17].cell(3 + chan * 14, 3).text = str(nominal_pow_max)
            self.pass_verdict(self.tables[17].rows[3 + chan * 14], nominal_pow_max)
            nominal_pow = nominal_pow_max
            for i in range(4, 9):
                self.csp.config_power_control(para='DOWN')
                # self.csp.RF_Power_settings(rx_level=-50, tx_power=10-3*(i-3), margin=8)
                # self.csp.trigger_settings(source='power', threshold=-20 - i, tout=1)
                res1 = self.csp.get_power_measure_res()
                nominal_pow_d = eval(res1[2])
                power_step = nominal_pow - nominal_pow_d
                nominal_pow = nominal_pow_d
                self.tables[17].cell(i + chan * 14, 3).text = str(power_step)
                self.pass_verdict(self.tables[17].rows[i + chan * 14], power_step)

            for i in range(10):
                self.csp.config_power_control(para='DOWN')
                res = self.csp.config_power_control_state()
                res = res.split(',')[1]
                logdebug(res)
                if res == 'MIN':
                    res1 = self.csp.get_power_measure_res()
                    nominal_pow_min = eval(res1[2])
                    break
            self.tables[17].cell(9 + chan * 14, 3).text = str(nominal_pow_min)
            self.pass_verdict(self.tables[17].rows[9 + chan * 14], nominal_pow_min)
            nominal_pow = nominal_pow_min

            for i in range(10, 15):
                self.csp.config_power_control(para='UP')
                # self.csp.RF_Power_settings(rx_level=-50, tx_power=-10 + 3 * (i-9), margin=8)
                # self.csp.trigger_settings(source='power', threshold=-35 + i, tout=1)
                res1 = self.csp.get_power_measure_res()
                nominal_pow_d = eval(res1[2])
                power_step = nominal_pow_d - nominal_pow
                nominal_pow = nominal_pow_d
                self.tables[17].cell(i + chan * 14, 3).text = str(power_step)
                self.pass_verdict(self.tables[17].rows[i + chan * 14], power_step)

            self.csp.config_power_control(para='MAX')
            # self.csp.RF_Power_settings(rx_level=-50, tx_power=10, margin=8)
            res1 = self.csp.get_power_measure_res()
            nominal_pow_max = eval(res1[2])
            peak_pow = eval(res1[3])
            self.tables[17].cell(15 + chan * 14, 3).text = str(nominal_pow_max)
            self.pass_verdict(self.tables[17].rows[15 + chan * 14], nominal_pow_max)
            # self.csp.trigger_settings(source='power', threshold=-20, tout=1)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

        self.csp.RF_Power_settings_autoranging(0)

    def BR_Sensitivity_single_slot_packets(self):
        self.csp.mode_set(mode='BR')
        self.csp.sig_btype(btype='BR')
        self.csp.config_rx_level(rxpwr=-80)
        self.csp.config_rxq_br_search_packets(num=2000)
        self.csp.config_rxq_br_search_step(step=0.5)
        self.csp.config_connect_br_packet_pattern(pattern='PRBS9')
        self.csp.config_connect_br_packet_ptype(ptype='DH1')
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan*39)
            self.csp.ber_search_meas_state(state=0)
            while 1:
                state_res = self.csp.get_ber_search_meas_state()  # ber search 完成则退出循环
                # logdebug(state_res)
                if state_res == 'RDY':
                    loginfo('search complete')
                    break
                elif state_res == 'OFF':
                    self.cmd_br_connect()
                    self.csp.ber_search_meas_state(state=0)
                else:
                    logdebug('searching')
                time.sleep(5)

            res = self.csp.meas_bt_ber_search_res(cmd_type='FETCh')
            ber = eval(res[1])
            per = eval(res[2])
            NAK = eval(res[5])
            hec_err = eval(res[6])
            crc_err = eval(res[7])
            packet_type_err = eval(res[8])
            pay_err = eval(res[9])
            sens = eval(res[11])
            loginfo('{} {}  {}  {}  {}  {}  {}  {}'.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err, sens))
            self.tables[18].cell(2 + chan * 1, 3).text = str(sens)
            self.pass_verdict(self.tables[18].rows[2 + chan * 1], sens)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def BR_Sensitivity_multi_slot_packets(self):
        self.csp.mode_set(mode='BR')
        self.csp.sig_btype(btype='BR')
        self.csp.config_rx_level(rxpwr=-80)
        self.csp.config_rxq_br_search_packets(num=2000)
        self.csp.config_rxq_br_search_step(step=0.5)
        self.csp.config_connect_br_packet_pattern(pattern='PRBS9')
        self.csp.config_connect_br_packet_ptype(ptype='DH5')
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan*39)
            self.csp.ber_search_meas_state(state=0)
            while 1:
                state_res = self.csp.get_ber_search_meas_state()  # ber search 完成则退出循环
                # logdebug(state_res)
                if state_res == 'RDY':
                    loginfo('search complete')
                    break
                elif state_res == 'OFF':
                    self.cmd_br_connect()
                    self.csp.ber_search_meas_state(state=0)
                else:
                    logdebug('searching')
                time.sleep(5)

            res = self.csp.meas_bt_ber_search_res(cmd_type='FETCh')
            ber = eval(res[1])
            per = eval(res[2])
            NAK = eval(res[5])
            hec_err = eval(res[6])
            crc_err = eval(res[7])
            packet_type_err = eval(res[8])
            pay_err = eval(res[9])
            sens = eval(res[11])
            loginfo('{} {}  {}  {}  {}  {}  {}  {}'.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err, sens))
            self.tables[19].cell(2 + chan * 1, 3).text = str(sens)
            self.pass_verdict(self.tables[19].rows[2 + chan * 1], sens)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def BR_CI_Performance(self):
        self.csp.mode_set(mode='BR')
        self.dict_pwr_inter_freqoffset = {
            0: range(-75, -40),
            1: range(-60, -30),
            2: range(-40, -10),
            3: range(-35, 0),
            4: range(-35, 0),
            -1: range(-60, -30),
            -2: range(-50, -10),
            -3: range(-35, 0),
            -4: range(-35, 0),
            }

        self.csp.sig_btype(btype='BR')
        self.csp.config_rx_level(rxpwr=-60)
        self.csp.config_rxq_br_packets(num=1000)
        self.csp.config_connect_br_packet_pattern(pattern='PRBS9')
        self.csp.config_connect_br_packet_ptype(ptype='DH5')
        self.tester_inter = mxg.MXG(device_name='N5182B', num_of_machine=1)
        self.tester_inter.arb_waveform(rate='DH1_PN15')
        self.tester_inter.trriger_para_set(type='CONT', count=2000)
        self.tester_inter.para_set(freq=2402, power=-100)
        for chan in range(1):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=3+chan*36)
            self.tester_inter.output_state(1, 1)
            self.tester_inter.arb_state(1)
            num = 0
            for freq_inter_offset in [0, 1, -1, 2, -2, 3, -3, 4, -4]:
                freq_inter = freq_inter_offset + 2405 + chan*36
                # if freq_inter_offset == 0 and btype == 'EDR':
                #     self.tester_inter.arb_waveform(rate=rate + '_PN15')
                # else:
                #     self.tester_inter.arb_waveform(rate='DH1_PN15')
                if -3 < freq_inter_offset < 3:
                    self.csp.config_rx_level(rxpwr=-60)
                    rxpwr_ref = -60
                else:
                    self.csp.config_rx_level(rxpwr=-67)
                    rxpwr_ref = -67
                pwr_list = self.dict_pwr_inter_freqoffset[freq_inter_offset]
                for power_inter in pwr_list:
                    self.tester_inter.para_set(freq=freq_inter, power=power_inter + self.cableloss)
                    time.sleep(0.2)
                    loginfo('freq_inter:    {}      pwr_inter:  {}'.format(freq_inter, power_inter))
                    timeout = 0
                    while 1:
                        if self.csp.get_bt_connect_state() == 'SBY':
                            self.cmd_br_connect()
                        self.csp.ber_meas_state(state=0)
                        while 1:
                            state = self.csp.get_ber_meas_state()
                            if state == 'RDY':
                                loginfo('ber run complte')
                                break
                            time.sleep(2)
                            timeout = timeout + 1
                            if timeout > 30:
                                loginfo('get result timeout')
                                break
                        if timeout > 30:
                            res = ['0', '1', '100', '100', '100', '100', '100', '100', '100', '100']
                            break
                        res = self.csp.meas_bt_ber_res(cmd_type='FETCh')
                        if eval(res[0]) == 0 or eval(res[0]) == 1:
                            break
                        else:
                            if self.csp.get_bt_connect_state() == 'SBY':
                                self.cmd_br_connect()
                    # while 1:
                    #     self.csp.ber_meas_state(state=0)
                    #     while 1:
                    #         state = self.csp.get_ber_meas_state()
                    #         if state == 'RDY':
                    #             loginfo('ber run complte')
                    #             break
                    #         time.sleep(2)
                    #     res = self.csp.meas_bt_ber_res(cmd_type='FETCh')
                    #     if eval(res[0]) == 0 or eval(res[0]) == 1:
                    #         break
                    #     else:
                    #         if self.csp.get_bt_connect_state() == 'SBY':
                    #             self.cmd_br_connect()

                    ber = eval(res[1])
                    per = eval(res[2])
                    NAK = eval(res[5])
                    hec_err = eval(res[6])
                    crc_err = eval(res[7])
                    packet_type_err = eval(res[8])
                    pay_err = eval(res[9])

                    loginfo('{} {}  {}  {}  {}  {}  {} '.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err))
                    if ber > 0.1 or per > 50 or NAK > 50:
                        self.tables[20].cell(4 + num + chan * 10, 3).text = str(rxpwr_ref-power_inter+1)
                        self.pass_verdict(self.tables[20].rows[4 + num + chan * 10], rxpwr_ref-power_inter+1)
                        break
                num = num + 1
            self.tester_inter.output_state(0, 0)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def BR_Blocking_Performance(self):
        self.csp.mode_set(mode='BR')
        self.csp.sig_btype(btype='BR')
        self.csp.config_rx_level(rxpwr=-60)
        self.csp.config_rxq_br_packets(num=500)
        self.csp.config_connect_br_packet_pattern(pattern='PRBS9')
        self.csp.config_connect_br_packet_ptype(ptype='DH1')
        self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=58)

        rxpwr = -13
        freq_range = range(30, 2401, 1) + range(2500, 6000, 1)
        freq_inter_list = []
        freq_inter_list2 = []
        for freq_inter in freq_range:
            freq_blocking_2 = 'NA'
            freq_blocking_1 = 'NA'

            if freq_inter < 2001 or freq_inter > 2999:
                pwr = rxpwr + 17
            else:
                pwr = rxpwr
            self.csp.config_rx_level(rxpwr=-67)
            timeout = 0
            while 1:
                if self.csp.get_bt_connect_state() == 'SBY':
                    self.cmd_br_connect()
                self.csp.ber_meas_state(state=0)
                while 1:
                    state = self.csp.get_ber_meas_state()
                    if state == 'RDY':
                        loginfo('ber run complte')
                        break
                    time.sleep(2)
                    timeout = timeout + 1
                    if timeout > 30:
                        loginfo('get result timeout')
                        break
                if timeout > 30:
                    res = ['0', '1', '100', '100', '100', '100', '100', '100', '100', '100']
                    break
                res = self.csp.meas_bt_ber_res(cmd_type='FETCh')
                if eval(res[0]) == 0 or eval(res[0]) == 1:
                    break
                else:
                    if self.csp.get_bt_connect_state() == 'SBY':
                        self.cmd_br_connect()
            # connect_num = 0
            # while 1:
            #     self.tester_inter.para_set(freq=freq_inter, power=pwr+self.cableloss)
            #
            #     loginfo('freq_inter:    {}'.format(freq_inter))
            #
            #     res = self.csp.meas_bt_ber_res(cmd_type='READ')
            #     if eval(res[0]) == 0 or eval(res[0]) == 1:
            #         break
            #     else:
            #         connect_num = connect_num + 1
            #         loginfo('reconnect num {}'.format(connect_num))
            #         if self.csp.get_bt_connect_state() == 'SBY':
            #             self.cmd_br_connect()
            #             if connect_num > 3:
            #                 res = ['0', '100', '100', '100', '100', '100', '100', '100', '100', '100', '100']
            #                 break
            ber = eval(res[1])
            per = eval(res[2])
            NAK = eval(res[5])
            hec_err = eval(res[6])
            crc_err = eval(res[7])
            packet_type_err = eval(res[8])
            pay_err = eval(res[9])
            loginfo('freq_blocking_1:   {}      level:  {}      ber:    {}      per:    {}'.format(freq_inter, pwr, ber, per))
            if ber > 0.1 or per > 90:
                loginfo('freq_blocking_1:   {}      level:  {}      ber:    {}      per:    {}'.format(freq_inter, pwr, ber, per))
                freq_blocking_1 = freq_inter
                freq_inter_list.append(freq_inter)

        while 1:
            loginfo('***********************************************************')
            loginfo(freq_inter_list)
            if len(freq_inter_list) > 24:
                rxpwr = rxpwr - 1
                freq_range = freq_inter_list
                freq_inter_list = []
                freq_inter_list2 = []
                for freq_inter in freq_range:
                    freq_blocking_2 = 'NA'
                    freq_blocking_1 = 'NA'

                    if freq_inter < 2001 or freq_inter > 2999:
                        pwr = rxpwr + 17
                    else:
                        pwr = rxpwr
                    self.csp.config_rx_level(rxpwr=-67)
                    timeout = 0
                    while 1:
                        if self.csp.get_bt_connect_state() == 'SBY':
                            self.cmd_br_connect()
                        self.csp.ber_meas_state(state=0)
                        while 1:
                            state = self.csp.get_ber_meas_state()
                            if state == 'RDY':
                                loginfo('ber run complte')
                                break
                            time.sleep(2)
                            timeout = timeout + 1
                            if timeout > 30:
                                loginfo('get result timeout')
                                break
                        if timeout > 30:
                            res = ['0', '1', '100', '100', '100', '100', '100', '100', '100', '100']
                            break
                        res = self.csp.meas_bt_ber_res(cmd_type='FETCh')
                        if eval(res[0]) == 0 or eval(res[0]) == 1:
                            break
                        else:
                            if self.csp.get_bt_connect_state() == 'SBY':
                                self.cmd_br_connect()
                    # connect_num = 0
                    # while 1:
                    #     self.tester_inter.para_set(freq=freq_inter, power=pwr + self.cableloss)
                    #
                    #     loginfo('freq_inter:    {}'.format(freq_inter))
                    #
                    #     res = self.csp.meas_bt_ber_res(cmd_type='READ')
                    #     if eval(res[0]) == 0 or eval(res[0]) == 1:
                    #         break
                    #     else:
                    #         connect_num = connect_num + 1
                    #         loginfo('reconnect num {}'.format(connect_num))
                    #         if self.csp.get_bt_connect_state() == 'SBY':
                    #             self.cmd_br_connect()
                    #             if connect_num > 3:
                    #                 res = ['0', '100', '100', '100', '100', '100', '100', '100', '100', '100', '100']
                    #                 break
                    ber = eval(res[1])
                    per = eval(res[2])
                    NAK = eval(res[5])
                    hec_err = eval(res[6])
                    crc_err = eval(res[7])
                    packet_type_err = eval(res[8])
                    pay_err = eval(res[9])
                    loginfo('freq_blocking_1:   {}      level:  {}      ber:    {}      per:    {}'.format(freq_inter, pwr, ber, per))
                    if ber > 0.1 or per > 50:
                        loginfo('freq_blocking_1:   {}      level:  {}      ber:    {}      per:    {}'.format(freq_inter, pwr, ber, per))
                        freq_blocking_1 = freq_inter
                        freq_inter_list.append(freq_inter)
            else:
                self.tables[21].cell(3, 3).text = str(rxpwr+17)
                self.pass_verdict(self.tables[21].rows[3], rxpwr+17)
                self.tables[21].cell(4, 3).text = str(rxpwr)
                self.pass_verdict(self.tables[21].rows[4], rxpwr)
                self.tables[21].cell(5, 3).text = str(rxpwr)
                self.pass_verdict(self.tables[21].rows[5], rxpwr)
                self.tables[21].cell(6, 3).text = str(rxpwr+17)
                self.pass_verdict(self.tables[21].rows[6], rxpwr+17)
                break
        self.tester_inter.output_state(0, 0)

        self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def BR_Maximum_Input_Level(self):
        self.csp.mode_set(mode='BR')
        rxpwr_range=range(-20,1)
        self.csp.sig_btype(btype='BR')
        self.csp.config_rx_level(rxpwr=-40)
        self.csp.config_rxq_br_packets(num=2000)
        self.csp.config_connect_br_packet_pattern(pattern='PRBS9')
        self.csp.config_connect_br_packet_ptype(ptype='DH1')
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan*39)
            self.csp.ber_meas_state(0)
            for rxpwr in rxpwr_range:
                self.csp.config_rx_level(rxpwr=rxpwr)
                # while 1:
                #     state_res = self.csp.get_ber_search_meas_state()
                #     if state_res == 'RDY':
                #         logdebug(state_res)
                #         break
                #     time.sleep(3)
                # delay = 0.001*packet_num*(rx_level+100)/step
                # time.sleep(delay)
                while 1:
                    self.csp.ber_meas_state(state=0)
                    while 1:
                        state = self.csp.get_ber_meas_state()
                        if state == 'RDY':
                            loginfo('ber run complte')
                            break
                        time.sleep(2)
                    res = self.csp.meas_bt_ber_res(cmd_type='FETCh')
                    if eval(res[0]) == 0 or eval(res[0]) == 1:
                        break
                    else:
                        if self.csp.get_bt_connect_state() == 'SBY':
                            self.cmd_br_connect()

                ber = eval(res[1])
                per = eval(res[2])
                NAK = eval(res[5])
                hec_err = eval(res[6])
                crc_err = eval(res[7])
                packet_type_err = eval(res[8])
                pay_err = eval(res[9])

                time.sleep(0.5)
                loginfo('{} {}  {}  {}  {}  {}  {}  {}'.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err, rxpwr))
                if ber > 0.1 or per > 90:
                    self.tables[23].cell(2 + chan * 1, 3).text = str(rxpwr)
                    self.pass_verdict(self.tables[23].rows[2 + chan * 1], rxpwr)
                    break
                if rxpwr == max(rxpwr_range):
                    self.tables[23].cell(2 + chan * 1, 3).text = str(rxpwr)
                    self.pass_verdict(self.tables[23].rows[2 + chan * 1], rxpwr)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))



    def EDR_Sensitivity(self):
        self.csp.mode_set(mode='EDR')
        self.csp.sig_btype(btype='EDR')
        self.csp.config_rx_level(rxpwr=-76)
        self.csp.config_rxq_br_search_packets(num=2000)
        self.csp.config_rxq_br_search_step(step=0.5)
        self.csp.config_connect_edr_packet_pattern(pattern='PRBS9')  # 设置EDR包 payload 数据类型
        self.csp.config_connect_edr_packet_ptype(ptype='E25P')
        for i in range(3):
            self.csp.config_power_control(para='DOWN')
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan*39)
            self.csp.ber_search_meas_state(state=0)
            while 1:
                state_res = self.csp.get_ber_search_meas_state()  # ber search 完成则退出循环
                # logdebug(state_res)
                if state_res == 'RDY':
                    loginfo('search complete')
                    break
                elif state_res == 'OFF':
                    self.cmd_br_connect()
                    self.csp.ber_search_meas_state(state=0)
                else:
                    logdebug('searching')
                time.sleep(5)

            res = self.csp.meas_bt_ber_search_res(cmd_type='FETCh')
            ber = eval(res[1])
            per = eval(res[2])
            NAK = eval(res[5])
            hec_err = eval(res[6])
            crc_err = eval(res[7])
            packet_type_err = eval(res[8])
            pay_err = eval(res[9])
            sens = eval(res[11])
            loginfo('{} {}  {}  {}  {}  {}  {}  {}'.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err, sens))
            self.tables[24].cell(3 + chan * 1, 3).text = str(sens)
            self.pass_verdict(self.tables[24].rows[3 + chan * 1], sens)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

        self.csp.config_connect_edr_packet_ptype(ptype='E35P')
        for i in range(3):
            self.csp.config_power_control(para='DOWN')
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan*39)
            self.csp.ber_search_meas_state(state=0)
            while 1:
                state_res = self.csp.get_ber_search_meas_state()  # ber search 完成则退出循环
                # logdebug(state_res)
                if state_res == 'RDY':
                    loginfo('search complete')
                    break
                elif state_res == 'OFF':
                    self.cmd_br_connect()
                    self.csp.ber_search_meas_state(state=0)
                else:
                    logdebug('searching')
                time.sleep(5)

            res = self.csp.meas_bt_ber_search_res(cmd_type='FETCh')
            ber = eval(res[1])
            per = eval(res[2])
            NAK = eval(res[5])
            hec_err = eval(res[6])
            crc_err = eval(res[7])
            packet_type_err = eval(res[8])
            pay_err = eval(res[9])
            sens = eval(res[11])
            loginfo('{} {}  {}  {}  {}  {}  {}  {}'.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err, sens))
            self.tables[24].cell(7 + chan * 1, 3).text = str(sens)
            self.pass_verdict(self.tables[24].rows[7 + chan * 1], sens)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))


    def EDR_BER_Floor_Performance(self):
        self.csp.mode_set(mode='EDR')
        self.csp.sig_btype(btype='EDR')
        self.csp.config_rx_level(rxpwr=-60)
        self.csp.config_rxq_br_search_packets(num=2000)
        self.csp.config_rxq_br_search_step(step=0.5)
        self.csp.config_connect_edr_packet_pattern(pattern='PRBS9')  # 设置EDR包 payload 数据类型
        self.csp.config_connect_edr_packet_ptype(ptype='E25P')
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan*39)
            time.sleep(1)
            self.csp.config_power_control(para='MAX')
            time.sleep(2)
            for i in range(3):
                self.csp.config_power_control(para='DOWN')
            while 1:
                self.csp.ber_meas_state(state=0)
                while 1:
                    state = self.csp.get_ber_meas_state()
                    if state == 'RDY':
                        loginfo('ber run complte')
                        break
                    time.sleep(2)
                res = self.csp.meas_bt_ber_res(cmd_type='FETCh')
                if eval(res[0]) == 0 or eval(res[0]) == 1:
                    break
                else:
                    if self.csp.get_bt_connect_state() == 'SBY':
                        self.cmd_br_connect()

            ber = eval(res[1])
            per = eval(res[2])
            NAK = eval(res[5])
            hec_err = eval(res[6])
            crc_err = eval(res[7])
            packet_type_err = eval(res[8])
            pay_err = eval(res[9])

            time.sleep(0.5)
            loginfo('{} {}  {}  {}  {}  {}  {}  {}'.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err, -60))

            self.tables[25].cell(3 + chan * 1, 3).text = str(ber)
            self.pass_verdict(self.tables[25].rows[3 + chan * 1], ber)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

        self.csp.config_connect_edr_packet_ptype(ptype='E35P')
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan*39)
            time.sleep(1)
            self.csp.config_power_control(para='MAX')
            time.sleep(2)
            for i in range(3):
                self.csp.config_power_control(para='DOWN')
            while 1:
                self.csp.ber_meas_state(state=0)
                while 1:
                    state = self.csp.get_ber_meas_state()
                    if state == 'RDY':
                        loginfo('ber run complte')
                        break
                    time.sleep(2)
                res = self.csp.meas_bt_ber_res(cmd_type='FETCh')
                if eval(res[0]) == 0 or eval(res[0]) == 1:
                    break
                else:
                    if self.csp.get_bt_connect_state() == 'SBY':
                        self.cmd_br_connect()

            ber = eval(res[1])
            per = eval(res[2])
            NAK = eval(res[5])
            hec_err = eval(res[6])
            crc_err = eval(res[7])
            packet_type_err = eval(res[8])
            pay_err = eval(res[9])

            time.sleep(0.5)
            loginfo('{} {}  {}  {}  {}  {}  {}  {}'.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err, -60))

            self.tables[25].cell(7 + chan * 1, 3).text = str(ber)
            self.pass_verdict(self.tables[25].rows[7 + chan * 1], ber)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def EDR_CI_Performance(self):
        self.csp.mode_set(mode='EDR')
        self.dict_pwr_inter_freqoffset = {
            0: range(-85, -40),
            1: range(-60, -30),
            2: range(-50, -10),
            3: range(-40, 0),
            4: range(-35, 0),
            -1: range(-60, -30),
            -2: range(-50, -10),
            -3: range(-40, 0),
            -4: range(-35, 0),
            }
        self.tester_inter = mxg.MXG(device_name='N5182B', num_of_machine=1)
        self.tester_inter.arb_waveform(rate='2_DH1_PN15')
        self.tester_inter.trriger_para_set(type='CONT', count=1000)
        self.tester_inter.para_set(freq=2402, power=-100)

        self.csp.sig_btype(btype='EDR')
        self.csp.config_rx_level(rxpwr=-60)
        self.csp.config_rxq_br_packets(num=1000)
        self.csp.config_connect_edr_packet_pattern(pattern='PRBS9')  # 设置EDR包 payload 数据类型
        self.csp.config_connect_edr_packet_ptype(ptype='E25P')

        for chan in range(1):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=3 + chan * 36)
            self.tester_inter.output_state(1, 1)
            self.tester_inter.arb_state(1)
            time.sleep(1)
            self.csp.config_power_control(para='MAX')
            time.sleep(2)
            for i in range(2):
                self.csp.config_power_control(para='DOWN')
            num = 0
            for freq_inter_offset in [0, 1, -1, 2, -2, 3, -3, 4, -4]:
                freq_inter = freq_inter_offset + 2405 + chan * 36
                if freq_inter_offset == 0:
                    self.tester_inter.arb_waveform(rate='2_DH1_PN15')
                else:
                    self.tester_inter.arb_waveform(rate='DH1_PN15')
                if -3 < freq_inter_offset < 3:
                    self.csp.config_rx_level(rxpwr=-60)
                    rxpwr_ref = -60
                else:
                    self.csp.config_rx_level(rxpwr=-67)
                    rxpwr_ref = -67
                pwr_list = self.dict_pwr_inter_freqoffset[freq_inter_offset]
                for power_inter in pwr_list:
                    self.tester_inter.para_set(freq=freq_inter, power=power_inter + self.cableloss)
                    time.sleep(0.2)
                    loginfo('freq_inter:    {}      pwr_inter:  {}'.format(freq_inter, power_inter))
                    timeout = 0
                    while 1:
                        if self.csp.get_bt_connect_state() == 'SBY':
                            self.cmd_br_connect()
                        self.csp.ber_meas_state(state=0)
                        while 1:
                            state = self.csp.get_ber_meas_state()
                            if state == 'RDY':
                                loginfo('ber run complte')
                                break
                            time.sleep(2)
                            timeout = timeout + 1
                            if timeout > 30:
                                loginfo('get result timeout')
                                break
                        if timeout > 30:
                            res = ['0', '1', '100', '100', '100', '100', '100', '100', '100', '100']
                            break
                        res = self.csp.meas_bt_ber_res(cmd_type='FETCh')
                        if eval(res[0]) == 0 or eval(res[0]) == 1:
                            break
                        else:
                            if self.csp.get_bt_connect_state() == 'SBY':
                                self.cmd_br_connect()

                    ber = eval(res[1])
                    per = eval(res[2])
                    NAK = eval(res[5])
                    hec_err = eval(res[6])
                    crc_err = eval(res[7])
                    packet_type_err = eval(res[8])
                    pay_err = eval(res[9])


                    if ber > 0.1 or per > 50 or NAK > 50:
                        loginfo('{} {}  {}  {}  {}  {}  {} '.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err))
                        self.tables[26].cell(5 + num + chan * 10, 3).text = str(rxpwr_ref-power_inter+1)
                        self.pass_verdict(self.tables[26].rows[5 + num + chan * 10], rxpwr_ref-power_inter+1)
                        break
                num = num + 1
            self.tester_inter.output_state(0, 0)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

        self.csp.config_connect_edr_packet_ptype(ptype='E35P')

        for chan in range(1):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=3 + chan * 36)
            self.tester_inter.output_state(1, 1)
            self.tester_inter.arb_state(1)
            time.sleep(1)
            self.csp.config_power_control(para='MAX')
            time.sleep(2)
            for i in range(2):
                self.csp.config_power_control(para='DOWN')
            num = 0
            for freq_inter_offset in [0, 1, -1, 2, -2, 3, -3, 4, -4]:
                freq_inter = freq_inter_offset + 2405 + chan * 36
                if freq_inter_offset == 0:
                    self.tester_inter.arb_waveform(rate='3_DH1_PN15')
                else:
                    self.tester_inter.arb_waveform(rate='DH1_PN15')
                if -3 < freq_inter_offset < 3:
                    self.csp.config_rx_level(rxpwr=-60)
                    rxpwr_ref = -60
                else:
                    self.csp.config_rx_level(rxpwr=-67)
                    rxpwr_ref = -67
                pwr_list = self.dict_pwr_inter_freqoffset[freq_inter_offset]
                for power_inter in pwr_list:
                    self.tester_inter.para_set(freq=freq_inter, power=power_inter + self.cableloss)
                    time.sleep(0.2)
                    loginfo('freq_inter:    {}      pwr_inter:  {}'.format(freq_inter, power_inter))
                    timeout = 0
                    while 1:
                        if self.csp.get_bt_connect_state() == 'SBY':
                            self.cmd_br_connect()
                        self.csp.ber_meas_state(state=0)
                        while 1:
                            state = self.csp.get_ber_meas_state()
                            if state == 'RDY':
                                loginfo('ber run complte')
                                break
                            time.sleep(2)
                            timeout = timeout + 1
                            if timeout > 30:
                                loginfo('get result timeout')
                                break
                        if timeout > 30:
                            res = ['0', '1', '100', '100', '100', '100', '100', '100', '100', '100']
                            break
                        res = self.csp.meas_bt_ber_res(cmd_type='FETCh')
                        if eval(res[0]) == 0 or eval(res[0]) == 1:
                            break
                        else:
                            if self.csp.get_bt_connect_state() == 'SBY':
                                self.cmd_br_connect()

                    ber = eval(res[1])
                    per = eval(res[2])
                    NAK = eval(res[5])
                    hec_err = eval(res[6])
                    crc_err = eval(res[7])
                    packet_type_err = eval(res[8])
                    pay_err = eval(res[9])


                    if ber > 0.1 or per > 50 or NAK > 50:
                        loginfo('{} {}  {}  {}  {}  {}  {} '.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err))
                        self.tables[27].cell(5 + num + chan * 10, 3).text = str(rxpwr_ref-power_inter+1)
                        self.pass_verdict(self.tables[27].rows[5 + num + chan * 10], rxpwr_ref-power_inter+1)
                        break
                num = num + 1
            self.tester_inter.output_state(0, 0)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def EDR_CI_debug(self,cableloss=5,ptype='E25P'):
        self.csp.mode_set(mode='EDR')
        self.dict_pwr_inter_freqoffset = {
            0: range(-85, -40),
            1: range(-60, -30),
            2: range(-35, -10),
            3: range(-35, 0),
            4: range(-35, 0),
            -1: range(-60, -30),
            -2: range(-50, -10),
            -3: range(-40, 0),
            -4: range(-35, 0),
            }
        self.tester_inter = mxg.MXG(device_name='N5182B', num_of_machine=1)
        self.tester_inter.arb_waveform(rate='2_DH1_PN15')
        self.tester_inter.trriger_para_set(type='CONT', count=1000)
        self.tester_inter.para_set(freq=2402, power=-100)
        self.tester_inter.output_state(0, 0)

        self.csp.sig_std('CLAS')
        self.config_classic_ber(rate='BR', chan=1, rx_level=-50, tx_power=10)
        self.csp.tx_measure_para(tout=10, repetition='SINGleshot', count=10, MOEXception='OFF')
        self.csp.sig_btype(btype='EDR')
        self.csp.config_rx_level(rxpwr=-60)
        self.csp.config_rxq_br_packets(num=1000)
        self.csp.config_connect_edr_packet_pattern(pattern='PRBS9')  # 设置EDR包 payload 数据类型
        self.csp.config_connect_edr_packet_ptype(ptype=ptype)


        ci_list=[]
        for chan in range(1):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=3 + chan * 36)
            self.tester_inter.output_state(1, 1)
            self.tester_inter.arb_state(1)
            time.sleep(1)
            self.csp.config_power_control(para='MAX')
            time.sleep(2)
            for i in range(2):
                self.csp.config_power_control(para='DOWN')
            num = 0
            for freq_inter_offset in [2, 3]:
                freq_inter = freq_inter_offset + 2405 + chan * 36
                if freq_inter_offset == 0:
                    self.tester_inter.arb_waveform(rate='2_DH1_PN15')
                else:
                    self.tester_inter.arb_waveform(rate='DH1_PN15')
                if -3 < freq_inter_offset < 3:
                    self.csp.config_rx_level(rxpwr=-60)
                    rxpwr_ref = -60
                else:
                    self.csp.config_rx_level(rxpwr=-67)
                    rxpwr_ref = -67
                pwr_list = self.dict_pwr_inter_freqoffset[freq_inter_offset]
                for power_inter in pwr_list:
                    self.tester_inter.para_set(freq=freq_inter, power=power_inter + cableloss)
                    time.sleep(0.2)
                    loginfo('freq_inter:    {}      pwr_inter:  {}'.format(freq_inter, power_inter))
                    timeout = 0
                    while 1:
                        if self.csp.get_bt_connect_state() == 'SBY':
                            self.cmd_br_connect()
                        self.csp.ber_meas_state(state=0)
                        while 1:
                            state = self.csp.get_ber_meas_state()
                            if state == 'RDY':
                                loginfo('ber run complte')
                                break
                            time.sleep(2)
                            timeout = timeout + 1
                            if timeout > 30:
                                loginfo('get result timeout')
                                break
                        if timeout > 30:
                            res = ['0', '1', '100', '100', '100', '100', '100', '100', '100', '100']
                            break
                        res = self.csp.meas_bt_ber_res(cmd_type='FETCh')
                        if eval(res[0]) == 0 or eval(res[0]) == 1:
                            break
                        else:
                            if self.csp.get_bt_connect_state() == 'SBY':
                                self.cmd_br_connect()

                    ber = eval(res[1])
                    per = eval(res[2])
                    NAK = eval(res[5])
                    hec_err = eval(res[6])
                    crc_err = eval(res[7])
                    packet_type_err = eval(res[8])
                    pay_err = eval(res[9])


                    if ber > 0.1 or per > 50 or NAK > 50:
                        loginfo('{} {}  {}  {}  {}  {}  {} '.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err))
                        loginfo('freq_offset: {} , CI: {}'.format(freq_inter_offset,rxpwr_ref-power_inter+1))
                        ci_list.append([freq_inter_offset,rxpwr_ref-power_inter+1])
                        # self.tables[26].cell(5 + num + chan * 10, 3).text = str(rxpwr_ref-power_inter+1)
                        # self.pass_verdict(self.tables[26].rows[5 + num + chan * 10], rxpwr_ref-power_inter+1)
                        break
                num = num + 1

            self.tester_inter.output_state(0, 0)

        loginfo(ci_list)


    def EDR_Maximum_Input_Level(self):
        self.csp.mode_set(mode='EDR')
        rxpwr_range = range(-20, 1)
        self.csp.sig_btype(btype='EDR')
        self.csp.config_rx_level(rxpwr=-40)
        self.csp.config_rxq_br_packets(num=2000)
        self.csp.config_connect_edr_packet_pattern(pattern='PRBS9')  # 设置EDR包 payload 数据类型
        self.csp.config_connect_edr_packet_ptype(ptype='E25P')
        for i in range(3):
            self.csp.config_power_control(para='DOWN')
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan * 39)
            self.csp.ber_meas_state(0)
            for rxpwr in rxpwr_range:
                self.csp.config_rx_level(rxpwr=rxpwr)
                while 1:
                    self.csp.ber_meas_state(state=0)
                    while 1:
                        state = self.csp.get_ber_meas_state()
                        if state == 'RDY':
                            loginfo('ber run complte')
                            break
                        time.sleep(2)
                    res = self.csp.meas_bt_ber_res(cmd_type='FETCh')
                    if eval(res[0]) == 0 or eval(res[0]) == 1:
                        break
                    else:
                        if self.csp.get_bt_connect_state() == 'SBY':
                            self.cmd_br_connect()

                ber = eval(res[1])
                per = eval(res[2])
                NAK = eval(res[5])
                hec_err = eval(res[6])
                crc_err = eval(res[7])
                packet_type_err = eval(res[8])
                pay_err = eval(res[9])

                time.sleep(0.5)
                loginfo('{} {}  {}  {}  {}  {}  {}  {}'.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err, rxpwr))
                if ber > 0.1 or per > 50 or NAK > 50:
                    self.tables[28].cell(3 + chan * 1, 3).text = str(rxpwr-1)
                    self.pass_verdict(self.tables[28].rows[3 + chan * 1], rxpwr-1)
                    break
                if rxpwr == max(rxpwr_range):
                    self.tables[28].cell(3 + chan * 1, 3).text = str(rxpwr)
                    self.pass_verdict(self.tables[28].rows[3 + chan * 1], rxpwr)
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

        self.csp.config_connect_edr_packet_ptype(ptype='E35P')
        time.sleep(1)
        self.csp.config_power_control(para='MAX')
        time.sleep(2)
        for i in range(3):
            self.csp.config_power_control(para='DOWN')
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan * 39)
            self.csp.ber_meas_state(0)
            for rxpwr in rxpwr_range:
                self.csp.config_rx_level(rxpwr=rxpwr)
                while 1:
                    self.csp.ber_meas_state(state=0)
                    while 1:
                        state = self.csp.get_ber_meas_state()
                        if state == 'RDY':
                            loginfo('ber run complte')
                            break
                        time.sleep(2)
                    res = self.csp.meas_bt_ber_res(cmd_type='FETCh')
                    if eval(res[0]) == 0 or eval(res[0]) == 1:
                        break
                    else:
                        if self.csp.get_bt_connect_state() == 'SBY':
                            self.cmd_br_connect()

                ber = eval(res[1])
                per = eval(res[2])
                NAK = eval(res[5])
                hec_err = eval(res[6])
                crc_err = eval(res[7])
                packet_type_err = eval(res[8])
                pay_err = eval(res[9])

                time.sleep(0.5)
                loginfo('{} {}  {}  {}  {}  {}  {}  {}'.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err, rxpwr))
                if ber > 0.1 or per > 50 or NAK > 50:
                    self.tables[28].cell(7 + chan * 1, 3).text = str(rxpwr-1)
                    self.pass_verdict(self.tables[28].rows[7 + chan * 1], rxpwr-1)
                    break
                if rxpwr == max(rxpwr_range):
                    self.tables[28].cell(7 + chan * 1, 3).text = str(rxpwr)
                    self.pass_verdict(self.tables[28].rows[7 + chan * 1], rxpwr)
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))


    def LE1M_Output_power(self):
        self.csp.mode_set(mode='LE1M')
        self.csp.sig_le_phy(phy='LE1M')
        # self.csp.trigger_settings(source='power', threshold=-35, tout=1)
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=1+chan * 19)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan * 19)
            self.csp.config_connect_le_packet_pattern(rate='LE1M', pattern='PRBS9')
            res1 = self.csp.get_power_measure_res(rate='LE1M')

            logdebug('{}'.format(res1))

            nominal_pow = eval(res1[2])
            peak_pow = eval(res1[3])
            leakage_pow = eval(res1[4])
            peak_avg_pwr = eval(res1[5])

            self.tables[29].cell(3 + chan * 3, 3).text = str(nominal_pow)
            self.pass_verdict(self.tables[29].rows[3 + chan * 3], nominal_pow)
            self.tables[29].cell(4 + chan * 3, 3).text = str(peak_avg_pwr)
            self.pass_verdict(self.tables[29].rows[4 + chan * 3], peak_avg_pwr)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))


    def LE1M_Inband_emissions(self):
        self.csp.mode_set(mode='LE1M')
        self.csp.sig_le_phy(phy='LE1M')
        self.csp.acp_meas_settings(opt='CH40')
        self.csp.trigger_settings(source='power', threshold=-40, tout=1)
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=3 + chan * 17)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=2+chan*17)
            self.csp.config_connect_le_packet_pattern(rate='LE1M', pattern='PRBS9')
            res2 = self.csp.get_acp_res()
            time.sleep(1)
            logdebug('{}'.format(res2))

            acp_list = [eval(i) for i in res2[1:]]
            for i in range(len(acp_list)):
                self.tables[30].cell(3 + i + chan * 82, 3).text = str(acp_list[i])
                self.pass_verdict(self.tables[30].rows[3 + i + chan * 82], acp_list[i])

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def LE1M_Modulation_Characteristics(self):
        self.csp.mode_set(mode='LE1M')
        self.csp.sig_le_phy(phy='LE1M')
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=1+chan * 19)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan * 19)

            delta_f1_99, delta_f2_99, delta_f1_avg, delta_f1_min, delta_f1_max, delta_f2_avg, delta_f2_min, delta_f2_max, mod_ratio = -999, -999, -999, -999, -999, -999, -999, -999, -999
            self.csp.config_connect_le_packet_pattern(rate='LE1M', pattern='P44')
            res = self.csp.get_modulation_measure_res(rate='LE1M')
            logdebug('{}'.format(res))

            delta_f1_99 = eval(res[2]) / 1000.00
            freq_accuracy = eval(res[3]) / 1000.00
            freq_drift = eval(res[4]) / 1000.00
            drift_rate = eval(res[5]) / 1000.00
            delta_f1_avg = eval(res[6]) / 1000.00
            delta_f1_min = eval(res[7]) / 1000.00
            delta_f1_max = eval(res[8]) / 1000.00

            time.sleep(0.5)

            self.csp.config_connect_le_packet_pattern(rate='LE1M', pattern='P11')
            res = self.csp.get_modulation_measure_res(rate='LE1M')
            logdebug('{}'.format(res))

            delta_f2_99 = eval(res[2]) / 1000.00
            freq_accuracy = eval(res[3]) / 1000.00
            freq_drift = eval(res[4]) / 1000.00
            drift_rate = eval(res[5]) / 1000.00
            delta_f2_avg = eval(res[9]) / 1000.00
            delta_f2_min = eval(res[10]) / 1000.00
            delta_f2_max = eval(res[11]) / 1000.00
            mod_ratio = delta_f2_avg / delta_f1_avg
            freq_offset = eval(res[14]) / 1000.00
            init_drift = eval(res[15]) / 1000.00

            self.tables[31].cell(3 + chan * 4, 3).text = str(delta_f1_avg)
            self.pass_verdict(self.tables[31].rows[3 + chan * 4], delta_f1_avg)

            self.tables[31].cell(4 + chan * 4, 3).text = str(delta_f2_99)
            self.pass_verdict(self.tables[31].rows[4 + chan * 4], delta_f2_99)

            self.tables[31].cell(5 + chan * 4, 3).text = str(mod_ratio)
            self.pass_verdict(self.tables[31].rows[5 + chan * 4], mod_ratio)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def LE1M_Carrier_frequency_offset_and_drift(self):
        self.csp.mode_set(mode='LE1M')
        self.csp.sig_le_phy(phy='LE1M')
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=1+chan * 19)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan * 19)

            delta_f1_99, delta_f2_99, delta_f1_avg, delta_f1_min, delta_f1_max, delta_f2_avg, delta_f2_min, delta_f2_max, mod_ratio = -999, -999, -999, -999, -999, -999, -999, -999, -999

            self.csp.config_connect_le_packet_pattern(rate='LE1M', pattern='P11')
            res = self.csp.get_modulation_measure_res(rate='LE1M')
            logdebug('{}'.format(res))

            delta_f2_99 = eval(res[2]) / 1000.00
            freq_accuracy = eval(res[3]) / 1000.00
            freq_drift = eval(res[4]) / 1000.00
            drift_rate = eval(res[5]) / 1000.00
            delta_f2_avg = eval(res[9]) / 1000.00
            delta_f2_min = eval(res[10]) / 1000.00
            delta_f2_max = eval(res[11]) / 1000.00
            mod_ratio = delta_f2_avg / delta_f1_avg
            freq_offset = eval(res[14]) / 1000.00
            init_drift = eval(res[15]) / 1000.00

            self.tables[32].cell(3 + chan * 5, 3).text = str(freq_accuracy)
            self.pass_verdict(self.tables[32].rows[3 + chan * 5], freq_accuracy)

            self.tables[32].cell(4 + chan * 5, 3).text = str(freq_drift)
            self.pass_verdict(self.tables[32].rows[4 + chan * 5], freq_drift)

            self.tables[32].cell(5 + chan * 5, 3).text = str(drift_rate)
            self.pass_verdict(self.tables[32].rows[5 + chan * 5], drift_rate)

            self.tables[32].cell(6 + chan * 5, 3).text = str(init_drift)
            self.pass_verdict(self.tables[32].rows[6 + chan * 5], init_drift)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def LE1M_Receiver_sensitivity(self):
        self.csp.mode_set(mode='LE1M')
        self.csp.sig_le_phy(phy='LE1M')
        self.csp.config_rxq_le_packets(rate='LE1M', num=2000)
        self.csp.config_connect_le_packet_pattern(rate='LE1M', pattern='PRBS9')
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=1+chan * 19)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan * 19)

            for i in range(41):
                rxpwr = -80 - i/2.0
                self.csp.config_rx_level(rxpwr=rxpwr)

                per_list = []
                for loop in range(3):
                    res = self.csp.meas_le_per_res(rate='LE1M')
                    time.sleep(0.5)
                    per = eval(res[1])
                    per_list.append(per)

                per = min(per_list)
                if per > 30.8:
                    self.tables[33].cell(2 + chan * 1, 3).text = str(rxpwr+0.5)
                    self.pass_verdict(self.tables[33].rows[2 + chan * 1], rxpwr+0.5)
                    break
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def LE1M_CI_Performance(self):
        self.csp.mode_set(mode='LE1M')
        self.csp.sig_le_phy(phy='LE1M')
        self.csp.config_rxq_le_packets(rate='LE1M', num=2000)
        self.csp.config_connect_le_packet_pattern(rate='LE1M', pattern='PRBS9')
        self.csp.config_rx_level(rxpwr=-67)
        self.dict_pwr_inter_freqoffset = {
            0: range(-90, -60),
            1: range(-85, -40),
            2: range(-52, -20),
            3: range(-41, -10),
            4: range(-41, -10),
            -1: range(-85, -40),
            -2: range(-52, -20),
            -3: range(-41, -10),
            -4: range(-41, -10)
            }
        self.tester_inter = mxg.MXG(device_name='N5182B', num_of_machine=1)
        self.tester_inter.arb_waveform(rate='LE_1M_PN15')
        self.tester_inter.trriger_para_set(type='CONT', count=2000)
        self.tester_inter.output_state(1, 1)
        self.tester_inter.arb_state(1)

        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=3 + chan * 17)
                cf = 3 + chan * 17
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=2 + chan * 17)
                cf = 2 + chan * 17
            time.sleep(1)
            num = 0
            for freq_inter_offset in[0,1,-1,2,-2,3,-3,4,-4]:
                freq_inter = freq_inter_offset+2402+cf*2
                pwr_list = self.dict_pwr_inter_freqoffset[freq_inter_offset]
                for power_inter in pwr_list:
                    self.tester_inter.para_set(freq=freq_inter, power=power_inter + self.cableloss)
                    time.sleep(0.2)
                    loginfo('freq_inter:    {}      pwr_inter:  {}'.format(freq_inter,power_inter))
                    per_list = []
                    for loop in range(3):
                        res = self.csp.meas_le_per_res(rate='LE1M')
                        time.sleep(0.2)
                        per = eval(res[1])
                        per_list.append(per)

                    per = min(per_list)
                    loginfo('{}     {}      {}      {}      {}'.format(chan, 'LE1M', freq_inter, power_inter,  per))
                    if per > 30.8:
                        self.tables[34].cell(4 + num + chan * 10, 3).text = str(-67-power_inter+1)
                        self.pass_verdict(self.tables[34].rows[4 + num + chan * 10], -67-power_inter+1)
                        break
                num = num + 1
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))
        self.tester_inter.output_state(0, 0)

    def LE1M_Blocking_Performance(self):
        self.csp.mode_set(mode='LE1M')
        self.csp.sig_le_phy(phy='LE1M')
        self.csp.config_rxq_le_packets(rate='LE1M', num=2000)
        self.csp.config_connect_le_packet_pattern(rate='LE1M', pattern='PRBS9')
        self.csp.config_rx_level(rxpwr=-67)
        self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=19)

        self.tester_inter = mxg.MXG(device_name='N5182B',num_of_machine=1)
        self.tester_inter.para_set(30,-30+self.cableloss)
        self.tester_inter.output_state(1, 0)

        freq_range= range(30,2001,10)+range(2003,2400,1)+range(2484,2998,1)+range(3000,6000,25)
        # freq_range = freq_range_le
        freq_blocking_2 = 'NA'
        freq_blocking_1 = 'NA'
        freq_inter_list = []
        freq_inter_list1 = []
        inter_rxpwr = -10
        # for rxpwr in inter_rxpwr:
        count = 0
        count1 = 0
        for freq_inter in freq_range:
            freq_blocking_2 = 'NA'
            freq_blocking_1 = 'NA'
            if freq_inter < 2001 or freq_inter > 2999:
                pwr =  inter_rxpwr + 5
                # self.tester_inter.para_set(freq=freq_inter, power=rxpwr + inter_loss)
            else:
                pwr = inter_rxpwr
            #     # self.tester_inter.para_set(freq=freq_inter, power=rxpwr-5 + inter_loss)
            self.tester_inter.para_set(freq=freq_inter, power=pwr+self.cableloss)
            loginfo('freq_inter:    {}'.format(freq_inter))
            self.csp.config_rx_level(rxpwr=-67)
            self.csp.per_meas_state(0)
            time.sleep(1)
            res = self.csp.meas_le_per_res(rate='LE1M')
            per = eval(res[1])
            if per > 30.8:
                loginfo('freq_blocking_1:   {}      level:  {}      per:    {}'.format(freq_inter,pwr,per))
                # freq_blocking_1 = freq_inter
                # if freq_inter > 1999 or freq_inter < 3001:
                #     # pwr = rxpwr - 5 + inter_loss
                #     rxpwr = rxpwr - 5
                freq_inter_list.append(freq_inter)
                # fw1.write_data([rate,rxpwr,freq_blocking_1,freq_blocking_2,per])
                count = count + 1

        for freq_inter_check in freq_inter_list:
            if freq_inter_check < 2001 or freq_inter_check > 2999:
                pwr =  inter_rxpwr + 5
                # self.tester_inter.para_set(freq=freq_inter, power=rxpwr + inter_loss)
            else:
                pwr = inter_rxpwr
            self.tester_inter.para_set(freq=freq_inter_check, power=pwr+self.cableloss)
            loginfo('freq_inter:    {}'.format(freq_inter_check))
            self.csp.per_meas_state(0)
            time.sleep(1)
            res = self.csp.meas_le_per_res(rate='LE1M')
            per = eval(res[1])
            if per > 30.8:
                loginfo('freq_blocking_1:   {}      level:  {}      per:    {}'.format(freq_inter_check, pwr, per))
                freq_inter_list1.append(freq_inter_check)
        while 1:
            loginfo('***********************************************************')
            loginfo(freq_inter_list1)
            if len(freq_inter_list1) >= 10:
                inter_rxpwr = inter_rxpwr -1
            elif len(freq_inter_list1) <= 5:
                inter_rxpwr = inter_rxpwr + 1
            else:
                self.tables[35].cell(3, 3).text = str(inter_rxpwr+5)
                self.pass_verdict(self.tables[35].rows[3], inter_rxpwr+5)
                self.tables[35].cell(4, 3).text = str(inter_rxpwr)
                self.pass_verdict(self.tables[35].rows[4], inter_rxpwr)
                self.tables[35].cell(5, 3).text = str(inter_rxpwr)
                self.pass_verdict(self.tables[35].rows[5], inter_rxpwr)
                self.tables[35].cell(6, 3).text = str(inter_rxpwr+5)
                self.pass_verdict(self.tables[35].rows[6], inter_rxpwr+5)

                break
            freq_range = freq_inter_list1
            freq_inter_list1 = []
            for freq_inter in freq_range:
                if freq_inter_check < 2001 or freq_inter_check > 2999:
                    pwr = inter_rxpwr + 5
                    # self.tester_inter.para_set(freq=freq_inter, power=rxpwr + inter_loss)
                else:
                    pwr = inter_rxpwr
                self.tester_inter.para_set(freq=freq_inter, power=pwr+self.cableloss)
                self.csp.per_meas_state(0)
                time.sleep(1)
                res = self.csp.meas_le_per_res(rate='LE1M')
                per = eval(res[1])
                if per > 30.8:
                    loginfo('freq_blocking_1:   {}      level:  {}      per:    {}'.format(freq_inter, pwr, per))
                    freq_inter_list1.append(freq_inter)

        self.tester_inter.output_state(0, 0)
        self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def LE1M_Maximum_input_signal_level(self):
        rxpwr_range = range(-20,1)
        self.csp.mode_set(mode='LE1M')
        self.csp.sig_le_phy(phy='LE1M')
        self.csp.config_rxq_le_packets(rate='LE1M', num=2000)
        self.csp.config_connect_le_packet_pattern(rate='LE1M', pattern='PRBS9')
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=1+chan * 19)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan * 19)

            for rxpwr in rxpwr_range:
                self.csp.config_rx_level(rxpwr=rxpwr)

                per_list = []
                for loop in range(3):
                    res = self.csp.meas_le_per_res(rate='LE1M')
                    time.sleep(0.5)
                    per = eval(res[1])
                    per_list.append(per)

                per = min(per_list)
                if per > 30.8:
                    self.tables[37].cell(2 + chan * 1, 3).text = str(rxpwr-1)
                    self.pass_verdict(self.tables[37].rows[2 + chan * 1], rxpwr-1)
                    break
                if rxpwr == max(rxpwr_range):
                    self.tables[37].cell(2 + chan * 1, 3).text = str(rxpwr)
                    self.pass_verdict(self.tables[37].rows[2 + chan * 1], rxpwr)
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))


    def LE1M_PER_Report_Integrity(self):
        self.csp.mode_set(mode='LE1M')
        self.csp.sig_le_phy(phy='LE1M')
        self.csp.config_rxq_le_packets(rate='LE1M', num=2000)
        self.csp.config_connect_le_packet_pattern(rate='LE1M', pattern='PRBS9')
        self.csp.config_rxq_le_integrity(rate='LE1M',en='ON')
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=1+chan * 19)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan * 19)


            self.csp.config_rx_level(rxpwr=-30)

            per_list = []
            for loop in range(3):
                res = self.csp.meas_le_per_res(rate='LE1M')
                time.sleep(0.5)
                per = eval(res[1])
                per_list.append(per)

            per = min(per_list)

            self.tables[38].cell(3 + chan * 1, 3).text = str(per)
            self.pass_verdict(self.tables[38].rows[3 + chan * 1], per)
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

        self.csp.config_rxq_le_integrity(rate='LE1M', en='OFF')


    def LE2M_Output_power(self):
        self.csp.mode_set(mode='LE2M')
        self.csp.sig_le_phy(phy='LE2M')
        self.csp.trigger_settings(source='power', threshold=-35, tout=1)
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=1+chan * 19)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan * 19)
            self.csp.config_connect_le_packet_pattern(rate='LE2M', pattern='PRBS9')
            res1 = self.csp.get_power_measure_res(rate='LE2M')

            logdebug('{}'.format(res1))

            nominal_pow = eval(res1[2])
            peak_pow = eval(res1[3])
            leakage_pow = eval(res1[4])
            peak_avg_pwr = eval(res1[5])

            self.tables[39].cell(3 + chan * 3, 3).text = str(nominal_pow)
            self.pass_verdict(self.tables[39].rows[3 + chan * 3], nominal_pow)
            self.tables[39].cell(4 + chan * 3, 3).text = str(peak_avg_pwr)
            self.pass_verdict(self.tables[39].rows[4 + chan * 3], peak_avg_pwr)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))
        self.csp.trigger_settings(source='power', threshold=-35, tout=1)

    def LE2M_Inband_emissions(self):
        self.csp.mode_set(mode='LE2M')
        self.csp.sig_le_phy(phy='LE2M')
        self.csp.acp_meas_settings(opt='CH40')
        self.csp.trigger_settings(source='power', threshold=-30, tout=1)
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=3 + chan * 17)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=2+chan*17)
            self.csp.config_connect_le_packet_pattern(rate='LE2M', pattern='PRBS9')
            res2 = self.csp.get_acp_res()
            time.sleep(1)
            logdebug('{}'.format(res2))

            acp_list = [eval(i) for i in res2[1:]]
            for i in range(len(acp_list)):
                self.tables[40].cell(3 + i + chan * 82, 3).text = str(acp_list[i])
                self.pass_verdict(self.tables[40].rows[3 + i + chan * 82], acp_list[i])

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))
        self.csp.trigger_settings(source='power', threshold=-35, tout=1)

    def LE2M_Modulation_Characteristics(self):
        self.csp.mode_set(mode='LE2M')
        self.csp.sig_le_phy(phy='LE2M')
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=1+chan * 19)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan * 19)

            delta_f1_99, delta_f2_99, delta_f1_avg, delta_f1_min, delta_f1_max, delta_f2_avg, delta_f2_min, delta_f2_max, mod_ratio = -999, -999, -999, -999, -999, -999, -999, -999, -999
            self.csp.config_connect_le_packet_pattern(rate='LE2M', pattern='P44')
            res = self.csp.get_modulation_measure_res(rate='LE2M')
            logdebug('{}'.format(res))

            delta_f1_99 = eval(res[2]) / 1000.00
            freq_accuracy = eval(res[3]) / 1000.00
            freq_drift = eval(res[4]) / 1000.00
            drift_rate = eval(res[5]) / 1000.00
            delta_f1_avg = eval(res[6]) / 1000.00
            delta_f1_min = eval(res[7]) / 1000.00
            delta_f1_max = eval(res[8]) / 1000.00

            time.sleep(0.5)

            self.csp.config_connect_le_packet_pattern(rate='LE2M', pattern='P11')
            res = self.csp.get_modulation_measure_res(rate='LE2M')
            logdebug('{}'.format(res))

            delta_f2_99 = eval(res[2]) / 1000.00
            freq_accuracy = eval(res[3]) / 1000.00
            freq_drift = eval(res[4]) / 1000.00
            drift_rate = eval(res[5]) / 1000.00
            delta_f2_avg = eval(res[9]) / 1000.00
            delta_f2_min = eval(res[10]) / 1000.00
            delta_f2_max = eval(res[11]) / 1000.00
            mod_ratio = delta_f2_avg / delta_f1_avg
            freq_offset = eval(res[14]) / 1000.00
            init_drift = eval(res[15]) / 1000.00

            self.tables[41].cell(3 + chan * 4, 3).text = str(delta_f1_avg)
            self.pass_verdict(self.tables[41].rows[3 + chan * 4], delta_f1_avg)

            self.tables[41].cell(4 + chan * 4, 3).text = str(delta_f2_99)
            self.pass_verdict(self.tables[41].rows[4 + chan * 4], delta_f2_99)

            self.tables[41].cell(5 + chan * 4, 3).text = str(mod_ratio)
            self.pass_verdict(self.tables[41].rows[5 + chan * 4], mod_ratio)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def LE2M_Carrier_frequency_offset_and_drift(self):
        self.csp.mode_set(mode='LE2M')
        self.csp.sig_le_phy(phy='LE2M')
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=1+chan * 19)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan * 19)

            delta_f1_99, delta_f2_99, delta_f1_avg, delta_f1_min, delta_f1_max, delta_f2_avg, delta_f2_min, delta_f2_max, mod_ratio = -999, -999, -999, -999, -999, -999, -999, -999, -999

            self.csp.config_connect_le_packet_pattern(rate='LE2M', pattern='P11')
            res = self.csp.get_modulation_measure_res(rate='LE2M')
            logdebug('{}'.format(res))

            delta_f2_99 = eval(res[2]) / 1000.00
            freq_accuracy = eval(res[3]) / 1000.00
            freq_drift = eval(res[4]) / 1000.00
            drift_rate = eval(res[5]) / 1000.00
            delta_f2_avg = eval(res[9]) / 1000.00
            delta_f2_min = eval(res[10]) / 1000.00
            delta_f2_max = eval(res[11]) / 1000.00
            mod_ratio = delta_f2_avg / delta_f1_avg
            freq_offset = eval(res[14]) / 1000.00
            init_drift = eval(res[15]) / 1000.00

            self.tables[42].cell(3 + chan * 5, 3).text = str(freq_accuracy)
            self.pass_verdict(self.tables[42].rows[3 + chan * 5], freq_accuracy)

            self.tables[42].cell(4 + chan * 5, 3).text = str(freq_drift)
            self.pass_verdict(self.tables[42].rows[4 + chan * 5], freq_drift)

            self.tables[42].cell(5 + chan * 5, 3).text = str(drift_rate)
            self.pass_verdict(self.tables[42].rows[5 + chan * 5], drift_rate)

            self.tables[42].cell(6 + chan * 5, 3).text = str(init_drift)
            self.pass_verdict(self.tables[42].rows[6 + chan * 5], init_drift)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def LE2M_Receiver_sensitivity(self):
        self.csp.mode_set(mode='LE2M')
        self.csp.sig_le_phy(phy='LE2M')
        self.csp.config_rxq_le_packets(rate='LE2M', num=2000)
        self.csp.config_connect_le_packet_pattern(rate='LE2M', pattern='PRBS9')
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=1+chan * 19)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan * 19)

            for i in range(41):
                rxpwr = -80 - i / 2.0
                self.csp.config_rx_level(rxpwr=rxpwr)

                per_list = []
                for loop in range(3):
                    res = self.csp.meas_le_per_res(rate='LE2M')
                    time.sleep(0.5)
                    per = eval(res[1])
                    per_list.append(per)

                per = min(per_list)
                if per > 30.8:
                    self.tables[43].cell(2 + chan * 1, 3).text = str(rxpwr+0.5)
                    self.pass_verdict(self.tables[43].rows[2 + chan * 1], rxpwr+0.5)
                    break
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def LE2M_CI_Performance(self):
        self.csp.mode_set(mode='LE2M')
        self.csp.sig_le_phy(phy='LE2M')
        self.csp.config_rxq_le_packets(rate='LE2M', num=2000)
        self.csp.config_connect_le_packet_pattern(rate='LE2M', pattern='PRBS9')
        self.csp.config_rx_level(rxpwr=-67)
        self.dict_pwr_inter_freqoffset = {
            0: range(-90, -60),
            1: range(-85, -40),
            2: range(-52, -20),
            3: range(-41, -10),
            4: range(-41, -10),
            -1: range(-85, -40),
            -2: range(-52, -20),
            -3: range(-41, -10),
            -4: range(-41, -10)
            }
        self.tester_inter = mxg.MXG(device_name='N5182B', num_of_machine=1)
        self.tester_inter.arb_waveform(rate='LE_2M_PN15')
        self.tester_inter.trriger_para_set(type='CONT', count=2000)
        self.tester_inter.output_state(1, 1)
        self.tester_inter.arb_state(1)

        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=3 + chan * 17)
                cf = 3 + chan * 17
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=2 + chan * 17)
                cf = 2 + chan * 17

            time.sleep(1)
            num = 0
            for freq_inter_offset in[0,1,-1,2,-2,3,-3,4,-4]:
                freq_inter = freq_inter_offset*2 + 2402 + cf*2
                pwr_list = self.dict_pwr_inter_freqoffset[freq_inter_offset]
                for power_inter in pwr_list:
                    self.tester_inter.para_set(freq=freq_inter, power=power_inter + self.cableloss)
                    time.sleep(0.2)
                    loginfo('freq_inter:    {}      pwr_inter:  {}'.format(freq_inter,power_inter))
                    per_list = []
                    for loop in range(3):
                        res = self.csp.meas_le_per_res(rate='LE2M')
                        time.sleep(0.2)
                        per = eval(res[1])
                        per_list.append(per)

                    per = min(per_list)
                    loginfo('{}     {}      {}      {}      {}'.format(chan, 'LE2M', freq_inter, power_inter,  per))
                    if per > 30.8:
                        self.tables[44].cell(4 + num + chan * 10, 3).text = str(-67-power_inter+1)
                        self.pass_verdict(self.tables[44].rows[4 + num + chan * 10], -67-power_inter+1)
                        break
                num = num + 1
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))
        self.tester_inter.output_state(0, 0)

    def LE2M_Blocking_Performance(self):
        self.csp.mode_set(mode='LE2M')
        self.csp.sig_le_phy(phy='LE2M')
        self.csp.config_rxq_le_packets(rate='LE2M', num=2000)
        self.csp.config_connect_le_packet_pattern(rate='LE2M', pattern='PRBS9')
        self.csp.config_rx_level(rxpwr=-67)
        self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=19)

        self.tester_inter = mxg.MXG(device_name='N5182B',num_of_machine=1)
        self.tester_inter.para_set(30,-30+self.cableloss)
        self.tester_inter.output_state(1, 0)

        freq_range= range(30,2001,10)+range(2003,2400,1)+range(2484,2998,1)+range(3000,6000,25)
        # freq_range = freq_range_le
        freq_blocking_2 = 'NA'
        freq_blocking_1 = 'NA'
        freq_inter_list = []
        freq_inter_list1 = []
        inter_rxpwr = -10
        # for rxpwr in inter_rxpwr:
        count = 0
        count1 = 0
        for freq_inter in freq_range:
            freq_blocking_2 = 'NA'
            freq_blocking_1 = 'NA'
            if freq_inter < 2001 or freq_inter > 2999:
                pwr =  inter_rxpwr + 5
                # self.tester_inter.para_set(freq=freq_inter, power=rxpwr + inter_loss)
            else:
                pwr = inter_rxpwr
            #     # self.tester_inter.para_set(freq=freq_inter, power=rxpwr-5 + inter_loss)
            self.tester_inter.para_set(freq=freq_inter, power=pwr+self.cableloss)
            loginfo('freq_inter:    {}'.format(freq_inter))
            self.csp.config_rx_level(rxpwr=-67)
            self.csp.per_meas_state(0)
            time.sleep(1)
            res = self.csp.meas_le_per_res(rate='LE2M')
            per = eval(res[1])
            if per > 30.8:
                loginfo('freq_blocking_1:   {}      level:  {}      per:    {}'.format(freq_inter,pwr,per))
                # freq_blocking_1 = freq_inter
                # if freq_inter > 1999 or freq_inter < 3001:
                #     # pwr = rxpwr - 5 + inter_loss
                #     rxpwr = rxpwr - 5
                freq_inter_list.append(freq_inter)
                # fw1.write_data([rate,rxpwr,freq_blocking_1,freq_blocking_2,per])
                count = count + 1

        for freq_inter_check in freq_inter_list:
            if freq_inter_check < 2001 or freq_inter_check > 2999:
                pwr =  inter_rxpwr + 5
                # self.tester_inter.para_set(freq=freq_inter, power=rxpwr + inter_loss)
            else:
                pwr = inter_rxpwr
            self.tester_inter.para_set(freq=freq_inter_check, power=pwr+self.cableloss)
            loginfo('freq_inter:    {}'.format(freq_inter_check))
            self.csp.per_meas_state(0)
            time.sleep(1)
            res = self.csp.meas_le_per_res(rate='LE2M')
            per = eval(res[1])
            if per > 30.8:
                loginfo('freq_blocking_1:   {}      level:  {}      per:    {}'.format(freq_inter_check, pwr, per))
                freq_inter_list1.append(freq_inter_check)
        while 1:
            loginfo('***********************************************************')
            loginfo(freq_inter_list1)
            if len(freq_inter_list1) >= 10:
                inter_rxpwr = inter_rxpwr -1
            elif len(freq_inter_list1) <= 5:
                inter_rxpwr = inter_rxpwr + 1
            else:
                self.tables[45].cell(3, 3).text = str(inter_rxpwr+5)
                self.pass_verdict(self.tables[45].rows[3], inter_rxpwr+5)
                self.tables[45].cell(4, 3).text = str(inter_rxpwr)
                self.pass_verdict(self.tables[45].rows[4], inter_rxpwr)
                self.tables[45].cell(5, 3).text = str(inter_rxpwr)
                self.pass_verdict(self.tables[45].rows[5], inter_rxpwr)
                self.tables[45].cell(6, 3).text = str(inter_rxpwr+5)
                self.pass_verdict(self.tables[45].rows[6], inter_rxpwr+5)

                break
            freq_range = freq_inter_list1
            freq_inter_list1 = []
            for freq_inter in freq_range:
                if freq_inter_check < 2001 or freq_inter_check > 2999:
                    pwr = inter_rxpwr + 5
                    # self.tester_inter.para_set(freq=freq_inter, power=rxpwr + inter_loss)
                else:
                    pwr = inter_rxpwr
                self.tester_inter.para_set(freq=freq_inter, power=pwr+self.cableloss)
                self.csp.per_meas_state(0)
                time.sleep(1)
                res = self.csp.meas_le_per_res(rate='LE2M')
                per = eval(res[1])
                if per > 30.8:
                    loginfo('freq_blocking_1:   {}      level:  {}      per:    {}'.format(freq_inter, pwr, per))
                    freq_inter_list1.append(freq_inter)

        self.tester_inter.output_state(0, 0)
        self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def LE2M_Maximum_input_signal_level(self):
        rxpwr_range = range(-20,1)
        self.csp.mode_set(mode='LE2M')
        self.csp.sig_le_phy(phy='LE2M')
        self.csp.config_rxq_le_packets(rate='LE2M', num=2000)
        self.csp.config_connect_le_packet_pattern(rate='LE2M', pattern='PRBS9')
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=1+chan * 19)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan * 19)

            for rxpwr in rxpwr_range:
                self.csp.config_rx_level(rxpwr=rxpwr)

                per_list = []
                for loop in range(3):
                    res = self.csp.meas_le_per_res(rate='LE2M')
                    time.sleep(0.5)
                    per = eval(res[1])
                    per_list.append(per)

                per = min(per_list)
                if per > 30.8:
                    self.tables[47].cell(2 + chan * 1, 3).text = str(rxpwr-1)
                    self.pass_verdict(self.tables[47].rows[2 + chan * 1], rxpwr-1)
                    break
                if rxpwr == max(rxpwr_range):
                    self.tables[37].cell(2 + chan * 1, 3).text = str(rxpwr)
                    self.pass_verdict(self.tables[37].rows[2 + chan * 1], rxpwr)
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))


    def LE2M_PER_Report_Integrity(self):
        self.csp.mode_set(mode='LE2M')
        self.csp.sig_le_phy(phy='LE2M')
        self.csp.config_rxq_le_packets(rate='LE2M', num=2000)
        self.csp.config_connect_le_packet_pattern(rate='LE2M', pattern='PRBS9')
        self.csp.config_rxq_le_integrity(rate='LE2M',en='ON')
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=1+chan * 19)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan * 19)


            self.csp.config_rx_level(rxpwr=-30)

            per_list = []
            for loop in range(3):
                res = self.csp.meas_le_per_res(rate='LE2M')
                time.sleep(0.5)
                per = eval(res[1])
                per_list.append(per)

            per = min(per_list)

            self.tables[48].cell(3 + chan * 1, 3).text = str(per)
            self.pass_verdict(self.tables[48].rows[3 + chan * 1], per)
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

        self.csp.config_rxq_le_integrity(rate='LE2M', en='OFF')
