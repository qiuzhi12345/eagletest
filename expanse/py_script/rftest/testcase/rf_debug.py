#!/usr/bin/python
# -*- coding:utf8 -*-

import re
from baselib.loglib.log_lib import *
from baselib.plot import *
from baselib.instrument import *
import numpy as np
import pandas as pd
import scipy
from math import *
from shutil import copy
import time
import csv
import pylab
import matplotlib.pyplot as plt
from baselib.test_channel import *
import xlrd
import sys
import random
import os
from baselib.instrument.cmw_bt import *
from baselib.instrument.spa import *
from rftest.rflib import *
from hal.common import *
from rftest.rflib.csv_report import csvreport
from rftest.testcase.bt_api import bt_api
from rftest.rflib import rfglobal
from rftest.testcase.bt_nosignal_test import tester_cmw

import serial
from sys import argv
import binascii
from binascii import b2a_hex, a2b_hex
import pylink
import docx
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_TAB_ALIGNMENT,WD_TAB_LEADER #设置制表符等
from docx.shared import Inches #设置图像大小
from docx.shared import Pt #设置像素、缩进等
from docx.shared import RGBColor #设置字体颜色
from docx.shared import Length #设置宽度
import win32api
import shutil
class rf_debug(object):
    def __init__(self, comport, chipv='TX232_MPW3', jlink='59610138',jlink_en=1,board_name=''):
        self.comport = comport
        self.chipv = chipv
        self.mem_ts = MEM_TS(self.comport)
        self.board_name = board_name
        self.jlink_en = jlink_en
        if jlink_en !=0:
            self.jlink = jlink
            self.btapi=bt_api(self.comport, chipv=self.chipv, jlink_en=jlink_en,jlink=self.jlink)
        else:
            self.btapi=bt_api(self.comport, chipv=self.chipv, jlink_en=jlink_en)
        self.br_len_dic = {
            'DH1': 27,
            'DH3': 183,
            'DH5': 339
        }
        self.br_rate_dic = {
            'DH1': 1,
            'DH3': 4,
            'DH5': 7
        }
        self.edr_len_dic = {
            '2_DH1': 31,
            '2_DH3': 356,
            '2_DH5': 656,
            '3_DH1': 83,
            '3_DH3': 552,
            '3_DH5': 1021,
        }
        self.edr_rate_dic = {
            '2_DH1': 2,
            '2_DH3': 5,
            '2_DH5': 8,
            '3_DH1': 3,
            '3_DH3': 6,
            '3_DH5': 9,
        }
        self.le_rate_dic = {
            'LE1M': 1,
            'LE2M': 2,
            'LE500k': 3,
            'LE125K': 4
        }

    def test232_tp_gain_cal(self,mode=1, freq=2441):
        '''
        mode:
        0:  le1m mode
        other:  le2m mode

        freq:   MHz
        '''
        if self.jlink_en==0:
            if mode != 0:
                self.mem_ts.wrm(0xa0421064, 12, 12, 1)  ##Frequency value manual enable
                self.mem_ts.wrm(0xa0421064, 11, 0, freq)  ##Manual value of frequency (unit: MHz)
                self.mem_ts.wrm(0xa0421010, 3, 3, 0)
                self.mem_ts.wrm(0xa0421024, 21, 18, 0xf)
                self.mem_ts.wrm(0xa0422014, 3, 0, 0x0)
                self.mem_ts.wrm(0xa0422014, 3, 0, 0xe)
                time.sleep(0.2)
                res = self.mem_ts.rdm(0xa04210f8, 10, 0)
                tp_gian_cal_le2m = res & 0x7ff
                logdebug('tp_gian_cal_le2m: {}'.format(tp_gian_cal_le2m))
            else:
                self.mem_ts.wrm(0xa0421064, 12, 12, 1)  ##Frequency value manual enable
                self.mem_ts.wrm(0xa0421064, 11, 0, freq)  ##Manual value of frequency (unit: MHz)
                self.mem_ts.wrm(0xa0421010, 3, 3, 0)
                self.mem_ts.wrm(0xa0421024, 21, 18, 0x0)
                self.mem_ts.wrm(0xa0422014, 3, 0, 0x0)
                self.mem_ts.wrm(0xa0422014, 3, 0, 0xe)
                time.sleep(0.2)
                res = self.mem_ts.rdm(0xa04210f8, 10, 0)
                tp_gian_cal_le1m = res & 0x7ff
                logdebug('tp_gian_cal_le1m: {}'.format(tp_gian_cal_le1m))
            self.mem_ts.wrm(0xa0422014, 3, 0, 0x0)
            self.mem_ts.wrm(0xa0421010, 3, 3, 1)
            self.mem_ts.wrm(0xa0421024, 21, 18, 0x0)
            self.mem_ts.wrm(0xa0421064, 12, 12, 0)  ##Frequency value auto
        else:
            if mode !=0:
                self.jlink.wrm(0xa0421064, 12, 12, 1)  ##Frequency value manual enable
                self.jlink.wrm(0xa0421064, 11, 0, freq)  ##Manual value of frequency (unit: MHz)
                self.jlink.wrm(0xa0421010, 3, 3, 0)
                self.jlink.wrm(0xa0421024,21,18,0xf)
                self.jlink.wrm(0xa0422014,3,0,0x0)
                self.jlink.wrm(0xa0422014, 3, 0, 0xe)
                time.sleep(0.2)
                res = self.jlink.rdm(0xa04210f8,10,0)
                tp_gian_cal_le2m = res & 0x7ff
                logdebug('tp_gian_cal_le2m: {}'.format(tp_gian_cal_le2m))
            else:
                self.jlink.wrm(0xa0421064, 12, 12, 1)  ##Frequency value manual enable
                self.jlink.wrm(0xa0421064, 11, 0, freq)  ##Manual value of frequency (unit: MHz)
                self.jlink.wrm(0xa0421010, 3, 3, 0)
                self.jlink.wrm(0xa0421024,21,18,0x0)
                self.jlink.wrm(0xa0422014,3,0,0x0)
                self.jlink.wrm(0xa0422014, 3, 0, 0xe)
                time.sleep(0.2)
                res = self.jlink.rdm(0xa04210f8,10,0)
                tp_gian_cal_le1m = res & 0x7ff
                logdebug('tp_gian_cal_le1m: {}'.format(tp_gian_cal_le1m))
            self.jlink.wrm(0xa0422014, 3, 0, 0x0)
            self.jlink.wrm(0xa0421010, 3, 3, 1)
            self.jlink.wrm(0xa0421024, 21, 18, 0x0)
            self.jlink.wrm(0xa0421064, 12, 12, 0)  ##Frequency value manual enable
        return res

    def test232_tp_gain_cal_scan(self):
        title = 'mode,channel,tp_gain_cal_value\n'
        fname = rfglobal.get_filename('ts_bt_test/', 'test232_tp_gain_cal_scan_{}'.format(self.board_name))
        fw1 = csvreport(fname, title)
        for mode in [0,1]:
            if mode == 0:
                rate = 'LE1M'
            else:
                rate = 'LE2M'
            for chan in range(0,40):
                tp_gain_cal_value = self.test232_tp_gain_cal(mode=mode,freq=2402+chan*2)
                logdebug('tp_gain_cal_value:    {}'.format(tp_gain_cal_value))
                fw1.write_data([rate,chan,tp_gain_cal_value])

    def test232_dig_gain_set(self, mode=0, gfsk_gain=0x7f,dpsk_gain=0x98):
        '''
        mode:
        0:  BR/LE
        other:  EDR
        '''
        if mode != 0:
            self.jlink.wrm(0xa0420018, 7, 0, gfsk_gain)
        else:
            self.jlink.wrm(0xa0420010, 7, 0, gfsk_gain)
        self.jlink.wrm(0xa042000c, 15, 8, dpsk_gain)

    def test232_cfo_est_test(self,chan_list=[0],signal_freq_oft=400):
        title = 'channel,signal_freq(khz),cfo_est, cfo_est_step\n'
        fname = rfglobal.get_filename('ts_bt_test/','test232_cfo_est_test_{}'.format(self.board_name))
        fw1 = csvreport(fname,title)
        txg = mxg.MXG()
        txg.arb_waveform(rate='LE_1M')
        txg.arb_state(1)
        txg.trriger_para_set(type='CONTinuous', count=1000)
        txg.output_state(1, 1)
        txg.set_power(-40)
        for chan in chan_list:
            self.btapi.cmdstop(0)
            time.sleep(0.5)
            self.btapi.LE_RX(chan=chan,phy=1,mod=0)
            signal_freq_base = 2402+2*chan
            cfo_est_init = 0
            for freq_oft in range(-signal_freq_oft,signal_freq_oft+1):
                while 1:
                    recv_pkt1 = self.jlink.rdm(0xa04008d8, 15, 0)
                    time.sleep(0.1)
                    recv_pkt2 = self.jlink.rdm(0xa04008d8, 15, 0)
                    recv_pkt = recv_pkt2 - recv_pkt1
                    if recv_pkt != 0:
                        freq = signal_freq_base+freq_oft/1000.000
                        txg.set_cfreq(freq)
                        time.sleep(0.5)
                        cfo_est = self.jlink.rdm(0xa04202f0,15,0)
                        if cfo_est>>15 ==1:
                            cfo_est = cfo_est - 65535
                        cfo_est_step = cfo_est - cfo_est_init
                        cfo_est_init = cfo_est
                        loginfo('chan:{}    signal_freq:{}      cfo_est:{}      step:{}'.format(chan, freq, cfo_est, cfo_est_step))
                        fw1.write_data([chan,freq,cfo_est,cfo_est_step])
                        break
                    else:
                        self.btapi.cmdstop(0)
                        time.sleep(0.5)
                        self.btapi.LE_RX(chan=chan, phy=1, mod=0)


    def test232_rxdcoc_code_rd(self):
        res1_i = self.mem_ts.rdm(0xa04210f4,11,6)
        res1_q = self.mem_ts.rdm(0xa04210f4, 17, 12)
        res2_i = self.mem_ts.rdm(0xa04210f4,23,18)
        res2_q = self.mem_ts.rdm(0xa04210f4, 29, 24)
        loginfo("rxdcoc1_i_code: {}     rxdcoc1_q_code: {}     rxdcoc2_i_code: {}     rxdcoc2_q_code: {}     ".format(res1_i,res1_q,res2_i,res2_q))
        return [res1_i,res1_q,res2_i,res2_q]

    def test232_rxdcoc_volt_rd(self):
        mydm = dm.dm(device_name='MY50180049', num_of_machine=0, comm='USB')
        self.mem_ts.wrm(0xa0421024, 17, 17, 1)  ##RF diag  soc enable
        self.mem_ts.wrm(0xa0421024, 3, 3, 0)  ##RF diag  enable
        self.mem_ts.wrm(0xa0421024, 11, 11, 1)

        self.mem_ts.wr(0xa0422014, 0x0)
        time.sleep(0.5)
        self.mem_ts.wr(0xa0422014, 0x3)
        time.sleep(0.5)
        self.mem_ts.wrm(0xa0421024, 2, 0, 1)  ##RF diag code select
        res_qn = eval(mydm.device.ask('MEAS:VOLT?'))
        time.sleep(0.5)
        self.mem_ts.wrm(0xa0421024, 2, 0, 2)  ##RF diag code select
        res_qp = eval(mydm.device.ask('MEAS:VOLT?'))
        time.sleep(0.5)
        self.mem_ts.wrm(0xa0421024, 2, 0, 3)  ##RF diag code select
        res_in = eval(mydm.device.ask('MEAS:VOLT?'))
        time.sleep(0.5)
        self.mem_ts.wrm(0xa0421024, 2, 0, 4)  ##RF diag code select
        res_ip = eval(mydm.device.ask('MEAS:VOLT?'))

        loginfo("rxdcoc_qn: {}     rxdcoc_qp: {}     rxdcoc_in: {}     rxdcoc_ip: {}     ".format(res_qn,res_qp,res_in,res_ip))
        return [res_qn,res_qp,res_in,res_ip]

    def test232_rxdcoc_offset_wr(self,i_offset=0,q_offset=0):
        if self.jlink_en != 0:
            self.jlink.wrm(0xa0421058, 12, 12, 1)
            self.jlink.wrm(0xa0421010, 6, 5, 3)
            self.jlink.wrm(0xa0421058, 18, 13, i_offset)
            self.jlink.wrm(0xa0421058, 24, 19, q_offset)
        else:
            self.mem_ts.wrm(0xa0421058,12,12,1)
            self.mem_ts.wrm(0xa0421010, 6, 5, 3)
            self.mem_ts.wrm(0xa0421058, 18, 13, i_offset)
            self.mem_ts.wrm(0xa0421058, 24, 19, q_offset)
        loginfo("wr rxdcoc offset:  i_offset {}     q_offset {}".format(i_offset,q_offset))

    def test232_rxdcoc_cal_bypass(self,en=1):
        self.mem_ts.wrm(0xa0421058, 12, 12, 0)
        if en == 0:
            self.mem_ts.wrm(0xa0421010, 6, 5, 0)
        else:
            self.mem_ts.wrm(0xa0421010, 6, 5, 3)

    def test232_rxdcoc_scan_lna_gain(self, device_name='MY50180049',idac_code=30,qdac_code=28):
        title = 'idac_code,qdac_code,hgain_code,lgain_code,atten,rxdcoc_qn(v),rxdcoc_qp(v),rxdcoc_in(v),rxdcoc_ip(v)\n'
        fname = rfglobal.get_filename('ts_bt_test/', 'test232_rxdcoc_scan_lna_gain_{}'.format(self.board_name))
        fw1 = csvreport(fname, title)
        mydm = dm.dm(device_name=device_name, num_of_machine=0, comm='USB')
        self.mem_ts.wrm(0xa0421024, 17, 17, 1)  ##RF diag  soc enable
        self.mem_ts.wrm(0xa0421024, 3, 3, 0)  ##RF diag  enable
        self.mem_ts.wrm(0xa0421024, 11, 11, 1)
        self.mem_ts.wrm(0xa04210a0, 0, 0, 1)  ##agc off
        self.mem_ts.wrm(0xa04210a0, 20, 17, 0xc)  ##agc index set 0
        self.mem_ts.wrm(0xa04210cc, 14, 13, 3)  ##cbpf gain1 on
        self.test232_rxdcoc_cal_bypass(1)
        self.test232_rxdcoc_offset_wr(idac_code, qdac_code)
        self.mem_ts.wrm(0xa0422014, 3, 0, 3)

        for hgain_code in [0x3f,0x20,0x10,0x8,0x4,0x2]:
            self.mem_ts.wrm(0xa04210b4, 24, 19, hgain_code)
            lgain_code = self.mem_ts.rdm(0xa04210b4, 18, 17)
            atten = self.mem_ts.rdm(0xa04210b4, 16, 15)
            self.mem_ts.wrm(0xa0421024, 2, 0, 1)  ##RF diag code select
            res_qn = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 2)  ##RF diag code select
            res_qp = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 3)  ##RF diag code select
            res_in = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 4)  ##RF diag code select
            res_ip = eval(mydm.device.ask('MEAS:VOLT?'))

            loginfo("rxdcoc_qn: {}     rxdcoc_qp: {}     rxdcoc_in: {}     rxdcoc_ip: {}     ".format(res_qn, res_qp, res_in, res_ip))
            rxdcoc_qn = res_qn
            rxdcoc_qp = res_qp
            rxdcoc_in = res_in
            rxdcoc_ip = res_ip

            fw1.write_data([idac_code, qdac_code, hgain_code, lgain_code, atten, rxdcoc_qn, rxdcoc_qp, rxdcoc_in, rxdcoc_ip ])
        for lgain_code in range(3,-1,-1):
            self.mem_ts.wrm(0xa04210b4, 25, 25, 0)
            self.mem_ts.wrm(0xa04210b4, 18, 17, lgain_code)
            hgain_code = self.mem_ts.rdm(0xa04210b4, 24, 19)
            atten = self.mem_ts.rdm(0xa04210b4, 16, 15)

            self.mem_ts.wrm(0xa0421024, 2, 0, 1)  ##RF diag code select
            res_qn = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 2)  ##RF diag code select
            res_qp = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 3)  ##RF diag code select
            res_in = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 4)  ##RF diag code select
            res_ip = eval(mydm.device.ask('MEAS:VOLT?'))

            loginfo("rxdcoc_qn: {}     rxdcoc_qp: {}     rxdcoc_in: {}     rxdcoc_ip: {}     ".format(res_qn, res_qp, res_in, res_ip))
            rxdcoc_qn = res_qn
            rxdcoc_qp = res_qp
            rxdcoc_in = res_in
            rxdcoc_ip = res_ip

            fw1.write_data([idac_code, qdac_code, hgain_code, lgain_code, atten, rxdcoc_qn, rxdcoc_qp, rxdcoc_in, rxdcoc_ip ])

        for atten in range(4):
            self.mem_ts.wrm(0xa04210b4, 16, 15, atten)
            hgain_code = self.mem_ts.rdm(0xa04210b4, 24, 19)
            lgain_code = self.mem_ts.rdm(0xa04210b4, 18, 17)

            # res = self.test232_rxdcoc_volt_rd()
            self.mem_ts.wrm(0xa0421024, 2, 0, 1)  ##RF diag code select
            res_qn = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 2)  ##RF diag code select
            res_qp = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 3)  ##RF diag code select
            res_in = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 4)  ##RF diag code select
            res_ip = eval(mydm.device.ask('MEAS:VOLT?'))

            loginfo("rxdcoc_qn: {}     rxdcoc_qp: {}     rxdcoc_in: {}     rxdcoc_ip: {}     ".format(res_qn, res_qp, res_in, res_ip))
            rxdcoc_qn = res_qn
            rxdcoc_qp = res_qp
            rxdcoc_in = res_in
            rxdcoc_ip = res_ip

            fw1.write_data([idac_code, qdac_code, hgain_code, lgain_code, atten, rxdcoc_qn, rxdcoc_qp, rxdcoc_in, rxdcoc_ip ])
        self.mem_ts.wrm(0xa04210b4, 16, 15, 0)
        self.mem_ts.wrm(0xa04210b4, 25, 17, 0x1ff)
    def test232_rxdcoc_cal_diff(self, device_name='MY50180049',board_name=''):
        mydm = dm.dm(device_name=device_name, num_of_machine=0, comm='USB')
        self.mem_ts.wrm(0xa0421024, 17, 17, 1)  ##RF diag  soc enable
        self.mem_ts.wrm(0xa0421024, 3, 3, 0)  ##RF diag  enable
        self.mem_ts.wrm(0xa0421024, 11, 11, 1)
        self.mem_ts.wrm(0xa04210a0, 0, 0, 1)  ##agc off
        self.mem_ts.wrm(0xa04210a0, 20, 17, 0)  ##agc index set 0
        self.mem_ts.wrm(0xa04210cc, 14, 13, 3)  ##cbpf gain1&2 on

        title = 'cbpf2_en,cbpf1_idac_code,cbpf1_qdac_code,cbpf2_idac_code,cbpf2_qdac_code,rxdcoc_qn(v),rxdcoc_qp(v),rxdcoc_in(v),rxdcoc_ip(v)\n'
        fname = rfglobal.get_filename('ts_bt_test/', 'test232_rxdcoc_cal_diff_{}'.format(board_name))
        fw1 = csvreport(fname, title)
        for cbpf_en in [0,1]:
            self.mem_ts.wrm(0xa0422014, 3, 0, 0)
            self.mem_ts.wrm(0xa0421020, 20, 20, cbpf_en)
            self.mem_ts.wrm(0xa0421010, 6, 5, 3)
            self.mem_ts.wrm(0xa0421058, 12, 12, 1)  ##cbpf1 rxdcoc code manual
            self.mem_ts.wrm(0xa042105c, 12, 12, 1)  ##cbpf2 rxdcoc code manual
            self.mem_ts.wrm(0xa042105c, 18, 13, 32)
            self.mem_ts.wrm(0xa042105c, 24, 19, 32)
            self.mem_ts.wrm(0xa0421058, 18, 13, 32)
            self.mem_ts.wrm(0xa0421058, 24, 19, 32)
            self.mem_ts.wrm(0xa0422014, 3, 0, 3)

            self.mem_ts.wrm(0xa0421024, 2, 0, 1)  ##RF diag code select
            res_qn = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 2)  ##RF diag code select
            res_qp = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 3)  ##RF diag code select
            res_in = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 4)  ##RF diag code select
            res_ip = eval(mydm.device.ask('MEAS:VOLT?'))

            loginfo("rxdcoc_qn: {}     rxdcoc_qp: {}     rxdcoc_in: {}     rxdcoc_ip: {}     ".format(res_qn, res_qp, res_in, res_ip))
            rxdcoc_qn = res_qn
            rxdcoc_qp = res_qp
            rxdcoc_in = res_in
            rxdcoc_ip = res_ip

            fw1.write_data([cbpf_en, 32, 32, 32, 32, rxdcoc_qn, rxdcoc_qp, rxdcoc_in, rxdcoc_ip])

            self.mem_ts.wrm(0xa0421058, 12, 12, 0)  ##cbpf1 rxdcoc code auto
            self.mem_ts.wrm(0xa042105c, 12, 12, 0)  ##cbpf2 rxdcoc code auto
            time.sleep(2)
            self.mem_ts.wrm(0xa0421010, 6, 5, 0)  ##rxdcoc cal bypass disable
            self.mem_ts.wrm(0xa0422014, 3, 0, 0)
            self.mem_ts.wrm(0xa0422014, 3, 0, 3)
            time.sleep(0.5)
            ro_rxdcoc1_q_code = self.mem_ts.rdm(0xa04210f4,17,12)
            ro_rxdcoc1_i_code = self.mem_ts.rdm(0xa04210f4,11,6)
            ro_rxdcoc2_q_code = self.mem_ts.rdm(0xa04210f4,23,18)
            ro_rxdcoc2_i_code = self.mem_ts.rdm(0xa04210f4,29,24)
            self.mem_ts.wrm(0xa0421024, 2, 0, 1)  ##RF diag code select
            res_qn = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 2)  ##RF diag code select
            res_qp = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 3)  ##RF diag code select
            res_in = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 4)  ##RF diag code select
            res_ip = eval(mydm.device.ask('MEAS:VOLT?'))

            loginfo("rxdcoc_qn: {}     rxdcoc_qp: {}     rxdcoc_in: {}     rxdcoc_ip: {}     ".format(res_qn, res_qp, res_in, res_ip))
            rxdcoc_qn = res_qn
            rxdcoc_qp = res_qp
            rxdcoc_in = res_in
            rxdcoc_ip = res_ip
            fw1.write_data([cbpf_en, ro_rxdcoc1_i_code, ro_rxdcoc1_q_code, ro_rxdcoc2_i_code, ro_rxdcoc2_q_code, rxdcoc_qn, rxdcoc_qp, rxdcoc_in, rxdcoc_ip])

    def test232_rxdcoc_cal_diff_loop(self, device_name='MY50180049',board_name='',loop=100):
        mydm = dm.dm(device_name=device_name, num_of_machine=0, comm='USB')
        self.mem_ts.wrm(0xa0421024, 17, 17, 1)  ##RF diag  soc enable
        self.mem_ts.wrm(0xa0421024, 3, 3, 0)  ##RF diag  enable
        self.mem_ts.wrm(0xa0421024, 11, 11, 1)
        self.mem_ts.wrm(0xa04210a0, 0, 0, 1)  ##agc off
        self.mem_ts.wrm(0xa04210a0, 20, 17, 0)  ##agc index set 0
        self.mem_ts.wrm(0xa04210cc, 14, 13, 3)  ##cbpf gain1&2 on
        self.mem_ts.wrm(0xa0421020, 20, 20, 1)

        title = 'cbpf2_en,cbpf1_idac_code,cbpf1_qdac_code,cbpf2_idac_code,cbpf2_qdac_code,rxdcoc_qn(v),rxdcoc_qp(v),rxdcoc_in(v),rxdcoc_ip(v)\n'
        fname = rfglobal.get_filename('ts_bt_test/', 'test232_rxdcoc_cal_diff_{}'.format(board_name))
        fw1 = csvreport(fname, title)
        for i in range(loop):
            # self.mem_ts.wrm(0xa0421010, 6, 5, 3)
            # self.mem_ts.wrm(0xa0421058, 12, 12, 1)  ##cbpf1 rxdcoc code manual
            # self.mem_ts.wrm(0xa042105c, 12, 12, 1)  ##cbpf2 rxdcoc code manual
            # self.mem_ts.wrm(0xa042105c, 18, 13, 32)
            # self.mem_ts.wrm(0xa042105c, 24, 19, 32)
            # self.mem_ts.wrm(0xa0421058, 18, 13, 32)
            # self.mem_ts.wrm(0xa0421058, 24, 19, 32)
            self.mem_ts.wrm(0xa0422014, 3, 0, 0)
            self.mem_ts.wrm(0xa0421058, 12, 12, 0)  ##cbpf1 rxdcoc code auto
            self.mem_ts.wrm(0xa042105c, 12, 12, 0)  ##cbpf2 rxdcoc code auto
            time.sleep(0.1)
            # self.mem_ts.wrm(0xa0421010, 6, 5, 0)  ##rxdcoc cal bypass disable
            self.mem_ts.wrm(0xa0422014, 3, 0, 3)
            time.sleep(0.5)
            ro_rxdcoc1_q_code = self.mem_ts.rdm(0xa04210f4,17,12)
            ro_rxdcoc1_i_code = self.mem_ts.rdm(0xa04210f4,11,6)
            ro_rxdcoc2_q_code = self.mem_ts.rdm(0xa04210f4,23,18)
            ro_rxdcoc2_i_code = self.mem_ts.rdm(0xa04210f4,29,24)
            self.mem_ts.wrm(0xa0421024, 2, 0, 1)  ##RF diag code select
            res_qn = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 2)  ##RF diag code select
            res_qp = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 3)  ##RF diag code select
            res_in = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 4)  ##RF diag code select
            res_ip = eval(mydm.device.ask('MEAS:VOLT?'))

            loginfo("rxdcoc_qn: {}     rxdcoc_qp: {}     rxdcoc_in: {}     rxdcoc_ip: {}     ".format(res_qn, res_qp, res_in, res_ip))
            rxdcoc_qn = res_qn
            rxdcoc_qp = res_qp
            rxdcoc_in = res_in
            rxdcoc_ip = res_ip
            fw1.write_data([1, ro_rxdcoc1_i_code, ro_rxdcoc1_q_code, ro_rxdcoc2_i_code, ro_rxdcoc2_q_code, rxdcoc_qn, rxdcoc_qp, rxdcoc_in, rxdcoc_ip])


    def test232_rxdcoc_scan(self, device_name='MY50180049', cbpf1_en=1, cbpf2_en=1):

        mydm = dm.dm(device_name=device_name, num_of_machine=0, comm='USB')
        self.mem_ts.wrm(0xa0421024, 17, 17, 1)  ##RF diag  soc enable
        self.mem_ts.wrm(0xa0421024, 3, 3, 0)  ##RF diag  enable
        self.mem_ts.wrm(0xa0421024, 11, 11, 1)
        self.mem_ts.wrm(0xa04210a0, 0, 0, 1)  ##agc off
        self.mem_ts.wrm(0xa04210a0, 20, 17, 0)  ##agc index set 0
        self.mem_ts.wrm(0xa04210cc, 14, 13, 2)  ##cbpf gain1 on
        self.test232_rxdcoc_cal_bypass(1)
        self.test232_rxdcoc_offset_wr(32, 32)
        self.mem_ts.wrm(0xa0422014, 3, 0, 3)

        if cbpf1_en!=0:
            title = 'idac_code,qdac_code,rxdcoc_qn(v),rxdcoc_qp(v),rxdcoc_in(v),rxdcoc_ip(v)\n'
            fname = rfglobal.get_filename('ts_bt_test/', 'test232_rxdcoc_scan_cbpf1_{}'.format(self.board_name))
            fw1 = csvreport(fname, title)
            self.mem_ts.wrm(0xa0421020, 20, 20, 0)  ##cbpf2 disable
            for idac_code in range(64):
                for qdac_code in range(64):
                    self.test232_rxdcoc_offset_wr(idac_code, qdac_code)
                    # res = self.test232_rxdcoc_volt_rd()
                    self.mem_ts.wrm(0xa0421024, 2, 0, 1)  ##RF diag code select
                    res_qn = eval(mydm.device.ask('MEAS:VOLT?'))

                    self.mem_ts.wrm(0xa0421024, 2, 0, 2)  ##RF diag code select
                    res_qp = eval(mydm.device.ask('MEAS:VOLT?'))

                    self.mem_ts.wrm(0xa0421024, 2, 0, 3)  ##RF diag code select
                    res_in = eval(mydm.device.ask('MEAS:VOLT?'))

                    self.mem_ts.wrm(0xa0421024, 2, 0, 4)  ##RF diag code select
                    res_ip = eval(mydm.device.ask('MEAS:VOLT?'))

                    loginfo("rxdcoc_qn: {}     rxdcoc_qp: {}     rxdcoc_in: {}     rxdcoc_ip: {}     ".format(res_qn, res_qp, res_in, res_ip))
                    rxdcoc_qn = res_qn
                    rxdcoc_qp = res_qp
                    rxdcoc_in = res_in
                    rxdcoc_ip = res_ip

                    fw1.write_data([idac_code, qdac_code, rxdcoc_qn, rxdcoc_qp, rxdcoc_in, rxdcoc_ip, ])
        if cbpf2_en!=0:
            self.mem_ts.wrm(0xa0421020, 20, 20, 0)  ##cbpf2 disable
            self.mem_ts.wrm(0xa0422014, 3, 0, 0)
            self.test232_rxdcoc_cal_bypass(0)
            self.mem_ts.wrm(0xa0421058, 12, 12, 0)
            self.mem_ts.wrm(0xa04210cc, 14, 13, 3)  ##cbpf gain1&2 on
            self.mem_ts.wrm(0xa0422014, 3, 0, 3)    ##rxon
            time.sleep(0.5)
            ro_rxdcoc1_q_code = self.mem_ts.rdm(0xa04210f4,17,12)
            ro_rxdcoc1_i_code = self.mem_ts.rdm(0xa04210f4,11,6)
            ro_rxdcoc2_q_code = self.mem_ts.rdm(0xa04210f4,23,18)
            ro_rxdcoc2_i_code = self.mem_ts.rdm(0xa04210f4,29,24)
            self.test232_rxdcoc_offset_wr(ro_rxdcoc1_i_code, ro_rxdcoc1_q_code)
            self.mem_ts.wrm(0xa0422014, 3, 0, 0)
            self.mem_ts.wrm(0xa0421020, 20, 20, 1)  ##cbpf2 enable
            self.mem_ts.wrm(0xa042105c, 12, 12, 1)  ##cbpf2 rxdcoc code manual
            self.mem_ts.wrm(0xa0422014, 3, 0, 3)  ##rxon

            title = 'ro_rxdcoc1_q_code={},ro_rxdcoc1_i_code={},ro_rxdcoc2_q_code={},ro_rxdcoc2_i_code={}\n'.format(ro_rxdcoc1_q_code,ro_rxdcoc1_i_code,ro_rxdcoc2_q_code,ro_rxdcoc2_i_code)
            title = title + 'idac_code,qdac_code,rxdcoc_qn(v),rxdcoc_qp(v),rxdcoc_in(v),rxdcoc_ip(v)\n'
            fname = rfglobal.get_filename('ts_bt_test/', 'test232_rxdcoc_scan_cbpf2')
            fw1 = csvreport(fname, title)
            for idac_code in range(64):
                self.mem_ts.wrm(0xa042105c, 18, 13, idac_code)
                for qdac_code in range(64):
                    self.mem_ts.wrm(0xa042105c, 24, 19, qdac_code)
                    # res = self.test232_rxdcoc_volt_rd()
                    self.mem_ts.wrm(0xa0421024, 2, 0, 1)  ##RF diag code select
                    res_qn = eval(mydm.device.ask('MEAS:VOLT?'))

                    self.mem_ts.wrm(0xa0421024, 2, 0, 2)  ##RF diag code select
                    res_qp = eval(mydm.device.ask('MEAS:VOLT?'))

                    self.mem_ts.wrm(0xa0421024, 2, 0, 3)  ##RF diag code select
                    res_in = eval(mydm.device.ask('MEAS:VOLT?'))

                    self.mem_ts.wrm(0xa0421024, 2, 0, 4)  ##RF diag code select
                    res_ip = eval(mydm.device.ask('MEAS:VOLT?'))

                    loginfo("rxdcoc_qn: {}     rxdcoc_qp: {}     rxdcoc_in: {}     rxdcoc_ip: {}     ".format(res_qn, res_qp, res_in, res_ip))
                    rxdcoc_qn = res_qn
                    rxdcoc_qp = res_qp
                    rxdcoc_in = res_in
                    rxdcoc_ip = res_ip

                    fw1.write_data([idac_code, qdac_code, rxdcoc_qn, rxdcoc_qp, rxdcoc_in, rxdcoc_ip, ])
            self.mem_ts.wrm(0xa042105c, 18, 13, ro_rxdcoc2_i_code)
            self.mem_ts.wrm(0xa042105c, 24, 19, ro_rxdcoc2_q_code)
            self.mem_ts.wrm(0xa0421020, 20, 20, 0)  ##cbpf2 disable
            self.mem_ts.wrm(0xa042105c, 12, 12, 0)  ##cbpf2 rxdcoc code manual


    def test232_rxdcoc_manual_cal(self, device_name='MY50180049'):
        title = 'idac_code,rxdcoc_qn(v),rxdcoc_qp(v),rxdcoc_in(v),rxdcoc_ip(v),cbpf_outi,cbpf_outq\n'
        fname = rfglobal.get_filename('ts_bt_test/', 'test232_rxdcoc_manual_cal_{}'.format(self.board_name))
        fw1 = csvreport(fname, title)
        mydm = dm.dm(device_name=device_name, num_of_machine=0, comm='USB')
        self.mem_ts.wrm(0xa0421024, 17, 17, 1)  ##RF diag  soc enable
        self.mem_ts.wrm(0xa0421024, 3, 3, 0)  ##RF diag  enable
        self.mem_ts.wrm(0xa0421024, 11, 11, 1)
        self.mem_ts.wrm(0xa04210a0, 0, 0, 1)    ##agc off
        self.mem_ts.wrm(0xa04210a0, 20, 17, 0)  ##agc index set 0
        self.mem_ts.wrm(0xa04210cc, 14, 13, 2)  ##cbpf gain1 on
        self.test232_rxdcoc_cal_bypass(1)
        self.test232_rxdcoc_offset_wr(32,32)
        self.mem_ts.wrm(0xa0422014, 3, 0, 3)

        for idac_code in range(64):
            self.test232_rxdcoc_offset_wr(idac_code, 32)
            # res = self.test232_rxdcoc_volt_rd()
            self.mem_ts.wrm(0xa0421024, 2, 0, 1)  ##RF diag code select
            res_qn = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 2)  ##RF diag code select
            res_qp = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 3)  ##RF diag code select
            res_in = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 4)  ##RF diag code select
            res_ip = eval(mydm.device.ask('MEAS:VOLT?'))

            loginfo("rxdcoc_qn: {}     rxdcoc_qp: {}     rxdcoc_in: {}     rxdcoc_ip: {}     ".format(res_qn, res_qp, res_in, res_ip))
            rxdcoc_qn = res_qn
            rxdcoc_qp = res_qp
            rxdcoc_in = res_in
            rxdcoc_ip = res_ip
            daci_outi = rxdcoc_ip-rxdcoc_in
            daci_outq = rxdcoc_qp-rxdcoc_qn
            fw1.write_data([idac_code,rxdcoc_qn,rxdcoc_qp,rxdcoc_in,rxdcoc_ip,daci_outi,daci_outq])
        fw1.write_string('qdac_code,rxdcoc_qn(v),rxdcoc_qp(v),rxdcoc_in(v),rxdcoc_ip(v),cbpf_outi,cbpf_outq\n')
        for qdac_code in range(64):
            self.test232_rxdcoc_offset_wr(32, qdac_code)
            # res = self.test232_rxdcoc_volt_rd()
            self.mem_ts.wrm(0xa0421024, 2, 0, 1)  ##RF diag code select
            res_qn = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 2)  ##RF diag code select
            res_qp = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 3)  ##RF diag code select
            res_in = eval(mydm.device.ask('MEAS:VOLT?'))

            self.mem_ts.wrm(0xa0421024, 2, 0, 4)  ##RF diag code select
            res_ip = eval(mydm.device.ask('MEAS:VOLT?'))

            loginfo("rxdcoc_qn: {}     rxdcoc_qp: {}     rxdcoc_in: {}     rxdcoc_ip: {}     ".format(res_qn, res_qp, res_in, res_ip))
            rxdcoc_qn = res_qn
            rxdcoc_qp = res_qp
            rxdcoc_in = res_in
            rxdcoc_ip = res_ip
            dacq_outi = rxdcoc_ip-rxdcoc_in
            dacq_outq = rxdcoc_qp-rxdcoc_qn
            fw1.write_data([qdac_code,rxdcoc_qn,rxdcoc_qp,rxdcoc_in,rxdcoc_ip,dacq_outi,dacq_outq])


    def test232_txdcoc_rd(self):
        res1_i = self.mem_ts.rdm(0xa042103c,25,20)
        res1_q = self.mem_ts.rdm(0xa042103c, 31, 26)
        loginfo("txdcoc_i_code: {}     txdcoc_q_code: {}".format(res1_i,res1_q))
        return [res1_i,res1_q]

    def test232_txdcoc_wr(self, manual_en=1, icode=32, qcode=32):
        self.mem_ts.wrm(0xa0421040, 6, 6, manual_en)
        self.mem_ts.wrm(0xa0421040, 14, 14, manual_en)
        self.mem_ts.wrm(0xa0421040, 5, 0, icode)
        self.mem_ts.wrm(0xa0421040, 13, 8, qcode)
        self.mem_ts.wr(0xa0422014, 0x0)
        time.sleep(1)
        self.mem_ts.wr(0xa0422014, 0xe)
        self.test232_txdcoc_rd()

    def test232_txdcoc_allchannel_rd(self):
        title = 'channel,txdcoc_i_code,txdcoc_q_code\n'
        fname = rfglobal.get_filename('ts_bt_test/', 'test232_txdcoc_allchannel_rd_{}'.format(self.board_name))
        fw1 = csvreport(fname, title)
        self.mem_ts.wrm(0xa0421064, 12, 12, 1)  ##Frequency value manual enable
        for channel in range(0,79):
            self.mem_ts.wrm(0xa0421064, 11, 0, 2402 + channel)
            self.mem_ts.wr(0xa0422014, 0x0)
            time.sleep(1)
            self.mem_ts.wr(0xa0422014, 0xe)
            time.sleep(1)
            res = self.test232_txdcoc_rd()
            fw1.write_data([channel, res[0],res[1]])

    def test232_tx_gain_scan(self, csv_save=True):
        if csv_save:
            title = 'tx pa tpr value,tx carrier power(dbm)\n'
            fname = rfglobal.get_filename('ts_bt_test/','test232_tx_gain_scan_{}'.format(self.board_name))
            fw1 = csvreport(fname,title)

        spa = Agilent()
        spa.set_param(2402,0.5,3,3)
        spa.set_reflvl(15)
        self.btapi.tx_carrier(freq=2402, dac_gain=0x9f, pa_gain=63)
        self.mem_ts.wrm(0xa0421004, 22, 22, 1)  ##Manual control of Tx PA target power enable

        for gain in range(63,0,-1):
            self.mem_ts.wrm(0xa042100c, 31, 26, gain)  ##Manual control of Tx PA target power value
            self.mem_ts.wr(0xa0422014, 0x0)
            time.sleep(0.2)
            self.mem_ts.wr(0xa0422014, 0xe)
            res = spa.pk_search_avg_timesleep(timesleep=1.5)
            pwr = res[0][1]
            loginfo('gain:  {}  tx carrier power:   {}'.format(gain,pwr))
            fw1.write_data([gain,pwr])
        self.btapi.tx_carrier_stop()


    def rxgain_scan_tx232(self, channel_list=[0], csv_save=True, rf_cableloss=1, if_cableloss=1):
        self.btapi.cmdstop(0)
        self.btapi.cmdstop(1)
        if csv_save:
            title = 'channel,lna_itrim,cbpf_bias_trim,LNA_HGAIN,LNA_LGAIN,LNA_ATTEN,FLT_GAIN1,FLT_GAIN2,RF_pwr(dbm),IF_pwr(dbm),dalta\n'
            fname = rfglobal.get_filename('ts_bt_test/','rxgain_scan_tx232_{}'.format(self.board_name))
            fw1 = csvreport(fname,title)

        spa = Agilent()
        spa.set_param(1,0.5,1,3)
        txg = mxg.MXG()
        # txg.arb_waveform(rate='LE_1M')
        txg.trriger_para_set(type='CONTinuous', count=1000)
        txg.output_state(1, 0)
        # txg.arb_state(1)
        self.mem_ts.wrm(0xa04210A0, 0, 0, 0x1)  ##AGC OFF
        self.mem_ts.wrm(0xa0421020, 20, 20, 0x0)
        self.mem_ts.wrm(0xa0120038, 11, 8, 0)
        self.mem_ts.wrm(0xa012005c, 31, 16, 0)  ##
        self.mem_ts.wrm(0xa0200240, 17, 6, 0)  ##PA5/4/3/2 高阻
        self.mem_ts.wrm(0xa01200a0, 29, 29, 1)  ##bbpll_bg_pup_ibg_tbuf
        self.mem_ts.wrm(0xa01200c8, 4, 4, 1)  ##test buf en
        ##self.mem_ts.wrm(0xa01200c8,3,2,2)	##tx IF test buf
        self.mem_ts.wrm(0xa01200c8, 3, 2, 1)  ##rx IF test buf
        lna_itrim_init = self.mem_ts.rdm(0xa0421034,8,7)
        cbpf_bias_trim_init = self.mem_ts.rdm(0xa0421034,18,16)
        for channel in channel_list:
            for lna_itrim in range(4):
                self.mem_ts.wrm(0xa0421034, 8, 7,lna_itrim)
                for cbpf_bias_trim in range(8):
                    self.mem_ts.wrm(0xa0421034, 18, 16, cbpf_bias_trim)
                    self.mem_ts.wrm(0xa04210b4, 25, 17, 0x1ff)
                    self.mem_ts.wrm(0xa04210b4, 16, 13, 0x3)
                    txg.para_set(freq=2402 + channel , power=-67)
                    self.mem_ts.wr(0xa0422014, 0x0)
                    self.mem_ts.wrm(0xa0421064, 12, 12, 1)  ##Frequency value manual enable
                    self.mem_ts.wrm(0xa0421064, 11, 0, 2402+channel)
                    self.mem_ts.wr(0xa0422014, 0x3)  ##manual rx on

                    for lna_hgain in (0x3f,0x20,0x10,0x8,0x4,0x2):

                        self.mem_ts.wrm(0xa04210b4, 24, 19, lna_hgain)
                        lna_lgain = 3
                        lna_atten = 0
                        flt_gain1 = 1
                        flt_gain2 = 1

                        pwr = -40
                        while 1:
                            txg.set_power(pwr+rf_cableloss)
                            time.sleep(0.2)
                            res = spa.pk_search_avg_timesleep(timesleep=1.5)
                            if_pwr = res[0][1]+if_cableloss
                            loginfo(if_pwr)
                            if if_pwr > 0:
                                pwr = pwr - 5
                            else:
                                break
                        loginfo('RF pwer:{}    IF power:{}'.format(pwr, if_pwr))
                        fw1.write_data([channel, lna_itrim, cbpf_bias_trim, lna_hgain, lna_lgain, lna_atten,  flt_gain1, flt_gain2, pwr, if_pwr, if_pwr - pwr])

                    self.mem_ts.wrm(0xa04210b4, 25, 25, 0)
                    for lna_lgain in (3,2):
                        self.mem_ts.wrm(0xa04210b4, 18, 17, lna_lgain)
                        lna_hgain = 2
                        lna_atten = 0
                        flt_gain1 = 1
                        flt_gain2 = 1

                        pwr = -40
                        while 1:
                            txg.set_power(pwr+rf_cableloss)
                            time.sleep(0.2)
                            res = spa.pk_search_avg_timesleep(timesleep=1.5)
                            if_pwr = res[0][1]+if_cableloss
                            print(if_pwr)
                            if if_pwr > 0:
                                pwr = pwr - 5
                            else:
                                break
                        loginfo('RF pwer:{}    IF power:{}'.format(pwr, if_pwr))
                        fw1.write_data([channel, lna_itrim, cbpf_bias_trim, lna_hgain, lna_lgain, lna_atten,  flt_gain1, flt_gain2, pwr, if_pwr, if_pwr - pwr])
                    self.mem_ts.wrm(0xa04210b4, 18, 17, 2)
                    for lna_atten in (1,2,3):
                        self.mem_ts.wrm(0xa04210b4, 16, 15, lna_atten)
                        lna_hgain = 2
                        lna_lgain = 2
                        flt_gain1 = 1
                        flt_gain2 = 1

                        pwr = -40
                        while 1:
                            txg.set_power(pwr+rf_cableloss)
                            time.sleep(0.2)
                            res = spa.pk_search_avg_timesleep(timesleep=1.5)
                            if_pwr = res[0][1]+if_cableloss
                            print(if_pwr)
                            if if_pwr > 0:
                                pwr = pwr - 5
                            else:
                                break
                        loginfo('RF pwer:{}    IF power:{}'.format(pwr, if_pwr))
                        fw1.write_data([channel, lna_itrim, cbpf_bias_trim, lna_hgain, lna_lgain, lna_atten,  flt_gain1, flt_gain2, pwr, if_pwr, if_pwr - pwr])
                    self.mem_ts.wrm(0xa04210b4, 16, 15, 3)
                    self.mem_ts.wrm(0xa04210b4, 14, 14, 0)
                    for flt_gain2 in (1,0) :
                        self.mem_ts.wrm(0xa04210b4, 13, 13, flt_gain2)
                        lna_hgain = 0
                        lna_lgain = 0
                        flt_gain1 = 0
                        lna_atten = 3

                        pwr = -40
                        while 1:
                            txg.set_power(pwr+rf_cableloss)
                            time.sleep(0.2)
                            res = spa.pk_search_avg_timesleep(timesleep=1.5)
                            if_pwr = res[0][1]+if_cableloss
                            print(if_pwr)
                            if if_pwr > 0:
                                pwr = pwr - 5
                            else:
                                break
                        loginfo('RF pwer:{}    IF power:{}'.format(pwr, if_pwr))
                        fw1.write_data([channel, lna_itrim, cbpf_bias_trim, lna_hgain, lna_lgain, lna_atten,  flt_gain1, flt_gain2, pwr, if_pwr, if_pwr - pwr])
        txg.output_state(0, 0)
        self.mem_ts.wrm(0xa0421034, 8, 7, lna_itrim_init)
        self.mem_ts.wrm(0xa0421034, 18, 16, cbpf_bias_trim_init)
        self.mem_ts.wrm(0xa04210A0, 0, 0, 0x0)
        self.mem_ts.wrm(0xa04210b4, 25, 13, 0x1ff3)
        self.mem_ts.wrm(0xa0421020, 20, 20, 0x0)


    def rxgain_scan_tx231(self, channel_list=[0], csv_save=True):
        self.btapi.cmdstop(0)
        self.btapi.cmdstop(1)
        if csv_save:
            title = 'channel,LNA_GAIN,LNA_ATTEN_R,LNA_ATTEN_C,FLT_GAIN1,FLT_GAIN2,RF_pwr(dbm),IF_pwr(dbm),dalta\n'
            fname = rfglobal.get_filename('ts_bt_test/','rxgain_scan_tx231_{}'.format(self.board_name))
            fw1 = csvreport(fname,title)

        spa = Agilent()
        spa.set_param(1,1,10,30)
        txg = mxg.MXG()
        # txg.arb_waveform(rate='LE_1M')
        txg.trriger_para_set(type='CONTinuous', count=1000)
        txg.output_state(1, 0)
        # txg.arb_state(1)
        for channel in channel_list:
            txg.para_set(freq=2402 + channel * 2, power=-67)
            self.mem_ts.wrm(0xa042107c,30,26,0x1f)  ##rx gain manual
            time.sleep(.5)
            self.mem_ts.wrm(0xa0421064, 0, 0, 0x1)  ##AGC OFF
            time.sleep(.5)
            self.mem_ts.wrm(0xa04210a8, 6, 6, 1)  ##Tx DAC output test enable signal
            time.sleep(.5)
            self.mem_ts.wrm(0xa04210a8, 0, 0, 1)  ##Test buffer enable signal
            time.sleep(.5)
            self.btapi.LE_RX(chan=channel,phy=1,mod=0,countMode=0)
            for lna_gain in range(3,-1,-1):
                if lna_gain == 0:
                    for lna_atten_c in range(0,4):
                        for lna_atten_r in range(0, 4):
                            for flt_gain1 in range(3, -1, -1):
                                for flt_gain2 in range(3, -1, -1):
                                    self.mem_ts.wrm(0xa042107c, 21, 20, lna_atten_r)
                                    time.sleep(.5)
                                    self.mem_ts.wrm(0xa042107c, 17, 16, lna_gain)
                                    time.sleep(.5)
                                    self.mem_ts.wrm(0xa042107c, 19, 18, lna_atten_c)
                                    time.sleep(.5)
                                    self.mem_ts.wrm(0xa042107c, 23, 22, flt_gain1)
                                    time.sleep(.5)
                                    self.mem_ts.wrm(0xa042107c, 25, 24, flt_gain2)
                                    pwr = -40
                                    while 1:
                                        txg.set_power(pwr)
                                        time.sleep(0.2)
                                        res = spa.pk_search_avg_timesleep(timesleep=1.5)
                                        if_pwr = res[0][1]
                                        print(if_pwr)
                                        if if_pwr > 0:
                                            pwr = pwr - 5
                                        else:
                                            break
                                    loginfo('RF pwer:{}    IF power:{}'.format(pwr, if_pwr))
                                    fw1.write_data([channel, lna_gain, lna_atten_r, lna_atten_c, flt_gain1, flt_gain2, pwr, if_pwr, if_pwr - pwr])
                else:
                    lna_atten_r = 0
                    lna_atten_c = 0
                    for flt_gain1 in range(3,-1,-1):
                        for flt_gain2 in range(3, -1, -1):
                            self.mem_ts.wrm(0xa042107c, 21, 20, lna_atten_r)
                            time.sleep(.5)
                            self.mem_ts.wrm(0xa042107c, 17, 16, lna_gain)
                            time.sleep(.5)
                            self.mem_ts.wrm(0xa042107c, 19, 18, lna_atten_c)
                            time.sleep(.5)
                            self.mem_ts.wrm(0xa042107c, 23, 22, flt_gain1)
                            time.sleep(.5)
                            self.mem_ts.wrm(0xa042107c, 25, 24, flt_gain2)
                            pwr = -40

                            while 1:
                                txg.set_power(pwr)
                                time.sleep(0.2)
                                res = spa.pk_search_avg_timesleep(timesleep=1.5)
                                if_pwr = res[0][1]
                                print(if_pwr)
                                if if_pwr > 0:
                                    pwr = pwr - 5
                                else:
                                    break
                            loginfo('RF pwer:{}    IF power:{}'.format(pwr,if_pwr))
                            fw1.write_data([channel,lna_gain,lna_atten_r,lna_atten_c,flt_gain1,flt_gain2,pwr,if_pwr,if_pwr-pwr])
            # for lna_gain in range(0,4):
            #     lna_atten_r = 0
            #     lna_atten_c = 0
            #     flt_gain1 = 3
            #     flt_gain2 = 3
            #     self.mem_ts.wrm(0xa042107c, 21, 20, lna_atten_r)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 17, 16, lna_gain)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 19, 18, lna_atten_c)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 23, 22, flt_gain1)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 25, 24, flt_gain2)
            #     pwr = -40
            #
            #     while 1:
            #         txg.set_power(pwr)
            #         time.sleep(0.2)
            #         res = spa.pk_search_avg_timesleep(timesleep=1.5)
            #         if_pwr = res[0][1]
            #         print(if_pwr)
            #         if if_pwr > 0:
            #             pwr = pwr - 5
            #         else:
            #             break
            #     loginfo('RF pwer:{}    IF power:{}'.format(pwr,if_pwr))
            #     fw1.write_data([channel,lna_gain,lna_atten_r,lna_atten_c,flt_gain1,flt_gain2,pwr,if_pwr,if_pwr-pwr])
            # for lna_atten_c in range(0,4):
            #     lna_gain = 0
            #     lna_atten_r = 0
            #     flt_gain1 = 3
            #     flt_gain2 = 3
            #     self.mem_ts.wrm(0xa042107c, 21, 20, lna_atten_r)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 17, 16, lna_gain)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 19, 18, lna_atten_c)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 23, 22, flt_gain1)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 25, 24, flt_gain2)
            #     pwr = -30
            #
            #     while 1:
            #         txg.set_power(pwr)
            #         time.sleep(0.2)
            #         res = spa.pk_search_avg_timesleep(timesleep=1.5)
            #         if_pwr = res[0][1]
            #         print(if_pwr)
            #         if if_pwr > 0 :
            #             pwr = pwr - 5
            #         else:
            #             break
            #     loginfo('RF pwer:{}    IF power:{}'.format(pwr,if_pwr))
            #     fw1.write_data([channel,lna_gain,lna_atten_r,lna_atten_c,flt_gain1,flt_gain2,pwr,if_pwr,if_pwr-pwr])
            # for lna_atten_r in range(0,4):
            #     lna_gain = 0
            #     lna_atten_c = 3
            #     flt_gain1 = 3
            #     flt_gain2 = 3
            #
            #     self.mem_ts.wrm(0xa042107c, 21, 20, lna_atten_r)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 17, 16, lna_gain)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 19, 18, lna_atten_c)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 23, 22, flt_gain1)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 25, 24, flt_gain2)
            #     pwr = -20
            #
            #     while 1:
            #         txg.set_power(pwr)
            #         time.sleep(0.2)
            #         res = spa.pk_search_avg_timesleep(timesleep=1.5)
            #         if_pwr = res[0][1]
            #         print(if_pwr)
            #         if if_pwr > 0 :
            #             pwr = pwr - 5
            #         else:
            #             break
            #     loginfo('RF pwer:{}    IF power:{}'.format(pwr,if_pwr))
            #     fw1.write_data([channel,lna_gain,lna_atten_r,lna_atten_c,flt_gain1,flt_gain2,pwr,if_pwr,if_pwr-pwr])
            # for flt_gain1 in range(0,4):
            #     lna_gain = 0
            #     lna_atten_r = 3
            #     lna_atten_c = 3
            #     flt_gain2 = 3
            #
            #     self.mem_ts.wrm(0xa042107c, 21, 20, lna_atten_r)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 17, 16, lna_gain)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 19, 18, lna_atten_c)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 23, 22, flt_gain1)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 25, 24, flt_gain2)
            #     pwr = -10
            #
            #     while 1:
            #         txg.set_power(pwr)
            #         time.sleep(0.2)
            #         res = spa.pk_search_avg_timesleep(timesleep=2)
            #         if_pwr = res[0][1]
            #         print(if_pwr)
            #         if if_pwr > 5 :
            #             pwr = pwr - 5
            #         else:
            #             break
            #     loginfo('RF pwer:{}    IF power:{}'.format(pwr,if_pwr))
            #     fw1.write_data([channel,lna_gain,lna_atten_r,lna_atten_c,flt_gain1,flt_gain2,pwr,if_pwr,if_pwr-pwr])
            # for flt_gain2 in range(0,4):
            #     lna_gain = 0
            #     lna_atten_r = 3
            #     lna_atten_c = 3
            #     flt_gain1 = 0
            #
            #     self.mem_ts.wrm(0xa042107c, 21, 20, lna_atten_r)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 17, 16, lna_gain)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 19, 18, lna_atten_c)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 23, 22, flt_gain1)
            #     time.sleep(.5)
            #     self.mem_ts.wrm(0xa042107c, 25, 24, flt_gain2)
            #     pwr = -10
            #
            #     while 1:
            #         txg.set_power(pwr)
            #         time.sleep(0.2)
            #         res = spa.pk_search_avg_timesleep(timesleep=2)
            #         if_pwr = res[0][1]
            #         print(if_pwr)
            #         if if_pwr > 0 :
            #             pwr = pwr - 5
            #         else:
            #             break
            #     loginfo('RF pwer:{}    IF power:{}'.format(pwr,if_pwr))
            #     fw1.write_data([channel,lna_gain,lna_atten_r,lna_atten_c,flt_gain1,flt_gain2,pwr,if_pwr,if_pwr-pwr])

    def test231_agc_scan(self,  flt_pkd_list=[0x40], agc_tgt_sat_list=[[407,580]], signal_cableloss=7, interference_cableloss=7,  if_cableloss=1, signal_freq=2440,
                         signal_level=-67, interference_en=1,
                         interference_offset_freq=3):

        title = 'flt_pkd_value,agc_tgt,agc_sat,interference_pwr(dbm)@offset {}MHz,signal_pwr(dbm),IF_pwr(dbm),dalta\n'.format(interference_offset_freq)
        fname = rfglobal.get_filename('ts_bt_test/','test231_agc_scan_{}'.format(self.board_name))
        fw1 = csvreport(fname,title)

        spa = Agilent()
        spa.set_param(1,1,10,30)
        txg = mxg.MXG()
        if interference_en == 0:
            txg_freq = signal_freq
        else:
            txg_freq = signal_freq + interference_offset_freq
            tester = tester_cmw(mode=1, rfport=1, cable_loss=signal_cableloss)
            tester.srx.trig_cw()
            tester.srx.para_settings(freq=signal_freq, level=signal_level)
            tester.srx.gen_switch('ON')
        txg.trriger_para_set(type='CONTinuous', count=1000)
        txg.output_state(1, 0)
        self.mem_ts.wrm(0xa04210a8, 6, 6, 1)  ##Tx DAC output test enable signal
        time.sleep(.5)
        self.mem_ts.wrm(0xa04210a8, 0, 0, 1)  ##Test buffer enable signal
        time.sleep(.5)
        self.mem_ts.wrm(0xa0421060, 12, 12, 1)
        self.mem_ts.wrm(0xa0421060, 11, 0, signal_freq)

        loginfo('txg_freq:  {}'.format(txg_freq))
        for flt_pkd_value in flt_pkd_list:
            self.mem_ts.wrm(0xa0421090, 20, 14, flt_pkd_value)
            for agc_tgt_sat in agc_tgt_sat_list:
                agc_tgt = agc_tgt_sat[0]
                agc_sat = agc_tgt_sat[1]
                self.mem_ts.wrm(0xa0421068, 9, 0, agc_tgt)
                self.mem_ts.wrm(0xa042106c, 25, 16, agc_sat)
                for rfpower in range(-80,10,3):
                    txg_rfpower = rfpower + interference_cableloss
                    txg.para_set(freq=txg_freq, power=txg_rfpower)
                    self.mem_ts.wr(0xa0422014, 0x0)
                    time.sleep(1)
                    self.mem_ts.wr(0xa0422014, 0x3)
                    res = spa.pk_search_avg_timesleep(timesleep=0.5)
                    if_pwr = res[0][1]
                    loginfo('flt_pkd_value:{}  agc_tgt:{}  agc_sat:{}  txg_freq:{}  RF pwer:{}  IF power:{}  gain:{}'.format(flt_pkd_value,agc_tgt,agc_sat,txg_freq,
                                                                                                                                          rfpower,
                                                                                                                                          if_pwr,
                                                                                                                                   if_pwr-rfpower))
                    if interference_en == 0:
                        fw1.write_data([flt_pkd_value,agc_tgt, agc_sat, -999, rfpower, if_pwr + if_cableloss, if_pwr - rfpower + if_cableloss])
                    else:
                        fw1.write_data([flt_pkd_value,agc_tgt, agc_sat, rfpower, signal_level, if_pwr + if_cableloss, if_pwr - signal_level + if_cableloss])
                    self.mem_ts.wr(0xa0422014, 0x0)

    def test232_agc_gain_table_set(self, file_path=r'E:\project_audio\TX232\report\NPW2\TX232_QFN68_A1\rxgain\tx232_rxgain_table_v1.xls'):
        fd = pd.read_excel(file_path,index_col=None,header=None)
        for i in range(len(fd.index)):
            reg_addr = eval(fd[0][i])
            addr_msb = fd[1][i]
            addr_lsb = fd[2][i]
            reg_value = eval(fd[7][i])
            self.mem_ts.wrm(reg_addr,addr_msb,addr_lsb,reg_value)


    def test232_agc_scan(self, agc_tgt_sat_list=[[407,580]], signal_cableloss=7, interference_cableloss=7, if_cableloss=1, signal_freq=2440, signal_level=-67,
                         interference_en=0,
                         interference_offset_freq=3):

        title = 'agc_tgt,agc_sat,interference_pwr(dbm)@offset {}MHz,signal_pwr(dbm),IF_pwr(dbm),dalta,agc_gain_idx\n'.format(interference_offset_freq)
        fname = rfglobal.get_filename('ts_bt_test/','test232_agc_scan_{}M'.format(interference_offset_freq))
        fw1 = csvreport(fname,title)
        self.test232_agc_gain_table_set()
        spa = Agilent()
        spa.set_param(1,1,10,30)
        txg = mxg.MXG()
        tester = tester_cmw(mode=1, rfport=1, cable_loss=signal_cableloss)
        if interference_en == 0:
            txg_freq = signal_freq
        else:
            txg_freq = signal_freq + interference_offset_freq
            tester.srx.trig_cw()
            tester.srx.para_settings(freq=signal_freq, level=signal_level)
            tester.srx.gen_switch('ON')
        txg.trriger_para_set(type='CONTinuous', count=1000)
        txg.output_state(1, 0)
        self.test232_rf_buf('rx')

        self.mem_ts.wrm(0xa0421064, 12, 12, 1)
        self.mem_ts.wrm(0xa0421064, 11, 0, signal_freq)

        for agc_tgt_sat in agc_tgt_sat_list:
            agc_tgt = agc_tgt_sat[0]
            agc_sat = agc_tgt_sat[1]
            self.mem_ts.wrm(0xa04210ac, 9, 0, agc_tgt)
            self.mem_ts.wrm(0xa04210a8, 11, 2, agc_sat)
            for rfpower in range(-90,10,3):
                txg_rfpower = rfpower + interference_cableloss
                txg.para_set(freq=txg_freq, power=txg_rfpower)
                self.mem_ts.wr(0xa0422014, 0x0)
                time.sleep(1)
                self.mem_ts.wr(0xa0422014, 0x3)
                res = spa.pk_search_avg_timesleep(timesleep=1.5)
                if_pwr = res[0][1]
                agc_gain_idx = self.mem_ts.rdm(0xa04210dc, 11, 8)
                loginfo('agc_tgt:{}  agc_sat:{}  txg_freq:{}  RF pwer:{}  IF power:{}  gain:{}  agc_gain_idx:{}'.format(agc_tgt,agc_sat,txg_freq,
                                                                                                                                      rfpower,
                                                                                                                                      if_pwr,
                                                                                                                               if_pwr-rfpower,agc_gain_idx))
                if interference_en == 0:
                    fw1.write_data([agc_tgt, agc_sat,-999, rfpower, if_pwr+if_cableloss, if_pwr - rfpower  + if_cableloss,agc_gain_idx])
                else:
                    fw1.write_data([agc_tgt, agc_sat, rfpower, signal_level, if_pwr + if_cableloss, if_pwr - signal_level + if_cableloss,agc_gain_idx])
                self.mem_ts.wr(0xa0422014, 0x0)
        txg.output_state(0, 0)
        tester.srx.gen_switch('OFF')

    def test232_agc_vs_gainindex(self, channel_list=[0], rf_cableloss=7):
        title = 'mode,channel,signal_pwr(dbm),agc_gain_idx\n'
        fname = rfglobal.get_filename('ts_bt_test/','test232_agc_vs_gainindex_{}'.format(self.board_name))
        fw1 = csvreport(fname,title)
        txg = mxg.MXG()
        txg.arb_waveform(rate='LE_1M')
        txg.arb_state(1)
        txg.trriger_para_set(type='CONTinuous', count=1000)
        txg.output_state(1, 1)
        for channel in channel_list:
            self.btapi.cmdstop(0)
            self.btapi.cmdstop(1)
            self.btapi.LE_RX(chan=channel, phy=1, mod=0, countMode=0)
            txg.para_set(freq=2402 + 2*channel, power=-67)
            for signal_pwr in range(-97,10):
                txg.set_power(signal_pwr + rf_cableloss)
                time.sleep(1)
                gain_index = self.jlink.rdm(0xa04210dc, 11, 8)

                fw1.write_data(['LE1M',channel,signal_pwr,gain_index])






    def le_rx_sniff(self, chan=0, rxpwr=-60):
        import matplotlib.pyplot as plt
        import numpy as np
        import time
        from math import *
        plt.ion()  # 开启interactive mode 成功的关键函数
        plt.figure(1)

        self.btapi.cmdstop(0)
        self.btapi.cmdstop(1)
        tester_signal = tester_cmw(mode=1, rfport=1, cable_loss=4.6)
        tester_signal.srx.gen_wave(repeat=1500, data_rate='LE_1M', dirty_en=0)
        tester_signal.srx.para_settings(freq=2402+chan*2, level=rxpwr)
        time.sleep(1)
        per_list=[]
        t=0
        t_list=[]
        while 1 :
            plt.clf()  # 清空画布上的所有内容
            self.btapi.LE_RX(chan=chan, phy=1, mod=0, countMode=0)
            tester_signal.srx.gen_switch('ON')
            time.sleep(1500 * 1e-3)
            res = self.btapi.cmdstop(0)
            res = res[0].split('bleRxCount=')[1].strip('\r\n')
            logdebug('rev_pkt:  {}'.format(res))
            rev_pkg = eval(res)
            per = (1500 - rev_pkg) * 100.0 / 1500
            # time.sleep(time_interval)
            t = t+1
            # per = eval(res[1])
            per_list.append(per)
            t_list.append(t)
            plt.plot(t_list,per_list, '-r')
            plt.pause(0.01)

    def test232_rssi_print(self):
        self.comport.ser.flushInput()
        self.comport.ser.flushOutput()
        self.rssi_print_end = 0
        while 1:

            self.rssi_print = self.comport.ser.readline()
            loginfo('{}'.format(self.rssi_print))
            if self.rssi_print_end == 1:
                break
            # elif self.rssi_print == '':
            #     break

    def test232_rssi_scan(self, freq=2476, mode=0, cable_loss=7, name_str=''):
        '''
        freq:   range2402--2480;unit:MHz
        mode:   0: br dh1;  1: LE1M
        '''

        title = 'rate,freq(MHz),instrument_tx_level(dBm),rssi/agc_gain_index\n'
        fname = rfglobal.get_filename('ts_bt_test/','test232_rssi_scan_{}'.format(name_str))
        fw1 = csvreport(fname,title)
        tester_signal = mxg.MXG('N5182B', 1)
        if mode!=0:
            tester_signal.arb_waveform(rate='LE_1M')
            rate = 'LE_1M'
        else:
            tester_signal.arb_waveform(rate='1M_DH1')
            rate = '1M_DH1'
        tester_signal.trriger_para_set(type='SINGle', count=100)
        tester_signal.output_state(1, 1)
        tester_signal.arb_state(1)
        tester_signal.para_set(freq=freq, power=-67 + cable_loss)
        time.sleep(2)

        # rssi_rd = threading.Thread(target=self.test232_rssi_print)
        # rssi_rd.start()
        for pwr in range(-97,0):
            tester_signal.set_power(pwr+cable_loss)
            time.sleep(1)
            if mode != 0:
                self.btapi.LE_RX(chan=(freq - 2402) / 2, phy=1, mod=0, countMode=0)
            else:
                self.btapi.BR_RX(chan=freq - 2402, len=27, ptype=5, rate=1, ltaddr=1, uap=0x0, lap=0x8)
            self.comport.ser.readlines()
            for num in range(10):
                tester_signal.trigger_on()

                time.sleep(1)
                rssi_print = self.comport.ser.readlines()
                # loginfo('{}'.format(rssi_print))
                if rssi_print != []:
                    break

            rssi_list=[]
            for i in range(1,len(rssi_print)):
                if rssi_print[i].find('rssi=')!=-1:
                    rssi_val = rssi_print[i].split('rssi=')[1].split(',')[0]
                    loginfo('{}'.format(rssi_val))
                    agc_gain_index = rssi_print[i].split('agc_gain_index=')[1]
                    loginfo('agc_gain_index:    {}'.format(agc_gain_index))
                    rssi_list.append(eval(rssi_val))
                    rssi_list.append(eval(agc_gain_index))

            # agc_gain_index = self.jlink.rdm(0xa04210dc,11,8)

            fw1.write_data([rate,freq,pwr]+rssi_list)

            if mode != 0:
                self.btapi.cmdstop(0)
            else:
                self.btapi.cmdstop(1)

    def rf_reg_read_all(self,file_path=''):
        # import docx
        # from docx import Document
        # title = 'reg_addr,reg_msb,reg_lsb,RW,reg_name,description,value_default,value_present_{}\n'.format(time.strftime("%Y-%m-%d %H:%M:%S", time.localtime()))
        # fname = self.get_filename('ts_bt_test/', 'rf_reg_read_all')
        # fw1 = csvreport(fname, title)
        # document = Document(file_path)
        # tables = document.tables
        # for i in range(1,len(tables)):
        #     table = tables[i]
        #     reg_addr = 0xa0421000 + (i-1)*4
        #     for table_row in range(1,len(table.rows)):
        #         reg_bit = table.cell(table_row,0).text.encode('utf-8')
        #         reg_bit_num = reg_bit.split(':')
        #         msb = reg_bit_num[0]
        #         if len(reg_bit_num) > 1 :
        #             lsb = reg_bit_num[1]
        #         elif len(reg_bit_num) > 2 :
        #             logerror('{} bit read fail'.format(reg_addr))
        #         else:
        #             lsb = msb
        #         value = self.mem_ts.rdm(reg_addr,eval(msb),eval(lsb))
        #
        #         reg_rw = table.cell(table_row,1).text.encode('utf-8')
        #         reg_value_default = table.cell(table_row,2).text.encode('utf-8')
        #         reg_name = table.cell(table_row,3).text.encode('utf-8')
        #         reg_description = table.cell(table_row,4).text.encode('utf-8')
        #         reg_description = reg_description.replace('\n', '; ').replace('\r', '; ')
        #         reg_description = r'{}'.format(reg_description)
        #         print(msb,lsb,reg_rw,reg_name,reg_description,reg_value_default)
        #         fw1.write_data([str(hex(reg_addr)),msb,lsb,reg_rw,reg_name,reg_description,reg_value_default,str(hex(value))],float_num=0)

        df=pd.read_excel(file_path,header=None,names=['a','b','c','d','e','f'])
        reg_addr_offset = 0
        for i in df['a'].index:
            offset = df.iloc[i,0]
            logdebug(offset)
            if str(offset).find('Offset:')!=-1:
                reg_addr_offset = offset.split('Offset:')[1]
                reg_addr_offset = eval(reg_addr_offset)
            reg_addr = 0xa0421000 + reg_addr_offset
            logdebug(reg_addr)
            if df.iloc[i,2] == 'RW' or df.iloc[i,2] == 'R':
                reg_bit = df.iloc[i,0]
                if reg_bit.find(':')!=-1:
                    bit_msb = reg_bit.split(':')[0]
                    bit_lsb = reg_bit.split(':')[1]
                else:
                    bit_msb = reg_bit
                    bit_lsb = reg_bit
                reg_value_pre = self.mem_ts.rdm(reg_addr,eval(bit_msb),eval(bit_lsb))
                df.iloc[i, 5] = hex(reg_value_pre)
        df.to_excel(r'E:\360MoveData\Users\123\Desktop\bt_rf_mpw3.xlsx',header=False,index=False)


    def rf_top_reg_to_code(self, file_path=''):
        import docx
        from docx import Document
        document = Document(file_path)
        tables = document.tables
        paragraph_list = []
        fname = rfglobal.get_filename('ts_bt_test/', 'rf_top_reg_to_code')
        logtime = time.strftime('%Y%m%d_%H%M%S', time.localtime(time.time()))
        with open('{}_{}.txt'.format(fname,logtime),'w+') as f:
            f.write('/*\n * DEFINES\n *****************************************************************************************\n */\n')
            # f.write('#define    RF_TOP_BASE_ADDR                0xA0421000\n')
            for paragraph in document.paragraphs:
                logdebug('{}'.format(paragraph.text))
                if paragraph.text.find('(offset:') != -1:
                    stru_name = paragraph.text.split('(offset:')[0].split(' ')[1]
                    offset = paragraph.text.split('(offset:')[1].split(')')[0].replace(' ','')
                    reg_addr = 0xa0421000 + eval(offset)
                    w_str = '#define    {}_REG'.format(stru_name)
                    w_str = w_str + (40-len(w_str))*' ' + '{}\n'.format(hex(reg_addr).replace('L',''))
                    f.write(w_str)


            for paragraph in document.paragraphs:
                logdebug('{}'.format(paragraph.text))
                if paragraph.text.find('(offset:') != -1:
                    stru_name = paragraph.text.split('(offset:')[0].split(' ')[1]
                    offset = paragraph.text.split('(offset:')[1].split(')')[0].replace(' ','')
                    table_num = eval(offset)/4 +1
                    logdebug('table_num: {}'.format(table_num))
                    table = tables[table_num]
                    struct_wstr1 = 'typedef union{\n    struct{\n'
                    struct_wstr2_list = ''
                    struct_reset_list = ''
                    for table_row in range(len(table.rows)-1,0,-1):
                        logdebug('lens: {} ;table_row: {}'.format(len(table.rows),table_row))
                        reg_bit = table.cell(table_row,0).text.encode('utf-8')
                        reg_bit_num = reg_bit.split(':')
                        msb = eval(reg_bit_num[0])
                        if len(reg_bit_num) > 1 :
                            lsb = eval(reg_bit_num[1])
                        elif len(reg_bit_num) > 2 :
                            logerror('{} bit read fail'.format(stru_name))
                        else:
                            lsb = msb
                        reg_rw = table.cell(table_row,1).text.encode('utf-8')
                        reg_value_default = table.cell(table_row,2).text.encode('utf-8')
                        reg_name = table.cell(table_row,3).text.encode('utf-8')
                        reg_description = table.cell(table_row,4).text.encode('utf-8')
                        reg_description = reg_description.replace('\n', '; ').replace('\r', '; ')
                        reg_description = r'{}'.format(reg_description)
                        struct_wstr2 = '        unsigned int {} : {};'.format(reg_name,(msb-lsb+1))
                        struct_wstr2 = struct_wstr2 + (60-len(struct_wstr2))*' ' + '/* bit[{}:{}]; {}; {} */\n'.format(lsb,msb,reg_rw,reg_description)
                        struct_wstr2_list = struct_wstr2_list + struct_wstr2
                        if reg_rw == 'RW':
                            struct_reset = '    {}->{}_BIT.{} = {};\n'.format(stru_name,stru_name,reg_name,reg_value_default)
                            struct_reset_list = struct_reset_list + struct_reset
                    # struct_wstr3 = '    }'+'{}_BIT;\n    uint32_t    {}_BYTE;\n'.format(stru_name,stru_name)+'}'+'_{};\n'.format(stru_name)
                    struct_wstr3 = '    }' + 'BF;\n    uint32_t    U32;\n' + '}' + '_{};\n'.format(stru_name)
                    struct_wstr4 = 'volatile _{} *{} = (_{}*){}_REG;\n'.format(stru_name,stru_name,stru_name,stru_name)
                    struct_wstr5 = 'extern void {}_reset(void)\n'.format(stru_name.lower()) +'{\n'+ struct_reset_list + '}\n'
                    f.write(struct_wstr1+struct_wstr2_list+struct_wstr3+struct_wstr4+struct_wstr5)


