#!/usr/bin/python
# -*- coding:utf8 -*-

import re
from exceptions import StandardError

from baselib.loglib.log_lib import *
from baselib.plot import *
from baselib.instrument import *
import numpy as np
import pandas as pd
import scipy
from math import *
from shutil import copy
import time
import csv
import pylab
import matplotlib.pyplot as plt
from baselib.test_channel import *
import xlrd
import sys
import random
import os
from baselib.instrument.cmw_bt import *
from baselib.instrument.spa import *
from rftest.rflib import *
from hal.common import *
from rftest.rflib.csv_report import csvreport
from rftest.testcase.bt_api import bt_api
from rftest.rflib import rfglobal
from rftest.testcase.rf_debug import rf_debug

import serial
from sys import argv
import binascii
from binascii import b2a_hex, a2b_hex
import pylink
import docx
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_TAB_ALIGNMENT,WD_TAB_LEADER #设置制表符等
from docx.shared import Inches #设置图像大小
from docx.shared import Pt #设置像素、缩进等
from docx.shared import RGBColor #设置字体颜色
from docx.shared import Length #设置宽度
import win32api
import shutil

class bt_signaling(object):

    def __init__(self, comport, chipv='TX212',board_name='', jlink='', rfport=1, atten=4.6, le_com_baudrate='B500K'):
        self.rfport = rfport
        self.atten = atten
        self.comport = comport
        self.chipv = chipv
        self.board_name = board_name
        self.jlink = jlink
        self.comport.Hexmd=1
        self.le_com_baudrate = le_com_baudrate
        # if jlink_en !=0:
        #     self.jlink = jlink(jlink_sn=jlink_sn)
            # self.jlink = pylink.JLink()
            # self.jlink_sn = jlink_sn
            # self.jlink.open()
            # self.jlink.set_tif(pylink.enums.JLinkInterfaces.SWD)
            # self.jlink.connect('ARMCM4_FP')

        cmw_bt().reset()
        self.csp = combined_signal_path()
        self.csp.signaling_switch(1)
        self.csp.rf_port(mode='signaling', rfport=self.rfport, atten=self.atten)
        self.br_rate_list = ['DH1','DH3','DH5']
        self.edr_rate_list = ['2_DH1','2_DH3','2_DH5','3_DH1','3_DH3','3_DH5']
        self.dic_edr_rate = {
            '2_DH1':    'E21P',
            '2_DH3':    'E23P',
            '2_DH5':    'E25P',
            '3_DH1':    'E31P',
            '3_DH3':    'E33P',
            '3_DH5':    'E35P'
            }

    # def jlink_connect(self):
    #     self.jlink.jlink_connect()
        # self.jlink.open(self.jlink_sn)
        # self.jlink.set_tif(pylink.enums.JLinkInterfaces.SWD)
        # self.jlink.connect('ARMCM4_FP')

    # def jlink_mem_rd(self, addr=0xa04210a0):
    #     res = self.jlink.memory_read(addr,4)
    #     res_int = 0x1000000*res[3] + 0x10000*res[2] + 0x100*res[1] + res[0]
    #     res = '0x%x'%(res_int)
    #     loginfo('addr:  {}  value:  {}'.format(hex(addr), res))
    #     return res
    #
    # def jlink_mem_rdm(self, addr=0xa04210a0, msb=1, lsb=0):
    #     result = self.jlink_mem_rd(addr=addr)
    #     mask = (1<<(msb+1)) - (1<<lsb);
    #     try:
    #         result = (eval(result) & mask) >> lsb;
    #         # logdebug("resm:" + "0x%x"%result)
    #     except:
    #         logerror("mem reads " + "%s"%result)
    #     return result
    #
    # def jlink_mem_wr(self, addr=0xa04210a0, value=0x18c619):
    #
    #     value_list = [value&0xff,(value&0xff00)>>8,(value&0xff0000)>>16,(value&0xff000000)>>24]
    #     self.jlink.memory_write(addr,value_list)
    #     res = self.jlink_mem_rd(addr)
    #
    # def jlink_mem_wrm(self, addr, msb, lsb, value):
    #     result = self.jlink_mem_rd(addr=addr)
    #     mask = (1<<(msb+1)) - (1<<lsb);
    #     try:
    #         result = (eval(result) & ~mask) | ((value << lsb) & mask);
    #         return self.jlink_mem_wr(addr,result);
    #     except:
    #         logerror("mem write fail, reads  %s"%result)
    #         return

    def hci_reset(self):
        self.comport.req_com(cmdstr='\x01\x03\x0c\x00',timeout=1,endstr='TS')

    def hci_read_buffer_size(self):
        self.comport.req_com(cmdstr='\x01\x05\x10\x00', timeout=1, endstr='TS')

    def hci_enable_under_test_mode(self):
        self.comport.req_com(cmdstr='\x01\x03\x18\x00', timeout=1, endstr='TS')

    def hci_enable_write_both_scan(self):
        self.comport.req_com(cmdstr='\x01\x1a\x0c\x01\x03', timeout=1, endstr='TS')

    def hci_auto_accept_all_connection(self):
        self.comport.req_com(cmdstr='\x01\x05\x0c\x03\x02\x00\x02', timeout=1, endstr='TS')

    def dut_signal_mode_set(self):
        res = self.comport.req_com(cmdstr='signaltest', wr_end='\r\n', timeout=5)
        if res.find('signal test')!=-1:
            self.comport.Hexmd = 1
            self.hci_reset()
            self.hci_read_buffer_size()
            self.hci_enable_under_test_mode()
            self.hci_enable_write_both_scan()
            self.hci_auto_accept_all_connection()
            self.comport.Hexmd = 0
        else:
            logwarn('can not enter signal test mode')

    def config_per_le_connect_usb(self, rate='LE1M', packet_len=37):
        self.csp.reset()
        self.csp.clean()
        time.sleep(1)
        self.csp.signaling_switch(1)
        time.sleep(2)
        self.csp.rf_port(mode='signaling', rfport=self.rfport, atten=self.atten)
        self.csp.sig_opmode(mode='RFT')
        self.csp.sig_std(std='LESignaling')
        self.csp.sig_btype('LE')
        self.csp.sig_le_phy(phy=rate)
        self.csp.RF_Frequency_Settings_rx(mode='DTM', ch_tx=0)
        self.csp.RF_Power_settings(rx_level=-40, tx_power=10,margin=8)
        self.csp.RF_Power_settings_autoranging(1)
        self.csp.config_connect_le_packet_len(rate=rate, len=packet_len)
        self.csp.config_connect_le_packet_pattern(rate=rate, pattern='PRBS9')
        self.csp.config_hwinterface('RS232')
        self.csp.config_comport()
        self.csp.config_comport_baudrate(self.le_com_baudrate)
        self.csp.check_le_connect_usb()
        time.sleep(1)

        self.csp.config_rxq_repetion(rep='SING')
        self.csp.config_rxq_le_packets(rate=rate, num=1500)
        self.csp.per_meas_state(0)
        time.sleep(1)

    def config_classic_ber(self, rate='BR', chan=0, rx_level=-60, tx_power=10):
        '''
        配置经典蓝牙连接CMW流程：
        复位仪器，打开信令模式，配置好RF参数和蓝牙包的参数，inquire EUT的BD，paging EUT并连接EUT。
        '''
        if rate not in ('BR','EDR'):
            raise StandardError('rate command wrong')
        # self.dut_signal_mode_set()  #设置dut进入BT信令测试模式
        self.csp.reset()    #复位cmw设置
        self.csp.clean()
        self.csp.signaling_switch(1)    #1：仪器的蓝牙信令模式打开，0：关闭
        self.csp.rf_port(mode='signaling', rfport=self.rfport, atten=self.atten)    #设置信令模式下 rfport和线损
        self.csp.sig_opmode(mode='RFT')
        self.csp.sig_std(std='CLAS')
        self.csp.sig_btype(btype=rate)
        self.csp.config_sig_testmode(testmode='LOOP')
        self.csp.RF_Frequency_Settings_rx(mode='LOOP', ch_tx=chan)
        self.csp.RF_Power_settings(rx_level=rx_level, tx_power=tx_power, margin=8)
        # self.csp.RF_Power_settings_autoranging()
        if rate == 'BR':
            self.csp.config_connect_br_packet_pattern(pattern='PRBS9')  #设置BR包 payload 数据类型
            self.csp.config_connect_br_packet_ptype(ptype='DH1')    #设置BR包类型
            # self.csp.config_connect_br_packet_len(len=packet_len)   #设置BR包长
        else:
            self.csp.config_connect_edr_packet_pattern(pattern='PRBS9') #设置EDR包 payload 数据类型
            self.csp.config_connect_edr_packet_ptype(ptype='E21P')  #设置EDR的包类型，E21P：2-DH1，E31P：3-DH1
            # self.csp.config_connect_edr_packet_len(len=packet_len)  #设置EDR包长
        self.csp.config_paging_classic_NOResponses(1)   #设置stop inquire的条件，当inquire到的BD数量为设定值时则stop inquire
        self.csp.bt_connect_action(action='INQuire')    #开始inquire命令
        time.sleep(2)
        for i in range(5):
            self.csp.bt_connect_action(action='INQuire')  # 开始inquire命令
            time.sleep(2)
            target = self.csp.get_paging_classic_target()   #返回inquire到的BD 列表
            self.csp.config_paging_classic_target(target=1) #设置inquire到的第一个BD用来paging
            self.csp.bt_connect_action(action='TMConnect')  #
            res = self.csp.get_bt_connect_state()
            if res == 'TCON':
                break

        time.sleep(1)

    def cmd_br_connect(self):
        for i in range(5):
            self.csp.bt_connect_action(action='INQuire')  # 开始inquire命令
            time.sleep(2)
            target = self.csp.get_paging_classic_target()   #返回inquire到的BD 列表
            self.csp.config_paging_classic_target(target=1) #设置inquire到的第一个BD用来paging
            self.csp.bt_connect_action(action='TMConnect')  #
            res = self.csp.get_bt_connect_state()
            if res == 'TCON':
                loginfo('*********** connect sucess ************')
                time.sleep(2)
                break
            else:
                logwarn('connect fail, state: {} , please try again'.format(res))

    def le_tx_tp_debug(self, chan_list=[0,19,39]):
        '''

        '''
        title = 'type,tp_gain_value,channel,nominal_pow(dBm),'
        title = title + 'delta_f1_avg(KHz),delta_f1_avg_max(KHz),delta_f2_avg(KHz),delta_f2_avg_max(KHz),mod_ratio,delta_f2_99(KHz),'
        title = title + '\n'
        fname = self.get_filename('ts_bt_test/','le_tx_tp_debug_{}'.format(self.board_name))
        fw1=csvreport(fname,title)
        self.config_per_le_connect_usb(rate='LE1M', packet_len=37)
        self.csp.config_rx_level(rxpwr=-67)
        self.csp.tx_measure_para(tout=5, repetition='SINGleshot', count=1000, MOEXception='OFF')
        for rate in ['LE1M','LE2M']:
            if rate == 'LE500k':
                _rate = 'LELR'
                coding = 'S2'
                mode = 'LRANge'
                self.csp.sig_le_lr_coding(coding=coding)
            elif rate == 'LE125K':
                _rate = 'LELR'
                coding = 'S8'
                mode = 'LRANge'
                self.csp.sig_le_lr_coding(coding=coding)
            else:
                mode = rate
                _rate = rate

            self.csp.mode_set(mode=mode)
            self.csp.sig_le_phy(phy=_rate)
            rfdebug = rf_debug(self.comport,self.chipv,self.jlink)
            self.jlink.wrm(0xa0421020,25,24,0x1)
            for tp_gain_val_chan in [0,19,39]:
                if rate == 'LE1M':
                    tp_gain_cal_value = rfdebug.test232_tp_gain_cal(0,2402+tp_gain_val_chan*2)
                else:
                    tp_gain_cal_value = rfdebug.test232_tp_gain_cal(1, 2402 + tp_gain_val_chan * 2)
                for chan in chan_list:
                    self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan)
                    self.csp.config_connect_le_packet_pattern(rate=_rate, pattern='P44')
                    res = self.csp.get_modulation_measure_res(rate=mode,data_type='AVER')
                    logdebug('{}'.format(res))
                    delta_f1_avg = eval(res[6]) / 1000.00
                    nominal_pow = eval(res[12])

                    res = self.csp.get_modulation_measure_res(rate=mode, data_type='MAX')
                    logdebug('{}'.format(res))
                    delta_f1_avg_max = eval(res[6]) / 1000.00


                    time.sleep(0.5)

                    self.csp.config_connect_le_packet_pattern(rate=_rate, pattern='P11')
                    res = self.csp.get_modulation_measure_res(rate=mode,data_type='AVER')
                    logdebug('{}'.format(res))

                    delta_f2_99 = eval(res[2]) / 1000.00
                    delta_f2_avg = eval(res[9]) / 1000.00
                    mod_ratio = delta_f2_avg / delta_f1_avg

                    res = self.csp.get_modulation_measure_res(rate=mode, data_type='MAX')
                    logdebug('{}'.format(res))
                    delta_f2_avg_max = eval(res[9]) / 1000.00

                    fw1.write_data([rate,tp_gain_cal_value,chan,nominal_pow,delta_f1_avg,delta_f1_avg_max,delta_f2_avg,delta_f2_avg_max,mod_ratio,delta_f2_99])



    def le_tx_meas(self, chan_list=[0], rate_list=['LE1M'], packet_len=37, csv_save=True, report_save=True, fig_en=0):
        '''
        rate_list:  LE1M、LE2M、LE500K、LE125K
        '''
        if csv_save:
            title = 'channel,type,nominal_pow(dBm),peak_pow(dBm),leakage_pow(dBm),freq_accuracy(KHz),freq_drift(KHz),drift_rate(Hz/50us),'
            title = title + 'delta_f1_avg(KHz),delta_f1_min(KHz),delta_f1_max(KHz),delta_f2_avg(KHz),delta_f2_min(KHz),delta_f2_max(KHz),mod_ratio,delta_f1_99(KHz),delta_f2_99(KHz),'
            title = title + 'acp_list_21ch\n'
            fname = self.get_filename('ts_bt_test/','test_le_tx_{}'.format(self.board_name))
            fw1=csvreport(fname,title)
        if report_save:
            title2 = 'BR_tx_performance'
            fname2 = self.get_filename('ts_bt_test/','test_le_tx_report_{}'.format(self.board_name))
            fw2=csvreport(fname2,title2)
            item = ['channel', 'rate', u'TP/TRM-LE/CA/BV-01-C - Output power  - Average Power', u' - Peak Power',
                    u'- Peak - Average Power',
                    u'- leakage Power',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -10',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -9',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -8',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -7',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -6',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -5',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -4',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -3',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -2',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   -1',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -    Ptx',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   1',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   2',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   3',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   4',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   5',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   6',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   7',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   8',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   9',
                    u'TP/TRM-LE/CA/BV-03-C - In-Band Emissions -   10',
                    u'TP/TRM-LE/CA/BV-05-C - Modulation Characteristics - Fre Accuracy',
                    u'- Delta F1 Avg',
                    u' - Delta F1 Min',
                    u' - Delta F1Max',
                    u' - Fre Accuracy',
                    u'- Delta F2 Avg',
                    u'- Delta F2 Min',
                    u'- Delta F2Max',
                    u' - Delta F2 Max 99%',
                    u'- Delta F2 Avg/Delta F1 Avg',
                    u'- Frequency Accuracy',
                    u' - Frequency Offset',
                    u'- Frequency Drift',
                    u'- Max Drift Rate - 50us',
                    u'- Intial Frequency Drift',
                    ]
            df = pd.DataFrame(item)
            df.to_csv(fw2.filename, index=False)
        res_list=[]
        self.config_per_le_connect_usb(rate='LE1M', packet_len=37)
        self.csp.config_rx_level(rxpwr=-67)
        self.csp.tx_measure_para(tout=5, repetition='SINGleshot', count=20, MOEXception='OFF')
        for chan in chan_list:
            self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan)
            for rate in rate_list:
                if rate == 'LE500k':
                    _rate = 'LELR'
                    coding = 'S2'
                    mode = 'LRANge'
                    self.csp.sig_le_lr_coding(coding=coding)
                elif rate == 'LE125K':
                    _rate = 'LELR'
                    coding = 'S8'
                    mode = 'LRANge'
                    self.csp.sig_le_lr_coding(coding=coding)
                else:
                    mode = rate
                    _rate = rate

                self.csp.mode_set(mode=mode)
                self.csp.sig_le_phy(phy=_rate)
                # self.config_per_le_connect_usb(rate=_rate, packet_len=packet_len)
                # self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan)
                # self.csp.config_rx_level(rxpwr=-67)
                time.sleep(1)
                delta_f1_99, delta_f2_99, delta_f1_avg, delta_f1_min, delta_f1_max,delta_f2_avg, delta_f2_min, delta_f2_max, mod_ratio = -999, -999, -999, -999, -999, -999, -999, -999, -999
                self.csp.config_connect_le_packet_pattern(rate=_rate,pattern='P44')
                res = self.csp.get_modulation_measure_res(rate=mode)
                logdebug('{}'.format(res))

                if mode == 'LRANge':
                    freq_accuracy = eval(res[3]) / 1000.00
                    freq_drift = eval(res[4]) / 1000.00
                    drift_rate = eval(res[5]) / 1000.00
                else:
                    delta_f1_99 = eval(res[2]) / 1000.00
                    freq_accuracy = eval(res[3]) / 1000.00
                    freq_drift = eval(res[4]) / 1000.00
                    drift_rate = eval(res[5]) / 1000.00
                    delta_f1_avg = eval(res[6]) / 1000.00
                    delta_f1_min = eval(res[7]) / 1000.00
                    delta_f1_max = eval(res[8]) / 1000.00

                time.sleep(0.5)

                self.csp.config_connect_le_packet_pattern(rate=_rate,pattern='P11')
                res = self.csp.get_modulation_measure_res(rate=mode)
                logdebug('{}'.format(res))

                if mode == 'LRANge':
                    freq_accuracy = eval(res[3]) / 1000.00
                    freq_drift = eval(res[4]) / 1000.00
                    drift_rate = eval(res[5]) / 1000.00
                else:
                    delta_f2_99 = eval(res[2]) / 1000.00
                    freq_accuracy = eval(res[3]) / 1000.00
                    freq_drift = eval(res[4]) / 1000.00
                    drift_rate = eval(res[5]) / 1000.00
                    delta_f2_avg = eval(res[9]) / 1000.00
                    delta_f2_min = eval(res[10]) / 1000.00
                    delta_f2_max = eval(res[11]) / 1000.00
                    mod_ratio = delta_f2_avg / delta_f1_avg
                    freq_offset = eval(res[14])/1000.00
                    init_drift = eval(res[15]) / 1000.00
                time.sleep(0.5)

                self.csp.config_connect_le_packet_pattern(rate=_rate,pattern='PRBS9')
                res1 = self.csp.get_power_measure_res(rate=mode)
                res2 = self.csp.get_acp_res()
                logdebug('{}'.format(res1))
                logdebug('{}'.format(res2))

                nominal_pow = eval(res1[2])
                peak_pow = eval(res1[3])
                leakage_pow = eval(res1[4])
                peak_avg_pwr = eval(res1[5])
                acp_list = [eval(i) for i in res2[1:]]
                acp_center = len(acp_list)/2
                acp_max_pwr = acp_list[acp_center]
                acp_r10 = acp_list[acp_center + 10]
                acp_r9 = acp_list[acp_center + 9]
                acp_r8 = acp_list[acp_center + 8]
                acp_r7 = acp_list[acp_center + 7]
                acp_r6 = acp_list[acp_center + 6]
                acp_r5 = acp_list[acp_center + 5]
                acp_r4 = acp_list[acp_center + 4]
                acp_r3 = acp_list[acp_center + 3]
                acp_r2 = acp_list[acp_center + 2]
                acp_r1 = acp_list[acp_center + 1]
                acp_l1 = acp_list[acp_center - 1]
                acp_l2 = acp_list[acp_center - 2]
                acp_l3 = acp_list[acp_center - 3]
                acp_l4 = acp_list[acp_center - 4]
                acp_l5 = acp_list[acp_center - 5]
                acp_l6 = acp_list[acp_center - 6]
                acp_l7 = acp_list[acp_center - 7]
                acp_l8 = acp_list[acp_center - 8]
                acp_l9 = acp_list[acp_center - 9]
                acp_l10 = acp_list[acp_center - 10]
                loginfo("channel: {}    packet type: {} ".format(chan, rate))
                loginfo("nominal power: {} dBm      peak power: {} dBm      leakage power: {} dBm ".format(nominal_pow, peak_pow, leakage_pow))
                loginfo("freq accuracy: {} Hz     freq drift: {} Hz     drift rate: {} Hz".format(freq_accuracy, freq_drift, drift_rate))
                loginfo("delta f1 avg: {}  Hz     delta f1 min: {} Hz      delta f1 max: {} Hz".format(delta_f1_avg, delta_f1_min, delta_f1_max))

                if fig_en != 0:
                    plt.ion()
                    x = [2402+chan*2+xi for xi in range(-10, 11, 1)]
                    fig = plt.figure('le_acp_channel{}_{}'.format(chan,rate))
                    fig.set_size_inches(13, 7)
                    ax = fig.add_subplot(111)
                    ax.bar(x, 80+np.array(acp_list), width=0.5, bottom=-80)
                    ax.set_xticks(np.array(x))
                    ax.set_xlabel('freq(MHz)')
                    ax.set_ylabel('txp(dBm)')
                    ax.set_title('%s_ACP'%rate)
                    for xi,yi in zip(x,acp_list):
                        ax.text(xi-0.25, yi+0.25, '%.2f'%yi, fontsize=8)
                    plt.savefig(fname + 'le_acp_channel{}_{}_{}'.format(chan, rate, time.strftime('%Y%m%d_%H%M%S',time.localtime(time.time()))))
                    # plt.show()
                if report_save:
                    df = pd.DataFrame(pd.read_csv(fw2.filename))
                    df_value = [chan,rate,nominal_pow,peak_pow,peak_avg_pwr,leakage_pow,acp_l10,acp_l9,acp_l8,acp_l7,acp_l6,acp_l5,acp_l4,acp_l3,acp_l2,acp_l1,acp_max_pwr,
                                acp_r1,acp_r2,acp_r3,acp_r4,acp_r5,acp_r6,acp_r7,acp_r8,acp_r9,acp_r10,freq_accuracy,delta_f1_avg,delta_f1_min,delta_f1_max,freq_accuracy,
                                delta_f2_avg,delta_f2_min,delta_f2_max,delta_f2_99,mod_ratio,freq_accuracy,freq_offset,freq_drift,drift_rate,init_drift]

                    df['channel_{}_{}'.format(chan,rate)] = df_value
                    df.to_csv(fw2.filename,index=False)
                if csv_save:
                    fw1.write_data([chan,rate,nominal_pow,peak_pow,leakage_pow,freq_accuracy,freq_drift,drift_rate,delta_f1_avg,delta_f1_min,delta_f1_max,delta_f2_avg,delta_f2_min,delta_f2_max,mod_ratio,delta_f1_99,delta_f2_99,acp_list])
                else:
                    res_list.append([chan,rate,nominal_pow,peak_pow,leakage_pow,freq_accuracy,freq_drift,drift_rate,delta_f1_avg,delta_f1_min,delta_f1_max,delta_f2_avg,delta_f2_min,delta_f2_max,mod_ratio,delta_f1_99,delta_f2_99,acp_list])

        return res_list

    def br_tx_meas(self, chan_list=[0], rate_list=['DH1'], rx_level=-50, tx_level=10, csv_save=True, report_save=True, fig_en=0):
        '''
        rate_list:  DH1\DH3\DH5
        '''
        if csv_save :
            title = 'channel,type,nominal_pow(dBm),peak_pow(dBm),leakage_pow(dBm),freq_accuracy(kHz),freq_drift(kHz),drift_rate(Hz/50 μs),'
            title = title + 'delta_f1_avg(kHz),delta_f1_min(kHz),delta_f1_max(kHz),delta_f2_avg(kHz),delta_f2_min(kHz),delta_f2_max(kHz),mod_ratio,delta_f2_99(kHz),'
            title = title + 'obw(kHz),frange_l(kHz),frange_h(kHz),acp_list_21ch\n'
            fname = self.get_filename('ts_bt_test/','test_br_tx_{}'.format(self.board_name))
            fw1=csvreport(fname,title)
        if report_save:
            title2 = 'BR_tx_performance'
            fname2 = self.get_filename('ts_bt_test/','test_br_tx_report_{}'.format(self.board_name))
            fw2=csvreport(fname2,title2)
            item = ['channel', 'DH', u'Tx/01 - Output Power  - Average', u'- Peak',
                    u'@   - Leakage',
                    u' @   - Packet Timing',
                    u'Tx/04 - Spectrum - Frequency Range - Left',
                    u'- Right',
                    u'Tx/05 - Spectrum - 20 dB Bandwidth - f(L)',
                    u'- f(H)',
                    u'- f(H) - f(L)',
                    u'Tx/06 - Spectrum - Adjacent channel power - Max Power',
                    u'- ACP - >3 MHz',
                    u'- ACP -3 MHz',
                    u'- ACP -2 MHz',
                    u'- ACP -1 MHz',
                    u' - ACP +1 MHz',
                    u'- ACP +2 MHz',
                    u'- ACP +3 MHz',
                    u'- ACP + >3 MHz',
                    u'Tx/07 - Modulation Characteristics - Delta F1 Avg',
                    u'Tx/07 - Modulation Characteristics - Delta F2 Avg',
                    u'- Delta F2 Max 99%',
                    u'- Delta F2 Avg/Delta F1 Avg',
                    u'Tx/08 - Initial Carrier Frequency Tolerance - Avg',
                    u'- Max',
                    u'Tx/09 - Carrier Frequency Drift - Max.Drift DH1',
                    u' - Max.Drift DH5',
                    u' - Max.Drift',
                    u' - Max.Drift Rate/50us DH1',
                    u' - Max.Drift Rate/50us DH5',
                    u'- Max.Drift Rate/50us',
                    u'Harmonic Spurs']

            df = pd.DataFrame(item)
            df.to_csv(fw2.filename, index=False)
        res_list=[]
        self.config_classic_ber(rate='BR', chan=1, rx_level=rx_level, tx_power=tx_level)
        self.csp.tx_measure_para(tout=10, repetition='SINGleshot', count=20, MOEXception='OFF')
        for chan in chan_list:
            # self.config_classic_ber(rate='BR', chan=1, rx_level=rx_level, tx_power=tx_level)
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan)
            self.csp.config_rxq_br_tout(tout=10)
            for rate in rate_list:
                self.csp.config_connect_br_packet_ptype(ptype=rate)  # 设置BR包类型

                time.sleep(1)
                delta_f1_99, delta_f2_99, delta_f2_avg, delta_f2_min, delta_f2_max, mod_ratio = -999, -999, -999, -999, -999, -999
                self.csp.config_connect_br_packet_pattern(pattern='P44')  # 设置BR包 payload 数据类型
                res = self.csp.get_modulation_measure_res()
                logdebug('{}'.format(res))

                delta_f1_99 = eval(res[2]) / 1000.00
                freq_accuracy = eval(res[3]) / 1000.00
                freq_drift = eval(res[4]) / 1000.00
                drift_rate = eval(res[5])
                delta_f1_avg = eval(res[6]) / 1000.00
                delta_f1_min = eval(res[7]) / 1000.00
                delta_f1_max = eval(res[8]) / 1000.00
                time.sleep(0.5)

                self.csp.config_connect_br_packet_pattern(pattern='P11')  # 设置BR包 payload 数据类型
                res = self.csp.get_modulation_measure_res()
                logdebug('{}'.format(res))

                delta_f2_99 = eval(res[2]) / 1000.00
                freq_accuracy = eval(res[3]) / 1000.00
                freq_drift = eval(res[4]) / 1000.00
                drift_rate = eval(res[5]) / 1000.00
                delta_f2_avg = eval(res[9]) / 1000.00
                delta_f2_min = eval(res[10]) / 1000.00
                delta_f2_max = eval(res[11]) / 1000.00
                mod_ratio = delta_f2_avg / delta_f1_avg
                time.sleep(0.5)

                self.csp.config_connect_br_packet_pattern(pattern='PRBS9')  # 设置BR包 payload 数据类型
                res1 = self.csp.get_power_measure_res()
                time.sleep(1)
                res2 = self.csp.get_acp_res()
                # time.sleep(1)
                res3 = self.csp.get_obw_res()
                time.sleep(1)
                logdebug('{}'.format(res1))
                logdebug('{}'.format(res2))

                nominal_pow = eval(res1[2])
                peak_pow = eval(res1[3])
                leakage_pow = eval(res1[4])
                PacketTiming = eval(res1[5])
                acp_list = [eval(i) for i in res2[1:]]
                # obw = eval(res3[6]) / 1000.00
                # obw = -999
                acp_center = 10
                acp_max_pwr = acp_list[acp_center]
                acp_r4 = acp_list[acp_center + 4]
                acp_r3 = acp_list[acp_center + 3]
                acp_r2 = acp_list[acp_center + 2]
                acp_r1 = acp_list[acp_center + 1]
                acp_l1 = acp_list[acp_center - 1]
                acp_l2 = acp_list[acp_center - 2]
                acp_l3 = acp_list[acp_center - 3]
                acp_l4 = acp_list[acp_center - 4]
                obw = eval(res3[6])/1000.00
                obw_l = eval(res3[4])/1000.00
                obw_h = eval(res3[5])/1000.00

                self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=0)
                res4 = self.csp.get_frange_res()
                frange_l = eval(res4[3]) / 1000000.00
                self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=78)
                res4 = self.csp.get_frange_res()
                frange_h = eval(res4[4]) / 1000000.00
                self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan)
                loginfo("channel: {}    packet type: {}   ".format(chan, rate))
                loginfo("nominal power: {}      peak power: {}      leakage power: {}   unit: dBm".format(nominal_pow, peak_pow, leakage_pow))
                loginfo("freq accuracy: {}      freq drift: {}      drift rate: {}  unit: Hz".format(freq_accuracy, freq_drift, drift_rate))
                loginfo("delta f1 avg: {}       delta f1 min: {}       delta f1 max: {}  unit: Hz".format(delta_f1_avg, delta_f1_min, delta_f1_max))
                loginfo("delta f2 avg: {}       delta f2 min: {}       delta f2 max: {}  unit: Hz".format(delta_f2_avg, delta_f2_min, delta_f2_max))
                loginfo("mod ratio: {}".format(mod_ratio))
                loginfo("delta f2 99.9%: {}  unit: Hz".format(delta_f2_99))
                loginfo("spectrum 20 dB bandwidth: {}".format(obw))

                if fig_en != 0:
                    plt.ion()
                    x = [2402 + chan + xi for xi in range(-10, 11, 1)]
                    fig = plt.figure('br_acp_channel{}_{}'.format(chan, rate))
                    fig.set_size_inches(13, 7)
                    ax = fig.add_subplot(111)
                    ax.bar(x, 80 + np.array(acp_list), width=0.5, bottom=-80)
                    ax.set_xticks(np.array(x))
                    ax.set_xlabel('freq(MHz)')
                    ax.set_ylabel('txp(dBm)')
                    ax.set_title('BR_ACP')
                    for xi, yi in zip(x, acp_list):
                        ax.text(xi - 0.25, yi + 0.25, '%.2f' % yi, fontsize=8)
                    plt.savefig(fname + 'br_acp_channel{}_{}_{}'.format(chan, rate, time.strftime('%Y%m%d_%H%M%S', time.localtime(time.time()))))
                    # plt.show()
                if report_save:
                    df = pd.DataFrame(pd.read_csv(fw2.filename))
                    df_value = [chan,rate,nominal_pow,peak_pow,leakage_pow,PacketTiming,frange_l,frange_h,obw_l,obw_h,obw,acp_max_pwr,acp_l4,acp_l3,
                                                           acp_l2,acp_l1,acp_r1,acp_r2,acp_r3,acp_r4,delta_f1_avg,delta_f2_avg,delta_f2_99,mod_ratio,freq_accuracy,'',freq_drift,
                                                           '','',drift_rate,'','','']
                    print (len(df_value))
                    df['channel_{}_{}'.format(chan,rate)] = df_value
                    df.to_csv(fw2.filename,index=False)
                if csv_save:
                    fw1.write_data(
                        [chan, rate, nominal_pow, peak_pow, leakage_pow, freq_accuracy, freq_drift, drift_rate, delta_f1_avg, delta_f1_min, delta_f1_max, delta_f2_avg,
                         delta_f2_min,
                         delta_f2_max, mod_ratio, delta_f2_99, obw, acp_list])
                else:
                    res_list.append(
                        [chan, rate, nominal_pow, peak_pow, leakage_pow, freq_accuracy, freq_drift, drift_rate, delta_f1_avg, delta_f1_min, delta_f1_max, delta_f2_avg,
                         delta_f2_min,
                         delta_f2_max, mod_ratio, delta_f2_99, obw,  acp_list])
        return res_list

    def edr_tx_meas(self, chan_list=[0], rate_list=['2_DH1'], rx_level=-50, tx_level=10, csv_save=True, report_save=True, fig_en=0):
        '''
        rate_list:  2_DH1\2_DH3\2_DH5\3_DH1\3_DH3\3_DH5
        '''
        if csv_save:
            title = 'rate,channel,nominal_pwr(dBm),peak_pwr(dBm),gfsk_pwr(dBm),dpsk_pwr(dBm),dpsk_gfsk_diff_pwr,guard_period(us),'
            title = title + 'wi(KHz),w0_wi(KHz),w0_max(KHz),DEVM_RMS(%),DEVM_peak(%),DEVM_P99(%),bit_error_rate,packet0error,'
            title = title + 'PTxRef(dBm),N26ChN1Abs(dBm),N26ChP1Abs(dBm),N26ChN1Rel(dBm),N26ChP1Rel(dBm),acp_list\n'
            fname = self.get_filename('ts_bt_test/', 'test_edr_tx_{}'.format(self.board_name))
            fw1 = csvreport(fname, title)
        if report_save:
            title2 = 'BR_tx_performance'
            fname2 = self.get_filename('ts_bt_test/','test_edr_tx_report_{}'.format(self.board_name))
            fw2=csvreport(fname2,title2)
            item = ['channel', 'rate', u'Tx/10 - EDR Relative Transmit Power - PGFSK - Max Power', u' - PDPSK',
                    u'- PDPSK - PGFSK',
                    u'@   - Nomimal Power',
                    u'@   - Guard Period',
                    u'@   - Packet Timing',
                    u'Tx/11 - CFS and MA - omega i',
                    u' - omega i + omega o',
                    u'- omega o',
                    u'- DEVM RMS',
                    u'- DEVM Peak',
                    u'- DEVM 99% - Max 0.30',
                    u'Tx/12 - EDR Differential Phase Encoding',
                    u'Tx/13 - In Band Spurious Emissins - Nominal Power',
                    u' - ACPower:-3',
                    u'- ACPower:-2',
                    u'- ACPower:-1,Ptx-26dB',
                    u' - ACPower:Center Ptxref',
                    u'- ACPower:1,Ptx-26dB',
                    u'- ACPower:+2',
                    u'- ACPower:+3',
                    u'- ACPower:+- >3',]
            df = pd.DataFrame(item)
            df.to_csv(fw2.filename, index=False)
        res_list=[]
        self.config_classic_ber(rate='EDR', chan=1, rx_level=rx_level, tx_power=tx_level)
        self.csp.tx_measure_para(tout=5, repetition='SINGleshot', count=20, MOEXception='OFF')
        for chan in chan_list:
            # self.config_classic_ber(rate='EDR', chan=1, rx_level=rx_level, tx_power=tx_level)
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan)
            self.csp.config_connect_edr_packet_pattern(pattern='PRBS9')  # 设置EDR包 payload 数据类型
            for rate in rate_list:
                _rate = self.dic_edr_rate[rate]
                self.csp.config_connect_edr_packet_ptype(ptype=_rate)  # 设置EDR包类型

                res = self.csp.get_modulation_measure_res()
                logdebug('{}'.format(res))
                res = [eval(i) for i in res[0:]]
                wi = res[2]/1000
                w0_wi = res[3]/1000
                w0_max = res[4]/1000
                DEVM_RMS = res[5]
                DEVM_peak = res[6]
                DEVM_P99 = res[7]

                res1 = self.csp.get_power_measure_res()
                logdebug('{}'.format(res1))
                res1 = [eval(i) for i in res1[0:]]
                nominal_pwr = res1[2]
                gfsk_pwr = res1[3]
                dpsk_pwr = res1[4]
                dpsk_gfsk_diff_pwr = res1[5]
                guard_period = res1[6]
                packet_timing = res1[7]
                peak_pwr = res1[8]

                self.csp.RF_Frequency_Settings_tx(mode='TXT', ch_tx=chan)
                res2 = self.csp.get_diff_phase_encoding_res()
                logdebug('{}'.format(res2))
                res2 = [eval(i) for i in res2[0:]]
                bit_error_rate = res2[2]
                packet0error = res2[3]

                self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan)
                self.csp.tx_measure_para(tout=5, repetition='SINGleshot', count=100, MOEXception='OFF')
                self.csp.RF_Power_settings(rx_level=rx_level, tx_power=nominal_pwr+3, margin=8)
                [res3, res4] = self.csp.get_acp_res_edr()
                self.csp.tx_measure_para(tout=5, repetition='SINGleshot', count=20, MOEXception='OFF')
                logdebug('{}'.format(res3))
                logdebug('{}'.format(res4))
                res3 = [eval(i) for i in res3[0:]]
                res4 = [eval(i) for i in res4[0:]]
                acp_nominal_pwr = res3[2]
                PTxRef = res3[4]
                N26ChN1Abs = res3[5]
                N26ChP1Abs = res3[6]
                N26ChN1Rel = res3[7]
                N26ChP1Rel = res3[8]

                acp_list = res4[1:]
                acp_center = len(acp_list)/2
                acp_max_pwr = acp_list[acp_center]
                acp_r4 = acp_list[acp_center + 4]
                acp_r3 = acp_list[acp_center + 3]
                acp_r2 = acp_list[acp_center + 2]
                acp_r1 = acp_list[acp_center + 1]
                acp_l1 = acp_list[acp_center - 1]
                acp_l2 = acp_list[acp_center - 2]
                acp_l3 = acp_list[acp_center - 3]
                acp_l4 = acp_list[acp_center - 4]

                loginfo("channel: {}    packet type: {}     payload length: ".format(chan, rate))
                loginfo("nominal power: {}      peak power: {}      gfsk_pwr: {}    dpsk_pwr: {}".format(nominal_pwr, peak_pwr, gfsk_pwr,dpsk_pwr))
                loginfo("dpsk_gfsk_diff_pwr: {}      guard_period: {}      ".format(dpsk_gfsk_diff_pwr, guard_period))
                loginfo("wi: {}       w0_wi: {}       w0_max: {}".format(wi, w0_wi, w0_max))
                loginfo("DEVM_RMS: {}       DEVM_peak: {}       DEVM_P99: {}".format(DEVM_RMS, DEVM_peak, DEVM_P99))
                loginfo("bit_error_rate: {}     packet0error: {}".format(bit_error_rate, packet0error))
                loginfo("PTxRef: {} N26ChN1Abs: {}  N26ChP1Abs: {}  N26ChN1Rel: {}  N26ChP1Rel:{}".format(PTxRef, N26ChN1Abs, N26ChP1Abs, N26ChN1Rel, N26ChP1Rel))
                if report_save:
                    df = pd.DataFrame(pd.read_csv(fw2.filename))
                    df_value = [chan,rate,peak_pwr,dpsk_pwr,dpsk_gfsk_diff_pwr,nominal_pwr,guard_period,packet_timing,wi,w0_wi,w0_max,DEVM_RMS,DEVM_peak,DEVM_P99,packet0error,
                                acp_nominal_pwr,acp_l3,acp_l2,N26ChN1Rel,PTxRef,N26ChP1Rel,acp_r2,acp_r3,min([acp_l4,acp_r4])]

                    df['channel_{}_{}'.format(chan,rate)] = df_value
                    df.to_csv(fw2.filename,index=False)
                if csv_save:
                    if fig_en != 0:
                        plt.ion()
                        x = [2402 + chan + xi for xi in range(-10, 11, 1)]
                        fig = plt.figure('edr_acp_channel{}_{}'.format(chan, rate))
                        fig.set_size_inches(13, 7)
                        ax = fig.add_subplot(111)
                        ax.bar(x, 80 + np.array(acp_list), width=0.5, bottom=-80)
                        ax.set_xticks(np.array(x))
                        ax.set_xlabel('freq(MHz)')
                        ax.set_ylabel('txp(dBm)')
                        ax.set_title('%s_ACP' % rate)
                        for xi, yi in zip(x, acp_list):
                            ax.text(xi - 0.25, yi + 0.25, '%.2f' % yi, fontsize=8)
                        plt.savefig(fname + 'edr_acp_channel{}_{}_{}'.format(chan, rate, time.strftime('%Y%m%d_%H%M%S',
                                                                                                       time.localtime(
                                                                                                           time.time()))))
                        # plt.show()
                    fw1.write_data([rate,chan,nominal_pwr,peak_pwr,gfsk_pwr,dpsk_pwr,dpsk_gfsk_diff_pwr,guard_period,wi,w0_wi,w0_max,DEVM_RMS,DEVM_peak,DEVM_P99,bit_error_rate,packet0error,PTxRef,N26ChN1Abs,N26ChP1Abs,N26ChN1Rel,N26ChP1Rel,acp_list])
                else:
                    res_list.append([rate,chan,nominal_pwr,peak_pwr,gfsk_pwr,dpsk_pwr,dpsk_gfsk_diff_pwr,guard_period,wi,w0_wi,w0_max,DEVM_RMS,DEVM_peak,DEVM_P99,bit_error_rate,packet0error,PTxRef,N26ChN1Abs,N26ChP1Abs,N26ChN1Rel,N26ChP1Rel,acp_list])

        return res_list


    def bedr_rx_sens(self, chan_list=[0,38,78], rate_list=['DH1'], rx_level=-60, tx_power=10, packet_num=2000, step=0.5, csv_save=True):
        if csv_save:
            title = 'rate,channel,ber,per,NAK,hec_err,crc_err,packet_type_err,pay_err,sens\n'
            fname = self.get_filename('ts_bt_test/', 'br_rx_sens_{}'.format(self.board_name))
            fw1 = csvreport(fname, title)
        for chan in chan_list:
            self.config_classic_ber(rate='BR',chan=chan, rx_level=rx_level, tx_power=tx_power)
            self.csp.config_rx_level(rxpwr=rx_level)
            self.csp.config_rxq_br_search_packets(num=packet_num)
            self.csp.config_rxq_br_search_step(step=step)
            # self.csp.config_rxq_br_search_tout(tout=20)
            for rate in rate_list:
                if rate in self.br_rate_list:
                    btype = 'BR'
                    self.csp.config_connect_br_packet_ptype(ptype=rate)  # 设置BR包类型
                elif rate in self.edr_rate_list:
                    btype = 'EDR'
                    rate = self.dic_edr_rate[rate]
                    self.csp.config_connect_edr_packet_ptype(ptype=rate)  # 设置EDR的包类型，E21P：2-DH1，E31P：3-DH1
                self.csp.sig_btype(btype=btype)
                self.csp.ber_search_meas_state(state=0)
                while 1:
                    state_res = self.csp.get_ber_search_meas_state()    #ber search 完成则退出循环
                    # logdebug(state_res)
                    if state_res == 'RDY':
                        loginfo('search complete')
                        break
                    elif state_res == 'OFF':
                        self.cmd_br_connect()
                        self.csp.ber_search_meas_state(state=0)
                    else:
                        logdebug('searching')
                    time.sleep(5)

                res = self.csp.meas_bt_ber_search_res(cmd_type='FETCh')
                ber = eval(res[1])
                per = eval(res[2])
                NAK = eval(res[5])
                hec_err = eval(res[6])
                crc_err = eval(res[7])
                packet_type_err = eval(res[8])
                pay_err = eval(res[9])
                sens = eval(res[11])
                loginfo('{} {}  {}  {}  {}  {}  {}  {}'.format(ber,per,NAK,hec_err,crc_err,packet_type_err,pay_err,sens))
                if csv_save:
                    fw1.write_data([rate,chan,ber,per,NAK,hec_err,crc_err,packet_type_err,pay_err,sens])

    def bedr_rx_range(self, chan_list=[0,38,78], rate_list=['DH1'], rx_level=-60, tx_power=10, packet_num=2000, rxpwr_range=range(-98,0),  csv_save=True):
        if csv_save:
            title = 'rate,channel,ber,per,NAK,hec_err,crc_err,packet_type_err,pay_err,sens\n'
            fname = self.get_filename('ts_bt_test/', 'bedr_rx_range_{}'.format(self.board_name))
            fw1 = csvreport(fname, title)

        self.config_classic_ber(rate='BR', chan=1, rx_level=rx_level, tx_power=tx_power)
        self.csp.config_rx_level(rxpwr=rx_level)
        self.csp.config_rxq_br_packets(num=packet_num)
        self.csp.config_rxq_br_tout(tout=10)
        for chan in chan_list:
            self.csp.config_rx_level(rxpwr=rx_level)
            self.csp.RF_Frequency_Settings_rx(ch_tx=chan)
            for rate in rate_list:
                if rate in self.br_rate_list:
                    btype = 'BR'
                    self.csp.config_connect_br_packet_ptype(ptype=rate)  # 设置BR包类型
                elif rate in self.edr_rate_list:
                    btype = 'EDR'
                    rate = self.dic_edr_rate[rate]
                    self.csp.config_connect_edr_packet_ptype(ptype=rate)  # 设置EDR的包类型，E21P：2-DH1，E31P：3-DH1
                self.csp.sig_btype(btype=btype)
                self.csp.ber_meas_state(0)
                for rxpwr in rxpwr_range:
                    self.csp.config_rx_level(rxpwr=rxpwr)
                    # while 1:
                    #     state_res = self.csp.get_ber_search_meas_state()
                    #     if state_res == 'RDY':
                    #         logdebug(state_res)
                    #         break
                    #     time.sleep(3)
                    # delay = 0.001*packet_num*(rx_level+100)/step
                    # time.sleep(delay)
                    while 1:
                        res = self.csp.meas_bt_ber_res(cmd_type='READ')
                        if eval(res[0]) == 0:
                            break
                        else:
                            if self.csp.get_bt_connect_state() == 'SBY':
                                self.cmd_br_connect()

                    ber = eval(res[1])
                    per = eval(res[2])
                    NAK = eval(res[5])
                    hec_err = eval(res[6])
                    crc_err = eval(res[7])
                    packet_type_err = eval(res[8])
                    pay_err = eval(res[9])

                    time.sleep(0.5)
                    loginfo('{} {}  {}  {}  {}  {}  {}  {}'.format(ber,per,NAK,hec_err,crc_err,packet_type_err,pay_err,rxpwr))
                    if csv_save:
                        fw1.write_data([rate,chan,ber,per,NAK,hec_err,crc_err,packet_type_err,pay_err,rxpwr])

    def bedr_rx_ci(self, chan_list=[0,38,78], rate_list=['DH1'], cable_loss=4.6,rx_level=-60, tx_power=10, packet_num=2000,  csv_save=True):
        self.dict_pwr_inter_freqoffset = {
            0: range(-85, -40),
            1: range(-65, -10),
            2: range(-50, 0),
            3: range(-50, 0),
            4: range(-50, 0),
            -1: range(-65, -10),
            -2: range(-60, 0),
            -3: range(-50, 0),
            -4: range(-50, 0),
            -5: range(-50, 0),
            -6: range(-50, 0),
            -7: range(-50, 0),
            -8: range(-50, 0)
            }

        if csv_save:
            title = 'signal_channel,rate,freq_interference,rxpwr_interference,ber(%), per(%),agc_gain_index,agc_tgt,agc_sat\n'
            fname = self.get_filename('ts_bt_test/','bedr_rx_ci_data_{}'.format(self.board_name))
            fw1 = csvreport(fname,title)

            title1 = 'signal_channel,rate,freq_interference,rxpwr_interference(dbm),ci(db),agc_tgt,agc_sat\n'
            fname1 = self.get_filename('ts_bt_test/','bedr_rx_ci_report_{}'.format(self.board_name))
            fw2 = csvreport(fname1,title1)

        self.config_classic_ber(rate='BR', chan=1, rx_level=rx_level, tx_power=tx_power)
        self.csp.config_rx_level(rxpwr=rx_level)
        self.csp.config_rxq_br_packets(num=packet_num)
        self.csp.config_rxq_br_tout(tout=10)

        for chan in chan_list:
            self.csp.config_rx_level(rxpwr=-60)
            self.csp.RF_Frequency_Settings_rx(mode='LOOP', ch_tx=chan)
            for rate in rate_list:
                self.tester_inter = mxg.MXG(device_name='N5182B', num_of_machine=1)
                self.tester_inter.arb_waveform(rate=rate+'_PN15')
                self.tester_inter.trriger_para_set(type='CONT', count=2000)
                self.tester_inter.output_state(1, 1)
                self.tester_inter.arb_state(1)

                if rate in self.br_rate_list:
                    btype = 'BR'
                    self.csp.config_connect_br_packet_ptype(ptype=rate)  # 设置BR包类型
                elif rate in self.edr_rate_list:
                    btype = 'EDR'
                    _rate = self.dic_edr_rate[rate]
                    self.csp.config_connect_edr_packet_ptype(ptype=_rate)  # 设置EDR的包类型，E21P：2-DH1，E31P：3-DH1
                self.csp.sig_btype(btype=btype)
                for freq_inter_offset in[0,1,-1,2,-2,3,-3,4,-4]:
                    freq_inter = freq_inter_offset+2402+chan
                    if freq_inter_offset == 0 and btype == 'EDR':
                        self.tester_inter.arb_waveform(rate=rate + '_PN15')
                    else:
                        self.tester_inter.arb_waveform(rate='DH1_PN15')
                    if -3 < freq_inter_offset < 3:
                        self.csp.config_rx_level(rxpwr=-60)
                    else:
                        self.csp.config_rx_level(rxpwr=-67)
                    pwr_list = self.dict_pwr_inter_freqoffset[freq_inter_offset]
                    for power_inter in pwr_list:
                        self.tester_inter.para_set(freq=freq_inter, power=power_inter + cable_loss)
                        # time.sleep(1)
                        loginfo('freq_inter:    {}      pwr_inter:  {}'.format(freq_inter,power_inter))

                        while 1:
                            res = self.csp.meas_bt_ber_res(cmd_type='READ')

                            if eval(res[0]) == 0 or eval(res[0]) == 1:
                                break
                            else:
                                if self.csp.get_bt_connect_state() == 'SBY':
                                    self.tester_inter.para_set(freq=freq_inter, power=power_inter - 1 + cable_loss)
                                    self.cmd_br_connect()
                                    time.sleep(10)
                                    res = [-999,-999,-999,-999,-999,-999,-999,-999,-999,-999,-999,-999]
                                    break
                        ber = eval(res[1])
                        per = eval(res[2])
                        NAK = eval(res[5])
                        hec_err = eval(res[6])
                        crc_err = eval(res[7])
                        packet_type_err = eval(res[8])
                        pay_err = eval(res[9])

                        loginfo('{} {}  {}  {}  {}  {}  {} '.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err))

                        fw1.write_data([chan, rate, freq_inter, power_inter,  ber, per])
                        ber_min = 0.1
                        if btype == 'EDR':
                            ber_min = 0.01
                        if ber > ber_min or per > 90 or ber == -999:
                            if -3 < freq_inter_offset < 3:
                                ci = -60 - power_inter +1
                            else:
                                ci = -67 - power_inter + 1
                            fw2.write_data([chan, rate, freq_inter, power_inter-1, ci])
                            break
        self.tester_inter.output_state(0, 0)

    def bedr_rx_ci_debug(self, chan_list=[0,38,78], rate_list=['DH1'], cable_loss=4.6,rx_level=-60, tx_power=10, packet_num=2000,  csv_save=True, name_str='cbpf1',
                   cbpf2_en=0,agc_tgt_sat_list=[[300,400]], loopmax=1):
        self.dict_pwr_inter_freqoffset = {
            0: range(-85, -40),
            1: range(-65, -10),
            2: range(-50, 0),
            3: range(-50, 0),
            4: range(-50, 0),
            -1: range(-65, -10),
            -2: range(-60, 0),
            -3: range(-50, 0),
            -4: range(-50, 0),
            -5: range(-50, 0),
            -6: range(-50, 0),
            -7: range(-50, 0),
            -8: range(-50, 0)
            }

        if csv_save:
            title = 'signal_channel,rate,freq_interference,rxpwr_interference,ber(%), per(%),agc_gain_index,agc_tgt,agc_sat\n'
            fname = self.get_filename('ts_bt_test/','bedr_rx_ci_data_{}'.format(name_str))
            fw1 = csvreport(fname,title)

            title1 = 'signal_channel,rate,freq_interference,rxpwr_interference(dbm),ci(db),agc_tgt,agc_sat\n'
            fname1 = self.get_filename('ts_bt_test/','bedr_rx_ci_report_{}'.format(name_str))
            fw2 = csvreport(fname1,title1)
        if self.chipv.find('TX232') != -1:
            if cbpf2_en !=0:
                self.jlink.wrm(0xa0421020,20,20,1)
            else:
                self.jlink.wrm(0xa0421020, 20, 20, 0)

        self.config_classic_ber(rate='BR', chan=1, rx_level=rx_level, tx_power=tx_power)
        self.csp.config_rx_level(rxpwr=rx_level)
        self.csp.config_rxq_br_packets(num=packet_num)
        self.csp.config_rxq_br_tout(tout=10)
        for loop in range(0,loopmax):
            loop = loop+1
            for cbpf_bw in range(3,-1,-1):
                self.jlink.wrm(0xa0421034, 13, 12, cbpf_bw)
                for cbpf_bias_trim in range(0,8):
                    self.jlink.wrm(0xa0421034, 18, 16, cbpf_bias_trim)
                    for chan in chan_list:
                        self.csp.config_rx_level(rxpwr=-60)
                        self.csp.RF_Frequency_Settings_rx(mode='LOOP', ch_tx=chan)
                        for rate in rate_list:
                            self.tester_inter = mxg.MXG(device_name='N5182B', num_of_machine=1)
                            self.tester_inter.arb_waveform(rate=rate+'_PN15')
                            self.tester_inter.trriger_para_set(type='CONT', count=2000)
                            self.tester_inter.output_state(1, 1)
                            self.tester_inter.arb_state(1)

                            if rate in self.br_rate_list:
                                btype = 'BR'
                                self.csp.config_connect_br_packet_ptype(ptype=rate)  # 设置BR包类型
                            elif rate in self.edr_rate_list:
                                btype = 'EDR'
                                _rate = self.dic_edr_rate[rate]
                                self.csp.config_connect_edr_packet_ptype(ptype=_rate)  # 设置EDR的包类型，E21P：2-DH1，E31P：3-DH1
                            self.csp.sig_btype(btype=btype)
                            for agc_tgt_sat in agc_tgt_sat_list:
                                agc_tgt = agc_tgt_sat[0]
                                agc_sat = agc_tgt_sat[1]
                                if self.chipv == 'TX232':
                                    self.jlink.wrm(0xa04210ac,9,0,agc_tgt)
                                    self.jlink.wrm(0xa04210a8,11,2,agc_sat)

                                for freq_inter_offset in[0,1,-1,2,-2,3,-3,4,-4]:
                                    freq_inter = freq_inter_offset+2402+chan
                                    if freq_inter_offset == 0 and btype == 'EDR':
                                        self.tester_inter.arb_waveform(rate=rate + '_PN15')
                                    else:
                                        self.tester_inter.arb_waveform(rate='DH1_PN15')
                                    if -3 < freq_inter_offset < 3:
                                        self.csp.config_rx_level(rxpwr=-60)
                                    else:
                                        self.csp.config_rx_level(rxpwr=-67)
                                    pwr_list = self.dict_pwr_inter_freqoffset[freq_inter_offset]
                                    for power_inter in pwr_list:
                                        self.tester_inter.para_set(freq=freq_inter, power=power_inter + cable_loss)
                                        # time.sleep(1)
                                        loginfo('freq_inter:    {}      pwr_inter:  {}'.format(freq_inter,power_inter))
                                        agc_gain_index = self.jlink.rdm(0xa04210dc, 11, 8)
                                        while 1:
                                            res = self.csp.meas_bt_ber_res(cmd_type='READ')

                                            if eval(res[0]) == 0 or eval(res[0]) == 1:
                                                break
                                            else:
                                                if self.csp.get_bt_connect_state() == 'SBY':
                                                    self.tester_inter.para_set(freq=freq_inter, power=power_inter - 1 + cable_loss)
                                                    self.cmd_br_connect()
                                                    time.sleep(10)
                                                    res = [-999,-999,-999,-999,-999,-999,-999,-999,-999,-999,-999,-999]
                                                    break
                                        ber = eval(res[1])
                                        per = eval(res[2])
                                        NAK = eval(res[5])
                                        hec_err = eval(res[6])
                                        crc_err = eval(res[7])
                                        packet_type_err = eval(res[8])
                                        pay_err = eval(res[9])

                                        loginfo('{} {}  {}  {}  {}  {}  {} '.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err))

                                        fw1.write_data([chan, rate, freq_inter, power_inter,  ber, per, agc_gain_index,agc_tgt,agc_sat])
                                        ber_min = 0.1
                                        if btype == 'EDR':
                                            ber_min = 0.01
                                        if ber > ber_min or per > 90 or ber == -999:
                                            if -3 < freq_inter_offset < 3:
                                                ci = -60 - power_inter +1
                                            else:
                                                ci = -67 - power_inter + 1
                                            fw2.write_data([chan, rate, freq_inter, power_inter-1, ci,agc_tgt,agc_sat])
                                            break
        self.tester_inter.output_state(0, 0)

    def bedr_rx_blocking(self, rate_list=['DH1'],  packet_num=1500, inter_loss=11, inter_rxpwr=[-30], freq_range=[]):
        title = 'rate,blocking_level,freq_blocking_1,freq_blocking_2,ber,per(%)\n'
        fname = self.get_filename('ts_bt_test/', 'bedr_rx_blocking_{}'.format(self.board_name))
        fw1 = csvreport(fname, title)
        self.tester_inter = mxg.MXG(device_name='N5182B',num_of_machine=1)
        self.tester_inter.para_set(30,-30+inter_loss)
        self.tester_inter.output_state(1, 0)

        self.config_classic_ber(rate='BR', chan=58, rx_level=-60, tx_power=10)
        self.csp.config_rxq_br_packets(num=packet_num)
        self.csp.config_rxq_br_tout(tout=10)
        self.csp.config_rx_level(rxpwr=-67)
        self.csp.RF_Frequency_Settings_rx(mode='LOOP', ch_tx=58)

        for rate in rate_list:
            self.csp.config_rx_level(rxpwr=-67)
            if rate in self.br_rate_list:
                btype = 'BR'
                self.csp.config_connect_br_packet_ptype(ptype=rate)  # 设置BR包类型
            elif rate in self.edr_rate_list:
                btype = 'EDR'
                _rate = self.dic_edr_rate[rate]
                self.csp.config_connect_edr_packet_ptype(ptype=_rate)  # 设置EDR的包类型，E21P：2-DH1，E31P：3-DH1
            self.csp.sig_btype(btype=btype)

            # freq_range= range(30,2401,1)+range(2500,6000,1)

            # for rxpwr in inter_rxpwr:
            rxpwr = inter_rxpwr
            freq_inter_list = []
            freq_inter_list2 = []
            for freq_inter in freq_range:
                freq_blocking_2 = 'NA'
                freq_blocking_1 = 'NA'

                if freq_inter < 2001 or freq_inter > 2999:
                    pwr = rxpwr + 17 + inter_loss
                else:
                    pwr = rxpwr + inter_loss
                self.csp.config_rx_level(rxpwr=-67)
                connect_num = 0
                while 1:
                    self.tester_inter.para_set(freq=freq_inter, power=pwr)

                    loginfo('freq_inter:    {}'.format(freq_inter))

                    res = self.csp.meas_bt_ber_res(cmd_type='READ')
                    if eval(res[0]) == 0 or eval(res[0]) == 1:
                        break
                    else:
                        connect_num = connect_num+1
                        loginfo('reconnect num {}'.format(connect_num))
                        if self.csp.get_bt_connect_state() == 'SBY':
                            self.cmd_br_connect()
                            time.sleep(10)
                            if connect_num > 3:
                                res = ['0','100','100','100','100','100','100','100','100','100','100']
                                break
                ber = eval(res[1])
                per = eval(res[2])
                NAK = eval(res[5])
                hec_err = eval(res[6])
                crc_err = eval(res[7])
                packet_type_err = eval(res[8])
                pay_err = eval(res[9])
                loginfo('freq_blocking_1:   {}      level:  {}      ber:    {}      per:    {}'.format(freq_inter, inter_rxpwr, ber, per))
                if ber > 0.1 or per > 90:
                    loginfo('freq_blocking_1:   {}      level:  {}      ber:    {}      per:    {}'.format(freq_inter,inter_rxpwr,ber,per))
                    freq_blocking_1 = freq_inter
                    freq_inter_list.append(freq_inter)
                    fw1.write_data([rate,pwr-inter_loss,freq_blocking_1,freq_blocking_2,ber,per])


            for freq_inter2 in freq_inter_list:
                freq_blocking_2 = 'NA'
                freq_blocking_1 = 'NA'

                if freq_inter2 < 2001 or freq_inter2 > 2999:
                    pwr = rxpwr + 17 + inter_loss
                else:
                    pwr = rxpwr + inter_loss
                self.csp.config_rx_level(rxpwr=-50)
                connect_num = 0
                while 1:
                    self.tester_inter.para_set(freq=freq_inter2, power=pwr)

                    loginfo('freq_inter:    {}'.format(freq_inter2))

                    res = self.csp.meas_bt_ber_res(cmd_type='READ')
                    if eval(res[0]) == 0 or eval(res[0]) == 1:
                        break
                    else:
                        connect_num = connect_num+1
                        loginfo('reconnect num {}'.format(connect_num))
                        if self.csp.get_bt_connect_state() == 'SBY':
                            self.cmd_br_connect()
                            time.sleep(10)
                            if connect_num > 3:
                                res = ['0','100','100','100','100','100','100','100','100','100','100']
                                break
                ber = eval(res[1])
                per = eval(res[2])
                NAK = eval(res[5])
                hec_err = eval(res[6])
                crc_err = eval(res[7])
                packet_type_err = eval(res[8])
                pay_err = eval(res[9])
                loginfo('freq_blocking_1:   {}      level:  {}      ber:    {}      per:    {}'.format(freq_inter2, inter_rxpwr, ber, per))
                if ber > 0.1 or per > 90:
                    loginfo('freq_blocking_1:   {}      level:  {}      ber:    {}      per:    {}'.format(freq_inter2,inter_rxpwr,ber,per))
                    freq_blocking_2 = freq_inter2
                    freq_inter_list2.append(freq_inter2)

                    fw1.write_data([rate,pwr-inter_loss,freq_blocking_1,freq_blocking_2,ber,per])
            #     self.csp.config_rx_level(rxpwr=-50)
            #     connect_num = 0
            #     while 1:
            #         self.tester_inter.para_set(freq=freq_inter, power=pwr)
            #
            #         loginfo('freq_inter:    {}'.format(freq_inter))
            #
            #         res = self.csp.meas_bt_ber_res(cmd_type='READ')
            #         if eval(res[0]) == 0 or eval(res[0]) == 1:
            #             break
            #         else:
            #             connect_num = connect_num+1
            #             loginfo('reconnect num {}'.format(connect_num))
            #             if self.csp.get_bt_connect_state() == 'SBY':
            #                 self.cmd_br_connect()
            #                 if connect_num > 3:
            #                     res = [0, 100, 100, 100, 100, 100, 100, 100, 100, 100, 100]
            #                     break
            #     ber = eval(res[1])
            #     per = eval(res[2])
            #     NAK = eval(res[5])
            #     hec_err = eval(res[6])
            #     crc_err = eval(res[7])
            #     packet_type_err = eval(res[8])
            #     pay_err = eval(res[9])
            #     loginfo('freq_blocking_2:   {}      level:  {}      ber:    {}      per:    {}'.format(freq_inter, pwr-inter_loss, ber, per))
            #     if ber > 0.1 or per > 90:
            #         loginfo('freq_blocking_2:   {}      level:  {}      ber:    {}      per:    {}'.format(freq_inter,pwr-inter_loss,ber,per))
            #         freq_blocking_2 = freq_inter
            #         # if freq_inter < 2001 or freq_inter > 2999:
            #         #     pwr = rxpwr + 17 + inter_loss
            #         fw1.write_data([rate,pwr,freq_blocking_1,freq_blocking_2,ber,per])
            #         count1 = count1 + 1
            #     # fw1.write_data([rate, pwr - inter_loss, freq_blocking_1, freq_blocking_2, ber, per])
            # if count > 24 or count1 > 5:
            #     self.tester_inter.output_state(0, 0)
            #     break
            # if freq_blocking_1 == 'NA':
            #     fw1.write_data([rate, pwr-inter_loss, freq_blocking_1, freq_blocking_2, ber,per])

        self.tester_inter.output_state(0, 0)
        return freq_inter_list+freq_inter_list2

    def le_rx_blocking(self, rate_list=['LE1M'], packet_len=37, packet_num=1500, inter_loss=11, inter_rxpwr=-10, freq_range_le=[]):
        title = 'rate,blocking_level,freq_blocking_1,freq_blocking_2,per(%)\n'
        fname = self.get_filename('ts_bt_test/', 'le_rx_blocking_{}'.format(self.board_name))
        fw1 = csvreport(fname, title)
        self.tester_inter = mxg.MXG(device_name='N5182B',num_of_machine=1)
        self.tester_inter.para_set(30,-30+inter_loss)
        self.tester_inter.output_state(1, 0)
        for rate in rate_list:
            if rate in ['LE_1M','LE1M']:
                _rate = 'LE1M'
            else:
                _rate = 'LE2M'
            self.config_per_le_connect_usb(rate=_rate,packet_len=packet_len)
            self.csp.RF_Frequency_Settings_rx(mode='DTM', ch_tx=19)
            self.csp.config_rxq_le_packets(rate=_rate, num=packet_num)
            self.csp.config_rx_level(rxpwr=-67)
            # freq_range= range(30,2001,10)+range(2003,2400,1)+range(2484,2998,1)+range(3000,6000,25)
            freq_range = freq_range_le
            freq_blocking_2 = 'NA'
            freq_blocking_1 = 'NA'
            freq_inter_list = []
            # for rxpwr in inter_rxpwr:
            count = 0
            count1 = 0
            for freq_inter in freq_range:
                freq_blocking_2 = 'NA'
                freq_blocking_1 = 'NA'
                # if freq_inter < 2001 or freq_inter > 2999:
                #     pwr =  + 5  + inter_loss
                #     # self.tester_inter.para_set(freq=freq_inter, power=rxpwr + inter_loss)
                # else:
                #     pwr = rxpwr + inter_loss
                #     # self.tester_inter.para_set(freq=freq_inter, power=rxpwr-5 + inter_loss)
                pwr = inter_rxpwr + inter_loss
                self.tester_inter.para_set(freq=freq_inter, power=pwr)
                loginfo('freq_inter:    {}'.format(freq_inter))
                self.csp.config_rx_level(rxpwr=-67)
                self.csp.per_meas_state(0)
                time.sleep(1)
                res = self.csp.meas_le_per_res(rate=_rate)
                per = eval(res[1])
                if per > 30.8:
                    loginfo('freq_blocking_1:   {}      level:  {}      per:    {}'.format(freq_inter,inter_rxpwr,per))
                    # freq_blocking_1 = freq_inter
                    # if freq_inter > 1999 or freq_inter < 3001:
                    #     # pwr = rxpwr - 5 + inter_loss
                    #     rxpwr = rxpwr - 5
                    freq_inter_list.append(freq_inter)
                    # fw1.write_data([rate,rxpwr,freq_blocking_1,freq_blocking_2,per])
                    count = count + 1
            for freq_inter_check in freq_inter_list:
                pwr = inter_rxpwr + inter_loss
                self.tester_inter.para_set(freq=freq_inter_check, power=pwr)
                loginfo('freq_inter:    {}'.format(freq_inter_check))
                self.csp.per_meas_state(0)
                time.sleep(1)
                res = self.csp.meas_le_per_res(rate=_rate)
                per = eval(res[1])
                if per > 30.8:
                    loginfo('freq_blocking_1:   {}      level:  {}      per:    {}'.format(freq_inter_check,inter_rxpwr,per))
                    freq_blocking_1 = freq_inter_check
                    fw1.write_data([rate, inter_rxpwr, freq_blocking_1, freq_blocking_2, per])
                    # self.csp.config_rx_level(rxpwr=-50)
                    # self.csp.per_meas_state(0)
                    # time.sleep(1)
                    # res = self.csp.meas_le_per_res(rate=_rate)
                    # per = eval(res[1])
                    # if per > 30.8:
                    #     loginfo('freq_blocking_2:   {}      level:  {}      per:    {}'.format(freq_inter_check,rxpwr,per))
                    #     freq_blocking_2 = freq_inter
                    #     # pwr = rxpwr
                    #     # if freq_inter > 1999 or freq_inter < 3001:
                    #     #     rxpwr = rxpwr - 5
                    #     #     pwr = rxpwr - 5 + inter_loss
                    #     fw1.write_data([rate,rxpwr,freq_blocking_1,freq_blocking_2,per])
                    #     count1 = count1 + 1

                # if count > 10 or count1 > 3:
                #     self.tester_inter.output_state(0, 0)
                #     break
            # if freq_blocking_1 == 'NA':
            #     fw1.write_data([rate, pwr-inter_loss, freq_blocking_1, freq_blocking_2, per])


    def le_rx_ci(self,  cable_loss=5, chan_list=[0], rate_list=['LE_1M'], packet_len=37, pkt_num=2000,  csv_save=True, cbpf2_en=0,agc_tgt_sat_list=[[300,400]]):
        self.dict_pwr_inter_freqoffset = {
            0: range(-80, -50),
            1: range(-75, -20),
            2: range(-52, -10),
            3: range(-41, 0),
            -1: range(-75, -20),
            -2: range(-52, -10),
            -3: range(-41, 0)
            }

        if csv_save:
            title = 'signal_channel,rate,freq_interference,rxpwr_interference,per(%),agc_gain_index,agc_tgt,agc_sat\n'
            fname = self.get_filename('ts_bt_test/','le_rx_ci_{}'.format(self.board_name))
            fw1 = csvreport(fname,title)

            title1 = 'signal_channel,rate,freq_interference,rxpwr_interference(dbm),ci(db),agc_tgt,agc_sat\n'
            fname1 = self.get_filename('ts_bt_test/','le_rx_ci_report_{}'.format(self.board_name))
            fw2 = csvreport(fname1,title1)
        if self.chipv == 'TX232':
            if cbpf2_en !=0:
                self.jlink.wrm(0xa0421020,20,20,1)
            else:
                self.jlink.wrm(0xa0421020, 20, 20, 0)
        for chan in chan_list:
            for rate in rate_list:
                self.tester_inter = mxg.MXG(device_name='N5182B', num_of_machine=1)
                self.tester_inter.arb_waveform(rate=rate+'_PN15')
                self.tester_inter.trriger_para_set(type='CONT', count=pkt_num)
                self.tester_inter.output_state(1, 1)
                self.tester_inter.arb_state(1)

                if rate == 'LE_1M':
                    _rate = 'LE1M'
                else:
                    _rate = 'LE2M'
                self.config_per_le_connect_usb(rate=_rate, packet_len=packet_len)
                self.csp.RF_Frequency_Settings_rx(mode='DTM', ch_tx=chan)
                self.csp.config_rxq_le_packets(rate=_rate, num=pkt_num)
                self.csp.config_rx_level(rxpwr=-67)
                time.sleep(1)
                for agc_tgt_sat in agc_tgt_sat_list:
                    agc_tgt = agc_tgt_sat[0]
                    agc_sat = agc_tgt_sat[1]
                    if self.chipv == 'TX232':
                        self.jlink.wrm(0xa04210ac,9,0,agc_tgt)
                        self.jlink.wrm(0xa04210a8,11,2,agc_sat)
                    for freq_inter_offset in[0,1,-1,2,-2,3,-3]:
                        if rate == 'LE_2M':
                            freq_inter = freq_inter_offset*2 + 2402 + chan * 2
                        else:
                            freq_inter = freq_inter_offset+2402+chan*2
                        pwr_list = self.dict_pwr_inter_freqoffset[freq_inter_offset]
                        for power_inter in pwr_list:
                            self.tester_inter.para_set(freq=freq_inter, power=power_inter + cable_loss)
                            time.sleep(0.2)
                            loginfo('freq_inter:    {}      pwr_inter:  {}'.format(freq_inter,power_inter))
                            agc_gain_index = self.jlink.rdm(0xa04210dc,11,8)
                            per_list=[]
                            for loop in range(10):
                                res = self.csp.meas_le_per_res(rate=_rate)
                                # time.sleep(1)
                                per = eval(res[1])
                                per_list.append(per)
                                time.sleep(0.5)
                            loginfo(per_list)
                            per = max(per_list)
                            loginfo('{}     {}      {}      {}      {}'.format(chan, rate, freq_inter, power_inter,  per))
                            fw1.write_data([chan, rate, freq_inter, power_inter,  per, agc_gain_index, agc_tgt, agc_sat])

                            if per > 30:
                                ci = -67 - power_inter +1
                                fw2.write_data([chan, rate, freq_inter, power_inter, ci, agc_tgt, agc_sat])
                                break


        self.tester_inter.output_state(0, 0)

    def test_le_sens(self, chan_list=[], rate_list=[], rx_pwr_range=[-99,-90], cable_loss=5, csv_save=True):
        if csv_save:
            title = 'rate,channel,sensitivity,per(%)\n'
            fname = self.get_filename('ts_bt_test/','test_le_per_{}'.format(self.board_name))
            fw1 = csvreport(fname,title)
        for rate in rate_list:
            if rate == 'LE_1M' or rate == 'LE1M':
                _rate = 'LE1M'
            else:
                _rate = 'LE2M'
            self.config_per_le_connect_usb(rate=_rate, packet_len=37)
            self.csp.RF_Frequency_Settings_rx(mode='DTM', ch_tx=0)
            self.csp.config_rxq_le_packets(rate=_rate, num=1500)
            self.csp.config_rx_level(rxpwr=-67)
            self.csp.RF_Power_settings_autoranging(0)  # 打开autorange会导致接收per 概率性100%
            time.sleep(1)
            for chan in chan_list:
                self.csp.RF_Frequency_Settings_rx(mode='DTM', ch_tx=chan)
                for rxpwr in range(rx_pwr_range[0],rx_pwr_range[1],1):
                    self.csp.config_rx_level(rxpwr=rxpwr)

                    per_list = []
                    for loop in range(5):
                        res = self.csp.meas_le_per_res(rate=_rate)
                        time.sleep(1)
                        per = eval(res[1])
                        per_list.append(per)

                    per = max(per_list)
                    if per < 30.8:
                        break
                fw1.write_data([rate,chan,rxpwr,per])



    def le_rx_intermod(self,  cable_loss=5, chan_list=[0], rate_list=['LE_1M'], packet_len=37, pkt_num=2000,  csv_save=True):
        self.dict_pwr_inter_freqoffset = {
            0: range(-90, -60),
            1: range(-85, -50),
            2: range(-52, -30),
            3: range(-41, -18),
            -1: range(-85, -50),
            -2: range(-52, -30),
            -3: range(-41, -18)
            }

        if csv_save:
            title = 'signal_channel,rate,freq_interference_tone,freq_interference_mod,rxpwr_interference,per(%)\n'
            fname = self.get_filename('ts_bt_test/','le_rx_ci_{}'.format(self.board_name))
            fw1 = csvreport(fname,title)
        for chan in chan_list:
            for rate in rate_list:
                if rate == 'LE_1M':
                    _rate = 'LE1M'
                else:
                    _rate = 'LE2M'
                self.tester_inter1 = mxg.MXG(device_name='N5182B', num_of_machine=1)
                self.tester_inter1.arb_waveform(rate=rate+'_PN15')
                self.tester_inter1.trriger_para_set(type='CONT', count=pkt_num)
                self.tester_inter1.output_state(1, 1)
                self.tester_inter1.arb_state(1)

                self.tester_inter2 = mxg.MXG(device_name='N5182B', num_of_machine=2)

                self.tester_inter2.output_state(1, 0)

                self.config_per_le_connect_usb(rate=_rate, packet_len=packet_len)
                self.csp.RF_Frequency_Settings_rx(mode='DTM', ch_tx=chan)
                self.csp.config_rxq_le_packets(rate=_rate, num=pkt_num)
                self.csp.config_rx_level(rxpwr=-67)

                freq_inter_offset_range = range(3,13) + range(-3,-13,-1)
                for freq_inter_offset in freq_inter_offset_range:
                    freq_inter2 = freq_inter_offset + 2402 + chan * 2
                    freq_inter1 = freq_inter_offset * 2 + 2402 + chan * 2
                    pwr_list = self.dict_pwr_inter_freqoffset[freq_inter_offset]
                    for power_inter in pwr_list:
                        self.tester_inter1.para_set(freq=freq_inter1, power=power_inter + cable_loss)
                        self.tester_inter2.para_set(freq=freq_inter2, power=power_inter + cable_loss)
                        time.sleep(0.2)
                        loginfo('freq_interference_tone:    {}      pwr_inter:  {}'.format(freq_inter2,power_inter))
                        per_list=[]
                        for loop in range(10):
                            res = self.csp.meas_le_per_res(rate=_rate)
                            per = eval(res[1])
                            per_list.append(per)
                            time.sleep(0.5)

                        per = max(per_list)
                        loginfo('{}     {}      {}      {}      {}      {}'.format(chan, rate, freq_inter2, freq_inter1, power_inter + cable_loss,  per))
                        fw1.write_data([chan, rate, freq_inter2, freq_inter1,  power_inter + cable_loss,  per])


    def le_rx_sniff(self, chan=0, rxpwr=-60, tx_power=15, time_interval=0.5):
        import matplotlib.pyplot as plt
        import numpy as np
        import time
        from math import *
        plt.ion()  # 开启interactive mode 成功的关键函数
        plt.figure(1)

        self.config_per_le_connect_usb(rate='LE1M', packet_len=37)
        self.csp.RF_Frequency_Settings_rx(mode='DTM', ch_tx=chan)
        self.csp.config_rxq_le_packets(rate='LE1M', num=1500)
        self.csp.RF_Power_settings(rx_level=rxpwr, tx_power=tx_power,margin=8)
        self.csp.RF_Power_settings_autoranging(0)   #打开autorange会导致接收per 概率性100%
        time.sleep(1)
        per_list=[]
        t=0
        t_list=[]
        while 1 :
            # plt.clf()  # 清空画布上的所有内容
            res = self.csp.meas_le_per_res(rate='LE1M')
            time.sleep(time_interval)
            t = t+time_interval
            per = eval(res[1])
            rxpkt_count = self.jlink.rd(0xa04008d8)
            loginfo('rxpkt_count:{}'.format(rxpkt_count))
            # per_list.append(per)
            # t_list.append(t)
            # plt.plot(t_list,per_list, '.r')
            plt.plot(t, per, '.r')
            plt.draw()
            plt.pause(0.01)

    def br_rx_sniff(self, chan=0, rxpwr=-60, time_interval=0.5, count=10000):
        import matplotlib.pyplot as plt
        import numpy as np
        import time
        from math import *
        plt.ion()  # 开启interactive mode 成功的关键函数
        plt.figure(200)
        plt.subplot(211)
        plt.subplot(212)


        self.config_classic_ber(rate='BR', chan=chan, rx_level=rxpwr, tx_power=15)
        self.csp.config_rx_level(rxpwr=rxpwr)
        self.csp.config_rxq_br_packets(num=1000)
        self.csp.config_rxq_br_tout(tout=10)
        time.sleep(1)
        ber_list=[]
        per_list=[]
        t=0
        t_list=[]
        for i in range(0,count) :
            # plt.clf()  # 清空画布上的所有内容
            while 1:
                res = self.csp.meas_bt_ber_res(cmd_type='READ')
                if eval(res[0]) == 0 or eval(res[0]) == 1:
                    break
                else:
                    if self.csp.get_bt_connect_state() == 'SBY':
                        self.cmd_br_connect()
                        time.sleep(10)
            time.sleep(time_interval)
            t = t+1
            ber = eval(res[1])
            per = eval(res[2])
            loginfo('ber  {}   per  {}'.format(ber,per))
            # ber_list.append(ber)
            # per_list.append(per)
            # t_list.append(t)
            plt.subplot(211)
            plt.plot(t,ber, '.r')
            plt.subplot(212)
            plt.plot( t, per, '.b')
            plt.pause(0.01)

    def get_filename(self, folder, file_name, sub_folder=''):
        '''
        :folder: file store folder
        :file_name:  file name
        :sub_folder: if not need, it may be default ""
        '''
        if rfglobal.file_folder=="":
            rfdata_path = './rftest/rfdata/'
        else:
            rfdata_path = './rftest/rfdata/%s/'%rfglobal.file_folder
            if os.path.exists(rfdata_path) == False:
                os.mkdir(rfdata_path)

        data_path1 = rfdata_path+'%s/'%(folder)
        if os.path.exists(data_path1) == False:
            os.mkdir(data_path1)

        filetime = time.strftime('%Y%m%d_%H%M%S',time.localtime(time.time()));
        # mac = self.read_mac()
        mac = ''

        gen_folder = '_%s'%(filetime[0:8])
        data_path2 = data_path1 +'%s/'%(gen_folder)
        if os.path.exists(data_path2) == False:
            os.mkdir(data_path2)

        fname = '%s'%(file_name)
        outfile_name = data_path2 + fname

        if sub_folder != '':
            gen_folder = '%s_%s'%(sub_folder,filetime[0:8])
            sub_path = data_path2+'%s/'%(gen_folder)
            if os.path.exists(sub_path) == False:
                os.mkdir(sub_path)

            outfile_name = sub_path + file_name

        return outfile_name


class BQB_autotest(bt_signaling):
    # def __init__(self, comport, chipv='TX212', rfport=1, atten=4.6):
    #     # bt_signaling(comport, chipv=chipv, rfport=rfport, atten=atten)
    #     self.rfport = rfport
    #     self.atten = atten
    #     self.comport = comport
    #     self.chipv = chipv
        # self.comport.Hexmd=1

    def pass_verdict(self, argu, meas):
        low_limit = str(argu.cells[1].text)
        upper_limit = str(argu.cells[2].text)
        loginfo(low_limit,upper_limit)
        argu.cells[5].text = ''
        if low_limit == '' and upper_limit != '':
            if meas < eval(upper_limit):
                run = argu.cells[5].paragraphs[0].add_run('Pass')
                run.font.color.rgb = RGBColor(0, 200, 0)
                res = 'Pass'
            else:
                run = argu.cells[5].paragraphs[0].add_run('Fail')
                run.font.color.rgb = RGBColor(200, 0, 0)
                res = 'Fail'
        elif upper_limit == '' and low_limit != '':
            if meas > eval(low_limit):
                run = argu.cells[5].paragraphs[0].add_run('Pass')
                run.font.color.rgb = RGBColor(0, 200, 0)
                res = 'Pass'
            else:
                run = argu.cells[5].paragraphs[0].add_run('Fail')
                run.font.color.rgb = RGBColor(200, 0, 0)
                res = 'Fail'
        elif low_limit != '' and upper_limit != '':
            if meas >= eval(low_limit) and meas <= eval(upper_limit):
                run = argu.cells[5].paragraphs[0].add_run('Pass')
                run.font.color.rgb = RGBColor(0, 200, 0)
                res = 'Pass'
            else:
                run = argu.cells[5].paragraphs[0].add_run('Fail')
                run.font.color.rgb = RGBColor(200, 0, 0)
                res = 'Fail'
        else:
            run = argu.cells[5].paragraphs[0].add_run('Pass')
            run.font.color.rgb = RGBColor(0, 200, 0)
            res = 'Pass'
        return res

    def bqb_autotest(self, inter_cableloss=5, signal_cableloss=6, dut_name='', temp=25):
        # title = 'reg_addr,reg_msb,reg_lsb,RW,reg_name,description,value_default,value_present_{}\n'.format(time.strftime("%Y-%m-%d %H:%M:%S", time.localtime()))
        # fname = self.get_filename('BQB_autotest', 'BQB_autotest')
        # fw1 = csvreport(fname, title)
        # shutil.copyfile(r'E:\chip\eagletest\py_script\rftest\test_script\BQB_autotest\BQB_RF_PHY_test_report.docx',os.path.join(fname,'BQB_RF_PHY_test_report_{}_{}.docx'.format(
        #     dut_name,temp)))
        self.document = Document(r'E:\chip\eagletest\py_script\rftest\test_script\BQB_autotest\BQB_RF_PHY_test_report_tws_L_1.docx')
        fname = 'E:/chip/eagletest/py_script/rftest/test_script/BQB_autotest/'
        # self.document = Document(r'E:\chip\eagletest\py_script\rftest\rfdata\BQB_autotest\_20210629\BQB_RF_PHY_test_report_QFN56_ECO_25.docx')

        self.tables = self.document.tables

        self.fname = fname
        self.dut_name = dut_name
        self.temp = temp
        self.cableloss = inter_cableloss
        self.atten =signal_cableloss

        ##BR/EDR test
        # self.csp.sig_std('CLAS')
        # self.config_classic_ber(rate='BR', chan=1, rx_level=-50, tx_power=10)
        # self.csp.tx_measure_para(tout=10, repetition='SINGleshot', count=10, MOEXception='OFF')
        # self.BR_Output_Power()
        # self.BR_Power_Density()
        # self.BR_Power_Control()
        # self.BR_Frequency_Range()
        # self.BR_20dB_Bandwidth()
        # self.BR_ACP()
        # self.BR_Modulation_Characteristics()
        # self.BR_Initial_Carrier_Frequency_Tolerance()
        # self.BR_Carrier_Frequency_Drift()
        # self.EDR_Relative_Transmit_Power()
        # self.EDR_Modulation_Accuracy()
        # self.EDR_Differential_Phase_Encoding()
        # self.EDR_Inband_Spurious_Emissions()
        # self.EDR_Power_Control()
        # self.BR_Sensitivity_single_slot_packets()
        # self.BR_Sensitivity_multi_slot_packets()
        # self.BR_CI_Performance()
        # self.BR_Blocking_Performance()
        # self.BR_Maximum_Input_Level()
        # self.EDR_Sensitivity()
        # self.EDR_BER_Floor_Performance()
        # self.EDR_CI_Performance()
        # self.EDR_Maximum_Input_Level()
        #
        # ##LE test
        # self.csp.sig_std('LE')
        # self.config_per_le_connect_usb(rate='LE1M', packet_len=37)
        # self.csp.config_rx_level(rxpwr=-67)
        # self.csp.trigger_settings(source='power', threshold=-35, tout=1)
        #
        # self.LE1M_Output_power()
        # self.LE1M_Inband_emissions()
        # self.LE1M_Modulation_Characteristics()
        # self.LE1M_Carrier_frequency_offset_and_drift()
        # self.LE1M_Receiver_sensitivity()
        # self.LE1M_Maximum_input_signal_level()
        # self.LE1M_PER_Report_Integrity()
        # self.LE1M_CI_Performance()
        # # self.LE1M_Blocking_Performance()

        # self.LE2M_Output_power()
        # self.LE2M_Inband_emissions()
        # self.LE2M_Modulation_Characteristics()
        # self.LE2M_Carrier_frequency_offset_and_drift()
        # self.LE2M_Receiver_sensitivity()
        # self.LE2M_Maximum_input_signal_level()
        # self.LE2M_PER_Report_Integrity()
        # self.LE2M_CI_Performance()
        # # self.LE2M_Blocking_Performance()

        # for i in range(3,len(self.tables)):
        #     logdebug(len(self.tables))
        #     res_list=[]
        #     for row in range(2,len(self.tables[i].rows)):
        #         self.tables[i].cell(row, 3).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
        #         self.tables[i].cell(row, 5).paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER
        #         res_list.append(self.tables[i].cell(row,5).text)
        #
        #     loginfo(res_list)
        #     for res in res_list:
        #         if res == 'Fail':
        #             sum_res = 'Fail'
        #             break
        #         elif res == 'Pass':
        #             sum_res = 'Pass'
        #         else:
        #             sum_res = ''
        #     if i <= 15:
        #         self.tables[2].cell(i - 2, 2).text = ''
        #         run = self.tables[2].cell(i - 2, 2).paragraphs[0].add_run(sum_res)
        #     elif i ==16:
        #         if self.tables[2].cell(i - 3, 2).text != 'Fail':
        #             self.tables[2].cell(i - 3, 2).text = ''
        #             run = self.tables[2].cell(i - 3, 2).paragraphs[0].add_run(sum_res)
        #     elif i > 16 and i <= 26:
        #         self.tables[2].cell(i - 3, 2).text = ''
        #         run = self.tables[2].cell(i - 3, 2).paragraphs[0].add_run(sum_res)
        #     elif i == 27:
        #         if self.tables[2].cell(i - 4, 2).text != 'Fail':
        #             self.tables[2].cell(i - 4, 2).text = ''
        #             run = self.tables[2].cell(i - 4, 2).paragraphs[0].add_run(sum_res)
        #     else:
        #         self.tables[2].cell(i - 4, 2).text = ''
        #         run = self.tables[2].cell(i - 4, 2).paragraphs[0].add_run(sum_res)
        #     if sum_res == 'Fail':
        #         run.font.color.rgb = RGBColor(200, 0, 0)
        #     else:
        #         run.font.color.rgb = RGBColor(0, 200, 0)
        # for i in range(2,len(self.tables[2].rows)):
        #     self.tables[2].cell(i,2).paragraphs[0].alginment = WD_TABLE_ALIGNMENT.CENTER
        #
        # self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))


    def BR_Output_Power(self):
        self.csp.mode_set(mode='BR')
        self.csp.sig_btype(btype='BR')
        self.csp.config_connect_br_packet_ptype(ptype='DH1')  # 设置BR包类型
        self.csp.config_connect_br_packet_pattern(pattern='PRBS9')  # 设置BR包 payload 数据类型
        self.csp.trigger_settings(source='power', threshold=-20, tout=1)
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan*39)
            res1 = self.csp.get_power_measure_res()
            nominal_pow = eval(res1[2])
            peak_pow = eval(res1[3])
            self.tables[3].cell(3+chan*3, 3).text = str(nominal_pow)
            self.tables[3].cell(4+chan*3, 3).text = str(peak_pow)
            self.pass_verdict(self.tables[3].rows[3+chan*3], nominal_pow)
            self.pass_verdict(self.tables[3].rows[4+chan*3], peak_pow)
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))
        self.csp.trigger_settings(source='power', threshold=-30, tout=1)

    def BR_Power_Density(self):
        self.csp.mode_set(mode='BR')
        self.csp.sig_btype(btype='BR')
        self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=18)
        self.csp.config_connect_br_packet_ptype(ptype='DH1')  # 设置BR包类型
        self.csp.config_connect_br_packet_pattern(pattern='PRBS9')  # 设置BR包 payload 数据类型
        self.csp.trigger_settings(source='power', threshold=-20, tout=1)
        res1 = self.csp.get_power_measure_res()
        nominal_pow = eval(res1[2])
        peak_pow = eval(res1[3])
        self.tables[4].cell(2, 3).text = str(peak_pow)
        self.pass_verdict(self.tables[4].rows[2], peak_pow)
        self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))
        self.csp.trigger_settings(source='power', threshold=-30, tout=1)

    def BR_Power_Control(self):
        self.csp.mode_set(mode='BR')
        self.csp.sig_btype(btype='BR')
        self.csp.config_connect_br_packet_ptype(ptype='DH1')  # 设置BR包类型
        self.csp.config_connect_br_packet_pattern(pattern='PRBS9')  # 设置BR包 payload 数据类型
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan*39)

            self.csp.config_power_control(para='MAX')
            time.sleep(5)
            res1 = self.csp.get_power_measure_res()
            nominal_pow_max = eval(res1[2])
            peak_pow = eval(res1[3])
            self.tables[5].cell(3+chan*14, 3).text = str(nominal_pow_max)
            self.pass_verdict(self.tables[5].rows[3+chan*14], nominal_pow_max)
            nominal_pow = nominal_pow_max
            for i in range(4,9):
                self.csp.config_power_control(para='DOWN')
                self.csp.trigger_settings(source='power', threshold=-20-i, tout=1)
                res1 = self.csp.get_power_measure_res()
                nominal_pow_d = eval(res1[2])
                power_step  = nominal_pow - nominal_pow_d
                nominal_pow = nominal_pow_d
                self.tables[5].cell(i+chan*14, 3).text = str(power_step)
                self.pass_verdict(self.tables[5].rows[i+chan*14], power_step)

            for i in range(10):
                self.csp.config_power_control(para='DOWN')
                res = self.csp.config_power_control_state()
                res = res.split(',')[1]
                logdebug(res)
                if res == 'MIN':
                    time.sleep(5)
                    res1 = self.csp.get_power_measure_res()
                    nominal_pow_min = eval(res1[2])
                    break
            self.tables[5].cell(9+chan*14, 3).text = str(nominal_pow_min)
            self.pass_verdict(self.tables[5].rows[9+chan*14], nominal_pow_min)
            nominal_pow = nominal_pow_min

            for i in range(10,15):
                self.csp.config_power_control(para='UP')
                self.csp.trigger_settings(source='power', threshold=-35 + i, tout=1)
                res1 = self.csp.get_power_measure_res()
                nominal_pow_d = eval(res1[2])
                power_step = nominal_pow_d - nominal_pow
                nominal_pow = nominal_pow_d
                self.tables[5].cell(i+chan*14, 3).text = str(power_step)
                self.pass_verdict(self.tables[5].rows[i+chan*14], power_step)

            self.csp.config_power_control(para='MAX')
            res1 = self.csp.get_power_measure_res()
            nominal_pow_max = eval(res1[2])
            peak_pow = eval(res1[3])
            self.tables[5].cell(15+chan*14, 3).text = str(nominal_pow_max)
            self.pass_verdict(self.tables[5].rows[15+chan*14], nominal_pow_max)
            self.csp.trigger_settings(source='power', threshold=-30, tout=1)
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))


    def BR_Frequency_Range(self):
        self.csp.mode_set(mode='BR')
        self.csp.sig_btype(btype='BR')
        self.csp.config_connect_br_packet_pattern(pattern='PRBS9')  # 设置BR包 payload 数据类型
        self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=0)
        res4 = self.csp.get_frange_res()
        frange_l = eval(res4[3]) / 1000000.00
        self.tables[6].cell(2, 3).text = str(frange_l)
        self.pass_verdict(self.tables[6].rows[2], frange_l)
        self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=78)
        res4 = self.csp.get_frange_res()
        frange_h = eval(res4[4]) / 1000000.00
        self.tables[6].cell(3, 3).text = str(frange_h)
        self.pass_verdict(self.tables[6].rows[3], frange_h)
        self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def BR_20dB_Bandwidth(self):
        self.csp.mode_set(mode='BR')
        self.csp.sig_btype(btype='BR')
        self.csp.config_connect_br_packet_pattern(pattern='PRBS9')  # 设置BR包 payload 数据类型
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan*39)
            res3 = self.csp.get_obw_res()
            obw = eval(res3[6]) / 1000.00
            obw_l = eval(res3[4]) / 1000.00
            obw_h = eval(res3[5]) / 1000.00
            self.tables[7].cell(3+chan*4, 3).text = str(obw_l)
            self.pass_verdict(self.tables[7].rows[3+chan*4], obw_l)
            self.tables[7].cell(4+chan*4, 3).text = str(obw_h)
            self.pass_verdict(self.tables[7].rows[4+chan*4], obw_h)
            self.tables[7].cell(5+chan*4, 3).text = str(obw)
            self.pass_verdict(self.tables[7].rows[5+chan*4], obw)
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def BR_ACP(self):
        self.csp.mode_set(mode='BR')
        self.csp.sig_btype(btype='BR')
        self.csp.config_connect_br_packet_pattern(pattern='PRBS9')  # 设置BR包 payload 数据类型
        self.csp.acp_meas_settings(opt='CH79')
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=3+chan*36)
            self.csp.tx_measure_states(state=0)
            res2 = self.csp.get_acp_res('READ')
            time.sleep(1)
            acp_list = [eval(i) for i in res2[1:]]
            loginfo(acp_list)
            for i in range(len(acp_list)):
                self.tables[8].cell(3 + i + chan * 80, 3).text = str(acp_list[i])
                self.pass_verdict(self.tables[8].rows[3 + i + chan * 80], acp_list[i])
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def BR_Modulation_Characteristics(self):
        self.csp.mode_set(mode='BR')
        self.csp.sig_btype(btype='BR')
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan * 39)
            delta_f1_99, delta_f2_99, delta_f2_avg, delta_f2_min, delta_f2_max, mod_ratio = -999, -999, -999, -999, -999, -999
            self.csp.config_connect_br_packet_pattern(pattern='P44')  # 设置BR包 payload 数据类型
            res = self.csp.get_modulation_measure_res()
            logdebug('{}'.format(res))

            delta_f1_99 = eval(res[2]) / 1000.00
            freq_accuracy = eval(res[3]) / 1000.00
            freq_drift = eval(res[4]) / 1000.00
            drift_rate = eval(res[5])
            delta_f1_avg = eval(res[6]) / 1000.00
            delta_f1_min = eval(res[7]) / 1000.00
            delta_f1_max = eval(res[8]) / 1000.00
            time.sleep(0.5)

            self.csp.config_connect_br_packet_pattern(pattern='P11')  # 设置BR包 payload 数据类型
            res = self.csp.get_modulation_measure_res()
            logdebug('{}'.format(res))

            delta_f2_99 = eval(res[2]) / 1000.00
            freq_accuracy = eval(res[3]) / 1000.00
            freq_drift = eval(res[4]) / 1000.00
            drift_rate = eval(res[5]) / 1000.00
            delta_f2_avg = eval(res[9]) / 1000.00
            delta_f2_min = eval(res[10]) / 1000.00
            delta_f2_max = eval(res[11]) / 1000.00
            mod_ratio = delta_f2_avg / delta_f1_avg
            time.sleep(0.5)

            self.tables[9].cell(3 + chan * 4, 3).text = str(delta_f1_avg)
            self.pass_verdict(self.tables[9].rows[3 + chan * 4], delta_f1_avg)

            self.tables[9].cell(4 + chan * 4, 3).text = str(delta_f2_99)
            self.pass_verdict(self.tables[9].rows[4 + chan * 4], delta_f2_99)

            self.tables[9].cell(5 + chan * 4, 3).text = str(mod_ratio)
            self.pass_verdict(self.tables[9].rows[5 + chan * 4], mod_ratio)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def BR_Initial_Carrier_Frequency_Tolerance(self):
        self.csp.mode_set(mode='BR')
        self.csp.sig_btype(btype='BR')
        self.csp.config_connect_br_packet_pattern(pattern='P11')  # 设置BR包 payload 数据类型
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan * 39)
            res = self.csp.get_modulation_measure_res()
            freq_accuracy = eval(res[3]) / 1000.00
            self.tables[10].cell(3 + chan * 2, 3).text = str(freq_accuracy)
            self.pass_verdict(self.tables[10].rows[3 + chan * 2], freq_accuracy)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def BR_Carrier_Frequency_Drift(self):
        self.csp.mode_set(mode='BR')
        self.csp.sig_btype(btype='BR')
        self.csp.config_connect_br_packet_pattern(pattern='P11')  # 设置BR包 payload 数据类型
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan * 39)
            self.csp.config_connect_br_packet_ptype(ptype='DH1')  # 设置BR包类型
            res = self.csp.get_modulation_measure_res()
            freq_drift = eval(res[4]) / 1000.00
            drift_rate = eval(res[5]) / 1000.00
            self.tables[11].cell(3 + chan * 7, 3).text = str(freq_drift)
            self.pass_verdict(self.tables[11].rows[3 + chan * 7], freq_drift)
            self.tables[11].cell(6 + chan * 7, 3).text = str(freq_drift)
            self.pass_verdict(self.tables[11].rows[6 + chan * 7], freq_drift)

            self.csp.config_connect_br_packet_ptype(ptype='DH3')  # 设置BR包类型
            res = self.csp.get_modulation_measure_res()
            freq_drift = eval(res[4]) / 1000.00
            drift_rate = eval(res[5]) / 1000.00
            self.tables[11].cell(4 + chan * 7, 3).text = str(freq_drift)
            self.pass_verdict(self.tables[11].rows[4 + chan * 7], freq_drift)
            self.tables[11].cell(7 + chan * 7, 3).text = str(freq_drift)
            self.pass_verdict(self.tables[11].rows[7 + chan * 7], freq_drift)

            self.csp.config_connect_br_packet_ptype(ptype='DH5')  # 设置BR包类型
            res = self.csp.get_modulation_measure_res()
            freq_drift = eval(res[4]) / 1000.00
            drift_rate = eval(res[5]) / 1000.00
            self.tables[11].cell(5 + chan * 7, 3).text = str(freq_drift)
            self.pass_verdict(self.tables[11].rows[5 + chan * 7], freq_drift)
            self.tables[11].cell(8 + chan * 7, 3).text = str(freq_drift)
            self.pass_verdict(self.tables[11].rows[8 + chan * 7], freq_drift)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def EDR_Relative_Transmit_Power(self):
        self.csp.mode_set(mode='EDR')
        self.csp.sig_btype(btype='EDR')
        self.csp.config_connect_edr_packet_pattern(pattern='PRBS9')  # 设置EDR包 payload 数据类型
        self.csp.trigger_settings(source='power', threshold=-20, tout=1)
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan * 39)
            self.csp.config_connect_edr_packet_ptype(ptype='E25P')  # 设置EDR包类型
            res1 = self.csp.get_power_measure_res()
            logdebug('{}'.format(res1))
            res1 = [eval(i) for i in res1[0:]]
            nominal_pwr = res1[2]
            gfsk_pwr = res1[3]
            dpsk_pwr = res1[4]
            dpsk_gfsk_diff_pwr = res1[5]
            guard_period = res1[6]
            packet_timing = res1[7]
            peak_pwr = res1[8]

            self.tables[12].cell(3 + chan * 7, 3).text = str(dpsk_gfsk_diff_pwr)
            self.pass_verdict(self.tables[12].rows[3 + chan * 7], dpsk_gfsk_diff_pwr)
            self.tables[12].cell(4 + chan * 7, 3).text = str(gfsk_pwr)
            self.pass_verdict(self.tables[12].rows[4 + chan * 7], gfsk_pwr)
            self.tables[12].cell(5 + chan * 7, 3).text = str(dpsk_pwr)
            self.pass_verdict(self.tables[12].rows[5 + chan * 7], dpsk_pwr)

            self.csp.config_connect_edr_packet_ptype(ptype='E35P')  # 设置EDR包类型
            res1 = self.csp.get_power_measure_res()
            logdebug('{}'.format(res1))
            res1 = [eval(i) for i in res1[0:]]
            nominal_pwr = res1[2]
            gfsk_pwr = res1[3]
            dpsk_pwr = res1[4]
            dpsk_gfsk_diff_pwr = res1[5]
            guard_period = res1[6]
            packet_timing = res1[7]
            peak_pwr = res1[8]

            self.tables[12].cell(6 + chan * 7, 3).text = str(dpsk_gfsk_diff_pwr)
            self.pass_verdict(self.tables[12].rows[6 + chan * 7], dpsk_gfsk_diff_pwr)
            self.tables[12].cell(7 + chan * 7, 3).text = str(gfsk_pwr)
            self.pass_verdict(self.tables[12].rows[7 + chan * 7], gfsk_pwr)
            self.tables[12].cell(8 + chan * 7, 3).text = str(dpsk_pwr)
            self.pass_verdict(self.tables[12].rows[8 + chan * 7], dpsk_pwr)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))
        self.csp.trigger_settings(source='power', threshold=-30, tout=1)

    def EDR_Modulation_Accuracy(self):
        self.csp.mode_set(mode='EDR')
        self.csp.sig_btype(btype='EDR')
        self.csp.config_connect_edr_packet_pattern(pattern='PRBS9')  # 设置EDR包 payload 数据类型
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan * 39)
            self.csp.config_connect_edr_packet_ptype(ptype='E25P')  # 设置EDR包类型
            for i in range(1):
                self.csp.config_power_control(para='DOWN')
            res = self.csp.get_modulation_measure_res()
            logdebug('{}'.format(res))
            res = [eval(i) for i in res[0:]]
            wi = res[2] / 1000
            w0_wi = res[3] / 1000
            w0_max = res[4] / 1000
            DEVM_RMS = res[5]*100.00
            DEVM_peak = res[6]*100.00
            DEVM_P99 = res[7]*100.00

            self.tables[13].cell(3 + chan * 13, 3).text = str(wi)
            self.pass_verdict(self.tables[13].rows[3 + chan * 13], wi)
            self.tables[13].cell(4 + chan * 13, 3).text = str(w0_wi)
            self.pass_verdict(self.tables[13].rows[4 + chan * 13], w0_wi)
            self.tables[13].cell(5 + chan * 13, 3).text = str(w0_max)
            self.pass_verdict(self.tables[13].rows[5 + chan * 13], w0_max)
            self.tables[13].cell(6 + chan * 13, 3).text = str(DEVM_RMS)
            self.pass_verdict(self.tables[13].rows[6 + chan * 13], DEVM_RMS)
            self.tables[13].cell(7 + chan * 13, 3).text = str(DEVM_peak)
            self.pass_verdict(self.tables[13].rows[7 + chan * 13], DEVM_peak)
            self.tables[13].cell(8 + chan * 13, 3).text = str(DEVM_P99)
            self.pass_verdict(self.tables[13].rows[8 + chan * 13], DEVM_P99)

            self.csp.config_connect_edr_packet_ptype(ptype='E35P')  # 设置EDR包类型
            for i in range(1):
                self.csp.config_power_control(para='DOWN')
            res = self.csp.get_modulation_measure_res()
            logdebug('{}'.format(res))
            res = [eval(i) for i in res[0:]]
            wi = res[2] / 1000
            w0_wi = res[3] / 1000
            w0_max = res[4] / 1000
            DEVM_RMS = res[5]*100.00
            DEVM_peak = res[6]*100.00
            DEVM_P99 = res[7]*100.00

            self.tables[13].cell(9 + chan * 13, 3).text = str(wi)
            self.pass_verdict(self.tables[13].rows[9 + chan * 13], wi)
            self.tables[13].cell(10 + chan * 13, 3).text = str(w0_wi)
            self.pass_verdict(self.tables[13].rows[10 + chan * 13], w0_wi)
            self.tables[13].cell(11 + chan * 13, 3).text = str(w0_max)
            self.pass_verdict(self.tables[13].rows[11 + chan * 13], w0_max)
            self.tables[13].cell(12 + chan * 13, 3).text = str(DEVM_RMS)
            self.pass_verdict(self.tables[13].rows[12 + chan * 13], DEVM_RMS)
            self.tables[13].cell(13 + chan * 13, 3).text = str(DEVM_peak)
            self.pass_verdict(self.tables[13].rows[13 + chan * 13], DEVM_peak)
            self.tables[13].cell(14 + chan * 13, 3).text = str(DEVM_P99)
            self.pass_verdict(self.tables[13].rows[14 + chan * 13], DEVM_P99)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def EDR_Differential_Phase_Encoding(self):
        self.csp.mode_set(mode='EDR')
        self.csp.sig_btype(btype='EDR')
        self.csp.config_connect_edr_packet_pattern(pattern='PRBS9')
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='TXT', ch_tx=chan * 39)
            self.csp.config_connect_edr_packet_ptype(ptype='E21P')  # 设置EDR包类型
            res2 = self.csp.get_diff_phase_encoding_res()
            logdebug('{}'.format(res2))
            res2 = [eval(i) for i in res2[0:]]
            bit_error_rate = res2[2]
            packet0error = res2[3]
            self.tables[14].cell(3 + chan * 3, 3).text = str(packet0error)
            self.pass_verdict(self.tables[14].rows[3 + chan * 3], packet0error)

            self.csp.config_connect_edr_packet_ptype(ptype='E31P')  # 设置EDR包类型
            res2 = self.csp.get_diff_phase_encoding_res()
            logdebug('{}'.format(res2))
            res2 = [eval(i) for i in res2[0:]]
            bit_error_rate = res2[2]
            packet0error = res2[3]
            self.tables[14].cell(4 + chan * 3, 3).text = str(packet0error)
            self.pass_verdict(self.tables[14].rows[4 + chan * 3], packet0error)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def EDR_Inband_Spurious_Emissions(self):
        self.csp.mode_set(mode='EDR')
        self.csp.sig_btype(btype='EDR')
        self.csp.config_connect_edr_packet_pattern(pattern='PRBS9')
        self.csp.acp_meas_settings(opt='CH79')
        self.csp.config_connect_edr_packet_ptype(ptype='E21P')  # 设置EDR包类型
        self.csp.trigger_settings(source='power', threshold=-40, tout=1)
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=3+chan * 36)
            time.sleep(1)
            self.csp.config_power_control(para='MAX')
            self.csp.RF_Power_settings(rx_level=-40, tx_power=15, margin=8)

            res1 = self.csp.get_power_measure_res()
            logdebug('{}'.format(res1))
            res1 = [eval(i) for i in res1[0:]]
            nominal_pwr = res1[2]
            self.csp.RF_Power_settings(rx_level=-40, tx_power=nominal_pwr+3, margin=8)
            # time.sleep(2)
            # # for i in range(1):
            # self.csp.config_power_control(para='DOWN')
            # self.csp.config_power_control(para='DOWN')
            time.sleep(1)
            [res3, res4] = self.csp.get_acp_res_edr('READ')
            logdebug('{}'.format(res3))
            logdebug('{}'.format(res4))
            res3 = [eval(i) for i in res3[0:]]
            res4 = [eval(i) for i in res4[0:]]
            acp_nominal_pwr = res3[2]
            PTxRef = res3[4]
            N26ChN1Abs = res3[5]
            N26ChP1Abs = res3[6]
            N26ChN1Rel = res3[7]
            N26ChP1Rel = res3[8]

            acp_list = res4[1:]
            for i in range(len(acp_list)):
                self.tables[15].cell(4 + i + chan * 80, 3).text = str(acp_list[i])
                self.pass_verdict(self.tables[15].rows[4 + i + chan * 80], acp_list[i])

            self.tables[15].cell(6 + chan*36 + chan * 80, 3).text = str(N26ChN1Abs)
            self.pass_verdict(self.tables[15].rows[6 + chan*36 + chan * 80], N26ChN1Abs)
            self.tables[15].cell(7 + chan*36 + chan * 80, 3).text = str(PTxRef)
            self.pass_verdict(self.tables[15].rows[7 + chan*36 + chan * 80], PTxRef)
            self.tables[15].cell(8 + chan*36 + chan * 80, 3).text = str(N26ChP1Abs)
            self.pass_verdict(self.tables[15].rows[8 + chan*36 + chan * 80], N26ChP1Abs)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

        self.csp.config_connect_edr_packet_ptype(ptype='E31P')  # 设置EDR包类型

        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=3+chan * 36)
            time.sleep(1)
            self.csp.config_power_control(para='MAX')
            # time.sleep(2)
            # # for i in range(1):
            # self.csp.config_power_control(para='DOWN')
            # self.csp.config_power_control(para='DOWN')
            time.sleep(1)
            [res3, res4] = self.csp.get_acp_res_edr('READ')
            logdebug('{}'.format(res3))
            logdebug('{}'.format(res4))
            res3 = [eval(i) for i in res3[0:]]
            res4 = [eval(i) for i in res4[0:]]
            acp_nominal_pwr = res3[2]
            PTxRef = res3[4]
            N26ChN1Abs = res3[5]
            N26ChP1Abs = res3[6]
            N26ChN1Rel = res3[7]
            N26ChP1Rel = res3[8]

            acp_list = res4[1:]
            for i in range(len(acp_list)):
                self.tables[16].cell(4 + i + chan * 80, 3).text = str(acp_list[i])
                self.pass_verdict(self.tables[16].rows[4 + i + chan * 80], acp_list[i])

            self.tables[16].cell(6 + chan*36 + chan * 80, 3).text = str(N26ChN1Abs)
            self.pass_verdict(self.tables[16].rows[6 + chan*36 + chan * 80], N26ChN1Abs)
            self.tables[16].cell(7 + chan*36 + chan * 80, 3).text = str(PTxRef)
            self.pass_verdict(self.tables[16].rows[7 + chan*36 + chan * 80], PTxRef)
            self.tables[16].cell(8 + chan*36 + chan * 80, 3).text = str(N26ChP1Abs)
            self.pass_verdict(self.tables[16].rows[8 + chan*36 + chan * 80], N26ChP1Abs)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))
        self.csp.trigger_settings(source='power', threshold=-30, tout=1)

    def EDR_Power_Control(self):
        self.csp.mode_set(mode='EDR')
        self.csp.sig_btype(btype='EDR')
        self.csp.config_connect_edr_packet_pattern(pattern='PRBS9')
        self.csp.config_connect_edr_packet_ptype(ptype='E25P')  # 设置EDR包类型
        self.csp.RF_Power_settings_autoranging(1)
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan*39)
            self.csp.config_power_control(para='MAX')
            res1 = self.csp.get_power_measure_res()
            logdebug('{}'.format(res1))
            res1 = [eval(i) for i in res1[0:]]
            nominal_pow_max = res1[2]
            self.tables[17].cell(3 + chan * 14, 3).text = str(nominal_pow_max)
            self.pass_verdict(self.tables[17].rows[3 + chan * 14], nominal_pow_max)
            nominal_pow = nominal_pow_max
            for i in range(4, 9):
                self.csp.config_power_control(para='DOWN')
                # self.csp.RF_Power_settings(rx_level=-50, tx_power=10-3*(i-3), margin=8)
                # self.csp.trigger_settings(source='power', threshold=-20 - i, tout=1)
                res1 = self.csp.get_power_measure_res()
                nominal_pow_d = eval(res1[2])
                power_step = nominal_pow - nominal_pow_d
                nominal_pow = nominal_pow_d
                self.tables[17].cell(i + chan * 14, 3).text = str(power_step)
                self.pass_verdict(self.tables[17].rows[i + chan * 14], power_step)

            for i in range(10):
                self.csp.config_power_control(para='DOWN')
                res = self.csp.config_power_control_state()
                res = res.split(',')[1]
                logdebug(res)
                if res == 'MIN':
                    res1 = self.csp.get_power_measure_res()
                    nominal_pow_min = eval(res1[2])
                    break
            self.tables[17].cell(9 + chan * 14, 3).text = str(nominal_pow_min)
            self.pass_verdict(self.tables[17].rows[9 + chan * 14], nominal_pow_min)
            nominal_pow = nominal_pow_min

            for i in range(10, 15):
                self.csp.config_power_control(para='UP')
                # self.csp.RF_Power_settings(rx_level=-50, tx_power=-10 + 3 * (i-9), margin=8)
                # self.csp.trigger_settings(source='power', threshold=-35 + i, tout=1)
                res1 = self.csp.get_power_measure_res()
                nominal_pow_d = eval(res1[2])
                power_step = nominal_pow_d - nominal_pow
                nominal_pow = nominal_pow_d
                self.tables[17].cell(i + chan * 14, 3).text = str(power_step)
                self.pass_verdict(self.tables[17].rows[i + chan * 14], power_step)

            self.csp.config_power_control(para='MAX')
            # self.csp.RF_Power_settings(rx_level=-50, tx_power=10, margin=8)
            res1 = self.csp.get_power_measure_res()
            nominal_pow_max = eval(res1[2])
            peak_pow = eval(res1[3])
            self.tables[17].cell(15 + chan * 14, 3).text = str(nominal_pow_max)
            self.pass_verdict(self.tables[17].rows[15 + chan * 14], nominal_pow_max)
            # self.csp.trigger_settings(source='power', threshold=-20, tout=1)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

        self.csp.RF_Power_settings_autoranging(0)

    def BR_Sensitivity_single_slot_packets(self):
        self.csp.mode_set(mode='BR')
        self.csp.sig_btype(btype='BR')
        self.csp.config_rx_level(rxpwr=-80)
        self.csp.config_rxq_br_search_packets(num=2000)
        self.csp.config_rxq_br_search_step(step=0.5)
        self.csp.config_connect_br_packet_pattern(pattern='PRBS9')
        self.csp.config_connect_br_packet_ptype(ptype='DH1')
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan*39)
            self.csp.ber_search_meas_state(state=0)
            while 1:
                state_res = self.csp.get_ber_search_meas_state()  # ber search 完成则退出循环
                # logdebug(state_res)
                if state_res == 'RDY':
                    loginfo('search complete')
                    break
                elif state_res == 'OFF':
                    self.cmd_br_connect()
                    self.csp.ber_search_meas_state(state=0)
                else:
                    logdebug('searching')
                time.sleep(5)

            res = self.csp.meas_bt_ber_search_res(cmd_type='FETCh')
            ber = eval(res[1])
            per = eval(res[2])
            NAK = eval(res[5])
            hec_err = eval(res[6])
            crc_err = eval(res[7])
            packet_type_err = eval(res[8])
            pay_err = eval(res[9])
            sens = eval(res[11])
            loginfo('{} {}  {}  {}  {}  {}  {}  {}'.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err, sens))
            self.tables[18].cell(2 + chan * 1, 3).text = str(sens)
            self.pass_verdict(self.tables[18].rows[2 + chan * 1], sens)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def BR_Sensitivity_multi_slot_packets(self):
        self.csp.mode_set(mode='BR')
        self.csp.sig_btype(btype='BR')
        self.csp.config_rx_level(rxpwr=-80)
        self.csp.config_rxq_br_search_packets(num=2000)
        self.csp.config_rxq_br_search_step(step=0.5)
        self.csp.config_connect_br_packet_pattern(pattern='PRBS9')
        self.csp.config_connect_br_packet_ptype(ptype='DH5')
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan*39)
            self.csp.ber_search_meas_state(state=0)
            while 1:
                state_res = self.csp.get_ber_search_meas_state()  # ber search 完成则退出循环
                # logdebug(state_res)
                if state_res == 'RDY':
                    loginfo('search complete')
                    break
                elif state_res == 'OFF':
                    self.cmd_br_connect()
                    self.csp.ber_search_meas_state(state=0)
                else:
                    logdebug('searching')
                time.sleep(5)

            res = self.csp.meas_bt_ber_search_res(cmd_type='FETCh')
            ber = eval(res[1])
            per = eval(res[2])
            NAK = eval(res[5])
            hec_err = eval(res[6])
            crc_err = eval(res[7])
            packet_type_err = eval(res[8])
            pay_err = eval(res[9])
            sens = eval(res[11])
            loginfo('{} {}  {}  {}  {}  {}  {}  {}'.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err, sens))
            self.tables[19].cell(2 + chan * 1, 3).text = str(sens)
            self.pass_verdict(self.tables[19].rows[2 + chan * 1], sens)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def BR_CI_Performance(self):
        self.csp.mode_set(mode='BR')
        self.dict_pwr_inter_freqoffset = {
            0: range(-75, -40),
            1: range(-60, -30),
            2: range(-40, -10),
            3: range(-35, 0),
            4: range(-35, 0),
            -1: range(-60, -30),
            -2: range(-50, -10),
            -3: range(-35, 0),
            -4: range(-35, 0),
            }

        self.csp.sig_btype(btype='BR')
        self.csp.config_rx_level(rxpwr=-60)
        self.csp.config_rxq_br_packets(num=1000)
        self.csp.config_connect_br_packet_pattern(pattern='PRBS9')
        self.csp.config_connect_br_packet_ptype(ptype='DH5')
        self.tester_inter = mxg.MXG(device_name='N5182B', num_of_machine=1)
        self.tester_inter.arb_waveform(rate='DH1_PN15')
        self.tester_inter.trriger_para_set(type='CONT', count=2000)
        self.tester_inter.para_set(freq=2402, power=-100)
        for chan in range(1):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=3+chan*36)
            self.tester_inter.output_state(1, 1)
            self.tester_inter.arb_state(1)
            num = 0
            for freq_inter_offset in [0, 1, -1, 2, -2, 3, -3, 4, -4]:
                freq_inter = freq_inter_offset + 2405 + chan*36
                # if freq_inter_offset == 0 and btype == 'EDR':
                #     self.tester_inter.arb_waveform(rate=rate + '_PN15')
                # else:
                #     self.tester_inter.arb_waveform(rate='DH1_PN15')
                if -3 < freq_inter_offset < 3:
                    self.csp.config_rx_level(rxpwr=-60)
                    rxpwr_ref = -60
                else:
                    self.csp.config_rx_level(rxpwr=-67)
                    rxpwr_ref = -67
                pwr_list = self.dict_pwr_inter_freqoffset[freq_inter_offset]
                for power_inter in pwr_list:
                    self.tester_inter.para_set(freq=freq_inter, power=power_inter + self.cableloss)
                    time.sleep(0.2)
                    loginfo('freq_inter:    {}      pwr_inter:  {}'.format(freq_inter, power_inter))
                    timeout = 0
                    while 1:
                        if self.csp.get_bt_connect_state() == 'SBY':
                            self.cmd_br_connect()
                        self.csp.ber_meas_state(state=0)
                        while 1:
                            state = self.csp.get_ber_meas_state()
                            if state == 'RDY':
                                loginfo('ber run complte')
                                break
                            time.sleep(2)
                            timeout = timeout + 1
                            if timeout > 30:
                                loginfo('get result timeout')
                                break
                        if timeout > 30:
                            res = ['0', '1', '100', '100', '100', '100', '100', '100', '100', '100']
                            break
                        res = self.csp.meas_bt_ber_res(cmd_type='FETCh')
                        if eval(res[0]) == 0 or eval(res[0]) == 1:
                            break
                        else:
                            if self.csp.get_bt_connect_state() == 'SBY':
                                self.cmd_br_connect()
                    # while 1:
                    #     self.csp.ber_meas_state(state=0)
                    #     while 1:
                    #         state = self.csp.get_ber_meas_state()
                    #         if state == 'RDY':
                    #             loginfo('ber run complte')
                    #             break
                    #         time.sleep(2)
                    #     res = self.csp.meas_bt_ber_res(cmd_type='FETCh')
                    #     if eval(res[0]) == 0 or eval(res[0]) == 1:
                    #         break
                    #     else:
                    #         if self.csp.get_bt_connect_state() == 'SBY':
                    #             self.cmd_br_connect()

                    ber = eval(res[1])
                    per = eval(res[2])
                    NAK = eval(res[5])
                    hec_err = eval(res[6])
                    crc_err = eval(res[7])
                    packet_type_err = eval(res[8])
                    pay_err = eval(res[9])

                    loginfo('{} {}  {}  {}  {}  {}  {} '.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err))
                    if ber > 0.1 or per > 50 or NAK > 50:
                        self.tables[20].cell(4 + num + chan * 10, 3).text = str(rxpwr_ref-power_inter+1)
                        self.pass_verdict(self.tables[20].rows[4 + num + chan * 10], rxpwr_ref-power_inter+1)
                        break
                num = num + 1
            self.tester_inter.output_state(0, 0)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def BR_Blocking_Performance(self):
        self.csp.mode_set(mode='BR')
        self.csp.sig_btype(btype='BR')
        self.csp.config_rx_level(rxpwr=-60)
        self.csp.config_rxq_br_packets(num=500)
        self.csp.config_connect_br_packet_pattern(pattern='PRBS9')
        self.csp.config_connect_br_packet_ptype(ptype='DH1')
        self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=58)

        rxpwr = -13
        freq_range = range(30, 2401, 1) + range(2500, 6000, 1)
        freq_inter_list = []
        freq_inter_list2 = []
        for freq_inter in freq_range:
            freq_blocking_2 = 'NA'
            freq_blocking_1 = 'NA'

            if freq_inter < 2001 or freq_inter > 2999:
                pwr = rxpwr + 17
            else:
                pwr = rxpwr
            self.csp.config_rx_level(rxpwr=-67)
            timeout = 0
            while 1:
                if self.csp.get_bt_connect_state() == 'SBY':
                    self.cmd_br_connect()
                self.csp.ber_meas_state(state=0)
                while 1:
                    state = self.csp.get_ber_meas_state()
                    if state == 'RDY':
                        loginfo('ber run complte')
                        break
                    time.sleep(2)
                    timeout = timeout + 1
                    if timeout > 30:
                        loginfo('get result timeout')
                        break
                if timeout > 30:
                    res = ['0', '1', '100', '100', '100', '100', '100', '100', '100', '100']
                    break
                res = self.csp.meas_bt_ber_res(cmd_type='FETCh')
                if eval(res[0]) == 0 or eval(res[0]) == 1:
                    break
                else:
                    if self.csp.get_bt_connect_state() == 'SBY':
                        self.cmd_br_connect()
            # connect_num = 0
            # while 1:
            #     self.tester_inter.para_set(freq=freq_inter, power=pwr+self.cableloss)
            #
            #     loginfo('freq_inter:    {}'.format(freq_inter))
            #
            #     res = self.csp.meas_bt_ber_res(cmd_type='READ')
            #     if eval(res[0]) == 0 or eval(res[0]) == 1:
            #         break
            #     else:
            #         connect_num = connect_num + 1
            #         loginfo('reconnect num {}'.format(connect_num))
            #         if self.csp.get_bt_connect_state() == 'SBY':
            #             self.cmd_br_connect()
            #             if connect_num > 3:
            #                 res = ['0', '100', '100', '100', '100', '100', '100', '100', '100', '100', '100']
            #                 break
            ber = eval(res[1])
            per = eval(res[2])
            NAK = eval(res[5])
            hec_err = eval(res[6])
            crc_err = eval(res[7])
            packet_type_err = eval(res[8])
            pay_err = eval(res[9])
            loginfo('freq_blocking_1:   {}      level:  {}      ber:    {}      per:    {}'.format(freq_inter, pwr, ber, per))
            if ber > 0.1 or per > 90:
                loginfo('freq_blocking_1:   {}      level:  {}      ber:    {}      per:    {}'.format(freq_inter, pwr, ber, per))
                freq_blocking_1 = freq_inter
                freq_inter_list.append(freq_inter)

        while 1:
            loginfo('***********************************************************')
            loginfo(freq_inter_list)
            if len(freq_inter_list) > 24:
                rxpwr = rxpwr - 1
                freq_range = freq_inter_list
                freq_inter_list = []
                freq_inter_list2 = []
                for freq_inter in freq_range:
                    freq_blocking_2 = 'NA'
                    freq_blocking_1 = 'NA'

                    if freq_inter < 2001 or freq_inter > 2999:
                        pwr = rxpwr + 17
                    else:
                        pwr = rxpwr
                    self.csp.config_rx_level(rxpwr=-67)
                    timeout = 0
                    while 1:
                        if self.csp.get_bt_connect_state() == 'SBY':
                            self.cmd_br_connect()
                        self.csp.ber_meas_state(state=0)
                        while 1:
                            state = self.csp.get_ber_meas_state()
                            if state == 'RDY':
                                loginfo('ber run complte')
                                break
                            time.sleep(2)
                            timeout = timeout + 1
                            if timeout > 30:
                                loginfo('get result timeout')
                                break
                        if timeout > 30:
                            res = ['0', '1', '100', '100', '100', '100', '100', '100', '100', '100']
                            break
                        res = self.csp.meas_bt_ber_res(cmd_type='FETCh')
                        if eval(res[0]) == 0 or eval(res[0]) == 1:
                            break
                        else:
                            if self.csp.get_bt_connect_state() == 'SBY':
                                self.cmd_br_connect()
                    # connect_num = 0
                    # while 1:
                    #     self.tester_inter.para_set(freq=freq_inter, power=pwr + self.cableloss)
                    #
                    #     loginfo('freq_inter:    {}'.format(freq_inter))
                    #
                    #     res = self.csp.meas_bt_ber_res(cmd_type='READ')
                    #     if eval(res[0]) == 0 or eval(res[0]) == 1:
                    #         break
                    #     else:
                    #         connect_num = connect_num + 1
                    #         loginfo('reconnect num {}'.format(connect_num))
                    #         if self.csp.get_bt_connect_state() == 'SBY':
                    #             self.cmd_br_connect()
                    #             if connect_num > 3:
                    #                 res = ['0', '100', '100', '100', '100', '100', '100', '100', '100', '100', '100']
                    #                 break
                    ber = eval(res[1])
                    per = eval(res[2])
                    NAK = eval(res[5])
                    hec_err = eval(res[6])
                    crc_err = eval(res[7])
                    packet_type_err = eval(res[8])
                    pay_err = eval(res[9])
                    loginfo('freq_blocking_1:   {}      level:  {}      ber:    {}      per:    {}'.format(freq_inter, pwr, ber, per))
                    if ber > 0.1 or per > 50:
                        loginfo('freq_blocking_1:   {}      level:  {}      ber:    {}      per:    {}'.format(freq_inter, pwr, ber, per))
                        freq_blocking_1 = freq_inter
                        freq_inter_list.append(freq_inter)
            else:
                self.tables[21].cell(3, 3).text = str(rxpwr+17)
                self.pass_verdict(self.tables[21].rows[3], rxpwr+17)
                self.tables[21].cell(4, 3).text = str(rxpwr)
                self.pass_verdict(self.tables[21].rows[4], rxpwr)
                self.tables[21].cell(5, 3).text = str(rxpwr)
                self.pass_verdict(self.tables[21].rows[5], rxpwr)
                self.tables[21].cell(6, 3).text = str(rxpwr+17)
                self.pass_verdict(self.tables[21].rows[6], rxpwr+17)
                break
        self.tester_inter.output_state(0, 0)

        self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def BR_Maximum_Input_Level(self):
        self.csp.mode_set(mode='BR')
        rxpwr_range=range(-20,1)
        self.csp.sig_btype(btype='BR')
        self.csp.config_rx_level(rxpwr=-40)
        self.csp.config_rxq_br_packets(num=2000)
        self.csp.config_connect_br_packet_pattern(pattern='PRBS9')
        self.csp.config_connect_br_packet_ptype(ptype='DH1')
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan*39)
            self.csp.ber_meas_state(0)
            for rxpwr in rxpwr_range:
                self.csp.config_rx_level(rxpwr=rxpwr)
                # while 1:
                #     state_res = self.csp.get_ber_search_meas_state()
                #     if state_res == 'RDY':
                #         logdebug(state_res)
                #         break
                #     time.sleep(3)
                # delay = 0.001*packet_num*(rx_level+100)/step
                # time.sleep(delay)
                while 1:
                    self.csp.ber_meas_state(state=0)
                    while 1:
                        state = self.csp.get_ber_meas_state()
                        if state == 'RDY':
                            loginfo('ber run complte')
                            break
                        time.sleep(2)
                    res = self.csp.meas_bt_ber_res(cmd_type='FETCh')
                    if eval(res[0]) == 0 or eval(res[0]) == 1:
                        break
                    else:
                        if self.csp.get_bt_connect_state() == 'SBY':
                            self.cmd_br_connect()

                ber = eval(res[1])
                per = eval(res[2])
                NAK = eval(res[5])
                hec_err = eval(res[6])
                crc_err = eval(res[7])
                packet_type_err = eval(res[8])
                pay_err = eval(res[9])

                time.sleep(0.5)
                loginfo('{} {}  {}  {}  {}  {}  {}  {}'.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err, rxpwr))
                if ber > 0.1 or per > 90:
                    self.tables[23].cell(2 + chan * 1, 3).text = str(rxpwr)
                    self.pass_verdict(self.tables[23].rows[2 + chan * 1], rxpwr)
                    break
                if rxpwr == max(rxpwr_range):
                    self.tables[23].cell(2 + chan * 1, 3).text = str(rxpwr)
                    self.pass_verdict(self.tables[23].rows[2 + chan * 1], rxpwr)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))



    def EDR_Sensitivity(self):
        self.csp.mode_set(mode='EDR')
        self.csp.sig_btype(btype='EDR')
        self.csp.config_rx_level(rxpwr=-76)
        self.csp.config_rxq_br_search_packets(num=2000)
        self.csp.config_rxq_br_search_step(step=0.5)
        self.csp.config_connect_edr_packet_pattern(pattern='PRBS9')  # 设置EDR包 payload 数据类型
        self.csp.config_connect_edr_packet_ptype(ptype='E25P')
        for i in range(3):
            self.csp.config_power_control(para='DOWN')
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan*39)
            self.csp.ber_search_meas_state(state=0)
            while 1:
                state_res = self.csp.get_ber_search_meas_state()  # ber search 完成则退出循环
                # logdebug(state_res)
                if state_res == 'RDY':
                    loginfo('search complete')
                    break
                elif state_res == 'OFF':
                    self.cmd_br_connect()
                    self.csp.ber_search_meas_state(state=0)
                else:
                    logdebug('searching')
                time.sleep(5)

            res = self.csp.meas_bt_ber_search_res(cmd_type='FETCh')
            ber = eval(res[1])
            per = eval(res[2])
            NAK = eval(res[5])
            hec_err = eval(res[6])
            crc_err = eval(res[7])
            packet_type_err = eval(res[8])
            pay_err = eval(res[9])
            sens = eval(res[11])
            loginfo('{} {}  {}  {}  {}  {}  {}  {}'.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err, sens))
            self.tables[24].cell(3 + chan * 1, 3).text = str(sens)
            self.pass_verdict(self.tables[24].rows[3 + chan * 1], sens)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

        self.csp.config_connect_edr_packet_ptype(ptype='E35P')
        for i in range(3):
            self.csp.config_power_control(para='DOWN')
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan*39)
            self.csp.ber_search_meas_state(state=0)
            while 1:
                state_res = self.csp.get_ber_search_meas_state()  # ber search 完成则退出循环
                # logdebug(state_res)
                if state_res == 'RDY':
                    loginfo('search complete')
                    break
                elif state_res == 'OFF':
                    self.cmd_br_connect()
                    self.csp.ber_search_meas_state(state=0)
                else:
                    logdebug('searching')
                time.sleep(5)

            res = self.csp.meas_bt_ber_search_res(cmd_type='FETCh')
            ber = eval(res[1])
            per = eval(res[2])
            NAK = eval(res[5])
            hec_err = eval(res[6])
            crc_err = eval(res[7])
            packet_type_err = eval(res[8])
            pay_err = eval(res[9])
            sens = eval(res[11])
            loginfo('{} {}  {}  {}  {}  {}  {}  {}'.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err, sens))
            self.tables[24].cell(7 + chan * 1, 3).text = str(sens)
            self.pass_verdict(self.tables[24].rows[7 + chan * 1], sens)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))


    def EDR_BER_Floor_Performance(self):
        self.csp.mode_set(mode='EDR')
        self.csp.sig_btype(btype='EDR')
        self.csp.config_rx_level(rxpwr=-60)
        self.csp.config_rxq_br_search_packets(num=2000)
        self.csp.config_rxq_br_search_step(step=0.5)
        self.csp.config_connect_edr_packet_pattern(pattern='PRBS9')  # 设置EDR包 payload 数据类型
        self.csp.config_connect_edr_packet_ptype(ptype='E25P')
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan*39)
            time.sleep(1)
            self.csp.config_power_control(para='MAX')
            time.sleep(2)
            for i in range(3):
                self.csp.config_power_control(para='DOWN')
            while 1:
                self.csp.ber_meas_state(state=0)
                while 1:
                    state = self.csp.get_ber_meas_state()
                    if state == 'RDY':
                        loginfo('ber run complte')
                        break
                    time.sleep(2)
                res = self.csp.meas_bt_ber_res(cmd_type='FETCh')
                if eval(res[0]) == 0 or eval(res[0]) == 1:
                    break
                else:
                    if self.csp.get_bt_connect_state() == 'SBY':
                        self.cmd_br_connect()

            ber = eval(res[1])
            per = eval(res[2])
            NAK = eval(res[5])
            hec_err = eval(res[6])
            crc_err = eval(res[7])
            packet_type_err = eval(res[8])
            pay_err = eval(res[9])

            time.sleep(0.5)
            loginfo('{} {}  {}  {}  {}  {}  {}  {}'.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err, -60))

            self.tables[25].cell(3 + chan * 1, 3).text = str(ber)
            self.pass_verdict(self.tables[25].rows[3 + chan * 1], ber)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

        self.csp.config_connect_edr_packet_ptype(ptype='E35P')
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan*39)
            time.sleep(1)
            self.csp.config_power_control(para='MAX')
            time.sleep(2)
            for i in range(3):
                self.csp.config_power_control(para='DOWN')
            while 1:
                self.csp.ber_meas_state(state=0)
                while 1:
                    state = self.csp.get_ber_meas_state()
                    if state == 'RDY':
                        loginfo('ber run complte')
                        break
                    time.sleep(2)
                res = self.csp.meas_bt_ber_res(cmd_type='FETCh')
                if eval(res[0]) == 0 or eval(res[0]) == 1:
                    break
                else:
                    if self.csp.get_bt_connect_state() == 'SBY':
                        self.cmd_br_connect()

            ber = eval(res[1])
            per = eval(res[2])
            NAK = eval(res[5])
            hec_err = eval(res[6])
            crc_err = eval(res[7])
            packet_type_err = eval(res[8])
            pay_err = eval(res[9])

            time.sleep(0.5)
            loginfo('{} {}  {}  {}  {}  {}  {}  {}'.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err, -60))

            self.tables[25].cell(7 + chan * 1, 3).text = str(ber)
            self.pass_verdict(self.tables[25].rows[7 + chan * 1], ber)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def EDR_CI_Performance(self):
        self.csp.mode_set(mode='EDR')
        self.dict_pwr_inter_freqoffset = {
            0: range(-85, -40),
            1: range(-60, -30),
            2: range(-50, -10),
            3: range(-40, 0),
            4: range(-35, 0),
            -1: range(-60, -30),
            -2: range(-50, -10),
            -3: range(-40, 0),
            -4: range(-35, 0),
            }
        self.tester_inter = mxg.MXG(device_name='N5182B', num_of_machine=1)
        self.tester_inter.arb_waveform(rate='2_DH1_PN15')
        self.tester_inter.trriger_para_set(type='CONT', count=1000)
        self.tester_inter.para_set(freq=2402, power=-100)

        self.csp.sig_btype(btype='EDR')
        self.csp.config_rx_level(rxpwr=-60)
        self.csp.config_rxq_br_packets(num=1000)
        self.csp.config_connect_edr_packet_pattern(pattern='PRBS9')  # 设置EDR包 payload 数据类型
        self.csp.config_connect_edr_packet_ptype(ptype='E25P')

        for chan in range(1):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=3 + chan * 36)
            self.tester_inter.output_state(1, 1)
            self.tester_inter.arb_state(1)
            time.sleep(1)
            self.csp.config_power_control(para='MAX')
            time.sleep(2)
            for i in range(2):
                self.csp.config_power_control(para='DOWN')
            num = 0
            for freq_inter_offset in [0, 1, -1, 2, -2, 3, -3, 4, -4]:
                freq_inter = freq_inter_offset + 2405 + chan * 36
                if freq_inter_offset == 0:
                    self.tester_inter.arb_waveform(rate='2_DH1_PN15')
                else:
                    self.tester_inter.arb_waveform(rate='DH1_PN15')
                if -3 < freq_inter_offset < 3:
                    self.csp.config_rx_level(rxpwr=-60)
                    rxpwr_ref = -60
                else:
                    self.csp.config_rx_level(rxpwr=-67)
                    rxpwr_ref = -67
                pwr_list = self.dict_pwr_inter_freqoffset[freq_inter_offset]
                for power_inter in pwr_list:
                    self.tester_inter.para_set(freq=freq_inter, power=power_inter + self.cableloss)
                    time.sleep(0.2)
                    loginfo('freq_inter:    {}      pwr_inter:  {}'.format(freq_inter, power_inter))
                    timeout = 0
                    while 1:
                        if self.csp.get_bt_connect_state() == 'SBY':
                            self.cmd_br_connect()
                        self.csp.ber_meas_state(state=0)
                        while 1:
                            state = self.csp.get_ber_meas_state()
                            if state == 'RDY':
                                loginfo('ber run complte')
                                break
                            time.sleep(2)
                            timeout = timeout + 1
                            if timeout > 30:
                                loginfo('get result timeout')
                                break
                        if timeout > 30:
                            res = ['0', '1', '100', '100', '100', '100', '100', '100', '100', '100']
                            break
                        res = self.csp.meas_bt_ber_res(cmd_type='FETCh')
                        if eval(res[0]) == 0 or eval(res[0]) == 1:
                            break
                        else:
                            if self.csp.get_bt_connect_state() == 'SBY':
                                self.cmd_br_connect()

                    ber = eval(res[1])
                    per = eval(res[2])
                    NAK = eval(res[5])
                    hec_err = eval(res[6])
                    crc_err = eval(res[7])
                    packet_type_err = eval(res[8])
                    pay_err = eval(res[9])


                    if ber > 0.1 or per > 50 or NAK > 50:
                        loginfo('{} {}  {}  {}  {}  {}  {} '.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err))
                        self.tables[26].cell(5 + num + chan * 10, 3).text = str(rxpwr_ref-power_inter+1)
                        self.pass_verdict(self.tables[26].rows[5 + num + chan * 10], rxpwr_ref-power_inter+1)
                        break
                num = num + 1
            self.tester_inter.output_state(0, 0)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

        self.csp.config_connect_edr_packet_ptype(ptype='E35P')

        for chan in range(1):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=3 + chan * 36)
            self.tester_inter.output_state(1, 1)
            self.tester_inter.arb_state(1)
            time.sleep(1)
            self.csp.config_power_control(para='MAX')
            time.sleep(2)
            for i in range(2):
                self.csp.config_power_control(para='DOWN')
            num = 0
            for freq_inter_offset in [0, 1, -1, 2, -2, 3, -3, 4, -4]:
                freq_inter = freq_inter_offset + 2405 + chan * 36
                if freq_inter_offset == 0:
                    self.tester_inter.arb_waveform(rate='3_DH1_PN15')
                else:
                    self.tester_inter.arb_waveform(rate='DH1_PN15')
                if -3 < freq_inter_offset < 3:
                    self.csp.config_rx_level(rxpwr=-60)
                    rxpwr_ref = -60
                else:
                    self.csp.config_rx_level(rxpwr=-67)
                    rxpwr_ref = -67
                pwr_list = self.dict_pwr_inter_freqoffset[freq_inter_offset]
                for power_inter in pwr_list:
                    self.tester_inter.para_set(freq=freq_inter, power=power_inter + self.cableloss)
                    time.sleep(0.2)
                    loginfo('freq_inter:    {}      pwr_inter:  {}'.format(freq_inter, power_inter))
                    timeout = 0
                    while 1:
                        if self.csp.get_bt_connect_state() == 'SBY':
                            self.cmd_br_connect()
                        self.csp.ber_meas_state(state=0)
                        while 1:
                            state = self.csp.get_ber_meas_state()
                            if state == 'RDY':
                                loginfo('ber run complte')
                                break
                            time.sleep(2)
                            timeout = timeout + 1
                            if timeout > 30:
                                loginfo('get result timeout')
                                break
                        if timeout > 30:
                            res = ['0', '1', '100', '100', '100', '100', '100', '100', '100', '100']
                            break
                        res = self.csp.meas_bt_ber_res(cmd_type='FETCh')
                        if eval(res[0]) == 0 or eval(res[0]) == 1:
                            break
                        else:
                            if self.csp.get_bt_connect_state() == 'SBY':
                                self.cmd_br_connect()

                    ber = eval(res[1])
                    per = eval(res[2])
                    NAK = eval(res[5])
                    hec_err = eval(res[6])
                    crc_err = eval(res[7])
                    packet_type_err = eval(res[8])
                    pay_err = eval(res[9])


                    if ber > 0.1 or per > 50 or NAK > 50:
                        loginfo('{} {}  {}  {}  {}  {}  {} '.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err))
                        self.tables[27].cell(5 + num + chan * 10, 3).text = str(rxpwr_ref-power_inter+1)
                        self.pass_verdict(self.tables[27].rows[5 + num + chan * 10], rxpwr_ref-power_inter+1)
                        break
                num = num + 1
            self.tester_inter.output_state(0, 0)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def EDR_CI_debug(self,cableloss=5,ptype='E25P'):
        self.csp.mode_set(mode='EDR')
        self.dict_pwr_inter_freqoffset = {
            0: range(-85, -40),
            1: range(-60, -30),
            2: range(-35, -10),
            3: range(-35, 0),
            4: range(-35, 0),
            -1: range(-60, -30),
            -2: range(-50, -10),
            -3: range(-40, 0),
            -4: range(-35, 0),
            }
        self.tester_inter = mxg.MXG(device_name='N5182B', num_of_machine=1)
        self.tester_inter.arb_waveform(rate='2_DH1_PN15')
        self.tester_inter.trriger_para_set(type='CONT', count=1000)
        self.tester_inter.para_set(freq=2402, power=-100)
        self.tester_inter.output_state(0, 0)

        self.csp.sig_std('CLAS')
        self.config_classic_ber(rate='BR', chan=1, rx_level=-50, tx_power=10)
        self.csp.tx_measure_para(tout=10, repetition='SINGleshot', count=10, MOEXception='OFF')
        self.csp.sig_btype(btype='EDR')
        self.csp.config_rx_level(rxpwr=-60)
        self.csp.config_rxq_br_packets(num=1000)
        self.csp.config_connect_edr_packet_pattern(pattern='PRBS9')  # 设置EDR包 payload 数据类型
        self.csp.config_connect_edr_packet_ptype(ptype=ptype)


        ci_list=[]
        for chan in range(1):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=3 + chan * 36)
            self.tester_inter.output_state(1, 1)
            self.tester_inter.arb_state(1)
            time.sleep(1)
            self.csp.config_power_control(para='MAX')
            time.sleep(2)
            for i in range(2):
                self.csp.config_power_control(para='DOWN')
            num = 0
            for freq_inter_offset in [2, 3]:
                freq_inter = freq_inter_offset + 2405 + chan * 36
                if freq_inter_offset == 0:
                    self.tester_inter.arb_waveform(rate='2_DH1_PN15')
                else:
                    self.tester_inter.arb_waveform(rate='DH1_PN15')
                if -3 < freq_inter_offset < 3:
                    self.csp.config_rx_level(rxpwr=-60)
                    rxpwr_ref = -60
                else:
                    self.csp.config_rx_level(rxpwr=-67)
                    rxpwr_ref = -67
                pwr_list = self.dict_pwr_inter_freqoffset[freq_inter_offset]
                for power_inter in pwr_list:
                    self.tester_inter.para_set(freq=freq_inter, power=power_inter + cableloss)
                    time.sleep(0.2)
                    loginfo('freq_inter:    {}      pwr_inter:  {}'.format(freq_inter, power_inter))
                    timeout = 0
                    while 1:
                        if self.csp.get_bt_connect_state() == 'SBY':
                            self.cmd_br_connect()
                        self.csp.ber_meas_state(state=0)
                        while 1:
                            state = self.csp.get_ber_meas_state()
                            if state == 'RDY':
                                loginfo('ber run complte')
                                break
                            time.sleep(2)
                            timeout = timeout + 1
                            if timeout > 30:
                                loginfo('get result timeout')
                                break
                        if timeout > 30:
                            res = ['0', '1', '100', '100', '100', '100', '100', '100', '100', '100']
                            break
                        res = self.csp.meas_bt_ber_res(cmd_type='FETCh')
                        if eval(res[0]) == 0 or eval(res[0]) == 1:
                            break
                        else:
                            if self.csp.get_bt_connect_state() == 'SBY':
                                self.cmd_br_connect()

                    ber = eval(res[1])
                    per = eval(res[2])
                    NAK = eval(res[5])
                    hec_err = eval(res[6])
                    crc_err = eval(res[7])
                    packet_type_err = eval(res[8])
                    pay_err = eval(res[9])


                    if ber > 0.1 or per > 50 or NAK > 50:
                        loginfo('{} {}  {}  {}  {}  {}  {} '.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err))
                        loginfo('freq_offset: {} , CI: {}'.format(freq_inter_offset,rxpwr_ref-power_inter+1))
                        ci_list.append([freq_inter_offset,rxpwr_ref-power_inter+1])
                        # self.tables[26].cell(5 + num + chan * 10, 3).text = str(rxpwr_ref-power_inter+1)
                        # self.pass_verdict(self.tables[26].rows[5 + num + chan * 10], rxpwr_ref-power_inter+1)
                        break
                num = num + 1

            self.tester_inter.output_state(0, 0)

        loginfo(ci_list)


    def EDR_Maximum_Input_Level(self):
        self.csp.mode_set(mode='EDR')
        rxpwr_range = range(-20, 1)
        self.csp.sig_btype(btype='EDR')
        self.csp.config_rx_level(rxpwr=-40)
        self.csp.config_rxq_br_packets(num=2000)
        self.csp.config_connect_edr_packet_pattern(pattern='PRBS9')  # 设置EDR包 payload 数据类型
        self.csp.config_connect_edr_packet_ptype(ptype='E25P')
        for i in range(3):
            self.csp.config_power_control(para='DOWN')
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan * 39)
            self.csp.ber_meas_state(0)
            for rxpwr in rxpwr_range:
                self.csp.config_rx_level(rxpwr=rxpwr)
                while 1:
                    self.csp.ber_meas_state(state=0)
                    while 1:
                        state = self.csp.get_ber_meas_state()
                        if state == 'RDY':
                            loginfo('ber run complte')
                            break
                        time.sleep(2)
                    res = self.csp.meas_bt_ber_res(cmd_type='FETCh')
                    if eval(res[0]) == 0 or eval(res[0]) == 1:
                        break
                    else:
                        if self.csp.get_bt_connect_state() == 'SBY':
                            self.cmd_br_connect()

                ber = eval(res[1])
                per = eval(res[2])
                NAK = eval(res[5])
                hec_err = eval(res[6])
                crc_err = eval(res[7])
                packet_type_err = eval(res[8])
                pay_err = eval(res[9])

                time.sleep(0.5)
                loginfo('{} {}  {}  {}  {}  {}  {}  {}'.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err, rxpwr))
                if ber > 0.1 or per > 50 or NAK > 50:
                    self.tables[28].cell(3 + chan * 1, 3).text = str(rxpwr-1)
                    self.pass_verdict(self.tables[28].rows[3 + chan * 1], rxpwr-1)
                    break
                if rxpwr == max(rxpwr_range):
                    self.tables[28].cell(3 + chan * 1, 3).text = str(rxpwr)
                    self.pass_verdict(self.tables[28].rows[3 + chan * 1], rxpwr)
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

        self.csp.config_connect_edr_packet_ptype(ptype='E35P')
        time.sleep(1)
        self.csp.config_power_control(para='MAX')
        time.sleep(2)
        for i in range(3):
            self.csp.config_power_control(para='DOWN')
        for chan in range(3):
            self.csp.RF_Frequency_Settings_tx(mode='LOOP', ch_tx=chan * 39)
            self.csp.ber_meas_state(0)
            for rxpwr in rxpwr_range:
                self.csp.config_rx_level(rxpwr=rxpwr)
                while 1:
                    self.csp.ber_meas_state(state=0)
                    while 1:
                        state = self.csp.get_ber_meas_state()
                        if state == 'RDY':
                            loginfo('ber run complte')
                            break
                        time.sleep(2)
                    res = self.csp.meas_bt_ber_res(cmd_type='FETCh')
                    if eval(res[0]) == 0 or eval(res[0]) == 1:
                        break
                    else:
                        if self.csp.get_bt_connect_state() == 'SBY':
                            self.cmd_br_connect()

                ber = eval(res[1])
                per = eval(res[2])
                NAK = eval(res[5])
                hec_err = eval(res[6])
                crc_err = eval(res[7])
                packet_type_err = eval(res[8])
                pay_err = eval(res[9])

                time.sleep(0.5)
                loginfo('{} {}  {}  {}  {}  {}  {}  {}'.format(ber, per, NAK, hec_err, crc_err, packet_type_err, pay_err, rxpwr))
                if ber > 0.1 or per > 50 or NAK > 50:
                    self.tables[28].cell(7 + chan * 1, 3).text = str(rxpwr-1)
                    self.pass_verdict(self.tables[28].rows[7 + chan * 1], rxpwr-1)
                    break
                if rxpwr == max(rxpwr_range):
                    self.tables[28].cell(7 + chan * 1, 3).text = str(rxpwr)
                    self.pass_verdict(self.tables[28].rows[7 + chan * 1], rxpwr)
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))


    def LE1M_Output_power(self):
        self.csp.mode_set(mode='LE1M')
        self.csp.sig_le_phy(phy='LE1M')
        # self.csp.trigger_settings(source='power', threshold=-35, tout=1)
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=1+chan * 19)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan * 19)
            self.csp.config_connect_le_packet_pattern(rate='LE1M', pattern='PRBS9')
            res1 = self.csp.get_power_measure_res(rate='LE1M')

            logdebug('{}'.format(res1))

            nominal_pow = eval(res1[2])
            peak_pow = eval(res1[3])
            leakage_pow = eval(res1[4])
            peak_avg_pwr = eval(res1[5])

            self.tables[29].cell(3 + chan * 3, 3).text = str(nominal_pow)
            self.pass_verdict(self.tables[29].rows[3 + chan * 3], nominal_pow)
            self.tables[29].cell(4 + chan * 3, 3).text = str(peak_avg_pwr)
            self.pass_verdict(self.tables[29].rows[4 + chan * 3], peak_avg_pwr)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))


    def LE1M_Inband_emissions(self):
        self.csp.mode_set(mode='LE1M')
        self.csp.sig_le_phy(phy='LE1M')
        self.csp.acp_meas_settings(opt='CH40')
        self.csp.trigger_settings(source='power', threshold=-40, tout=1)
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=3 + chan * 17)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=2+chan*17)
            self.csp.config_connect_le_packet_pattern(rate='LE1M', pattern='PRBS9')
            res2 = self.csp.get_acp_res()
            time.sleep(1)
            logdebug('{}'.format(res2))

            acp_list = [eval(i) for i in res2[1:]]
            for i in range(len(acp_list)):
                self.tables[30].cell(3 + i + chan * 82, 3).text = str(acp_list[i])
                self.pass_verdict(self.tables[30].rows[3 + i + chan * 82], acp_list[i])

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def LE1M_Modulation_Characteristics(self):
        self.csp.mode_set(mode='LE1M')
        self.csp.sig_le_phy(phy='LE1M')
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=1+chan * 19)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan * 19)

            delta_f1_99, delta_f2_99, delta_f1_avg, delta_f1_min, delta_f1_max, delta_f2_avg, delta_f2_min, delta_f2_max, mod_ratio = -999, -999, -999, -999, -999, -999, -999, -999, -999
            self.csp.config_connect_le_packet_pattern(rate='LE1M', pattern='P44')
            res = self.csp.get_modulation_measure_res(rate='LE1M')
            logdebug('{}'.format(res))

            delta_f1_99 = eval(res[2]) / 1000.00
            freq_accuracy = eval(res[3]) / 1000.00
            freq_drift = eval(res[4]) / 1000.00
            drift_rate = eval(res[5]) / 1000.00
            delta_f1_avg = eval(res[6]) / 1000.00
            delta_f1_min = eval(res[7]) / 1000.00
            delta_f1_max = eval(res[8]) / 1000.00

            time.sleep(0.5)

            self.csp.config_connect_le_packet_pattern(rate='LE1M', pattern='P11')
            res = self.csp.get_modulation_measure_res(rate='LE1M')
            logdebug('{}'.format(res))

            delta_f2_99 = eval(res[2]) / 1000.00
            freq_accuracy = eval(res[3]) / 1000.00
            freq_drift = eval(res[4]) / 1000.00
            drift_rate = eval(res[5]) / 1000.00
            delta_f2_avg = eval(res[9]) / 1000.00
            delta_f2_min = eval(res[10]) / 1000.00
            delta_f2_max = eval(res[11]) / 1000.00
            mod_ratio = delta_f2_avg / delta_f1_avg
            freq_offset = eval(res[14]) / 1000.00
            init_drift = eval(res[15]) / 1000.00

            self.tables[31].cell(3 + chan * 4, 3).text = str(delta_f1_avg)
            self.pass_verdict(self.tables[31].rows[3 + chan * 4], delta_f1_avg)

            self.tables[31].cell(4 + chan * 4, 3).text = str(delta_f2_99)
            self.pass_verdict(self.tables[31].rows[4 + chan * 4], delta_f2_99)

            self.tables[31].cell(5 + chan * 4, 3).text = str(mod_ratio)
            self.pass_verdict(self.tables[31].rows[5 + chan * 4], mod_ratio)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def LE1M_Carrier_frequency_offset_and_drift(self):
        self.csp.mode_set(mode='LE1M')
        self.csp.sig_le_phy(phy='LE1M')
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=1+chan * 19)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan * 19)

            delta_f1_99, delta_f2_99, delta_f1_avg, delta_f1_min, delta_f1_max, delta_f2_avg, delta_f2_min, delta_f2_max, mod_ratio = -999, -999, -999, -999, -999, -999, -999, -999, -999

            self.csp.config_connect_le_packet_pattern(rate='LE1M', pattern='P11')
            res = self.csp.get_modulation_measure_res(rate='LE1M')
            logdebug('{}'.format(res))

            delta_f2_99 = eval(res[2]) / 1000.00
            freq_accuracy = eval(res[3]) / 1000.00
            freq_drift = eval(res[4]) / 1000.00
            drift_rate = eval(res[5]) / 1000.00
            delta_f2_avg = eval(res[9]) / 1000.00
            delta_f2_min = eval(res[10]) / 1000.00
            delta_f2_max = eval(res[11]) / 1000.00
            mod_ratio = delta_f2_avg / delta_f1_avg
            freq_offset = eval(res[14]) / 1000.00
            init_drift = eval(res[15]) / 1000.00

            self.tables[32].cell(3 + chan * 5, 3).text = str(freq_accuracy)
            self.pass_verdict(self.tables[32].rows[3 + chan * 5], freq_accuracy)

            self.tables[32].cell(4 + chan * 5, 3).text = str(freq_drift)
            self.pass_verdict(self.tables[32].rows[4 + chan * 5], freq_drift)

            self.tables[32].cell(5 + chan * 5, 3).text = str(drift_rate)
            self.pass_verdict(self.tables[32].rows[5 + chan * 5], drift_rate)

            self.tables[32].cell(6 + chan * 5, 3).text = str(init_drift)
            self.pass_verdict(self.tables[32].rows[6 + chan * 5], init_drift)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def LE1M_Receiver_sensitivity(self):
        self.csp.mode_set(mode='LE1M')
        self.csp.sig_le_phy(phy='LE1M')
        self.csp.config_rxq_le_packets(rate='LE1M', num=2000)
        self.csp.config_connect_le_packet_pattern(rate='LE1M', pattern='PRBS9')
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=1+chan * 19)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan * 19)

            for i in range(41):
                rxpwr = -80 - i/2.0
                self.csp.config_rx_level(rxpwr=rxpwr)

                per_list = []
                for loop in range(3):
                    res = self.csp.meas_le_per_res(rate='LE1M')
                    time.sleep(0.5)
                    per = eval(res[1])
                    per_list.append(per)

                per = min(per_list)
                if per > 30.8:
                    self.tables[33].cell(2 + chan * 1, 3).text = str(rxpwr+0.5)
                    self.pass_verdict(self.tables[33].rows[2 + chan * 1], rxpwr+0.5)
                    break
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def LE1M_CI_Performance(self):
        self.csp.mode_set(mode='LE1M')
        self.csp.sig_le_phy(phy='LE1M')
        self.csp.config_rxq_le_packets(rate='LE1M', num=2000)
        self.csp.config_connect_le_packet_pattern(rate='LE1M', pattern='PRBS9')
        self.csp.config_rx_level(rxpwr=-67)
        self.dict_pwr_inter_freqoffset = {
            0: range(-90, -60),
            1: range(-85, -40),
            2: range(-52, -20),
            3: range(-41, -10),
            4: range(-41, -10),
            -1: range(-85, -40),
            -2: range(-52, -20),
            -3: range(-41, -10),
            -4: range(-41, -10)
            }
        self.tester_inter = mxg.MXG(device_name='N5182B', num_of_machine=1)
        self.tester_inter.arb_waveform(rate='LE_1M_PN15')
        self.tester_inter.trriger_para_set(type='CONT', count=2000)
        self.tester_inter.output_state(1, 1)
        self.tester_inter.arb_state(1)

        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=3 + chan * 17)
                cf = 3 + chan * 17
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=2 + chan * 17)
                cf = 2 + chan * 17
            time.sleep(1)
            num = 0
            for freq_inter_offset in[0,1,-1,2,-2,3,-3,4,-4]:
                freq_inter = freq_inter_offset+2402+cf*2
                pwr_list = self.dict_pwr_inter_freqoffset[freq_inter_offset]
                for power_inter in pwr_list:
                    self.tester_inter.para_set(freq=freq_inter, power=power_inter + self.cableloss)
                    time.sleep(0.2)
                    loginfo('freq_inter:    {}      pwr_inter:  {}'.format(freq_inter,power_inter))
                    per_list = []
                    for loop in range(3):
                        res = self.csp.meas_le_per_res(rate='LE1M')
                        time.sleep(0.2)
                        per = eval(res[1])
                        per_list.append(per)

                    per = min(per_list)
                    loginfo('{}     {}      {}      {}      {}'.format(chan, 'LE1M', freq_inter, power_inter,  per))
                    if per > 30.8:
                        self.tables[34].cell(4 + num + chan * 10, 3).text = str(-67-power_inter+1)
                        self.pass_verdict(self.tables[34].rows[4 + num + chan * 10], -67-power_inter+1)
                        break
                num = num + 1
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))
        self.tester_inter.output_state(0, 0)

    def LE1M_Blocking_Performance(self):
        self.csp.mode_set(mode='LE1M')
        self.csp.sig_le_phy(phy='LE1M')
        self.csp.config_rxq_le_packets(rate='LE1M', num=2000)
        self.csp.config_connect_le_packet_pattern(rate='LE1M', pattern='PRBS9')
        self.csp.config_rx_level(rxpwr=-67)
        self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=19)

        self.tester_inter = mxg.MXG(device_name='N5182B',num_of_machine=1)
        self.tester_inter.para_set(30,-30+self.cableloss)
        self.tester_inter.output_state(1, 0)

        freq_range= range(30,2001,10)+range(2003,2400,1)+range(2484,2998,1)+range(3000,6000,25)
        # freq_range = freq_range_le
        freq_blocking_2 = 'NA'
        freq_blocking_1 = 'NA'
        freq_inter_list = []
        freq_inter_list1 = []
        inter_rxpwr = -10
        # for rxpwr in inter_rxpwr:
        count = 0
        count1 = 0
        for freq_inter in freq_range:
            freq_blocking_2 = 'NA'
            freq_blocking_1 = 'NA'
            if freq_inter < 2001 or freq_inter > 2999:
                pwr =  inter_rxpwr + 5
                # self.tester_inter.para_set(freq=freq_inter, power=rxpwr + inter_loss)
            else:
                pwr = inter_rxpwr
            #     # self.tester_inter.para_set(freq=freq_inter, power=rxpwr-5 + inter_loss)
            self.tester_inter.para_set(freq=freq_inter, power=pwr+self.cableloss)
            loginfo('freq_inter:    {}'.format(freq_inter))
            self.csp.config_rx_level(rxpwr=-67)
            self.csp.per_meas_state(0)
            time.sleep(1)
            res = self.csp.meas_le_per_res(rate='LE1M')
            per = eval(res[1])
            if per > 30.8:
                loginfo('freq_blocking_1:   {}      level:  {}      per:    {}'.format(freq_inter,pwr,per))
                # freq_blocking_1 = freq_inter
                # if freq_inter > 1999 or freq_inter < 3001:
                #     # pwr = rxpwr - 5 + inter_loss
                #     rxpwr = rxpwr - 5
                freq_inter_list.append(freq_inter)
                # fw1.write_data([rate,rxpwr,freq_blocking_1,freq_blocking_2,per])
                count = count + 1

        for freq_inter_check in freq_inter_list:
            if freq_inter_check < 2001 or freq_inter_check > 2999:
                pwr =  inter_rxpwr + 5
                # self.tester_inter.para_set(freq=freq_inter, power=rxpwr + inter_loss)
            else:
                pwr = inter_rxpwr
            self.tester_inter.para_set(freq=freq_inter_check, power=pwr+self.cableloss)
            loginfo('freq_inter:    {}'.format(freq_inter_check))
            self.csp.per_meas_state(0)
            time.sleep(1)
            res = self.csp.meas_le_per_res(rate='LE1M')
            per = eval(res[1])
            if per > 30.8:
                loginfo('freq_blocking_1:   {}      level:  {}      per:    {}'.format(freq_inter_check, pwr, per))
                freq_inter_list1.append(freq_inter_check)
        while 1:
            loginfo('***********************************************************')
            loginfo(freq_inter_list1)
            if len(freq_inter_list1) >= 10:
                inter_rxpwr = inter_rxpwr -1
            elif len(freq_inter_list1) <= 5:
                inter_rxpwr = inter_rxpwr + 1
            else:
                self.tables[35].cell(3, 3).text = str(inter_rxpwr+5)
                self.pass_verdict(self.tables[35].rows[3], inter_rxpwr+5)
                self.tables[35].cell(4, 3).text = str(inter_rxpwr)
                self.pass_verdict(self.tables[35].rows[4], inter_rxpwr)
                self.tables[35].cell(5, 3).text = str(inter_rxpwr)
                self.pass_verdict(self.tables[35].rows[5], inter_rxpwr)
                self.tables[35].cell(6, 3).text = str(inter_rxpwr+5)
                self.pass_verdict(self.tables[35].rows[6], inter_rxpwr+5)

                break
            freq_range = freq_inter_list1
            freq_inter_list1 = []
            for freq_inter in freq_range:
                if freq_inter_check < 2001 or freq_inter_check > 2999:
                    pwr = inter_rxpwr + 5
                    # self.tester_inter.para_set(freq=freq_inter, power=rxpwr + inter_loss)
                else:
                    pwr = inter_rxpwr
                self.tester_inter.para_set(freq=freq_inter, power=pwr+self.cableloss)
                self.csp.per_meas_state(0)
                time.sleep(1)
                res = self.csp.meas_le_per_res(rate='LE1M')
                per = eval(res[1])
                if per > 30.8:
                    loginfo('freq_blocking_1:   {}      level:  {}      per:    {}'.format(freq_inter, pwr, per))
                    freq_inter_list1.append(freq_inter)

        self.tester_inter.output_state(0, 0)
        self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def LE1M_Maximum_input_signal_level(self):
        rxpwr_range = range(-20,1)
        self.csp.mode_set(mode='LE1M')
        self.csp.sig_le_phy(phy='LE1M')
        self.csp.config_rxq_le_packets(rate='LE1M', num=2000)
        self.csp.config_connect_le_packet_pattern(rate='LE1M', pattern='PRBS9')
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=1+chan * 19)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan * 19)

            for rxpwr in rxpwr_range:
                self.csp.config_rx_level(rxpwr=rxpwr)

                per_list = []
                for loop in range(3):
                    res = self.csp.meas_le_per_res(rate='LE1M')
                    time.sleep(0.5)
                    per = eval(res[1])
                    per_list.append(per)

                per = min(per_list)
                if per > 30.8:
                    self.tables[37].cell(2 + chan * 1, 3).text = str(rxpwr-1)
                    self.pass_verdict(self.tables[37].rows[2 + chan * 1], rxpwr-1)
                    break
                if rxpwr == max(rxpwr_range):
                    self.tables[37].cell(2 + chan * 1, 3).text = str(rxpwr)
                    self.pass_verdict(self.tables[37].rows[2 + chan * 1], rxpwr)
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))


    def LE1M_PER_Report_Integrity(self):
        self.csp.mode_set(mode='LE1M')
        self.csp.sig_le_phy(phy='LE1M')
        self.csp.config_rxq_le_packets(rate='LE1M', num=2000)
        self.csp.config_connect_le_packet_pattern(rate='LE1M', pattern='PRBS9')
        self.csp.config_rxq_le_integrity(rate='LE1M',en='ON')
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=1+chan * 19)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan * 19)


            self.csp.config_rx_level(rxpwr=-30)

            per_list = []
            for loop in range(3):
                res = self.csp.meas_le_per_res(rate='LE1M')
                time.sleep(0.5)
                per = eval(res[1])
                per_list.append(per)

            per = min(per_list)

            self.tables[38].cell(3 + chan * 1, 3).text = str(per)
            self.pass_verdict(self.tables[38].rows[3 + chan * 1], per)
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

        self.csp.config_rxq_le_integrity(rate='LE1M', en='OFF')


    def LE2M_Output_power(self):
        self.csp.mode_set(mode='LE2M')
        self.csp.sig_le_phy(phy='LE2M')
        self.csp.trigger_settings(source='power', threshold=-35, tout=1)
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=1+chan * 19)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan * 19)
            self.csp.config_connect_le_packet_pattern(rate='LE2M', pattern='PRBS9')
            res1 = self.csp.get_power_measure_res(rate='LE2M')

            logdebug('{}'.format(res1))

            nominal_pow = eval(res1[2])
            peak_pow = eval(res1[3])
            leakage_pow = eval(res1[4])
            peak_avg_pwr = eval(res1[5])

            self.tables[39].cell(3 + chan * 3, 3).text = str(nominal_pow)
            self.pass_verdict(self.tables[39].rows[3 + chan * 3], nominal_pow)
            self.tables[39].cell(4 + chan * 3, 3).text = str(peak_avg_pwr)
            self.pass_verdict(self.tables[39].rows[4 + chan * 3], peak_avg_pwr)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))
        self.csp.trigger_settings(source='power', threshold=-35, tout=1)

    def LE2M_Inband_emissions(self):
        self.csp.mode_set(mode='LE2M')
        self.csp.sig_le_phy(phy='LE2M')
        self.csp.acp_meas_settings(opt='CH40')
        self.csp.trigger_settings(source='power', threshold=-30, tout=1)
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=3 + chan * 17)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=2+chan*17)
            self.csp.config_connect_le_packet_pattern(rate='LE2M', pattern='PRBS9')
            res2 = self.csp.get_acp_res()
            time.sleep(1)
            logdebug('{}'.format(res2))

            acp_list = [eval(i) for i in res2[1:]]
            for i in range(len(acp_list)):
                self.tables[40].cell(3 + i + chan * 82, 3).text = str(acp_list[i])
                self.pass_verdict(self.tables[40].rows[3 + i + chan * 82], acp_list[i])

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))
        self.csp.trigger_settings(source='power', threshold=-35, tout=1)

    def LE2M_Modulation_Characteristics(self):
        self.csp.mode_set(mode='LE2M')
        self.csp.sig_le_phy(phy='LE2M')
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=1+chan * 19)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan * 19)

            delta_f1_99, delta_f2_99, delta_f1_avg, delta_f1_min, delta_f1_max, delta_f2_avg, delta_f2_min, delta_f2_max, mod_ratio = -999, -999, -999, -999, -999, -999, -999, -999, -999
            self.csp.config_connect_le_packet_pattern(rate='LE2M', pattern='P44')
            res = self.csp.get_modulation_measure_res(rate='LE2M')
            logdebug('{}'.format(res))

            delta_f1_99 = eval(res[2]) / 1000.00
            freq_accuracy = eval(res[3]) / 1000.00
            freq_drift = eval(res[4]) / 1000.00
            drift_rate = eval(res[5]) / 1000.00
            delta_f1_avg = eval(res[6]) / 1000.00
            delta_f1_min = eval(res[7]) / 1000.00
            delta_f1_max = eval(res[8]) / 1000.00

            time.sleep(0.5)

            self.csp.config_connect_le_packet_pattern(rate='LE2M', pattern='P11')
            res = self.csp.get_modulation_measure_res(rate='LE2M')
            logdebug('{}'.format(res))

            delta_f2_99 = eval(res[2]) / 1000.00
            freq_accuracy = eval(res[3]) / 1000.00
            freq_drift = eval(res[4]) / 1000.00
            drift_rate = eval(res[5]) / 1000.00
            delta_f2_avg = eval(res[9]) / 1000.00
            delta_f2_min = eval(res[10]) / 1000.00
            delta_f2_max = eval(res[11]) / 1000.00
            mod_ratio = delta_f2_avg / delta_f1_avg
            freq_offset = eval(res[14]) / 1000.00
            init_drift = eval(res[15]) / 1000.00

            self.tables[41].cell(3 + chan * 4, 3).text = str(delta_f1_avg)
            self.pass_verdict(self.tables[41].rows[3 + chan * 4], delta_f1_avg)

            self.tables[41].cell(4 + chan * 4, 3).text = str(delta_f2_99)
            self.pass_verdict(self.tables[41].rows[4 + chan * 4], delta_f2_99)

            self.tables[41].cell(5 + chan * 4, 3).text = str(mod_ratio)
            self.pass_verdict(self.tables[41].rows[5 + chan * 4], mod_ratio)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def LE2M_Carrier_frequency_offset_and_drift(self):
        self.csp.mode_set(mode='LE2M')
        self.csp.sig_le_phy(phy='LE2M')
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=1+chan * 19)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan * 19)

            delta_f1_99, delta_f2_99, delta_f1_avg, delta_f1_min, delta_f1_max, delta_f2_avg, delta_f2_min, delta_f2_max, mod_ratio = -999, -999, -999, -999, -999, -999, -999, -999, -999

            self.csp.config_connect_le_packet_pattern(rate='LE2M', pattern='P11')
            res = self.csp.get_modulation_measure_res(rate='LE2M')
            logdebug('{}'.format(res))

            delta_f2_99 = eval(res[2]) / 1000.00
            freq_accuracy = eval(res[3]) / 1000.00
            freq_drift = eval(res[4]) / 1000.00
            drift_rate = eval(res[5]) / 1000.00
            delta_f2_avg = eval(res[9]) / 1000.00
            delta_f2_min = eval(res[10]) / 1000.00
            delta_f2_max = eval(res[11]) / 1000.00
            mod_ratio = delta_f2_avg / delta_f1_avg
            freq_offset = eval(res[14]) / 1000.00
            init_drift = eval(res[15]) / 1000.00

            self.tables[42].cell(3 + chan * 5, 3).text = str(freq_accuracy)
            self.pass_verdict(self.tables[42].rows[3 + chan * 5], freq_accuracy)

            self.tables[42].cell(4 + chan * 5, 3).text = str(freq_drift)
            self.pass_verdict(self.tables[42].rows[4 + chan * 5], freq_drift)

            self.tables[42].cell(5 + chan * 5, 3).text = str(drift_rate)
            self.pass_verdict(self.tables[42].rows[5 + chan * 5], drift_rate)

            self.tables[42].cell(6 + chan * 5, 3).text = str(init_drift)
            self.pass_verdict(self.tables[42].rows[6 + chan * 5], init_drift)

            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def LE2M_Receiver_sensitivity(self):
        self.csp.mode_set(mode='LE2M')
        self.csp.sig_le_phy(phy='LE2M')
        self.csp.config_rxq_le_packets(rate='LE2M', num=2000)
        self.csp.config_connect_le_packet_pattern(rate='LE2M', pattern='PRBS9')
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=1+chan * 19)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan * 19)

            for i in range(41):
                rxpwr = -80 - i / 2.0
                self.csp.config_rx_level(rxpwr=rxpwr)

                per_list = []
                for loop in range(3):
                    res = self.csp.meas_le_per_res(rate='LE2M')
                    time.sleep(0.5)
                    per = eval(res[1])
                    per_list.append(per)

                per = min(per_list)
                if per > 30.8:
                    self.tables[43].cell(2 + chan * 1, 3).text = str(rxpwr+0.5)
                    self.pass_verdict(self.tables[43].rows[2 + chan * 1], rxpwr+0.5)
                    break
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def LE2M_CI_Performance(self):
        self.csp.mode_set(mode='LE2M')
        self.csp.sig_le_phy(phy='LE2M')
        self.csp.config_rxq_le_packets(rate='LE2M', num=2000)
        self.csp.config_connect_le_packet_pattern(rate='LE2M', pattern='PRBS9')
        self.csp.config_rx_level(rxpwr=-67)
        self.dict_pwr_inter_freqoffset = {
            0: range(-90, -60),
            1: range(-85, -40),
            2: range(-52, -20),
            3: range(-41, -10),
            4: range(-41, -10),
            -1: range(-85, -40),
            -2: range(-52, -20),
            -3: range(-41, -10),
            -4: range(-41, -10)
            }
        self.tester_inter = mxg.MXG(device_name='N5182B', num_of_machine=1)
        self.tester_inter.arb_waveform(rate='LE_2M_PN15')
        self.tester_inter.trriger_para_set(type='CONT', count=2000)
        self.tester_inter.output_state(1, 1)
        self.tester_inter.arb_state(1)

        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=3 + chan * 17)
                cf = 3 + chan * 17
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=2 + chan * 17)
                cf = 2 + chan * 17

            time.sleep(1)
            num = 0
            for freq_inter_offset in[0,1,-1,2,-2,3,-3,4,-4]:
                freq_inter = freq_inter_offset*2 + 2402 + cf*2
                pwr_list = self.dict_pwr_inter_freqoffset[freq_inter_offset]
                for power_inter in pwr_list:
                    self.tester_inter.para_set(freq=freq_inter, power=power_inter + self.cableloss)
                    time.sleep(0.2)
                    loginfo('freq_inter:    {}      pwr_inter:  {}'.format(freq_inter,power_inter))
                    per_list = []
                    for loop in range(3):
                        res = self.csp.meas_le_per_res(rate='LE2M')
                        time.sleep(0.2)
                        per = eval(res[1])
                        per_list.append(per)

                    per = min(per_list)
                    loginfo('{}     {}      {}      {}      {}'.format(chan, 'LE2M', freq_inter, power_inter,  per))
                    if per > 30.8:
                        self.tables[44].cell(4 + num + chan * 10, 3).text = str(-67-power_inter+1)
                        self.pass_verdict(self.tables[44].rows[4 + num + chan * 10], -67-power_inter+1)
                        break
                num = num + 1
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))
        self.tester_inter.output_state(0, 0)

    def LE2M_Blocking_Performance(self):
        self.csp.mode_set(mode='LE2M')
        self.csp.sig_le_phy(phy='LE2M')
        self.csp.config_rxq_le_packets(rate='LE2M', num=2000)
        self.csp.config_connect_le_packet_pattern(rate='LE2M', pattern='PRBS9')
        self.csp.config_rx_level(rxpwr=-67)
        self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=19)

        self.tester_inter = mxg.MXG(device_name='N5182B',num_of_machine=1)
        self.tester_inter.para_set(30,-30+self.cableloss)
        self.tester_inter.output_state(1, 0)

        freq_range= range(30,2001,10)+range(2003,2400,1)+range(2484,2998,1)+range(3000,6000,25)
        # freq_range = freq_range_le
        freq_blocking_2 = 'NA'
        freq_blocking_1 = 'NA'
        freq_inter_list = []
        freq_inter_list1 = []
        inter_rxpwr = -10
        # for rxpwr in inter_rxpwr:
        count = 0
        count1 = 0
        for freq_inter in freq_range:
            freq_blocking_2 = 'NA'
            freq_blocking_1 = 'NA'
            if freq_inter < 2001 or freq_inter > 2999:
                pwr =  inter_rxpwr + 5
                # self.tester_inter.para_set(freq=freq_inter, power=rxpwr + inter_loss)
            else:
                pwr = inter_rxpwr
            #     # self.tester_inter.para_set(freq=freq_inter, power=rxpwr-5 + inter_loss)
            self.tester_inter.para_set(freq=freq_inter, power=pwr+self.cableloss)
            loginfo('freq_inter:    {}'.format(freq_inter))
            self.csp.config_rx_level(rxpwr=-67)
            self.csp.per_meas_state(0)
            time.sleep(1)
            res = self.csp.meas_le_per_res(rate='LE2M')
            per = eval(res[1])
            if per > 30.8:
                loginfo('freq_blocking_1:   {}      level:  {}      per:    {}'.format(freq_inter,pwr,per))
                # freq_blocking_1 = freq_inter
                # if freq_inter > 1999 or freq_inter < 3001:
                #     # pwr = rxpwr - 5 + inter_loss
                #     rxpwr = rxpwr - 5
                freq_inter_list.append(freq_inter)
                # fw1.write_data([rate,rxpwr,freq_blocking_1,freq_blocking_2,per])
                count = count + 1

        for freq_inter_check in freq_inter_list:
            if freq_inter_check < 2001 or freq_inter_check > 2999:
                pwr =  inter_rxpwr + 5
                # self.tester_inter.para_set(freq=freq_inter, power=rxpwr + inter_loss)
            else:
                pwr = inter_rxpwr
            self.tester_inter.para_set(freq=freq_inter_check, power=pwr+self.cableloss)
            loginfo('freq_inter:    {}'.format(freq_inter_check))
            self.csp.per_meas_state(0)
            time.sleep(1)
            res = self.csp.meas_le_per_res(rate='LE2M')
            per = eval(res[1])
            if per > 30.8:
                loginfo('freq_blocking_1:   {}      level:  {}      per:    {}'.format(freq_inter_check, pwr, per))
                freq_inter_list1.append(freq_inter_check)
        while 1:
            loginfo('***********************************************************')
            loginfo(freq_inter_list1)
            if len(freq_inter_list1) >= 10:
                inter_rxpwr = inter_rxpwr -1
            elif len(freq_inter_list1) <= 5:
                inter_rxpwr = inter_rxpwr + 1
            else:
                self.tables[45].cell(3, 3).text = str(inter_rxpwr+5)
                self.pass_verdict(self.tables[45].rows[3], inter_rxpwr+5)
                self.tables[45].cell(4, 3).text = str(inter_rxpwr)
                self.pass_verdict(self.tables[45].rows[4], inter_rxpwr)
                self.tables[45].cell(5, 3).text = str(inter_rxpwr)
                self.pass_verdict(self.tables[45].rows[5], inter_rxpwr)
                self.tables[45].cell(6, 3).text = str(inter_rxpwr+5)
                self.pass_verdict(self.tables[45].rows[6], inter_rxpwr+5)

                break
            freq_range = freq_inter_list1
            freq_inter_list1 = []
            for freq_inter in freq_range:
                if freq_inter_check < 2001 or freq_inter_check > 2999:
                    pwr = inter_rxpwr + 5
                    # self.tester_inter.para_set(freq=freq_inter, power=rxpwr + inter_loss)
                else:
                    pwr = inter_rxpwr
                self.tester_inter.para_set(freq=freq_inter, power=pwr+self.cableloss)
                self.csp.per_meas_state(0)
                time.sleep(1)
                res = self.csp.meas_le_per_res(rate='LE2M')
                per = eval(res[1])
                if per > 30.8:
                    loginfo('freq_blocking_1:   {}      level:  {}      per:    {}'.format(freq_inter, pwr, per))
                    freq_inter_list1.append(freq_inter)

        self.tester_inter.output_state(0, 0)
        self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

    def LE2M_Maximum_input_signal_level(self):
        rxpwr_range = range(-20,1)
        self.csp.mode_set(mode='LE2M')
        self.csp.sig_le_phy(phy='LE2M')
        self.csp.config_rxq_le_packets(rate='LE2M', num=2000)
        self.csp.config_connect_le_packet_pattern(rate='LE2M', pattern='PRBS9')
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=1+chan * 19)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan * 19)

            for rxpwr in rxpwr_range:
                self.csp.config_rx_level(rxpwr=rxpwr)

                per_list = []
                for loop in range(3):
                    res = self.csp.meas_le_per_res(rate='LE2M')
                    time.sleep(0.5)
                    per = eval(res[1])
                    per_list.append(per)

                per = min(per_list)
                if per > 30.8:
                    self.tables[47].cell(2 + chan * 1, 3).text = str(rxpwr-1)
                    self.pass_verdict(self.tables[47].rows[2 + chan * 1], rxpwr-1)
                    break
                if rxpwr == max(rxpwr_range):
                    self.tables[37].cell(2 + chan * 1, 3).text = str(rxpwr)
                    self.pass_verdict(self.tables[37].rows[2 + chan * 1], rxpwr)
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))


    def LE2M_PER_Report_Integrity(self):
        self.csp.mode_set(mode='LE2M')
        self.csp.sig_le_phy(phy='LE2M')
        self.csp.config_rxq_le_packets(rate='LE2M', num=2000)
        self.csp.config_connect_le_packet_pattern(rate='LE2M', pattern='PRBS9')
        self.csp.config_rxq_le_integrity(rate='LE2M',en='ON')
        for chan in range(3):
            if chan == 2:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=1+chan * 19)
            else:
                self.csp.RF_Frequency_Settings_tx(mode='DTM', ch_tx=chan * 19)


            self.csp.config_rx_level(rxpwr=-30)

            per_list = []
            for loop in range(3):
                res = self.csp.meas_le_per_res(rate='LE2M')
                time.sleep(0.5)
                per = eval(res[1])
                per_list.append(per)

            per = min(per_list)

            self.tables[48].cell(3 + chan * 1, 3).text = str(per)
            self.pass_verdict(self.tables[48].rows[3 + chan * 1], per)
            self.document.save(self.fname + 'BQB_RF_PHY_test_report_{}_{}.docx'.format(self.dut_name, self.temp))

        self.csp.config_rxq_le_integrity(rate='LE2M', en='OFF')
